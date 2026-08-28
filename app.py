# ============================================================
#  School Bus Tracker — D205 School District
#  Flask + SQLAlchemy + APScheduler
# ============================================================

from flask import (Flask, render_template, request, redirect, url_for,
                   jsonify, flash, make_response, send_file, session, g, abort,
                   has_request_context)
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.middleware.proxy_fix import ProxyFix
from datetime import datetime, date, timedelta, timezone
from contextlib import contextmanager
from functools import wraps
from types import SimpleNamespace
from sqlalchemy import and_, case, func, or_
from sqlalchemy.orm import joinedload, selectinload
from sqlalchemy.exc import IntegrityError
from urllib.parse import urlsplit
import click
import hashlib, hmac, os, json, csv, io, pytz, re, time, secrets, html, math, tempfile, unicodedata
from powerschool_import import (
    DEFAULT_MAPPING_V1, NORMALIZER_REVISION, TRANSPORTATION_V2_CONTRACT,
    ImportValidationError,
    build_normalized_plan, normalize_route,
    canonical_plan_hash, safe_csv_cell,
)
from email_service import EmailTransportError, SMTPSettings, send_email, verify_connection


def _utcnow():
    """Return naive UTC for compatibility with the application's existing DB columns."""
    return datetime.now(timezone.utc).replace(tzinfo=None)

try:
    from apscheduler.schedulers.background import BackgroundScheduler
    SCHEDULER_AVAILABLE = True
except ImportError:
    SCHEDULER_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from twilio.rest import Client as TwilioClient
    from twilio.base.exceptions import TwilioRestException
    TWILIO_AVAILABLE = True
except ImportError:
    TWILIO_AVAILABLE = False


# ── APP SETUP ────────────────────────────────────────────────────────────────

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INSTANCE_DIR = os.path.realpath(os.environ.get('INSTANCE_DIR', os.path.join(BASE_DIR, 'instance')))

# Load instance-level config (.env written by the install wizard or admin)
try:
    from dotenv import load_dotenv
    load_dotenv(os.path.join(INSTANCE_DIR, '.env'), override=False)
except ImportError:
    pass
app = Flask(__name__)

def _env_int(name, default, minimum=0, maximum=None):
    """Read a bounded integer setting without accepting malformed configuration."""
    raw = os.environ.get(name, str(default)).strip()
    try:
        value = int(raw)
    except ValueError as exc:
        raise RuntimeError(f'{name} must be an integer') from exc
    if value < minimum or (maximum is not None and value > maximum):
        raise RuntimeError(f'{name} is outside the allowed range')
    return value


# ProxyFix is disabled unless deployment explicitly declares trusted proxy hops.
# This prevents arbitrary X-Forwarded-* headers from becoming security inputs.
app.wsgi_app = ProxyFix(
    app.wsgi_app,
    x_for=_env_int('TRUSTED_PROXY_X_FOR', 0, 0, 5),
    x_proto=_env_int('TRUSTED_PROXY_X_PROTO', 0, 0, 5),
    x_host=_env_int('TRUSTED_PROXY_X_HOST', 0, 0, 5),
)
_secret = os.environ.get('SECRET_KEY', '')
_secret_generated = False
if not _secret or _secret == 'changeme-set-in-env':
    if os.environ.get('FLASK_ENV') == 'production':
        raise RuntimeError('SECRET_KEY must be configured in production.')
    _secret = secrets.token_hex(32)   # ephemeral development-only key
    _secret_generated = True
app.config['SECRET_KEY'] = _secret
_db_url = os.environ.get('DATABASE_URL', f'sqlite:///{os.path.join(BASE_DIR, "bustrack.db")}')
if _db_url.startswith('postgres://'):
    _db_url = _db_url.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = _db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.realpath(
    os.environ.get('UPLOAD_FOLDER', os.path.join(BASE_DIR, 'static', 'uploads')))
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024
# Secure session cookies
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE']  = 'Lax'
app.config['SESSION_COOKIE_SECURE']    = os.environ.get('FLASK_ENV') == 'production'
app.config['PERMANENT_SESSION_LIFETIME'] = 86400  # 24 h
app.config['LOGIN_RATE_LIMIT_ATTEMPTS'] = _env_int('LOGIN_RATE_LIMIT_ATTEMPTS', 5, 1, 100)
app.config['LOGIN_RATE_LIMIT_WINDOW_SECONDS'] = _env_int('LOGIN_RATE_LIMIT_WINDOW_SECONDS', 300, 30, 86400)
app.config['LOGIN_RATE_LIMIT_LOCK_SECONDS'] = _env_int('LOGIN_RATE_LIMIT_LOCK_SECONDS', 300, 30, 86400)
app.config['RESTORE_JOB_TTL_SECONDS'] = _env_int('RESTORE_JOB_TTL_SECONDS', 1800, 60, 86400)
app.config['RESTORE_SNAPSHOT_RETENTION_DAYS'] = _env_int(
    'RESTORE_SNAPSHOT_RETENTION_DAYS', 30, 1, 3650)
app.config['BROADCAST_JOB_TTL_SECONDS'] = _env_int(
    'BROADCAST_JOB_TTL_SECONDS', 86400, 300, 604800)
app.config['EMAIL_OUTBOX_MAX_ATTEMPTS'] = _env_int(
    'EMAIL_OUTBOX_MAX_ATTEMPTS', 5, 1, 20)
app.config['EMAIL_OUTBOX_RETRY_BASE_SECONDS'] = _env_int(
    'EMAIL_OUTBOX_RETRY_BASE_SECONDS', 60, 5, 3600)
app.config['EMAIL_OUTBOX_RETRY_MAX_SECONDS'] = _env_int(
    'EMAIL_OUTBOX_RETRY_MAX_SECONDS', 3600, 60, 86400)
app.config['EMAIL_OUTBOX_BATCH_SIZE'] = _env_int(
    'EMAIL_OUTBOX_BATCH_SIZE', 50, 1, 500)
app.config['IMPORT_MAX_ROWS'] = _env_int('IMPORT_MAX_ROWS', 25000, 1, 250000)
app.config['IMPORT_MAX_COLUMNS'] = _env_int('IMPORT_MAX_COLUMNS', 64, 1, 512)
app.config['IMPORT_STAGE_TTL_HOURS'] = _env_int('IMPORT_STAGE_TTL_HOURS', 24, 1, 720)
app.config['POWERSCHOOL_ROLLBACK_RETENTION_DAYS'] = _env_int(
    'POWERSCHOOL_ROLLBACK_RETENTION_DAYS', 30, 1, 365)
app.config['POWERSCHOOL_IMPORT_ENABLED'] = (
    os.environ.get('POWERSCHOOL_IMPORT_ENABLED', '0').strip().lower() in {'1', 'true', 'yes'})
app.config['CSP_ENFORCE'] = (
    os.environ.get('CSP_ENFORCE', '1').strip().lower() in {'1', 'true', 'yes'})
app.config['CSP_REPORT_ONLY'] = (
    os.environ.get('CSP_REPORT_ONLY', '0').strip().lower() in {'1', 'true', 'yes'})

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(BASE_DIR, 'static', 'exports'), exist_ok=True)
os.makedirs(INSTANCE_DIR, exist_ok=True)
IMPORT_STAGE_DIR = os.path.join(INSTANCE_DIR, 'imports')
os.makedirs(IMPORT_STAGE_DIR, mode=0o700, exist_ok=True)
os.chmod(IMPORT_STAGE_DIR, 0o700)

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# ── INSTALLATION LOCK ─────────────────────────────────────────────────────────
INSTALLED_FILE = os.path.join(INSTANCE_DIR, '.installed')

def is_installed():
    if os.path.exists(INSTALLED_FILE):
        return True
    # The database is authoritative if the filesystem marker is lost. This
    # keeps setup closed instead of exposing a second administrator bootstrap.
    try:
        return db.session.query(User.id).first() is not None
    except Exception:
        db.session.rollback()
        return False

def _mark_installed():
    os.makedirs(INSTANCE_DIR, exist_ok=True)
    with open(INSTALLED_FILE, 'w', encoding='utf-8') as f:
        f.write('installed')
    os.chmod(INSTALLED_FILE, 0o600)

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'ico'}
def allowed_file(fn): return '.' in fn and fn.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def _csv_safe_cell(value):
    """Neutralize spreadsheet formulas while preserving ordinary CSV values."""
    if not isinstance(value, str):
        return value
    probe = value.lstrip(' \t\r\n\v\f')
    return "'" + value if probe.startswith(('=', '+', '-', '@')) else value


def _csv_safe_row(values):
    return [_csv_safe_cell(value) for value in values]


def _password_error(value):
    """One password policy shared by installation, user admin, and profile."""
    if not isinstance(value, str) or len(value) < 12:
        return 'Password must be at least 12 characters.'
    if len(value) > 1024:
        return 'Password must not exceed 1024 characters.'
    return None


def _normalize_text(value, maximum=None):
    normalized = unicodedata.normalize('NFKC', str(value or '')).strip()
    return normalized[:maximum] if maximum else normalized


def _normalize_email(value):
    values = []
    for item in _normalize_text(value, 500).split(','):
        email = item.strip().lower()
        if email:
            values.append(email)
    return ','.join(values)


def _normalize_phone(value):
    phone = _normalize_text(value, 40)
    if phone.endswith('.0') and phone[:-2].lstrip('+-').isdigit():
        phone = phone[:-2]
    compact = re.sub(r'[\s().-]', '', phone)
    if compact and not compact.startswith('+') and compact.isdigit():
        compact = '+' + compact
    return compact[:30]


def _normalize_language(value):
    """Normalize common labels to a compact language tag; blank defaults to English."""
    raw = _normalize_text(value, 32).casefold().replace('_', '-')
    aliases = {
        'english': 'en', 'ingles': 'en', 'inglés': 'en',
        'spanish': 'es', 'espanol': 'es', 'español': 'es',
    }
    raw = aliases.get(raw, raw)
    if not raw:
        return 'en'
    if not re.fullmatch(r'[a-z]{2,3}(?:-[a-z0-9]{2,8})?', raw):
        return 'en'
    parts = raw.split('-', 1)
    return parts[0] + (f'-{parts[1].upper()}' if len(parts) > 1 else '')


def _valid_csv_upload(file):
    if not file or not file.filename or not file.filename.lower().endswith('.csv'):
        return False
    mime = (file.mimetype or '').lower().split(';', 1)[0]
    return mime in {'text/csv', 'text/plain', 'application/csv',
                    'application/vnd.ms-excel', 'application/octet-stream'}


# ── MODELS ───────────────────────────────────────────────────────────────────

class Configuration(db.Model):
    id                  = db.Column(db.Integer, primary_key=True)
    # Identity
    app_name            = db.Column(db.String(100), default='School Bus Tracker')
    app_subtitle        = db.Column(db.String(200), default='D205 School District')
    logo_path           = db.Column(db.String(255), default='')
    icon_path           = db.Column(db.String(255), default='')
    # Theme
    theme_mode          = db.Column(db.String(10), default='light')
    color_bg            = db.Column(db.String(20), default='#f1f5f9')
    color_nav           = db.Column(db.String(20), default='#1e293b')
    color_card          = db.Column(db.String(20), default='#ffffff')
    color_text          = db.Column(db.String(20), default='#0f172a')
    color_accent        = db.Column(db.String(20), default='#3b82f6')
    color_nav_text      = db.Column(db.String(20), default='#f8fafc')
    # Operational
    timezone            = db.Column(db.String(50), default='America/Chicago')
    daily_reset_time    = db.Column(db.String(5), default='05:00')
    commit_delay_min    = db.Column(db.Integer, default=5)
    offline_message     = db.Column(db.Text, default='Bus service is currently offline. Check back during operational hours.')
    show_always         = db.Column(db.Boolean, default=True)
    # Language
    lang_frontend       = db.Column(db.String(10), default='en')
    lang_backend        = db.Column(db.String(10), default='en')
    time_format         = db.Column(db.String(4), default='12h')
    # Email
    mail_provider       = db.Column(db.String(20), default='custom')
    mail_server         = db.Column(db.String(100), default='')
    mail_port           = db.Column(db.Integer, default=587)
    mail_use_tls        = db.Column(db.Boolean, default=True)
    mail_use_ssl        = db.Column(db.Boolean, default=False)
    mail_username       = db.Column(db.String(320), default='')
    mail_password       = db.Column(db.String(1000), default='')
    mail_from_email     = db.Column(db.String(320), default='')
    mail_from_name      = db.Column(db.String(100), default='Bus Tracker')
    mail_last_verified_at = db.Column(db.DateTime)
    mail_last_verification_status = db.Column(db.String(20), default='unverified')
    mail_last_error_code = db.Column(db.String(80), default='')
    mail_last_verified_identity = db.Column(db.String(64), default='')
    # SMS / Twilio
    twilio_enabled          = db.Column(db.Boolean, default=False)
    twilio_account_sid      = db.Column(db.String(60), default='')
    twilio_auth_token       = db.Column(db.String(60), default='')
    twilio_from_number      = db.Column(db.String(20), default='')
    twilio_sms_cost_per_seg = db.Column(db.Float, default=0.0079)


class OperationalSchedule(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    name        = db.Column(db.String(100), nullable=False)
    days        = db.Column(db.String(50), default='mon-fri')   # mon-fri, all, weekend, custom
    start_time  = db.Column(db.String(5), nullable=False)       # HH:MM
    end_time    = db.Column(db.String(5), nullable=False)
    is_active   = db.Column(db.Boolean, default=True)
    created_at  = db.Column(db.DateTime, default=_utcnow)


class Holiday(db.Model):
    id              = db.Column(db.Integer, primary_key=True)
    name            = db.Column(db.String(100), nullable=False)
    holiday_type    = db.Column(db.String(50), default='school')  # federal, state, school, local
    holiday_date    = db.Column(db.Date, nullable=False)
    is_recurring    = db.Column(db.Boolean, default=False)
    is_active       = db.Column(db.Boolean, default=True)
    custom_message  = db.Column(db.Text)   # displayed on public page on the holiday day
    created_at      = db.Column(db.DateTime, default=_utcnow)


class UserGroup(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    name        = db.Column(db.String(100), unique=True, nullable=False)
    description = db.Column(db.String(255))
    is_admin    = db.Column(db.Boolean, default=False)
    created_at  = db.Column(db.DateTime, default=_utcnow)
    users       = db.relationship('User', backref='group', lazy=True)
    permissions = db.relationship('GroupPermission', backref='group', lazy=True, cascade='all, delete-orphan')


class GroupPermission(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    group_id     = db.Column(db.Integer, db.ForeignKey('user_group.id'), nullable=False)
    module_key   = db.Column(db.String(50), nullable=False)
    access_level = db.Column(db.String(10), default='none')  # none | limited | full
    __table_args__ = (db.UniqueConstraint('group_id', 'module_key'),)


class GroupCapability(db.Model):
    """Explicit privileged capability grants, separate from visual modules."""
    id             = db.Column(db.Integer, primary_key=True)
    group_id       = db.Column(db.Integer, db.ForeignKey('user_group.id'), nullable=False)
    capability_key = db.Column(db.String(80), nullable=False)
    granted        = db.Column(db.Boolean, nullable=False, default=True)
    created_at     = db.Column(db.DateTime, default=_utcnow)
    __table_args__ = (db.UniqueConstraint('group_id', 'capability_key'),)


MODULES = [
    {'key': 'buses',         'label': 'Buses',          'icon': 'fa-bus'},
    {'key': 'incidents',     'label': 'Status Types',    'icon': 'fa-exclamation-circle'},
    {'key': 'statistics',    'label': 'Statistics',      'icon': 'fa-chart-bar'},
    {'key': 'users',         'label': 'Users',           'icon': 'fa-users'},
    {'key': 'notifications', 'label': 'Notifications',   'icon': 'fa-bell'},
    {'key': 'config',        'label': 'Configuration',   'icon': 'fa-cog'},
    {'key': 'logs',          'label': 'System Logs',     'icon': 'fa-scroll'},
]

# One registry defines every authorization decision. Ordinary read/write
# capabilities retain compatibility with the existing module matrix; sensitive
# capabilities require an explicit GroupCapability grant or Administrator.
CAPABILITIES = {
    'buses.read':                  {'module': 'buses', 'level': 'limited'},
    'buses.write':                 {'module': 'buses', 'level': 'full'},
    'incidents.read':              {'module': 'incidents', 'level': 'limited'},
    'incidents.write':             {'module': 'incidents', 'level': 'full'},
    'statistics.read':             {'module': 'statistics', 'level': 'limited'},
    'statistics.write':            {'module': 'statistics', 'level': 'full'},
    'users.read':                  {'module': 'users', 'level': 'limited'},
    'user.manage':                 {'module': 'users', 'level': 'full'},
    'notifications.read':          {'module': 'notifications', 'level': 'limited'},
    'notifications.write':         {'module': 'notifications', 'level': 'full'},
    'config.read':                 {'module': 'config', 'level': 'limited'},
    'config.write':                {'module': 'config', 'level': 'full'},
    'audit.read':                  {'module': 'logs', 'level': 'limited'},
    'backup.export_operational':   {'module': 'config', 'level': 'full'},
    # Explicit-only capabilities below are never inferred from module access.
    'user.assign_admin':           {},
    'backup.export_sensitive':     {},
    'restore.operational':         {},
    'restore.identity':            {},
    'smtp.diagnose':               {},
    'audit.export':                {},
    'notifications.pii':           {},
    'notifications.export_pii':    {},
    'notifications.broadcast':     {},
    'import.legacy':               {},
    'import.powerschool':          {},
    'import.rollback':             {},
}

MODULE_CAPABILITIES = {
    ('buses', 'limited'): 'buses.read', ('buses', 'full'): 'buses.write',
    ('incidents', 'limited'): 'incidents.read', ('incidents', 'full'): 'incidents.write',
    ('statistics', 'limited'): 'statistics.read', ('statistics', 'full'): 'statistics.write',
    ('users', 'limited'): 'users.read', ('users', 'full'): 'user.manage',
    ('notifications', 'limited'): 'notifications.read',
    ('notifications', 'full'): 'notifications.write',
    ('config', 'limited'): 'config.read', ('config', 'full'): 'config.write',
    ('logs', 'limited'): 'audit.read', ('logs', 'full'): 'audit.read',
}

EXPLICIT_CAPABILITIES_BY_MODULE = {
    ('config', 'full'): {'backup.export_operational', 'smtp.diagnose'},
    ('logs', 'full'): {'audit.export'},
    ('notifications', 'full'): {
        'notifications.pii', 'notifications.export_pii',
        'notifications.broadcast', 'import.legacy',
    },
}


class AuditLog(db.Model):
    __tablename__ = 'audit_log'
    id         = db.Column(db.Integer, primary_key=True)
    user_id    = db.Column(db.Integer, db.ForeignKey('user.id', ondelete='SET NULL'), nullable=True)
    username   = db.Column(db.String(80))
    action     = db.Column(db.String(100))
    module     = db.Column(db.String(50))
    target     = db.Column(db.String(200))
    details    = db.Column(db.Text)
    ip_address = db.Column(db.String(45))
    created_at = db.Column(db.DateTime, default=_utcnow)


class LoginThrottle(db.Model):
    """Database-backed login throttling shared by every application worker."""
    id                = db.Column(db.Integer, primary_key=True)
    throttle_key      = db.Column(db.String(64), unique=True, nullable=False, index=True)
    failed_count      = db.Column(db.Integer, nullable=False, default=0)
    window_started_at = db.Column(db.DateTime, nullable=False, default=_utcnow)
    locked_until      = db.Column(db.DateTime)
    updated_at        = db.Column(db.DateTime, nullable=False, default=_utcnow,
                                  onupdate=_utcnow, index=True)


class User(UserMixin, db.Model):
    id                  = db.Column(db.Integer, primary_key=True)
    username            = db.Column(db.String(80), unique=True, nullable=False)
    email               = db.Column(db.String(120), unique=True, nullable=True)
    password_hash       = db.Column(db.String(256), nullable=False)
    first_name          = db.Column(db.String(80))
    last_name           = db.Column(db.String(80))
    phone               = db.Column(db.String(30))
    workplace           = db.Column(db.String(150))
    job_title           = db.Column(db.String(100))
    group_id            = db.Column(db.Integer, db.ForeignKey('user_group.id'))
    use_email_auth      = db.Column(db.Boolean, default=False)
    receive_notifications = db.Column(db.Boolean, default=True)
    avatar_initials     = db.Column(db.String(4))
    active              = db.Column(db.Boolean, default=True)
    session_version     = db.Column(db.Integer, nullable=False, default=1)
    created_at          = db.Column(db.DateTime, default=_utcnow)
    last_login          = db.Column(db.DateTime)

    @property
    def is_active(self): return self.active

    @property
    def full_name(self):
        return f"{self.first_name or ''} {self.last_name or ''}".strip() or self.username

    @property
    def is_admin(self): return bool(self.group and self.group.is_admin)

    def set_password(self, pwd): self.password_hash = generate_password_hash(pwd)
    def check_password(self, pwd): return check_password_hash(self.password_hash, pwd)

    def has_access(self, module_key, level='limited'):
        if self.is_admin: return True
        if not self.group: return False
        perm = GroupPermission.query.filter_by(group_id=self.group_id, module_key=module_key).first()
        if not perm or perm.access_level == 'none': return False
        if level == 'limited': return perm.access_level in ('limited', 'full')
        if level == 'full': return perm.access_level == 'full'
        return False

    def has_capability(self, capability_key):
        policy = CAPABILITIES.get(capability_key)
        if policy is None:
            return False
        if self.is_admin:
            return True
        if not self.group_id:
            return False
        explicit = GroupCapability.query.filter_by(
            group_id=self.group_id, capability_key=capability_key).first()
        if explicit is not None:
            return bool(explicit.granted)
        module_key = policy.get('module')
        return bool(module_key and self.has_access(
            module_key, policy.get('level', 'limited')))

    def accessible_modules(self):
        if self.is_admin: return MODULES
        if not self.group: return []
        return [m for m in MODULES if self.has_access(m['key'])]


class BusScheduleType(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    name         = db.Column(db.String(50), unique=True, nullable=False)
    time_label   = db.Column(db.String(20))   # e.g. "7:00 AM"
    sort_order   = db.Column(db.Integer, default=0)
    window_start = db.Column(db.String(5))    # HH:MM display window begins
    window_end   = db.Column(db.String(5))    # HH:MM display window ends


class IncidentType(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    name        = db.Column(db.String(100), unique=True, nullable=False)
    color       = db.Column(db.String(20), default='#6b7280')
    icon        = db.Column(db.String(50), default='fa-circle')
    description = db.Column(db.String(255))
    is_default  = db.Column(db.Boolean, default=False)   # On Time = default
    is_system   = db.Column(db.Boolean, default=False)   # Cannot delete
    sort_order  = db.Column(db.Integer, default=0)
    operational_priority = db.Column(db.Integer, nullable=False, default=50)
    created_at  = db.Column(db.DateTime, default=_utcnow)


class Bus(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    identifier   = db.Column(db.String(20), nullable=False)  # TR, TRS, TT — not unique alone
    name         = db.Column(db.String(150), nullable=False)
    route        = db.Column(db.String(200))
    capacity     = db.Column(db.Integer)
    description  = db.Column(db.Text)
    active       = db.Column(db.Boolean, default=True)
    created_at   = db.Column(db.DateTime, default=_utcnow)
    schedule_assignments = db.relationship('BusScheduleAssignment', backref='bus', lazy=True, cascade='all, delete-orphan')
    incident_records     = db.relationship('BusIncidentRecord', backref='bus', lazy=True)
    __table_args__ = (db.UniqueConstraint('identifier', 'name', name='uq_bus_identifier_name'),)

    @property
    def display_name(self): return f"{self.identifier} — {self.name}"


class BusScheduleAssignment(db.Model):
    id               = db.Column(db.Integer, primary_key=True)
    bus_id           = db.Column(db.Integer, db.ForeignKey('bus.id'), nullable=False)
    schedule_type_id = db.Column(db.Integer, db.ForeignKey('bus_schedule_type.id'), nullable=False)
    departure_time   = db.Column(db.String(5))   # HH:MM specific to this bus
    schedule_type    = db.relationship('BusScheduleType')
    __table_args__   = (db.UniqueConstraint('bus_id', 'schedule_type_id'),)


class DelayReason(db.Model):
    id         = db.Column(db.Integer, primary_key=True)
    reason     = db.Column(db.String(200), unique=True, nullable=False)
    sort_order = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=_utcnow)


class BusIncidentRecord(db.Model):
    id                = db.Column(db.Integer, primary_key=True)
    bus_id            = db.Column(db.Integer, db.ForeignKey('bus.id'), nullable=False)
    incident_type_id  = db.Column(db.Integer, db.ForeignKey('incident_type.id'), nullable=False)
    schedule_type_id  = db.Column(db.Integer, db.ForeignKey('bus_schedule_type.id'), nullable=True)
    delay_minutes     = db.Column(db.Integer, default=0)
    eta               = db.Column(db.String(5))    # HH:MM estimated arrival
    delay_reason_id   = db.Column(db.Integer, db.ForeignKey('delay_reason.id'), nullable=True)
    delay_reason_text = db.Column(db.String(200))  # free-text if no preset chosen
    notes             = db.Column(db.Text)
    incident_date     = db.Column(db.Date, default=lambda: district_today())
    is_pending        = db.Column(db.Boolean, default=True)
    committed_at      = db.Column(db.DateTime)
    created_by_id     = db.Column(db.Integer, db.ForeignKey('user.id'))
    created_at        = db.Column(db.DateTime, default=_utcnow)
    updated_at        = db.Column(db.DateTime, default=_utcnow, onupdate=_utcnow)
    request_token     = db.Column(db.String(64), unique=True, nullable=True)
    incident_type     = db.relationship('IncidentType')
    schedule_type     = db.relationship('BusScheduleType')
    delay_reason      = db.relationship('DelayReason')
    created_by        = db.relationship('User')


class SubscriberGroup(db.Model):
    id              = db.Column(db.Integer, primary_key=True)
    name            = db.Column(db.String(100), unique=True, nullable=False)
    description     = db.Column(db.String(200), default='')
    color           = db.Column(db.String(20), default='blue')
    created_at      = db.Column(db.DateTime, default=_utcnow)
    subscribers     = db.relationship('NotificationSubscriber', backref='group', lazy=True)
    bus_assignments = db.relationship('GroupBusAssignment', backref='group',
                                      lazy=True, cascade='all, delete-orphan')


class GroupBusAssignment(db.Model):
    __tablename__ = 'group_bus_assignment'
    id               = db.Column(db.Integer, primary_key=True)
    group_id         = db.Column(db.Integer, db.ForeignKey('subscriber_group.id'), nullable=False)
    bus_id           = db.Column(db.Integer, db.ForeignKey('bus.id'), nullable=False)
    schedule_type_id = db.Column(db.Integer, db.ForeignKey('bus_schedule_type.id'), nullable=True)
    bus              = db.relationship('Bus')
    schedule_type    = db.relationship('BusScheduleType')
    __table_args__ = (db.UniqueConstraint('group_id', 'bus_id', 'schedule_type_id'),)


class NotificationSubscriber(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    notes       = db.Column(db.String(200))        # optional household label
    active      = db.Column(db.Boolean, default=True)
    created_at  = db.Column(db.DateTime, default=_utcnow)
    group_id    = db.Column(db.Integer, db.ForeignKey('subscriber_group.id'), nullable=True)
    school      = db.Column(db.String(100))
    # Legacy columns — kept for DB compat, migrated to SubscriberContact on startup
    first_name  = db.Column(db.String(80))
    last_name   = db.Column(db.String(80))
    email       = db.Column(db.String(120))
    phone       = db.Column(db.String(30))
    bus_assignments = db.relationship('NotificationBusAssignment', backref='subscriber',
                                      lazy=True, cascade='all, delete-orphan')
    contacts    = db.relationship('SubscriberContact', backref='subscriber',
                                  lazy=True, cascade='all, delete-orphan',
                                  order_by='SubscriberContact.sort_order')

    @property
    def full_name(self):
        if self.contacts:
            name = self.contacts[0].full_name
            if name: return name
        if self.notes: return self.notes
        legacy = f"{self.first_name or ''} {self.last_name or ''}".strip()
        return legacy or 'Unnamed'


class SubscriberContact(db.Model):
    __tablename__ = 'subscriber_contact'
    id            = db.Column(db.Integer, primary_key=True)
    subscriber_id = db.Column(db.Integer, db.ForeignKey('notification_subscriber.id'), nullable=False)
    first_name    = db.Column(db.String(80))
    last_name     = db.Column(db.String(80))
    email         = db.Column(db.String(500))
    phone         = db.Column(db.String(30))
    role          = db.Column(db.String(20), default='parent')  # 'parent' | 'student'
    preferred_language = db.Column(db.String(10), nullable=False, default='en')
    sort_order    = db.Column(db.Integer, default=0)

    @property
    def full_name(self): return f"{self.first_name or ''} {self.last_name or ''}".strip()


class NotificationBusAssignment(db.Model):
    id            = db.Column(db.Integer, primary_key=True)
    subscriber_id = db.Column(db.Integer, db.ForeignKey('notification_subscriber.id'), nullable=False)
    bus_id        = db.Column(db.Integer, db.ForeignKey('bus.id'), nullable=False)
    bus           = db.relationship('Bus')
    __table_args__ = (db.UniqueConstraint('subscriber_id', 'bus_id'),)


class NotificationLog(db.Model):
    __tablename__ = 'notification_log'
    id                 = db.Column(db.Integer, primary_key=True)
    incident_record_id = db.Column(db.Integer, db.ForeignKey('bus_incident_record.id'), nullable=True)
    sent_at            = db.Column(db.DateTime, default=_utcnow, index=True)
    channel            = db.Column(db.String(10), nullable=False)   # 'email' | 'sms'
    recipient_name     = db.Column(db.String(160))
    recipient_address  = db.Column(db.String(500))                  # email or phone
    subscriber_id      = db.Column(db.Integer, db.ForeignKey('notification_subscriber.id'), nullable=True)
    group_id           = db.Column(db.Integer, db.ForeignKey('subscriber_group.id'), nullable=True)
    group_name         = db.Column(db.String(100))
    bus_id             = db.Column(db.Integer, db.ForeignKey('bus.id'), nullable=True)
    bus_label          = db.Column(db.String(80))
    status             = db.Column(db.String(10), default='sent')   # 'sent' | 'failed'
    error_message      = db.Column(db.Text)
    sms_sid            = db.Column(db.String(50))
    sms_segments       = db.Column(db.Integer)
    sms_cost_usd       = db.Column(db.Float)


class BroadcastJob(db.Model):
    __tablename__ = 'broadcast_job'
    public_id     = db.Column(db.String(64), primary_key=True)
    owner_id      = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    status        = db.Column(db.String(20), nullable=False, default='queued', index=True)
    total         = db.Column(db.Integer, nullable=False, default=0)
    sent          = db.Column(db.Integer, nullable=False, default=0)
    failed        = db.Column(db.Integer, nullable=False, default=0)
    errors_json   = db.Column(db.Text, nullable=False, default='[]')
    created_at    = db.Column(db.DateTime, nullable=False, default=_utcnow, index=True)
    updated_at    = db.Column(db.DateTime, nullable=False, default=_utcnow, onupdate=_utcnow)
    expires_at    = db.Column(db.DateTime, nullable=False, index=True)

    @property
    def done(self):
        return self.status in {'completed', 'failed', 'expired'}

    @property
    def errors(self):
        try:
            value = json.loads(self.errors_json or '[]')
            return value if isinstance(value, list) else []
        except (TypeError, ValueError):
            return []


class EmailOutbox(db.Model):
    __tablename__ = 'email_outbox'
    id                 = db.Column(db.Integer, primary_key=True)
    dedupe_key         = db.Column(db.String(128), unique=True, nullable=False, index=True)
    kind               = db.Column(db.String(30), nullable=False, index=True)
    recipient_name     = db.Column(db.String(160), default='')
    recipient_address  = db.Column(db.String(320), nullable=False)
    subject            = db.Column(db.String(300), nullable=False)
    body               = db.Column(db.Text, nullable=False)
    status             = db.Column(db.String(20), nullable=False, default='pending', index=True)
    attempts           = db.Column(db.Integer, nullable=False, default=0)
    max_attempts       = db.Column(db.Integer, nullable=False, default=5)
    available_at       = db.Column(db.DateTime, nullable=False, default=_utcnow, index=True)
    locked_at          = db.Column(db.DateTime)
    sent_at            = db.Column(db.DateTime)
    last_error_code    = db.Column(db.String(80), default='')
    incident_record_id = db.Column(db.Integer, db.ForeignKey('bus_incident_record.id'), index=True)
    subscriber_id      = db.Column(db.Integer, db.ForeignKey('notification_subscriber.id'))
    group_id           = db.Column(db.Integer, db.ForeignKey('subscriber_group.id'))
    group_name         = db.Column(db.String(100), default='')
    bus_id             = db.Column(db.Integer, db.ForeignKey('bus.id'))
    bus_label          = db.Column(db.String(80), default='')
    broadcast_job_id   = db.Column(db.String(64), db.ForeignKey('broadcast_job.public_id'), index=True)
    created_at         = db.Column(db.DateTime, nullable=False, default=_utcnow)
    updated_at         = db.Column(db.DateTime, nullable=False, default=_utcnow, onupdate=_utcnow)


class CommunicationEvent(db.Model):
    """Provider-neutral, PII-free event for future communication adapters."""
    __tablename__ = 'communication_event'
    id                 = db.Column(db.Integer, primary_key=True)
    event_key          = db.Column(db.String(128), unique=True, nullable=False, index=True)
    event_type         = db.Column(db.String(40), nullable=False, index=True)
    incident_record_id = db.Column(db.Integer, db.ForeignKey('bus_incident_record.id'),
                                   nullable=False, unique=True, index=True)
    payload_json       = db.Column(db.Text, nullable=False, default='{}')
    status             = db.Column(db.String(20), nullable=False, default='ready', index=True)
    created_at         = db.Column(db.DateTime, nullable=False, default=_utcnow, index=True)


class ImportMappingProfile(db.Model):
    __tablename__ = 'import_mapping_profile'
    id             = db.Column(db.Integer, primary_key=True)
    name           = db.Column(db.String(100), nullable=False)
    source_type    = db.Column(db.String(40), nullable=False)
    schema_version = db.Column(db.String(30), nullable=False)
    mapping_json   = db.Column(db.Text, nullable=False, default='{}')
    active         = db.Column(db.Boolean, nullable=False, default=True)
    created_at     = db.Column(db.DateTime, nullable=False, default=_utcnow)
    __table_args__ = (db.UniqueConstraint('source_type', 'schema_version'),)


class ImportBatch(db.Model):
    __tablename__ = 'import_batch'
    id             = db.Column(db.Integer, primary_key=True)
    public_id      = db.Column(db.String(64), unique=True, nullable=False, index=True)
    source_type    = db.Column(db.String(40), nullable=False, index=True)
    schema_version = db.Column(db.String(30), nullable=False)
    status         = db.Column(db.String(24), nullable=False, default='staged', index=True)
    snapshot_type  = db.Column(db.String(20), nullable=False, default='delta')
    school_year    = db.Column(db.String(20))
    uploaded_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False, index=True)
    file_sha256    = db.Column(db.String(64), nullable=False, index=True)
    analysis_context_sha256 = db.Column(db.String(64), index=True)
    plan_hash      = db.Column(db.String(64), nullable=False)
    total_rows     = db.Column(db.Integer, nullable=False, default=0)
    selected_rows  = db.Column(db.Integer, nullable=False, default=0)
    rejected_rows  = db.Column(db.Integer, nullable=False, default=0)
    excluded_rows  = db.Column(db.Integer, nullable=False, default=0)
    metadata_json  = db.Column(db.Text, nullable=False, default='{}')
    created_at     = db.Column(db.DateTime, nullable=False, default=_utcnow, index=True)
    expires_at     = db.Column(db.DateTime, nullable=False, index=True)
    applied_at     = db.Column(db.DateTime)


class ImportFile(db.Model):
    __tablename__ = 'import_file'
    id             = db.Column(db.Integer, primary_key=True)
    batch_id       = db.Column(db.Integer, db.ForeignKey('import_batch.id', ondelete='CASCADE'), nullable=False, index=True)
    file_type      = db.Column(db.String(30), nullable=False)
    original_name  = db.Column(db.String(255), nullable=False)
    sha256         = db.Column(db.String(64), nullable=False)
    byte_size      = db.Column(db.Integer, nullable=False)
    storage_path   = db.Column(db.String(500), nullable=False)
    headers_json   = db.Column(db.Text, nullable=False, default='[]')
    created_at     = db.Column(db.DateTime, nullable=False, default=_utcnow)
    __table_args__ = (db.UniqueConstraint('batch_id', 'file_type'),)


class ImportRow(db.Model):
    __tablename__ = 'import_row'
    id              = db.Column(db.Integer, primary_key=True)
    batch_id        = db.Column(db.Integer, db.ForeignKey('import_batch.id', ondelete='CASCADE'), nullable=False, index=True)
    row_number      = db.Column(db.Integer, nullable=False)
    external_key    = db.Column(db.String(160), index=True)
    classification  = db.Column(db.String(30), nullable=False, index=True)
    selected        = db.Column(db.Boolean, nullable=False, default=True)
    normalized_json = db.Column(db.Text, nullable=False)
    errors_json     = db.Column(db.Text, nullable=False, default='[]')
    row_hash        = db.Column(db.String(64), nullable=False)
    __table_args__  = (db.UniqueConstraint('batch_id', 'row_number'),)


class ExternalIdentity(db.Model):
    __tablename__ = 'external_identity'
    id              = db.Column(db.Integer, primary_key=True)
    source_type     = db.Column(db.String(40), nullable=False)
    entity_type     = db.Column(db.String(40), nullable=False)
    external_key    = db.Column(db.String(160), nullable=False)
    local_table     = db.Column(db.String(80), nullable=False)
    local_id        = db.Column(db.Integer, nullable=False)
    created_at      = db.Column(db.DateTime, nullable=False, default=_utcnow)
    updated_at      = db.Column(db.DateTime, nullable=False, default=_utcnow, onupdate=_utcnow)
    __table_args__  = (db.UniqueConstraint('source_type', 'entity_type', 'external_key'),)


class ImportChange(db.Model):
    __tablename__ = 'import_change'
    id           = db.Column(db.Integer, primary_key=True)
    batch_id     = db.Column(db.Integer, db.ForeignKey('import_batch.id', ondelete='CASCADE'), nullable=False, index=True)
    row_id       = db.Column(db.Integer, db.ForeignKey('import_row.id', ondelete='SET NULL'), index=True)
    operation    = db.Column(db.String(30), nullable=False)
    target_table = db.Column(db.String(80), nullable=False)
    target_id    = db.Column(db.Integer)
    before_json  = db.Column(db.Text)
    after_json   = db.Column(db.Text)
    created_at   = db.Column(db.DateTime, nullable=False, default=_utcnow)


@login_manager.user_loader
def load_user(uid):
    try:
        user = db.session.get(User, int(uid))
    except (TypeError, ValueError):
        return None
    return user if user and user.active else None


# ── JINJA2 GLOBALS ───────────────────────────────────────────────────────────

_cfg_cache = {}

def get_config():
    if has_request_context() and hasattr(g, '_busroute_configuration'):
        return g._busroute_configuration
    cfg = Configuration.query.first()
    if not cfg:
        cfg = Configuration()
        db.session.add(cfg)
        db.session.commit()
    if has_request_context():
        g._busroute_configuration = cfg
    return cfg


def district_timezone(cfg=None):
    """Return the configured district timezone with a safe district-local fallback."""
    timezone_name = (getattr(cfg or get_config(), 'timezone', '') or 'America/Chicago').strip()
    try:
        return pytz.timezone(timezone_name)
    except pytz.UnknownTimeZoneError:
        app.logger.error('Invalid configured district timezone %r; using America/Chicago.',
                         timezone_name)
        return pytz.timezone('America/Chicago')


def district_now(cfg=None, now_utc=None):
    """Return an aware datetime in the configured district timezone."""
    instant = now_utc or datetime.now(timezone.utc)
    if instant.tzinfo is None:
        instant = pytz.utc.localize(instant)
    return instant.astimezone(district_timezone(cfg))


def district_today(cfg=None, now_utc=None):
    """Return the district's business date, independent of the server timezone."""
    return district_now(cfg, now_utc).date()


def district_date_utc_bounds(date_from, date_to, cfg=None):
    """Convert inclusive district-local dates to naive UTC database boundaries."""
    tz = district_timezone(cfg)
    local_start = tz.localize(datetime.combine(date_from, datetime.min.time()))
    local_end = tz.localize(datetime.combine(
        date_to + timedelta(days=1), datetime.min.time()))
    return (
        local_start.astimezone(timezone.utc).replace(tzinfo=None),
        local_end.astimezone(timezone.utc).replace(tzinfo=None),
    )


def format_district_datetime(value, cfg=None, fmt='%b %d, %Y %I:%M %p'):
    """Render a naive-UTC database timestamp in the configured district timezone."""
    if not value:
        return ''
    instant = value
    if instant.tzinfo is None:
        instant = pytz.utc.localize(instant)
    return instant.astimezone(district_timezone(cfg)).strftime(fmt).replace(' 0', ' ')

def hex_to_text_class(hex_color):
    """Return 'text-white' or 'text-gray-900' based on luminance"""
    h = hex_color.lstrip('#')
    if len(h) == 6:
        r, g, b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
        lum = (0.299*r + 0.587*g + 0.114*b) / 255
        return 'text-white' if lum < 0.5 else 'text-gray-900'
    return 'text-white'

TRANSLATIONS = {
    'en': {
        'bus_legend': 'Bus Legend', 'filters': 'Filters', 'search': 'Search buses, routes, status…',
        'all_status': 'All Status', 'favorites': 'Favorites First', 'on_time': 'On Time',
        'delayed': 'Delayed', 'delay': 'min delay', 'no_incidents': 'No incidents today',
        'favorite': 'Favorite', 'remove_fav': 'Remove favorite',
        'loading': 'Loading…', 'route': 'Route', 'capacity': 'Capacity',
        'schedule': 'Schedule', 'morning': 'Morning', 'midday': 'Midday', 'afternoon': 'Afternoon',
        'all_schedules': 'All schedules', 'live': 'Live', 'admin': 'Admin',
        'no_bus_service_on': 'No bus service on', 'tomorrow': 'Tomorrow',
        'in_days': 'In {count} days', 'service_offline': 'Service Offline',
        'no_service_today': 'No bus service today.', 'holiday': 'Holiday',
        'period': 'Period', 'showing_period': 'Showing buses for this period only',
        'attention_title': 'Service attention',
        'attention_count': '{count} {bus_word} with a status update',
        'all_on_time': 'All buses are currently on time',
        'show_affected': 'Show affected buses', 'show_all': 'Show all buses',
        'search_label': 'Search buses', 'status_label': 'Filter by status',
        'schedule_label': 'Filter by schedule', 'clear_search': 'Clear search',
        'results_count': '{count} {bus_word}', 'bus': 'bus', 'buses': 'buses',
        'no_matches': 'No buses match your filters', 'no_buses': 'No buses registered yet',
        'todays_updates': "Today's Updates", 'service_schedule': 'Service schedule:',
        'updated': 'Updated', 'updated_just_now': 'Updated just now',
        'updated_seconds': 'Updated {count}s ago', 'updated_minutes': 'Updated {count}m ago',
        'reconnecting': 'Reconnecting…', 'connection_interrupted': 'Live updates interrupted',
        'theme_toggle': 'Toggle dark mode', 'favorite_bus': 'Favorite {bus}',
        'remove_favorite_bus': 'Remove {bus} from favorites',
        'e_learning': 'E-Learning', 'combined': 'Combined', 'double_back': 'Double-back',
        'out_of_service': 'Out of Service', 'combined_delayed': 'Combined/Delayed',
        'home': 'Home', 'alerts': 'Alerts', 'filters_nav': 'Filters',
        'affected_buses': 'Buses needing attention', 'favorite_buses': 'Favorite buses',
        'other_buses': 'Other buses', 'filters_title': 'Filters and legend',
        'filters_hint': 'Narrow the list by status or schedule.', 'close': 'Close',
        'reset_filters': 'Reset filters', 'apply_filters': 'View results',
        'install_app': 'Install Bus Tracker',
        'install_app_hint': 'Add this portal to your home screen for faster access.',
        'install': 'Install', 'offline_title': 'You are offline',
        'offline_message': 'Current bus statuses require an internet connection. Reconnect and try again for the latest information.',
        'offline_retry': 'Try again', 'current_alerts': 'Current service alerts',
        'no_current_alerts': 'There are no active service alerts.',
        'open_filters': 'Open filters and legend', 'mobile_search': 'Search a bus',
    },
    'es': {
        'bus_legend': 'Leyenda de Buses', 'filters': 'Filtros', 'search': 'Buscar buses, rutas, estado…',
        'all_status': 'Todos los estados', 'favorites': 'Favoritos primero', 'on_time': 'A Tiempo',
        'delayed': 'Retrasado', 'delay': 'min de retraso', 'no_incidents': 'Sin incidencias hoy',
        'favorite': 'Favorito', 'remove_fav': 'Quitar favorito',
        'loading': 'Cargando…', 'route': 'Ruta', 'capacity': 'Capacidad',
        'schedule': 'Horario', 'morning': 'Mañana', 'midday': 'Medio día', 'afternoon': 'Tarde',
        'all_schedules': 'Todos los horarios', 'live': 'En vivo', 'admin': 'Administración',
        'no_bus_service_on': 'No habrá servicio de buses el', 'tomorrow': 'Mañana',
        'in_days': 'En {count} días', 'service_offline': 'Servicio fuera de horario',
        'no_service_today': 'No hay servicio de buses hoy.', 'holiday': 'Día feriado',
        'period': 'Período', 'showing_period': 'Mostrando únicamente los buses de este período',
        'attention_title': 'Atención de servicio',
        'attention_count': '{count} {bus_word} con un cambio de estado',
        'all_on_time': 'Todos los buses están a tiempo',
        'show_affected': 'Mostrar buses afectados', 'show_all': 'Mostrar todos los buses',
        'search_label': 'Buscar buses', 'status_label': 'Filtrar por estado',
        'schedule_label': 'Filtrar por horario', 'clear_search': 'Borrar búsqueda',
        'results_count': '{count} {bus_word}', 'bus': 'bus', 'buses': 'buses',
        'no_matches': 'Ningún bus coincide con los filtros', 'no_buses': 'Aún no hay buses registrados',
        'todays_updates': 'Actualizaciones de hoy', 'service_schedule': 'Horario de servicio:',
        'updated': 'Actualizado', 'updated_just_now': 'Actualizado ahora',
        'updated_seconds': 'Actualizado hace {count}s', 'updated_minutes': 'Actualizado hace {count}m',
        'reconnecting': 'Reconectando…', 'connection_interrupted': 'Actualizaciones en vivo interrumpidas',
        'theme_toggle': 'Cambiar modo oscuro', 'favorite_bus': 'Marcar {bus} como favorito',
        'remove_favorite_bus': 'Quitar {bus} de favoritos',
        'e_learning': 'Aprendizaje virtual', 'combined': 'Combinado', 'double_back': 'Doble recorrido',
        'out_of_service': 'Fuera de servicio', 'combined_delayed': 'Combinado/Retrasado',
        'home': 'Inicio', 'alerts': 'Alertas', 'filters_nav': 'Filtros',
        'affected_buses': 'Buses que requieren atención', 'favorite_buses': 'Buses favoritos',
        'other_buses': 'Otros buses', 'filters_title': 'Filtros y leyenda',
        'filters_hint': 'Limita la lista por estado u horario.', 'close': 'Cerrar',
        'reset_filters': 'Restablecer filtros', 'apply_filters': 'Ver resultados',
        'install_app': 'Instalar Bus Tracker',
        'install_app_hint': 'Agrega este portal a la pantalla de inicio para acceder más rápido.',
        'install': 'Instalar', 'offline_title': 'No tienes conexión',
        'offline_message': 'Los estados actuales de los buses requieren conexión a internet. Reconéctate e inténtalo nuevamente para obtener la información más reciente.',
        'offline_retry': 'Intentar nuevamente', 'current_alerts': 'Alertas de servicio actuales',
        'no_current_alerts': 'No hay alertas de servicio activas.',
        'open_filters': 'Abrir filtros y leyenda', 'mobile_search': 'Buscar un bus',
    }
}

def t(key, lang=None):
    try:
        lang = lang or get_config().lang_frontend
    except Exception:
        lang = 'en'
    return TRANSLATIONS.get(lang, TRANSLATIONS['en']).get(key, key)

def t_admin(key):
    try:
        cfg = get_config()
        lang = cfg.lang_backend
    except Exception:
        lang = 'en'
    return TRANSLATIONS.get(lang, TRANSLATIONS['en']).get(key, key)

def fmt_time(time_str, fmt='12h'):
    """Convert HH:MM string to 12h (7:30 AM) or 24h (07:30) display format."""
    if not time_str:
        return ''
    try:
        from datetime import datetime as _dt
        t_obj = _dt.strptime(str(time_str)[:5], '%H:%M')
        if fmt == '12h':
            h, m = t_obj.hour, t_obj.minute
            period = 'AM' if h < 12 else 'PM'
            h12 = h % 12 or 12
            return f'{h12}:{m:02d} {period}'
        return f'{t_obj.hour:02d}:{t_obj.minute:02d}'
    except Exception:
        return str(time_str)


def _parse_clock_value(value):
    text = str(value or '').strip()
    if not re.fullmatch(r'\d{2}:\d{2}', text):
        return None
    try:
        parsed = datetime.strptime(text, '%H:%M')
    except ValueError:
        return None
    return parsed.hour * 60 + parsed.minute


def _time_is_in_window(value, window_start, window_end):
    selected = _parse_clock_value(value)
    start = _parse_clock_value(window_start)
    end = _parse_clock_value(window_end)
    if selected is None or start is None or end is None:
        return False
    if start <= end:
        return start <= selected <= end
    return selected >= start or selected <= end


def schedule_assignment_warning(assignment):
    """Describe an invalid persisted departure time without modifying existing data."""
    value = getattr(assignment, 'departure_time', None)
    schedule = getattr(assignment, 'schedule_type', None)
    if not value:
        return None
    if _parse_clock_value(value) is None:
        return f'Invalid departure time: {value}'
    if (schedule and schedule.window_start and schedule.window_end and
            not _time_is_in_window(value, schedule.window_start, schedule.window_end)):
        return (f'{fmt_time(value)} is outside the {schedule.name} window '
                f'({fmt_time(schedule.window_start)}–{fmt_time(schedule.window_end)}).')
    return None


_PUBLIC_STATUS_KEYS = {
    'On Time': 'on_time', 'Delayed': 'delayed', 'E-Learning': 'e_learning',
    'Combined': 'combined', 'Double-back': 'double_back',
    'Out of Service': 'out_of_service', 'Combined/Delayed': 'combined_delayed',
}
_PUBLIC_SCHEDULE_KEYS = {
    'Morning': 'morning', 'Midday': 'midday', 'Afternoon': 'afternoon',
}


def public_status_label(name, lang=None):
    """Translate fixed system statuses while preserving administrator-defined names."""
    return t(_PUBLIC_STATUS_KEYS[name], lang) if name in _PUBLIC_STATUS_KEYS else name


def public_schedule_label(name, lang=None):
    """Translate fixed schedule names while preserving administrator-defined names."""
    return t(_PUBLIC_SCHEDULE_KEYS[name], lang) if name in _PUBLIC_SCHEDULE_KEYS else name


def format_public_date(value, lang='en', include_year=False):
    """Locale-independent date display for the two supported portal languages."""
    if not value:
        return ''
    weekdays = {
        'en': ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'],
        'es': ['lunes', 'martes', 'miércoles', 'jueves', 'viernes', 'sábado', 'domingo'],
    }
    months = {
        'en': ['January', 'February', 'March', 'April', 'May', 'June', 'July',
               'August', 'September', 'October', 'November', 'December'],
        'es': ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio', 'julio',
               'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'],
    }
    selected = lang if lang in weekdays else 'en'
    if selected == 'es':
        rendered = f'{weekdays[selected][value.weekday()]}, {value.day} de {months[selected][value.month - 1]}'
        return f'{rendered} de {value.year}' if include_year else rendered
    rendered = f'{weekdays[selected][value.weekday()]}, {months[selected][value.month - 1]} {value.day}'
    return f'{rendered}, {value.year}' if include_year else rendered

def _csrf_token():
    """Generate (or retrieve) per-session CSRF token, stored in Flask session."""
    if '_csrf' not in session:
        session['_csrf'] = secrets.token_hex(32)
    return session['_csrf']

app.jinja_env.globals.update(
    get_config=get_config,
    MODULES=MODULES,
    hex_to_text_class=hex_to_text_class,
    t=t, t_admin=t_admin,
    fmt_time=fmt_time,
    public_status_label=public_status_label,
    public_schedule_label=public_schedule_label,
    format_public_date=format_public_date,
    schedule_assignment_warning=schedule_assignment_warning,
    format_district_datetime=format_district_datetime,
    csrf_token=_csrf_token,
)


# ── DB INITIALIZATION ────────────────────────────────────────────────────────

def _migrate_bus_table():
    """Migrate bus table: replace unique(identifier) with unique(identifier, name)."""
    try:
        from sqlalchemy import inspect as sa_inspect, text
        insp = sa_inspect(db.engine)
        if 'bus' not in insp.get_table_names():
            return  # table doesn't exist yet, create_all will handle it
        unique_cols = [
            set(c['column_names'])
            for c in insp.get_unique_constraints('bus')
        ]
        # Check if old constraint (only on identifier) still exists
        if {'identifier'} in unique_cols:
            with db.engine.connect() as conn:
                conn.execute(text('ALTER TABLE bus RENAME TO bus_old'))
                conn.commit()
            db.create_all()  # creates bus with new schema
            with db.engine.connect() as conn:
                conn.execute(text('INSERT INTO bus SELECT * FROM bus_old'))
                conn.execute(text('DROP TABLE bus_old'))
                conn.commit()
            print('[Migration] bus table: unique constraint updated to (identifier, name)')
    except Exception as e:
        print(f'[Migration] bus table skipped: {e}')


def _migrate_add_columns():
    """Add new columns to existing tables (safe: ignores if already exists)."""
    from sqlalchemy import inspect as sa_inspect, text
    inspector = sa_inspect(db.engine)
    incident_priority_missing = (
        'incident_type' in inspector.get_table_names() and
        'operational_priority' not in {
            column['name'] for column in inspector.get_columns('incident_type')
        }
    )
    cols = [
        ('bus_schedule_assignment', 'departure_time', 'VARCHAR(5)'),
        ('bus_incident_record',     'eta',             'VARCHAR(5)'),
        ('bus_incident_record',     'delay_reason_id', 'INTEGER'),
        ('bus_incident_record',     'delay_reason_text', 'VARCHAR(200)'),
        ('configuration',           'mail_use_ssl',    'BOOLEAN DEFAULT 0'),
        ('configuration',           'mail_last_verified_at', 'TIMESTAMP'),
        ('configuration',           'mail_last_verification_status', "VARCHAR(20) DEFAULT 'unverified'"),
        ('configuration',           'mail_last_error_code', "VARCHAR(80) DEFAULT ''"),
        ('configuration',           'mail_last_verified_identity', "VARCHAR(64) DEFAULT ''"),
        ('configuration',           'time_format',     "VARCHAR(4) DEFAULT '12h'"),
        ('notification_subscriber', 'group_id',        'INTEGER'),
        ('notification_subscriber', 'notes',           'VARCHAR(200)'),
        ('notification_subscriber', 'school',          'VARCHAR(100)'),
        ('subscriber_contact',      'preferred_language',
                                                   "VARCHAR(10) NOT NULL DEFAULT 'en'"),
        ('bus_schedule_type',       'window_start',    'VARCHAR(5)'),
        ('bus_schedule_type',       'window_end',      'VARCHAR(5)'),
        ('holiday',                 'custom_message',        'TEXT'),
        ('configuration',           'twilio_enabled',         'BOOLEAN DEFAULT 0'),
        ('configuration',           'twilio_account_sid',     "VARCHAR(60) DEFAULT ''"),
        ('configuration',           'twilio_auth_token',      "VARCHAR(60) DEFAULT ''"),
        ('configuration',           'twilio_from_number',     "VARCHAR(20) DEFAULT ''"),
        ('configuration',           'twilio_sms_cost_per_seg','REAL DEFAULT 0.0079'),
        ('group_bus_assignment',    'schedule_type_id',       'INTEGER'),
        ('user',                    'session_version',        'INTEGER NOT NULL DEFAULT 1'),
        ('import_batch',            'analysis_context_sha256','VARCHAR(64)'),
        ('incident_type',           'operational_priority',  'INTEGER NOT NULL DEFAULT 50'),
        ('bus_incident_record',     'request_token',         'VARCHAR(64)'),
    ]
    # Use a separate connection per column so a failed ALTER TABLE (column already
    # exists) never leaves a shared connection in an aborted-transaction state.
    for table, col, coltype in cols:
        try:
            with db.engine.connect() as conn:
                conn.execute(text(f'ALTER TABLE "{table}" ADD COLUMN "{col}" {coltype}'))
                conn.commit()
        except Exception:
            pass  # column already exists — safe to ignore
    try:
        with db.engine.connect() as conn:
            conn.execute(text(
                'CREATE INDEX IF NOT EXISTS '
                'ix_import_batch_analysis_context_sha256 '
                'ON import_batch (analysis_context_sha256)'))
            conn.commit()
    except Exception:
        pass
    if incident_priority_missing:
        priorities = {
            'On Time': 0, 'E-Learning': 20, 'Combined': 40,
            'Double-back': 60, 'Delayed': 70,
            'Combined/Delayed': 85, 'Out of Service': 100,
        }
        try:
            with db.engine.connect() as conn:
                for name, priority in priorities.items():
                    conn.execute(text(
                        'UPDATE incident_type SET operational_priority = :priority '
                        'WHERE name = :name'), {'priority': priority, 'name': name})
                conn.commit()
        except Exception:
            pass
    try:
        with db.engine.connect() as conn:
            conn.execute(text(
                'CREATE UNIQUE INDEX IF NOT EXISTS '
                'uq_bus_incident_record_request_token '
                'ON bus_incident_record (request_token)'))
            conn.commit()
    except Exception:
        pass


def _migrate_email_column():
    """Widen email and encrypted-secret columns on PostgreSQL."""
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    if not db_url.startswith('postgresql'):
        return  # SQLite doesn't enforce VARCHAR length — no action needed
    from sqlalchemy import text
    statements = (
        'ALTER TABLE subscriber_contact ALTER COLUMN email TYPE VARCHAR(500)',
        'ALTER TABLE configuration ALTER COLUMN mail_password TYPE VARCHAR(1000)',
        'ALTER TABLE configuration ALTER COLUMN mail_username TYPE VARCHAR(320)',
        'ALTER TABLE configuration ALTER COLUMN mail_from_email TYPE VARCHAR(320)',
        'ALTER TABLE notification_log ALTER COLUMN recipient_address TYPE VARCHAR(500)',
    )
    for statement in statements:
        try:
            with db.engine.connect() as conn:
                conn.execute(text(statement))
                conn.commit()
        except Exception:
            pass  # already wide enough or column doesn't exist yet


def _migrate_group_bus_period():
    """Change group_bus_assignment unique constraint to (group_id, bus_id, schedule_type_id)."""
    from sqlalchemy import text
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    is_pg  = db_url.startswith('postgresql')
    try:
        if is_pg:
            with db.engine.connect() as conn:
                # Find and drop any unique constraint that covers group_id+bus_id but NOT schedule_type_id
                rows = conn.execute(text("""
                    SELECT tc.constraint_name
                    FROM information_schema.table_constraints tc
                    JOIN information_schema.key_column_usage kcu
                      ON tc.constraint_name = kcu.constraint_name
                     AND tc.table_name = kcu.table_name
                    WHERE tc.table_name = 'group_bus_assignment'
                      AND tc.constraint_type = 'UNIQUE'
                    GROUP BY tc.constraint_name
                    HAVING COUNT(*) = 2
                       AND SUM(CASE WHEN kcu.column_name = 'group_id' THEN 1 ELSE 0 END) = 1
                       AND SUM(CASE WHEN kcu.column_name = 'bus_id' THEN 1 ELSE 0 END) = 1
                """)).fetchall()
                for row in rows:
                    try:
                        conn.execute(text(f'ALTER TABLE group_bus_assignment DROP CONSTRAINT "{row[0]}"'))
                        conn.commit()
                    except Exception:
                        conn.rollback()
                try:
                    conn.execute(text(
                        'ALTER TABLE group_bus_assignment ADD CONSTRAINT uq_gba_grp_bus_period '
                        'UNIQUE (group_id, bus_id, schedule_type_id)'
                    ))
                    conn.commit()
                except Exception:
                    conn.rollback()  # constraint already exists
        else:
            # SQLite: check if the current unique index covers schedule_type_id
            with db.engine.connect() as conn:
                indexes = conn.execute(text("PRAGMA index_list('group_bus_assignment')")).fetchall()
                needs_migration = False
                for idx in indexes:
                    if idx[2]:  # unique flag
                        cols = [r[2] for r in conn.execute(
                            text(f"PRAGMA index_info('{idx[1]}')")).fetchall()]
                        if 'group_id' in cols and 'bus_id' in cols and 'schedule_type_id' not in cols:
                            needs_migration = True
                            break
                if needs_migration:
                    conn.execute(text("""
                        CREATE TABLE group_bus_assignment_new (
                            id INTEGER NOT NULL,
                            group_id INTEGER NOT NULL,
                            bus_id INTEGER NOT NULL,
                            schedule_type_id INTEGER,
                            PRIMARY KEY (id),
                            FOREIGN KEY(group_id) REFERENCES subscriber_group(id),
                            FOREIGN KEY(bus_id) REFERENCES bus(id),
                            FOREIGN KEY(schedule_type_id) REFERENCES bus_schedule_type(id),
                            UNIQUE (group_id, bus_id, schedule_type_id)
                        )
                    """))
                    conn.execute(text("""
                        INSERT INTO group_bus_assignment_new (id, group_id, bus_id, schedule_type_id)
                        SELECT id, group_id, bus_id,
                               CASE WHEN EXISTS (
                                   SELECT 1 FROM pragma_table_info('group_bus_assignment')
                                   WHERE name='schedule_type_id'
                               ) THEN schedule_type_id ELSE NULL END
                        FROM group_bus_assignment
                    """))
                    conn.execute(text('DROP TABLE group_bus_assignment'))
                    conn.execute(text('ALTER TABLE group_bus_assignment_new RENAME TO group_bus_assignment'))
                    conn.commit()
                    print('[Migration] group_bus_assignment: unique constraint updated to include schedule_type_id')
    except Exception as e:
        print(f'[Migration] group_bus_period skipped: {e}')


def _migrate_subscriber_contacts():
    """One-time: convert legacy subscriber person fields → SubscriberContact records."""
    try:
        subs = NotificationSubscriber.query.all()
        changed = False
        for sub in subs:
            if not sub.contacts and (sub.email or sub.first_name):
                db.session.add(SubscriberContact(
                    subscriber_id=sub.id,
                    first_name=sub.first_name, last_name=sub.last_name,
                    email=sub.email, phone=sub.phone,
                    role='parent', sort_order=0,
                ))
                changed = True
        if changed:
            db.session.commit()
    except Exception as e:
        print(f'[Migration] subscriber_contacts skipped: {e}')
        db.session.rollback()


def _seed_phase2_security_and_imports():
    for group in UserGroup.query.filter_by(is_admin=False).all():
        _sync_group_capabilities(group.id)
    profiles = [
        ('Legacy CSV v1', 'legacy_csv', '1', {
            'identity': 'additive-household',
            'columns': ['schema_version', 'subscriber_id', 'household_label',
                        'group', 'active', 'role', 'first_name', 'last_name',
                        'email', 'phone'],
        }),
        ('PowerSchool Import v1', 'powerschool', '1', {
            **DEFAULT_MAPPING_V1,
            'enabled': True,
        }),
    ]
    for name, source_type, version, mapping in profiles:
        row = ImportMappingProfile.query.filter_by(
            source_type=source_type, schema_version=version).first()
        if not row:
            db.session.add(ImportMappingProfile(
                name=name, source_type=source_type, schema_version=version,
                mapping_json=json.dumps(mapping, sort_keys=True)))
        elif source_type == 'powerschool':
            try:
                existing_mapping = json.loads(row.mapping_json or '{}')
            except (TypeError, ValueError):
                existing_mapping = {}
            # Upgrade only the Phase 2 placeholder.  A complete operator-defined
            # profile is never overwritten during startup. Managed aliases are
            # merged additively so existing profiles can read newer saved exports
            # without discarding operator-defined aliases.
            if not isinstance(existing_mapping.get('files'), dict):
                row.mapping_json = json.dumps(mapping, sort_keys=True)
            else:
                original_mapping_json = json.dumps(existing_mapping, sort_keys=True)
                changed = False
                for file_key, managed_file in DEFAULT_MAPPING_V1['files'].items():
                    existing_file = existing_mapping['files'].get(file_key)
                    if existing_file is None:
                        existing_file = {}
                        existing_mapping['files'][file_key] = existing_file
                        changed = True
                    if not isinstance(existing_file, dict):
                        continue
                    existing_columns = existing_file.get('columns')
                    if existing_columns is None:
                        existing_columns = {}
                        existing_file['columns'] = existing_columns
                        changed = True
                    if not isinstance(existing_columns, dict):
                        continue
                    for canonical, managed_aliases in managed_file['columns'].items():
                        existing_aliases = existing_columns.get(canonical)
                        if existing_aliases is None:
                            existing_aliases = []
                            existing_columns[canonical] = existing_aliases
                            changed = True
                        if not isinstance(existing_aliases, list):
                            continue
                        for alias in managed_aliases:
                            if alias not in existing_aliases:
                                existing_aliases.append(alias)
                                changed = True
                    existing_required = existing_file.get('required')
                    if existing_required is None:
                        existing_required = []
                        existing_file['required'] = existing_required
                        changed = True
                    if isinstance(existing_required, list):
                        for canonical in managed_file.get('required', []):
                            if canonical not in existing_required:
                                existing_required.append(canonical)
                                changed = True
                existing_periods = existing_mapping.get('period_aliases')
                if existing_periods is None:
                    existing_periods = {}
                    existing_mapping['period_aliases'] = existing_periods
                    changed = True
                if isinstance(existing_periods, dict):
                    for canonical, managed_aliases in (
                            DEFAULT_MAPPING_V1['period_aliases'].items()):
                        existing_aliases = existing_periods.get(canonical)
                        if existing_aliases is None:
                            existing_aliases = []
                            existing_periods[canonical] = existing_aliases
                            changed = True
                        if not isinstance(existing_aliases, list):
                            continue
                        for alias in managed_aliases:
                            if alias not in existing_aliases:
                                existing_aliases.append(alias)
                                changed = True
                if changed or json.dumps(
                        existing_mapping, sort_keys=True) != original_mapping_json:
                    row.mapping_json = json.dumps(existing_mapping, sort_keys=True)

    now = _utcnow()
    BroadcastJob.query.filter(
        BroadcastJob.status.in_(['queued', 'running']),
        BroadcastJob.updated_at < now - timedelta(minutes=15),
    ).update({'status': 'failed', 'updated_at': now}, synchronize_session=False)
    BroadcastJob.query.filter(
        BroadcastJob.expires_at <= now,
        ~BroadcastJob.status.in_(['completed', 'failed', 'expired']),
    ).update({'status': 'expired', 'updated_at': now}, synchronize_session=False)
    db.session.commit()


@contextmanager
def _database_init_lock():
    """Serialize schema initialization across Gunicorn workers and processes."""
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    if db_url.startswith('postgresql'):
        from sqlalchemy import text
        database_name = db.engine.url.database or 'default'
        lock_material = f'{database_name}:bustrack-schema-init:v1'.encode('utf-8')
        lock_key = int.from_bytes(
            hashlib.sha256(lock_material).digest()[:8], 'big', signed=True)
        lock_connection = db.engine.connect()
        try:
            lock_connection.execute(
                text('SELECT pg_advisory_lock(:lock_key)'), {'lock_key': lock_key})
            yield
        finally:
            try:
                lock_connection.execute(
                    text('SELECT pg_advisory_unlock(:lock_key)'), {'lock_key': lock_key})
            finally:
                lock_connection.close()
        return

    lock_path = os.path.join(INSTANCE_DIR, '.database-init.lock')
    lock_fd = os.open(lock_path, os.O_RDWR | os.O_CREAT, 0o600)
    try:
        try:
            import fcntl
            fcntl.flock(lock_fd, fcntl.LOCK_EX)
        except ImportError:
            pass  # Single-process development fallback on platforms without fcntl.
        yield
    finally:
        try:
            import fcntl
            fcntl.flock(lock_fd, fcntl.LOCK_UN)
        except ImportError:
            pass
        os.close(lock_fd)


@contextmanager
def _roster_import_lock():
    """Try to serialize roster-changing imports across every web worker."""
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    if db_url.startswith('postgresql'):
        from sqlalchemy import text
        database_name = db.engine.url.database or 'default'
        lock_material = f'{database_name}:bustrack-roster-import:v1'.encode(
            'utf-8')
        lock_key = int.from_bytes(
            hashlib.sha256(lock_material).digest()[:8], 'big', signed=True)
        lock_connection = db.engine.connect()
        acquired = False
        try:
            acquired = bool(lock_connection.execute(
                text('SELECT pg_try_advisory_lock(:lock_key)'),
                {'lock_key': lock_key}).scalar())
            yield acquired
        finally:
            try:
                if acquired:
                    lock_connection.execute(
                        text('SELECT pg_advisory_unlock(:lock_key)'),
                        {'lock_key': lock_key})
            finally:
                lock_connection.close()
        return

    lock_path = os.path.join(INSTANCE_DIR, '.roster-import.lock')
    lock_fd = os.open(lock_path, os.O_RDWR | os.O_CREAT, 0o600)
    acquired = False
    try:
        try:
            import fcntl
            fcntl.flock(lock_fd, fcntl.LOCK_EX | fcntl.LOCK_NB)
            acquired = True
        except BlockingIOError:
            acquired = False
        except ImportError:
            acquired = True  # Single-process development fallback.
        yield acquired
    finally:
        try:
            if acquired:
                import fcntl
                fcntl.flock(lock_fd, fcntl.LOCK_UN)
        except ImportError:
            pass
        os.close(lock_fd)


def _serialized_roster_mutation(response_kind='json'):
    """Fail fast when another import or rollback owns the roster mutex."""
    def decorator(view):
        @wraps(view)
        def wrapped(*args, **kwargs):
            with _roster_import_lock() as acquired:
                if acquired:
                    return view(*args, **kwargs)
                message = (
                    'Another roster import or rollback is in progress. '
                    'Wait for it to finish and try again.')
                if response_kind == 'html':
                    flash(message, 'warning')
                    return redirect(url_for('notifications'))
                return jsonify({'ok': False, 'message': message}), 409
        return wrapped
    return decorator


def _initialize_database_unlocked():
    _migrate_bus_table()
    db.create_all()
    _migrate_add_columns()
    _migrate_email_column()
    _migrate_group_bus_period()
    _seed_defaults()
    _migrate_subscriber_contacts()
    _seed_phase2_security_and_imports()

    from sqlalchemy import inspect as sa_inspect
    inspector = sa_inspect(db.engine)
    user_columns = {column['name'] for column in inspector.get_columns('user')}
    import_batch_columns = {
        column['name'] for column in inspector.get_columns('import_batch')}
    throttle_columns = ({column['name'] for column in inspector.get_columns('login_throttle')}
                        if 'login_throttle' in inspector.get_table_names() else set())
    required_phase2_tables = {
        'group_capability', 'broadcast_job', 'import_mapping_profile',
        'import_batch', 'import_file', 'import_row', 'external_identity',
        'import_change', 'email_outbox',
    }
    if 'session_version' not in user_columns or not {
            'throttle_key', 'failed_count', 'window_started_at', 'locked_until'
    }.issubset(throttle_columns):
        raise RuntimeError('Security schema migration did not complete; refusing to start.')
    if 'analysis_context_sha256' not in import_batch_columns:
        raise RuntimeError('PowerSchool import schema migration did not complete; refusing to start.')
    if not required_phase2_tables.issubset(set(inspector.get_table_names())):
        raise RuntimeError('Phase 2 additive schema migration did not complete; refusing to start.')

    # Auto-detect existing installations: if users exist but no lock file, create it.
    if not os.path.exists(INSTALLED_FILE) and User.query.count() > 0:
        _mark_installed()


def init_db():
    with _database_init_lock():
        _initialize_database_unlocked()


# ── HELPERS ──────────────────────────────────────────────────────────────────

def get_current_period(cfg=None):
    """Returns the active BusScheduleType based on current local time, or None."""
    try:
        cfg = cfg or get_config()
        now = district_now(cfg)
        current_time = now.strftime('%H:%M')
        periods = BusScheduleType.query.filter(
            BusScheduleType.window_start != None,
            BusScheduleType.window_end   != None,
        ).order_by(BusScheduleType.sort_order).all()
        for p in periods:
            if p.window_start and p.window_end and p.window_start <= current_time <= p.window_end:
                return p
    except Exception:
        pass
    return None


def get_bus_status(bus_id, for_date=None, schedule_type_id=None):
    """Returns (IncidentType, delay_minutes) for a bus on a given date/period."""
    if for_date is None: for_date = district_today()
    q = BusIncidentRecord.query.filter_by(bus_id=bus_id, incident_date=for_date)
    if schedule_type_id:
        q = q.filter_by(schedule_type_id=schedule_type_id)
    rec = q.order_by(BusIncidentRecord.created_at.desc()).first()
    if rec:
        return rec.incident_type, rec.delay_minutes
    default = IncidentType.query.filter_by(is_default=True).first()
    return default, 0


def _schedule_assignments_from_form():
    """Validate and return selected schedule/departure pairs from an admin form."""
    raw_ids = request.form.getlist('schedule_ids')
    try:
        schedule_ids = list(dict.fromkeys(int(value) for value in raw_ids))
    except (TypeError, ValueError):
        return None, 'One or more selected schedule periods are invalid.'
    if not schedule_ids:
        return [], None

    schedules = BusScheduleType.query.filter(BusScheduleType.id.in_(schedule_ids)).all()
    schedule_by_id = {schedule.id: schedule for schedule in schedules}
    if len(schedule_by_id) != len(schedule_ids):
        return None, 'One or more selected schedule periods no longer exist.'

    validated = []
    for schedule_id in schedule_ids:
        schedule = schedule_by_id[schedule_id]
        departure_time = request.form.get(f'departure_time_{schedule_id}', '').strip() or None
        if departure_time and _parse_clock_value(departure_time) is None:
            return None, f'Enter a valid departure time for {schedule.name}.'
        if (departure_time and schedule.window_start and schedule.window_end and
                not _time_is_in_window(
                    departure_time, schedule.window_start, schedule.window_end)):
            return None, (
                f'{fmt_time(departure_time)} is outside the {schedule.name} window '
                f'({fmt_time(schedule.window_start)}–{fmt_time(schedule.window_end)}).')
        validated.append((schedule_id, departure_time))
    return validated, None

def is_operational():
    """Check current time against operational schedules. Returns (bool, message)."""
    cfg = get_config()
    if cfg.show_always:
        return True, None
    try:
        now = district_now(cfg)
        today_str = now.strftime('%A').lower()[:3]  # mon, tue…
        current_time = now.strftime('%H:%M')
        # Check holidays
        holiday = Holiday.query.filter_by(holiday_date=now.date(), is_active=True).first()
        if holiday:
            msg = holiday.custom_message or f"No bus service today — {holiday.name}"
            return False, msg
        # Check schedules
        schedules = OperationalSchedule.query.filter_by(is_active=True).all()
        for s in schedules:
            days = s.days
            applies = (
                days == 'all' or
                (days == 'mon-fri' and today_str in ('mon','tue','wed','thu','fri')) or
                (days == 'weekend' and today_str in ('sat','sun')) or
                today_str in days
            )
            if applies and s.start_time <= current_time <= s.end_time:
                return True, None
        return False, cfg.offline_message
    except Exception:
        return True, None

def bus_list_today(period=None, admin=False):
    """Return bus status list for today.

    admin=True  → all active buses, all today's incidents (no period filter).
    admin=False → public view: only buses assigned to current period.
    """
    today = district_today()
    current_period = period
    if current_period is None:
        current_period = get_current_period()

    if admin:
        buses = Bus.query.filter_by(active=True).order_by(Bus.identifier).all()
    elif current_period is not None:
        assigned_ids = {a.bus_id for a in BusScheduleAssignment.query.filter_by(
            schedule_type_id=current_period.id).all()}
        buses = Bus.query.filter(
            Bus.active == True,
            Bus.id.in_(assigned_ids),
        ).order_by(Bus.identifier).all()
    else:
        buses = Bus.query.filter_by(active=True).order_by(Bus.identifier).all()

    period_id = current_period.id if current_period else None
    result = []
    for bus in buses:
        status, delay = get_bus_status(bus.id, today, schedule_type_id=period_id)
        q = BusIncidentRecord.query.filter_by(bus_id=bus.id, incident_date=today)
        if not admin and period_id:
            q = q.filter_by(schedule_type_id=period_id)
        incidents = q.order_by(BusIncidentRecord.created_at.desc()).all()
        schedules = [a.schedule_type for a in bus.schedule_assignments]
        latest = incidents[0] if incidents else None
        eta = latest.eta if latest else None
        if latest and latest.delay_reason_id and latest.delay_reason:
            delay_reason = latest.delay_reason.reason
        elif latest and latest.delay_reason_text:
            delay_reason = latest.delay_reason_text
        else:
            delay_reason = None
        result.append({'bus': bus, 'status': status, 'delay': delay,
                       'incidents': incidents, 'schedules': schedules,
                       'schedule_assignments': bus.schedule_assignments,
                       'eta': eta, 'delay_reason': delay_reason,
                       'current_period': current_period})
    return result

SMTP_PROVIDER_PRESETS = {
    'office365': {
        'server': 'smtp.office365.com',
        'port': 587,
        'use_tls': True,
        'use_ssl': False,
        'allowed_transports': [(587, True, False)],
        'banner': 'Server: smtp.office365.com — Port: 587 — STARTTLS',
        'note': 'SMTP AUTH must be enabled for the mailbox. OAuth migration is recommended.',
    },
    'google': {
        'server': 'smtp.gmail.com',
        'port': 587,
        'use_tls': True,
        'use_ssl': False,
        'allowed_transports': [(587, True, False), (465, False, True)],
        'banner': 'Server: smtp.gmail.com — Port: 587 — STARTTLS',
        'note': 'Requires a Google App Password when password authentication is used.',
    },
}
SMTP_PRESET_SERVERS = {
    provider: preset['server'] for provider, preset in SMTP_PROVIDER_PRESETS.items()
}
_ENCRYPTED_SECRET_PREFIX = 'enc:v1:'
_EMAIL_RE = re.compile(r'^[^\s@]+@[^\s@]+\.[^\s@]+$')


def _smtp_allowed_hosts():
    configured = {host.strip().lower().rstrip('.') for host in
                  os.environ.get('SMTP_ALLOWED_HOSTS', '').split(',') if host.strip()}
    return configured | set(SMTP_PRESET_SERVERS.values())


def _validated_smtp_host(provider, server):
    if provider not in {'custom', *SMTP_PRESET_SERVERS.keys()}:
        raise ValueError('Unsupported mail provider.')
    host = (SMTP_PRESET_SERVERS.get(provider) or server or '').strip().lower().rstrip('.')
    if not host or not re.fullmatch(r'(?=.{1,253}$)[A-Za-z0-9.-]+', host):
        raise ValueError('A valid SMTP server hostname is required.')
    if host not in _smtp_allowed_hosts():
        raise ValueError('SMTP server is not in SMTP_ALLOWED_HOSTS.')
    return host


def _validated_smtp_port(value):
    try:
        port = int(value)
    except (TypeError, ValueError) as exc:
        raise ValueError('SMTP port must be numeric.') from exc
    if not 1 <= port <= 65535:
        raise ValueError('SMTP port is outside the valid range.')
    return port


def _credential_fernet():
    """Return the credential cipher without embedding key material in the database."""
    from cryptography.fernet import Fernet
    raw_key = (os.environ.get('EMAIL_CREDENTIAL_ENCRYPTION_KEY', '').strip() or
               os.environ.get('BACKUP_ENCRYPTION_KEY', '').strip())
    if not raw_key:
        raise RuntimeError(
            'EMAIL_CREDENTIAL_ENCRYPTION_KEY or BACKUP_ENCRYPTION_KEY is required.')
    try:
        return Fernet(raw_key.encode('ascii', 'strict'))
    except (ValueError, TypeError, UnicodeError) as exc:
        raise RuntimeError('The configured email credential encryption key is invalid.') from exc


def _encrypt_mail_password(password):
    if not password:
        return ''
    if password.startswith(_ENCRYPTED_SECRET_PREFIX):
        _decrypt_mail_password(password)
        return password
    token = _credential_fernet().encrypt(password.encode('utf-8')).decode('ascii')
    return _ENCRYPTED_SECRET_PREFIX + token


def _decrypt_mail_password(stored):
    if not stored:
        return ''
    if not stored.startswith(_ENCRYPTED_SECRET_PREFIX):
        return stored  # Backward-compatible legacy plaintext until first save/migration.
    from cryptography.fernet import InvalidToken
    token = stored[len(_ENCRYPTED_SECRET_PREFIX):]
    try:
        return _credential_fernet().decrypt(token.encode('ascii')).decode('utf-8')
    except (InvalidToken, ValueError, TypeError, UnicodeError) as exc:
        raise RuntimeError('The saved email credential cannot be decrypted.') from exc


def _mail_password_is_encrypted(cfg):
    return bool((cfg.mail_password or '').startswith(_ENCRYPTED_SECRET_PREFIX))


def _canonical_smtp_transport(provider, server, port, use_tls, use_ssl):
    validated_server = _validated_smtp_host(provider, server)
    validated_port = _validated_smtp_port(port)
    tls = bool(use_tls)
    ssl_enabled = bool(use_ssl)
    if tls and ssl_enabled:
        raise ValueError('STARTTLS and SSL/TLS cannot both be enabled.')

    preset = SMTP_PROVIDER_PRESETS.get(provider)
    if preset:
        allowed = {tuple(item) for item in preset['allowed_transports']}
        requested = (validated_port, tls, ssl_enabled)
        if provider == 'office365':
            requested = next(iter(allowed))
        elif requested not in allowed:
            raise ValueError('The selected provider does not support this port/security combination.')
        validated_port, tls, ssl_enabled = requested
    return validated_server, validated_port, tls, ssl_enabled


def _smtp_identity(provider, server, port, use_tls, use_ssl, username):
    material = json.dumps({
        'provider': provider,
        'server': server,
        'port': int(port),
        'use_tls': bool(use_tls),
        'use_ssl': bool(use_ssl),
        'username': (username or '').strip().lower(),
    }, sort_keys=True, separators=(',', ':')).encode('utf-8')
    return hashlib.sha256(material).hexdigest()


def _current_smtp_identity(cfg):
    server, port, tls, ssl_enabled = _canonical_smtp_transport(
        cfg.mail_provider, cfg.mail_server, cfg.mail_port or 587,
        cfg.mail_use_tls, getattr(cfg, 'mail_use_ssl', False))
    return _smtp_identity(
        cfg.mail_provider, server, port, tls, ssl_enabled, cfg.mail_username)


def _mail_verification_identity(connection_identity, from_email):
    material = f'{connection_identity}\n{(from_email or "").strip().lower()}'.encode('utf-8')
    return hashlib.sha256(material).hexdigest()


def _validated_email(value, label, *, required=False):
    address = (value or '').strip()
    if not address and not required:
        return ''
    if not address or len(address) > 320 or not _EMAIL_RE.fullmatch(address):
        raise ValueError(f'Enter a valid {label}.')
    return address


def _smtp_settings_from_payload(cfg, payload, *, allow_saved_password):
    provider = (payload.get('provider', cfg.mail_provider) or 'custom').strip().lower()
    server, port, tls, ssl_enabled = _canonical_smtp_transport(
        provider,
        payload.get('server', cfg.mail_server),
        payload.get('port', cfg.mail_port or 587),
        payload.get('use_tls', cfg.mail_use_tls),
        payload.get('use_ssl', getattr(cfg, 'mail_use_ssl', False)),
    )
    username = (payload.get('username', cfg.mail_username) or '').strip()
    from_email = _validated_email(
        payload.get('from_email', cfg.mail_from_email), 'From email address', required=True)
    from_name = (payload.get('from_name', cfg.mail_from_name) or 'Bus Tracker').strip()
    if len(username) > 320 or len(from_name) > 100:
        raise ValueError('Email account fields exceed the allowed length.')
    if provider in SMTP_PROVIDER_PRESETS and not username:
        raise ValueError('Username / Email is required for the selected provider.')

    requested_identity = _smtp_identity(
        provider, server, port, tls, ssl_enabled, username)
    supplied_password = payload.get('password', '')
    supplied_password = supplied_password if isinstance(supplied_password, str) else ''
    if supplied_password:
        password = supplied_password
    elif (allow_saved_password and cfg.mail_password and
          requested_identity == _current_smtp_identity(cfg)):
        password = _decrypt_mail_password(cfg.mail_password or '')
    elif username:
        raise EmailTransportError(
            'password_required_for_changes',
            'Re-enter the SMTP password because the connection settings changed.',
        )
    else:
        password = ''

    return SMTPSettings(
        provider=provider,
        server=server,
        port=port,
        use_tls=tls,
        use_ssl=ssl_enabled,
        username=username,
        password=password,
        from_email=from_email,
        from_name=from_name,
    ), requested_identity


def _smtp_settings_from_config(cfg):
    return _smtp_settings_from_payload(cfg, {}, allow_saved_password=True)[0]


def configure_mail(cfg, override=None):
    """Compatibility wrapper returning immutable settings instead of mutating Flask."""
    if override is None:
        return _smtp_settings_from_config(cfg)
    return _smtp_settings_from_payload(
        cfg, override, allow_saved_password=True)[0]


def _record_mail_verification(cfg, connection_identity, from_email, status, error_code=''):
    cfg.mail_last_verified_at = _utcnow()
    cfg.mail_last_verified_identity = _mail_verification_identity(
        connection_identity, from_email)
    cfg.mail_last_verification_status = status
    cfg.mail_last_error_code = error_code
    db.session.commit()


def _matches_saved_mail_identity(cfg, connection_identity, from_email):
    try:
        return bool(
            connection_identity and
            connection_identity == _current_smtp_identity(cfg) and
            (from_email or '').strip().lower() ==
            (cfg.mail_from_email or '').strip().lower()
        )
    except (ValueError, RuntimeError):
        return False


def _mail_configuration_status(cfg):
    status = {
        'has_server': False,
        'has_credentials': bool(cfg.mail_username and cfg.mail_password),
        'has_from': bool(cfg.mail_from_email),
        'valid': False,
        'connection_verified': False,
        'verified': False,
        'encrypted': _mail_password_is_encrypted(cfg),
        'last_verified_at': cfg.mail_last_verified_at,
        'last_error_code': cfg.mail_last_error_code or '',
    }
    try:
        server, port, tls, ssl_enabled = _canonical_smtp_transport(
            cfg.mail_provider, cfg.mail_server, cfg.mail_port or 587,
            cfg.mail_use_tls, getattr(cfg, 'mail_use_ssl', False))
        _validated_email(cfg.mail_from_email, 'From email address', required=True)
        identity = _smtp_identity(
            cfg.mail_provider, server, port, tls, ssl_enabled, cfg.mail_username)
        status['has_server'] = bool(server)
        status['valid'] = bool(server and status['has_credentials'] and status['has_from'])
        verification_identity = _mail_verification_identity(identity, cfg.mail_from_email)
        status['connection_verified'] = bool(
            status['valid'] and cfg.mail_last_verification_status in {
                'connection_verified', 'delivery_verified'} and
            cfg.mail_last_verified_identity == verification_identity)
        status['verified'] = bool(
            status['valid'] and cfg.mail_last_verification_status == 'delivery_verified' and
            cfg.mail_last_verified_identity == verification_identity)
    except (ValueError, RuntimeError):
        pass
    return status


def _smtp_public_settings(cfg):
    try:
        server, port, tls, ssl_enabled = _canonical_smtp_transport(
            cfg.mail_provider, cfg.mail_server, cfg.mail_port or 587,
            cfg.mail_use_tls, getattr(cfg, 'mail_use_ssl', False))
    except ValueError:
        server = cfg.mail_server or ''
        port = cfg.mail_port or 587
        tls = bool(cfg.mail_use_tls)
        ssl_enabled = bool(getattr(cfg, 'mail_use_ssl', False))
    return {
        'provider': cfg.mail_provider,
        'server': server,
        'port': port,
        'use_tls': tls,
        'use_ssl': ssl_enabled,
    }


@app.cli.command('migrate-email-config')
def migrate_email_config_command():
    """Encrypt a legacy credential and normalize provider-owned transport values."""
    cfg = Configuration.query.first()
    if not cfg:
        raise click.ClickException('Application configuration does not exist.')

    changed = []
    try:
        if cfg.mail_password and not _mail_password_is_encrypted(cfg):
            cfg.mail_password = _encrypt_mail_password(cfg.mail_password)
            changed.append('credential_encrypted')
        elif cfg.mail_password:
            _decrypt_mail_password(cfg.mail_password)

        if cfg.mail_provider in SMTP_PROVIDER_PRESETS:
            server, port, use_tls, use_ssl = _canonical_smtp_transport(
                cfg.mail_provider,
                cfg.mail_server,
                cfg.mail_port or 587,
                cfg.mail_use_tls,
                getattr(cfg, 'mail_use_ssl', False),
            )
            current = (
                cfg.mail_server, cfg.mail_port,
                bool(cfg.mail_use_tls), bool(getattr(cfg, 'mail_use_ssl', False)),
            )
            normalized = (server, port, use_tls, use_ssl)
            if current != normalized:
                cfg.mail_server, cfg.mail_port = server, port
                cfg.mail_use_tls, cfg.mail_use_ssl = use_tls, use_ssl
                cfg.mail_last_verification_status = 'unverified'
                cfg.mail_last_verified_identity = ''
                cfg.mail_last_error_code = ''
                changed.append('transport_normalized')
        db.session.commit()
    except (ValueError, RuntimeError) as exc:
        db.session.rollback()
        raise click.ClickException(str(exc)) from exc

    click.echo(','.join(changed) if changed else 'already_current')


# ── PERMISSION DECORATOR ─────────────────────────────────────────────────────

def _sync_group_capabilities(group_id, overwrite_existing=False):
    """Backfill legacy grants without overwriting an explicit policy decision."""
    desired = set()
    for permission in GroupPermission.query.filter_by(group_id=group_id).all():
        if permission.access_level == 'full':
            desired.update(EXPLICIT_CAPABILITIES_BY_MODULE.get(
                (permission.module_key, 'full'), set()))
    managed = set().union(*EXPLICIT_CAPABILITIES_BY_MODULE.values())
    existing = {row.capability_key: row for row in GroupCapability.query.filter_by(
        group_id=group_id).filter(GroupCapability.capability_key.in_(managed)).all()}
    for capability in managed:
        row = existing.get(capability)
        should_grant = capability in desired
        if row and overwrite_existing:
            row.granted = should_grant
        elif not row:
            db.session.add(GroupCapability(
                group_id=group_id, capability_key=capability, granted=should_grant))


def require_capability(capability_key):
    if capability_key not in CAPABILITIES:
        raise RuntimeError(f'Unknown capability: {capability_key}')
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login', next=request.url))
            if not current_user.has_capability(capability_key):
                abort(403)
            return f(*args, **kwargs)
        return decorated
    return decorator

def require_module(module_key, level='limited'):
    capability = MODULE_CAPABILITIES.get((module_key, level))
    if not capability:
        raise RuntimeError(f'No capability mapping for {module_key}:{level}')
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login', next=request.url))
            if not current_user.has_capability(capability):
                flash('You do not have permission to access this section.', 'error')
                return redirect(url_for('dashboard'))
            return f(*args, **kwargs)
        return decorated
    return decorator


def require_admin(f):
    """Reserve identity, secret, and disaster-recovery operations for admins."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login', next=request.url))
        if not current_user.is_admin:
            abort(403)
        return f(*args, **kwargs)
    return decorated


# ── AUDIT HELPER ─────────────────────────────────────────────────────────────

def _audit(action, module, target='', details=''):
    try:
        uid   = current_user.id if current_user.is_authenticated else None
        uname = current_user.username if current_user.is_authenticated else 'system'
        ip    = request.remote_addr or '0.0.0.0'
        db.session.add(AuditLog(user_id=uid, username=uname, action=action,
                                module=module, target=target or '',
                                details=details or '', ip_address=ip))
        db.session.commit()
    except Exception:
        pass


# In-memory store for DB import jobs


# ── APSCHEDULER ──────────────────────────────────────────────────────────────

def _email_dedupe_key(kind, *parts):
    material = '\n'.join([kind, *(str(part) for part in parts)]).encode('utf-8')
    return f'{kind}:{hashlib.sha256(material).hexdigest()}'


def _enqueue_email(*, dedupe_key, kind, recipient_name, recipient_address,
                   subject, body, available_at=None, incident_record_id=None,
                   subscriber_id=None, group_id=None, group_name='', bus_id=None,
                   bus_label='', broadcast_job_id=None):
    """Add one durable delivery without duplicating an already-known message."""
    address = _validated_email(
        recipient_address, 'recipient email address', required=True).lower()
    existing = EmailOutbox.query.filter_by(dedupe_key=dedupe_key).first()
    if existing:
        return existing
    row = EmailOutbox(
        dedupe_key=dedupe_key,
        kind=kind,
        recipient_name=(recipient_name or '')[:160],
        recipient_address=address,
        subject=(subject or '')[:300],
        body=body or '',
        status='pending',
        attempts=0,
        max_attempts=app.config['EMAIL_OUTBOX_MAX_ATTEMPTS'],
        available_at=available_at or _utcnow(),
        incident_record_id=incident_record_id,
        subscriber_id=subscriber_id,
        group_id=group_id,
        group_name=(group_name or '')[:100],
        bus_id=bus_id,
        bus_label=(bus_label or '')[:80],
        broadcast_job_id=broadcast_job_id,
    )
    try:
        with db.session.begin_nested():
            db.session.add(row)
            db.session.flush()
        return row
    except IntegrityError:
        return EmailOutbox.query.filter_by(dedupe_key=dedupe_key).first()


def _complete_broadcast_if_ready(job_id):
    if not job_id:
        return
    job = db.session.get(BroadcastJob, job_id)
    if not job:
        return
    if (job.status != 'completed' and
            int(job.sent or 0) + int(job.failed or 0) >= int(job.total or 0)):
        job.status = 'completed'
        job.updated_at = _utcnow()
        owner = db.session.get(User, job.owner_id)
        db.session.add(AuditLog(
            user_id=job.owner_id,
            username=owner.username if owner else 'system',
            action='broadcast_completed',
            module='notifications',
            target=job.public_id,
            details=f'{job.sent} sent; {job.failed} failed',
            ip_address='background',
        ))


def _requeue_configuration_failures():
    """Make configuration-related terminal deliveries eligible after a config fix."""
    retryable_codes = {
        'authentication_rejected', 'sender_rejected', 'configuration_invalid',
        'password_required', 'tls_failed', 'connection_failed',
        'connection_timeout', 'server_disconnected', 'smtp_feature_unsupported',
        'smtp_rejected', 'smtp_temporary_failure', 'delivery_failed',
    }
    rows = EmailOutbox.query.filter(
        EmailOutbox.status == 'failed',
        EmailOutbox.last_error_code.in_(retryable_codes),
    ).all()
    now = _utcnow()
    for row in rows:
        row.status = 'retry'
        row.attempts = 0
        row.available_at = now
        row.locked_at = None
        row.updated_at = now
        if row.broadcast_job_id:
            job = db.session.get(BroadcastJob, row.broadcast_job_id)
            if job:
                job.failed = max(0, int(job.failed or 0) - 1)
                job.status = 'running'
                job.updated_at = now
    return len(rows)


def _record_outbox_notification(row, status, error_code=''):
    if not row.incident_record_id:
        return
    existing = NotificationLog.query.filter_by(
        incident_record_id=row.incident_record_id,
        channel='email',
        recipient_address=row.recipient_address,
    ).first()
    if existing:
        existing.status = status
        existing.error_message = error_code or None
        existing.sent_at = _utcnow()
        return
    db.session.add(NotificationLog(
        incident_record_id=row.incident_record_id,
        channel='email',
        recipient_name=row.recipient_name,
        recipient_address=row.recipient_address,
        subscriber_id=row.subscriber_id,
        group_id=row.group_id,
        group_name=row.group_name,
        bus_id=row.bus_id,
        bus_label=row.bus_label,
        status=status,
        error_message=error_code or None,
    ))


def _claim_due_email_ids():
    now = _utcnow()
    stale_before = now - timedelta(minutes=10)
    query = EmailOutbox.query.filter(
        db.or_(
            db.and_(EmailOutbox.status.in_(['pending', 'retry']),
                    EmailOutbox.available_at <= now),
            db.and_(EmailOutbox.status == 'processing',
                    EmailOutbox.locked_at < stale_before),
        )
    ).order_by(EmailOutbox.available_at, EmailOutbox.id)
    if str(db.engine.url).startswith('postgresql'):
        query = query.with_for_update(skip_locked=True)
    rows = query.limit(app.config['EMAIL_OUTBOX_BATCH_SIZE']).all()
    ids = []
    for row in rows:
        row.status = 'processing'
        row.locked_at = now
        row.updated_at = now
        ids.append(row.id)
    db.session.commit()
    return ids


def process_email_outbox():
    """Deliver a bounded durable batch and retain retry state across restarts."""
    with app.app_context():
        try:
            row_ids = _claim_due_email_ids()
        except Exception as exc:
            db.session.rollback()
            print(f'[EmailOutbox] claim error: {type(exc).__name__}')
            return

        for row_id in row_ids:
            row = db.session.get(EmailOutbox, row_id)
            if not row or row.status != 'processing':
                continue
            try:
                cfg = Configuration.query.first()
                if not cfg:
                    raise EmailTransportError(
                        'configuration_missing', 'Email configuration is missing.')
                settings = _smtp_settings_from_config(cfg)
                send_email(
                    settings,
                    subject=row.subject,
                    recipients=[row.recipient_address],
                    body=row.body,
                )
                row.attempts = int(row.attempts or 0) + 1
                row.status = 'sent'
                row.sent_at = _utcnow()
                row.locked_at = None
                row.last_error_code = ''
                _record_outbox_notification(row, 'sent')
                if row.broadcast_job_id:
                    job = db.session.get(BroadcastJob, row.broadcast_job_id)
                    if job:
                        job.status = 'running'
                        job.sent = int(job.sent or 0) + 1
                        job.updated_at = _utcnow()
                        _complete_broadcast_if_ready(job.public_id)
                db.session.commit()
            except Exception as exc:
                db.session.rollback()
                row = db.session.get(EmailOutbox, row_id)
                if not row:
                    continue
                if isinstance(exc, EmailTransportError):
                    failure = exc
                elif isinstance(exc, (ValueError, RuntimeError)):
                    failure = EmailTransportError(
                        'configuration_invalid', 'The saved email configuration is invalid.')
                else:
                    failure = EmailTransportError(
                        'delivery_failed',
                        'Email delivery failed for an unexpected internal reason.',
                        retryable=True,
                    )
                row.attempts = int(row.attempts or 0) + 1
                row.last_error_code = failure.code
                row.locked_at = None
                terminal = not failure.retryable or row.attempts >= row.max_attempts
                if terminal:
                    row.status = 'failed'
                    _record_outbox_notification(row, 'failed', failure.code)
                    if row.broadcast_job_id:
                        job = db.session.get(BroadcastJob, row.broadcast_job_id)
                        if job:
                            job.status = 'running'
                            job.failed = int(job.failed or 0) + 1
                            errors = job.errors
                            if len(errors) < 100:
                                errors.append(
                                    f'{_mask_email(row.recipient_address)}: {failure.code}')
                            job.errors_json = json.dumps(errors)
                            job.updated_at = _utcnow()
                            _complete_broadcast_if_ready(job.public_id)
                else:
                    delay = min(
                        app.config['EMAIL_OUTBOX_RETRY_MAX_SECONDS'],
                        app.config['EMAIL_OUTBOX_RETRY_BASE_SECONDS'] *
                        (2 ** max(0, row.attempts - 1)),
                    )
                    row.status = 'retry'
                    row.available_at = _utcnow() + timedelta(seconds=delay)
                row.updated_at = _utcnow()
                db.session.commit()


def _record_communication_event(rec):
    event_key = f'bus_status:{rec.id}'
    existing = CommunicationEvent.query.filter_by(event_key=event_key).first()
    if existing:
        return existing
    payload = {
        'schema_version': 1,
        'incident_record_id': rec.id,
        'bus_id': rec.bus_id,
        'bus_label': rec.bus.display_name,
        'incident_type_id': rec.incident_type_id,
        'status': rec.incident_type.name,
        'schedule_type_id': rec.schedule_type_id,
        'delay_minutes': int(rec.delay_minutes or 0),
        'eta': rec.eta or '',
        'incident_date': rec.incident_date.isoformat(),
    }
    event = CommunicationEvent(
        event_key=event_key, event_type='bus_status_committed',
        incident_record_id=rec.id,
        payload_json=json.dumps(payload, sort_keys=True), status='ready')
    try:
        with db.session.begin_nested():
            db.session.add(event)
            db.session.flush()
        return event
    except IntegrityError:
        return CommunicationEvent.query.filter_by(event_key=event_key).first()


def _commit_pending_incident_once(record_id):
    """Atomically claim and commit one pending incident across all workers."""
    committed_at = _utcnow()
    claimed = BusIncidentRecord.query.filter(
        BusIncidentRecord.id == record_id,
        BusIncidentRecord.is_pending.is_(True),
    ).update({
        BusIncidentRecord.is_pending: False,
        BusIncidentRecord.committed_at: committed_at,
        BusIncidentRecord.updated_at: committed_at,
    }, synchronize_session=False)
    if claimed != 1:
        db.session.rollback()
        return None
    rec = BusIncidentRecord.query.options(
        joinedload(BusIncidentRecord.bus),
        joinedload(BusIncidentRecord.incident_type),
    ).filter_by(id=record_id).one()
    _record_communication_event(rec)
    _send_bus_notifications(rec)
    db.session.commit()
    return rec


def commit_pending_incidents():
    with app.app_context():
        try:
            cfg = Configuration.query.first()
            delay = cfg.commit_delay_min if cfg else 5
            cutoff = _utcnow() - timedelta(minutes=delay)
            pending_ids = [record_id for (record_id,) in db.session.query(
                BusIncidentRecord.id).filter(
                    BusIncidentRecord.is_pending.is_(True),
                    BusIncidentRecord.created_at <= cutoff,
                ).order_by(BusIncidentRecord.created_at).all()]
            for record_id in pending_ids:
                try:
                    _commit_pending_incident_once(record_id)
                except Exception as exc:
                    db.session.rollback()
                    print(f'[Scheduler] incident {record_id} commit error: {type(exc).__name__}')
        except Exception as e:
            db.session.rollback()
            print(f'[Scheduler] commit error: {type(e).__name__}')

def _send_bus_notifications(rec):
    try:
        cfg = Configuration.query.first()
        if not cfg:
            return
        bus, it = rec.bus, rec.incident_type

        # ── Email ────────────────────────────────────────────────────────────
        email_enabled = bool(cfg.mail_server or cfg.mail_provider in SMTP_PROVIDER_PRESETS)

        subject = f"Bus Update: {bus.display_name}"
        email_body = (f"Bus {bus.display_name} — Status Update\n\n"
                      f"Status: {it.name}\n"
                      f"Delay: {rec.delay_minutes} minutes\n"
                      f"Notes: {rec.notes or 'N/A'}\n\n"
                      f"Sent by {cfg.app_name}")

        # ── SMS body (concise ≤160 chars) ────────────────────────────────────
        delay_part = f' +{rec.delay_minutes}min' if rec.delay_minutes else ''
        eta_part   = f' ETA {rec.eta}' if rec.eta else ''
        sms_body   = f"[{cfg.app_name}] {bus.display_name}: {it.name}{delay_part}{eta_part}"
        if len(sms_body) > 160:
            sms_body = sms_body[:157] + '...'
        sms_segments  = math.ceil(len(sms_body) / 160)
        cost_per_seg  = getattr(cfg, 'twilio_sms_cost_per_seg', 0.0079) or 0.0079

        twilio_on = (getattr(cfg, 'twilio_enabled', False) and TWILIO_AVAILABLE
                     and getattr(cfg, 'twilio_account_sid', '') and getattr(cfg, 'twilio_auth_token', '')
                     and getattr(cfg, 'twilio_from_number', ''))

        sent_emails = set()
        sent_phones = set()

        def _log(channel, name, address, sub, grp_id, grp_name, status, error=None,
                 sms_sid=None, segs=None, cost=None):
            try:
                db.session.add(NotificationLog(
                    incident_record_id=rec.id,
                    channel=channel,
                    recipient_name=name or '',
                    recipient_address=address or '',
                    subscriber_id=sub.id if sub else None,
                    group_id=grp_id,
                    group_name=grp_name or '',
                    bus_id=bus.id,
                    bus_label=bus.display_name,
                    status=status,
                    error_message=error,
                    sms_sid=sms_sid,
                    sms_segments=segs,
                    sms_cost_usd=cost,
                ))
                db.session.commit()
            except Exception as le:
                db.session.rollback()
                print(f'[NotifLog] log error: {le}')

        def _try_email(name, email_field, sub, grp_id, grp_name):
            if not email_field or not email_enabled: return
            for email in [e.strip() for e in email_field.split(',') if e.strip()]:
                if email in sent_emails: continue
                sent_emails.add(email)
                try:
                    _enqueue_email(
                        dedupe_key=_email_dedupe_key(
                            'bus_notification', rec.id, email.strip().lower()),
                        kind='bus_notification',
                        recipient_name=name,
                        recipient_address=email,
                        subject=subject,
                        body=email_body,
                        incident_record_id=rec.id,
                        subscriber_id=sub.id if sub else None,
                        group_id=grp_id,
                        group_name=grp_name,
                        bus_id=bus.id,
                        bus_label=bus.display_name,
                    )
                except ValueError:
                    _log('email', name, email, sub, grp_id, grp_name, 'failed',
                         error='invalid_recipient')

        def _try_sms(name, phone, sub, grp_id, grp_name):
            if not phone or phone in sent_phones or not twilio_on: return
            sent_phones.add(phone)
            try:
                tw = TwilioClient(cfg.twilio_account_sid, cfg.twilio_auth_token)
                msg = tw.messages.create(to=phone, from_=cfg.twilio_from_number, body=sms_body)
                cost = round(sms_segments * cost_per_seg, 6)
                _log('sms', name, phone, sub, grp_id, grp_name, 'sent',
                     sms_sid=msg.sid, segs=sms_segments, cost=cost)
            except Exception as e:
                print(f'[Notifications] SMS error to {phone}: {e}')
                _log('sms', name, phone, sub, grp_id, grp_name, 'failed', error=str(e))

        # Primary path: group-level bus assignment → contacts
        # Only notify groups whose period matches the incident's period (or has no period = all periods)
        group_ids = {a.group_id for a in
                     GroupBusAssignment.query.filter(
                         GroupBusAssignment.bus_id == rec.bus_id,
                         db.or_(
                             GroupBusAssignment.schedule_type_id == None,
                             GroupBusAssignment.schedule_type_id == rec.schedule_type_id
                         )
                     ).all()}
        if group_ids:
            groups_map = {g.id: g for g in
                          SubscriberGroup.query.filter(SubscriberGroup.id.in_(group_ids)).all()}
            subs = NotificationSubscriber.query.filter(
                NotificationSubscriber.active == True,
                NotificationSubscriber.group_id.in_(group_ids)
            ).all()
            for sub in subs:
                grp = groups_map.get(sub.group_id)
                grp_id   = grp.id   if grp else None
                grp_name = grp.name if grp else ''
                if sub.contacts:
                    for contact in sub.contacts:
                        _try_email(contact.full_name, contact.email, sub, grp_id, grp_name)
                        _try_sms(contact.full_name, contact.phone, sub, grp_id, grp_name)
                else:
                    _try_email(sub.full_name, sub.email, sub, grp_id, grp_name)
                    _try_sms(sub.full_name, sub.phone, sub, grp_id, grp_name)

        # Backward compat: direct NotificationBusAssignment (legacy records)
        for a in NotificationBusAssignment.query.filter_by(bus_id=rec.bus_id).all():
            s = a.subscriber
            if s.active:
                _try_email(s.full_name, s.email, s, None, '')
                _try_sms(s.full_name, s.phone, s, None, '')

    except Exception as e:
        print(f'[Notifications] send error: {e}')

_sched_lock_fh = None   # module-level ref keeps file lock alive for process lifetime

def _start_scheduler_once():
    """Start the scheduler in only ONE gunicorn worker using a file lock.
    Falls back to unconditional start on Windows (dev) where fcntl is unavailable."""
    global _sched_lock_fh
    if not SCHEDULER_AVAILABLE or os.environ.get('DISABLE_SCHEDULER') == '1':
        return
    sched = BackgroundScheduler(daemon=True)
    sched.add_job(commit_pending_incidents, 'interval', minutes=1,
                  id='commit_pending', coalesce=True, max_instances=1)
    sched.add_job(process_email_outbox, 'interval', seconds=10,
                  id='process_email_outbox', coalesce=True, max_instances=1)
    try:
        import fcntl
        _sched_lock_fh = open('/tmp/bustrack_sched.lock', 'w')
        fcntl.flock(_sched_lock_fh, fcntl.LOCK_EX | fcntl.LOCK_NB)
        sched.start()   # we got the lock → this worker owns the scheduler
        print('[Scheduler] started (pid=%d)' % os.getpid())
    except ImportError:
        sched.start()   # Windows / dev mode — no fcntl, just start it
    except (IOError, OSError):
        pass            # another worker already holds the lock → skip

_start_scheduler_once()


# ── SECURITY MIDDLEWARE ───────────────────────────────────────────────────────

_WIZARD_ENDPOINTS = {
    'install_wizard', 'install_test_db', 'install_run', 'static', 'health',
    'web_manifest', 'service_worker', 'offline_portal',
}
_PUBLIC_ENDPOINTS = {
    'index', 'api_buses', 'web_manifest', 'service_worker', 'offline_portal',
}

@app.before_request
def pre_request_checks():
    ep = request.endpoint
    if ep is None:
        return  # 404 — handled by error handler

    # 1. Redirect to wizard if not yet installed
    if ep not in _WIZARD_ENDPOINTS:
        if not is_installed():
            return redirect(url_for('install_wizard'))

    # 2. Revalidate account state and the session generation on every request.
    if current_user.is_authenticated:
        expected_version = int(current_user.session_version or 1)
        if (not current_user.active or
                session.get('session_version') != expected_version):
            logout_user()
            session.clear()
            if request.path.startswith('/admin/'):
                flash('Your session is no longer valid. Please sign in again.', 'error')
                return redirect(url_for('login'))

    # 3. CSRF validation on all admin state-changing requests
    if request.path.startswith('/admin/') and request.method == 'POST':
        token  = request.form.get('_csrf') or request.headers.get('X-CSRF-Token', '')
        stored = session.get('_csrf', '')
        if not (token and stored and secrets.compare_digest(str(token), str(stored))):
            abort(403)


@app.after_request
def security_headers(resp):
    resp.headers['X-Content-Type-Options']  = 'nosniff'
    resp.headers['X-Frame-Options']         = 'SAMEORIGIN'
    resp.headers['X-XSS-Protection']        = '1; mode=block'
    resp.headers['Referrer-Policy']         = 'strict-origin-when-cross-origin'
    resp.headers['Permissions-Policy']      = 'geolocation=(), microphone=(), camera=()'
    resp.headers['Cross-Origin-Opener-Policy'] = 'same-origin'
    resp.headers['Cross-Origin-Resource-Policy'] = 'same-origin'
    csp = (
        "default-src 'self'; "
        "base-uri 'self'; object-src 'none'; frame-ancestors 'self'; form-action 'self'; "
        "script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com "
        "https://cdn.jsdelivr.net; "
        "style-src 'self' 'unsafe-inline' https://cdnjs.cloudflare.com; "
        "font-src 'self' data: https://cdnjs.cloudflare.com; "
        "img-src 'self' data:; connect-src 'self'"
    )
    if app.config['CSP_REPORT_ONLY']:
        resp.headers['Content-Security-Policy-Report-Only'] = csp
    if app.config['CSP_ENFORCE']:
        resp.headers['Content-Security-Policy'] = csp
    if request.path.startswith('/admin/'):
        resp.headers['Cache-Control'] = 'no-store, private, max-age=0'
        resp.headers['Pragma'] = 'no-cache'
    if os.environ.get('FLASK_ENV') == 'production' and request.is_secure:
        resp.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    if request.path.startswith('/static/uploads/') and request.path.lower().endswith('.svg'):
        resp.headers['Content-Disposition'] = 'attachment'
        resp.headers['Content-Security-Policy'] = "sandbox; default-src 'none'"
    return resp


# ── ERROR HANDLERS ────────────────────────────────────────────────────────────

@app.errorhandler(403)
def err_403(e):
    return render_template('errors/403.html'), 403

@app.errorhandler(404)
def err_404(e):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def err_500(e):
    return render_template('errors/500.html'), 500


# ── HEALTH CHECK ──────────────────────────────────────────────────────────────

@app.route('/health')
def health():
    return 'ok', 200


# ── INSTALL WIZARD ────────────────────────────────────────────────────────────

@app.route('/install')
def install_wizard():
    if is_installed():
        abort(404)
    database_label = db.engine.url.get_backend_name().replace('postgresql', 'PostgreSQL').replace('sqlite', 'SQLite')
    return render_template('install/wizard.html',
                           install_ready=bool(os.environ.get('INSTALL_TOKEN', '').strip()),
                           database_label=database_label)


def _install_token_valid(data=None):
    """Fail closed unless the operator supplied the server-side bootstrap token."""
    expected = os.environ.get('INSTALL_TOKEN', '').strip()
    supplied = request.headers.get('X-Install-Token', '').strip()
    if not supplied and isinstance(data, dict):
        supplied = str(data.get('install_token', '')).strip()
    return bool(expected and supplied and
                secrets.compare_digest(expected.encode('utf-8'), supplied.encode('utf-8')))


def _install_guard(data=None):
    if is_installed():
        abort(404)
    if not os.environ.get('INSTALL_TOKEN', '').strip():
        return jsonify({'ok': False, 'message':
                        'Installation is locked. Configure INSTALL_TOKEN and restart the service.'}), 503
    if not _install_token_valid(data):
        return jsonify({'ok': False, 'message': 'Invalid installation authorization.'}), 403
    return None


@app.route('/install/test-db', methods=['POST'])
def install_test_db():
    """Test only the server-configured database; caller input cannot select a destination."""
    from sqlalchemy import text as sa_text
    data   = request.get_json(silent=True) or {}
    denied = _install_guard(data)
    if denied:
        return denied
    try:
        with db.engine.connect() as conn:
            conn.execute(sa_text('SELECT 1'))
        return jsonify({'ok': True, 'message': 'Connection successful.'})
    except Exception:
        return jsonify({'ok': False, 'message':
                        'Connection failed. Verify the approved database settings and server logs.'}), 400


@app.route('/install/run', methods=['POST'])
def install_run():
    data     = request.get_json(silent=True) or {}
    denied = _install_guard(data)
    if denied:
        return denied
    username_value = data.get('username', '')
    password_value = data.get('password', '')
    email_value = data.get('email', '')
    username = username_value.strip() if isinstance(username_value, str) else ''
    password = password_value if isinstance(password_value, str) else ''
    email = email_value.strip() if isinstance(email_value, str) else ''
    email = email or None

    # Validate admin credentials
    if not 3 <= len(username) <= 80 or not re.fullmatch(r'[A-Za-z0-9_.-]+', username):
        return jsonify({'ok': False, 'message':
                        'Username must be 3–80 characters using letters, numbers, dots, hyphens or underscores.'}), 400
    password_error = _password_error(password)
    if password_error:
        return jsonify({'ok': False, 'message': password_error}), 400
    if email and (len(email) > 120 or not re.fullmatch(r'[^\s@]+@[^\s@]+\.[^\s@]+', email)):
        return jsonify({'ok': False, 'message': 'Enter a valid email address.'}), 400

    install_lock = os.path.join(INSTANCE_DIR, '.installing')
    try:
        lock_fd = os.open(install_lock, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
        os.write(lock_fd, str(os.getpid()).encode('ascii'))
    except FileExistsError:
        return jsonify({'ok': False, 'message': 'Installation is already in progress.'}), 409

    try:
        # Persist only a development-generated session key. Explicit deployment
        # configuration remains owned by the operator and is never rewritten here.
        if _secret_generated:
            _write_instance_env(app.config['SECRET_KEY'])

        # Create all tables and default data
        db.create_all()
        _migrate_add_columns()
        _seed_defaults()

        # Create admin user
        ag = UserGroup.query.filter_by(is_admin=True).first()
        if not ag:
            ag = UserGroup(name='Administrator', description='Full system access', is_admin=True)
            db.session.add(ag); db.session.commit()
        u = User(username=username, email=email,
                 first_name='Admin', group_id=ag.id, active=True)
        u.set_password(password)
        db.session.add(u)
        db.session.commit()

        _mark_installed()
        return jsonify({'ok': True, 'message': 'Installation complete. Redirecting to login…'})
    except (ValueError, TypeError) as e:
        db.session.rollback()
        return jsonify({'ok': False, 'message': str(e)}), 400
    except Exception:
        db.session.rollback()
        return jsonify({'ok': False, 'message':
                        'Installation failed. Review the server log; no installation lock was created.'}), 500
    finally:
        os.close(lock_fd)
        try:
            os.remove(install_lock)
        except FileNotFoundError:
            pass


def _write_instance_env(secret_key):
    """Atomically persist one setting without discarding existing instance config."""
    os.makedirs(INSTANCE_DIR, exist_ok=True)
    env_path = os.path.join(INSTANCE_DIR, '.env')
    existing_lines = []
    try:
        with open(env_path, 'r', encoding='utf-8') as current:
            existing_lines = current.readlines()
    except FileNotFoundError:
        pass

    replacement = f'SECRET_KEY={json.dumps(secret_key)}\n'
    updated_lines = []
    replaced = False
    setting_pattern = re.compile(r'^\s*(?:export\s+)?SECRET_KEY\s*=')
    for line in existing_lines:
        if setting_pattern.match(line):
            if not replaced:
                updated_lines.append(replacement)
                replaced = True
            continue
        updated_lines.append(line)
    if not replaced:
        if updated_lines and not updated_lines[-1].endswith(('\n', '\r')):
            updated_lines[-1] += '\n'
        updated_lines.append(replacement)

    temp_fd, temp_path = tempfile.mkstemp(prefix='.env.', dir=INSTANCE_DIR, text=True)
    try:
        os.fchmod(temp_fd, 0o600)
        with os.fdopen(temp_fd, 'w', encoding='utf-8') as temp_file:
            temp_fd = None
            temp_file.writelines(updated_lines)
            temp_file.flush()
            os.fsync(temp_file.fileno())
        os.replace(temp_path, env_path)
        os.chmod(env_path, 0o600)
        dir_fd = os.open(INSTANCE_DIR, os.O_RDONLY)
        try:
            os.fsync(dir_fd)
        finally:
            os.close(dir_fd)
    except Exception:
        if temp_fd is not None:
            os.close(temp_fd)
        try:
            os.remove(temp_path)
        except FileNotFoundError:
            pass
        raise


def _seed_defaults():
    """Insert default groups, schedule types, incident types, delay reasons, schedule and config."""
    # Groups
    if not UserGroup.query.filter_by(name='Administrator').first():
        db.session.add(UserGroup(name='Administrator', description='Full system access', is_admin=True))
        db.session.commit()
    if not UserGroup.query.filter_by(name='Staff Member').first():
        sg = UserGroup(name='Staff Member', description='Limited operational access', is_admin=False)
        db.session.add(sg); db.session.commit()
        sg = UserGroup.query.filter_by(name='Staff Member').first()
        for mod in ['buses', 'incidents', 'statistics']:
            db.session.add(GroupPermission(group_id=sg.id, module_key=mod, access_level='full'))
        for mod in ['users', 'notifications', 'config']:
            db.session.add(GroupPermission(group_id=sg.id, module_key=mod, access_level='none'))
        db.session.commit()
    # Config singleton
    if not Configuration.query.first():
        db.session.add(Configuration()); db.session.commit()
    # Schedule types (with default time windows)
    for name, label, order, w_start, w_end in [
        ('Morning',   '7:00 AM',  0, '06:00', '11:30'),
        ('Midday',    '12:00 PM', 1, '11:30', '14:00'),
        ('Afternoon', '3:00 PM',  2, '14:00', '19:00'),
    ]:
        existing = BusScheduleType.query.filter_by(name=name).first()
        if not existing:
            db.session.add(BusScheduleType(name=name, time_label=label, sort_order=order,
                                           window_start=w_start, window_end=w_end))
        elif not existing.window_start:
            existing.window_start = w_start
            existing.window_end   = w_end
    db.session.commit()
    # Incident types
    for name, color, icon, is_def, is_sys, order, priority in [
        ('On Time','#10b981','fa-check-circle',True,True,0,0),
        ('Delayed','#f59e0b','fa-clock',False,True,1,70),
        ('E-Learning','#8b5cf6','fa-laptop',False,True,2,20),
        ('Combined','#3b82f6','fa-link',False,True,3,40),
        ('Double-back','#06b6d4','fa-redo',False,True,4,60),
        ('Out of Service','#ef4444','fa-ban',False,True,5,100),
        ('Combined/Delayed','#f97316','fa-exclamation-triangle',False,True,6,85),
    ]:
        if not IncidentType.query.filter_by(name=name).first():
            db.session.add(IncidentType(name=name, color=color, icon=icon,
                                        is_default=is_def, is_system=is_sys,
                                        sort_order=order,
                                        operational_priority=priority))
    db.session.commit()
    # Delay reasons
    for reason, order in [('Traffic congestion',0),('Road construction',1),('Weather conditions',2),
                           ('Mechanical issue',3),('Driver delay',4),('Student boarding delay',5),
                           ('Accident on route',6),('Detour required',7)]:
        if not DelayReason.query.filter_by(reason=reason).first():
            db.session.add(DelayReason(reason=reason, sort_order=order))
    db.session.commit()
    # Operational schedule
    if not OperationalSchedule.query.first():
        db.session.add(OperationalSchedule(name='Weekday Service', days='mon-fri',
                                           start_time='06:30', end_time='18:00', is_active=True))
        db.session.commit()


# ── PUBLIC ROUTES ─────────────────────────────────────────────────────────────

@app.route('/manifest.webmanifest')
def web_manifest():
    response = send_file(
        os.path.join(BASE_DIR, 'static', 'manifest.webmanifest'),
        mimetype='application/manifest+json',
        conditional=True,
    )
    response.headers['Cache-Control'] = 'public, max-age=3600'
    return response


@app.route('/service-worker.js')
def service_worker():
    response = send_file(
        os.path.join(BASE_DIR, 'static', 'js', 'service-worker.js'),
        mimetype='application/javascript',
        conditional=True,
    )
    response.headers['Cache-Control'] = 'no-cache, max-age=0'
    response.headers['Service-Worker-Allowed'] = '/'
    return response


@app.route('/offline')
def offline_portal():
    cfg = get_config()
    response = make_response(render_template('public/offline.html', cfg=cfg))
    response.headers['Cache-Control'] = 'public, max-age=300'
    return response

def _public_bus_payload(item, cfg):
    bus = item['bus']
    status = item['status']
    status_name = status.name if status else 'On Time'
    incidents = []
    for incident in item['incidents']:
        incidents.append({
            'id': incident.id,
            'type': incident.incident_type.name,
            'type_label': public_status_label(incident.incident_type.name, cfg.lang_frontend),
            'color': incident.incident_type.color,
            'icon': incident.incident_type.icon,
            'delay': incident.delay_minutes or 0,
            'eta': incident.eta or '',
            'reason': (incident.delay_reason.reason
                       if incident.delay_reason_id and incident.delay_reason
                       else incident.delay_reason_text or ''),
            'notes': incident.notes or '',
            'created_at': incident.created_at.isoformat() if incident.created_at else '',
            'schedule': incident.schedule_type.name if incident.schedule_type else '',
        })
    assignments = [{
        'id': assignment.schedule_type_id,
        'name': assignment.schedule_type.name,
        'label': public_schedule_label(assignment.schedule_type.name, cfg.lang_frontend),
        'departure_time': assignment.departure_time or '',
    } for assignment in item['schedule_assignments']]
    return {
        'id': bus.id,
        'identifier': bus.identifier,
        'name': bus.name,
        'display_name': bus.display_name,
        'route': bus.route or '',
        'capacity': bus.capacity,
        'description': bus.description or '',
        'status': {
            'name': status_name,
            'label': public_status_label(status_name, cfg.lang_frontend),
            'color': status.color if status else '#10b981',
            'icon': status.icon if status else 'fa-check-circle',
            'is_default': bool(status.is_default) if status else True,
        },
        'delay_minutes': item['delay'] or 0,
        'incidents': incidents,
        'schedules': assignments,
    }


def _public_state(operational, current_period, buses_data, today_value, cfg):
    buses = [_public_bus_payload(item, cfg) for item in buses_data]
    period = None
    if current_period:
        period = {
            'id': current_period.id,
            'name': current_period.name,
            'label': public_schedule_label(current_period.name, cfg.lang_frontend),
            'window_start': current_period.window_start or '',
            'window_end': current_period.window_end or '',
        }
    stable = {
        'operational': bool(operational),
        'district_date': today_value.isoformat(),
        'current_period': period,
        'buses': buses,
    }
    revision = hashlib.sha256(json.dumps(
        stable, sort_keys=True, separators=(',', ':'), ensure_ascii=False
    ).encode('utf-8')).hexdigest()
    stable['attention_count'] = sum(
        1 for bus in buses if not bus['status']['is_default'])
    stable['revision'] = revision
    return stable

@app.route('/')
def index():
    cfg = get_config()
    operational, offline_msg = is_operational()
    current_period = get_current_period() if operational else None
    buses_data     = bus_list_today(period=current_period) if operational else []
    incident_types = IncidentType.query.order_by(IncidentType.sort_order).all()
    schedule_types = BusScheduleType.query.order_by(BusScheduleType.sort_order).all()
    today_dt = district_today(cfg)
    # Holiday for today (for richer offline display)
    today_holiday = Holiday.query.filter_by(
        holiday_date=today_dt, is_active=True).first() if not operational else None
    # Upcoming holidays in the next 7 days (for advance announcement)
    upcoming_holidays = Holiday.query.filter(
        Holiday.is_active == True,
        Holiday.holiday_date > today_dt,
        Holiday.holiday_date <= today_dt + timedelta(days=7)
    ).order_by(Holiday.holiday_date).all()
    public_state = _public_state(
        operational, current_period, buses_data, today_dt, cfg)
    return render_template('public/index.html',
                           buses_data=buses_data, incident_types=incident_types,
                           schedule_types=schedule_types, cfg=cfg,
                           current_period=current_period,
                           operational=operational, offline_msg=offline_msg,
                           today=today_dt,
                           today_holiday=today_holiday,
                           upcoming_holidays=upcoming_holidays,
                           public_state=public_state,
                           portal_i18n=TRANSLATIONS.get(
                               cfg.lang_frontend, TRANSLATIONS['en']),
                           portal_time_zone=district_timezone(cfg).zone)

@app.route('/api/buses')
def api_buses():
    cfg = get_config()
    operational, _ = is_operational()
    current_period = get_current_period() if operational else None
    buses_data = bus_list_today(period=current_period) if operational else []
    state = _public_state(
        operational, current_period, buses_data, district_today(cfg), cfg)
    revision = state['revision']
    if request.if_none_match.contains(revision):
        response = make_response('', 304)
    else:
        state['generated_at'] = district_now(cfg).isoformat()
        if request.args.get('render') == '1':
            state['cards_html'] = render_template(
                'public/_bus_cards.html', buses_data=buses_data, cfg=cfg)
        response = jsonify(state)
    response.set_etag(revision)
    response.headers['Cache-Control'] = 'no-cache, max-age=0'
    return response


# ── AUTH ROUTES ───────────────────────────────────────────────────────────────

def _login_throttle_key(ip, identifier):
    material = f'{ip}\0{identifier.casefold()}'.encode('utf-8', 'replace')
    return hmac.new(app.config['SECRET_KEY'].encode('utf-8'), material,
                    hashlib.sha256).hexdigest()


def _login_throttle_state(ip, identifier):
    key = _login_throttle_key(ip, identifier)
    record = LoginThrottle.query.filter_by(throttle_key=key).first()
    now = _utcnow()
    if record and record.locked_until and record.locked_until > now:
        return record, True
    return record, False


def _record_login_failure(ip, identifier):
    key = _login_throttle_key(ip, identifier)
    now = _utcnow()
    window = timedelta(seconds=app.config['LOGIN_RATE_LIMIT_WINDOW_SECONDS'])
    lock_for = timedelta(seconds=app.config['LOGIN_RATE_LIMIT_LOCK_SECONDS'])
    stale_before = now - max(window, lock_for) * 2
    LoginThrottle.query.filter(LoginThrottle.updated_at < stale_before).delete(
        synchronize_session=False)
    record = LoginThrottle.query.filter_by(throttle_key=key).with_for_update().first()
    if not record:
        record = LoginThrottle(throttle_key=key, failed_count=0,
                               window_started_at=now)
        db.session.add(record)
    elif now - record.window_started_at >= window:
        record.failed_count = 0
        record.window_started_at = now
        record.locked_until = None
    record.failed_count += 1
    if record.failed_count >= app.config['LOGIN_RATE_LIMIT_ATTEMPTS']:
        record.locked_until = now + lock_for
    try:
        db.session.commit()
    except IntegrityError:
        # A concurrent worker created the same key. Retry against that row.
        db.session.rollback()
        record = LoginThrottle.query.filter_by(throttle_key=key).with_for_update().one()
        if now - record.window_started_at >= window:
            record.failed_count = 0
            record.window_started_at = now
        record.failed_count += 1
        if record.failed_count >= app.config['LOGIN_RATE_LIMIT_ATTEMPTS']:
            record.locked_until = now + lock_for
        db.session.commit()


def _clear_login_failures(ip, identifier):
    key = _login_throttle_key(ip, identifier)
    LoginThrottle.query.filter_by(throttle_key=key).delete()
    db.session.commit()


def _safe_local_redirect(target):
    if not target:
        return None
    parsed = urlsplit(target)
    if parsed.scheme or parsed.netloc or not target.startswith('/') or target.startswith('//'):
        return None
    return target


def _issue_incident_request_token():
    """Issue a short-lived, session-bound idempotency token for one incident form."""
    token = secrets.token_hex(24)
    tokens = [value for value in session.get('_incident_request_tokens', [])
              if isinstance(value, str) and re.fullmatch(r'[0-9a-f]{48}', value)]
    tokens.append(token)
    session['_incident_request_tokens'] = tokens[-12:]
    session.modified = True
    return token


def _consume_incident_request_token(token):
    if not isinstance(token, str) or not re.fullmatch(r'[0-9a-f]{48}', token):
        return False
    tokens = list(session.get('_incident_request_tokens', []))
    matched = next((value for value in tokens
                    if isinstance(value, str) and secrets.compare_digest(value, token)), None)
    if matched is None:
        return False
    tokens.remove(matched)
    session['_incident_request_tokens'] = tokens
    session.modified = True
    return True

@app.route('/admin/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        ip  = request.remote_addr or '0.0.0.0'
        identifier = request.form.get('username', '').strip()
        _, locked = _login_throttle_state(ip, identifier)
        if locked:
            wait_minutes = math.ceil(app.config['LOGIN_RATE_LIMIT_LOCK_SECONDS'] / 60)
            flash(f'Too many failed attempts. Please wait {wait_minutes} minute(s) and try again.', 'error')
            return render_template('admin/login.html'), 429

        password   = request.form.get('password', '')
        user = User.query.filter_by(username=identifier).first()
        if not user:
            user = User.query.filter_by(email=identifier, use_email_auth=True).first()
        if user and user.check_password(password) and user.active:
            _clear_login_failures(ip, identifier)
            login_user(user)
            session['session_version'] = int(user.session_version or 1)
            session.permanent = True
            user.last_login = _utcnow()
            db.session.commit()
            _audit('login', 'auth', user.username)
            # Prevent open-redirect: only allow relative next URLs
            next_url = _safe_local_redirect(request.args.get('next', ''))
            if next_url:
                return redirect(next_url)
            return redirect(url_for('dashboard'))
        _record_login_failure(ip, identifier)
        _audit('login_failed', 'auth', identifier or '(empty)')
        flash('Invalid credentials. Please try again.', 'error')
    return render_template('admin/login.html')

@app.route('/admin/logout', methods=['POST'])
@login_required
def logout():
    _audit('logout', 'auth', current_user.username)
    logout_user()
    session.clear()
    return redirect(url_for('index'))


# ── DASHBOARD ─────────────────────────────────────────────────────────────────

def _dashboard_natural_key(*parts):
    text_value = ' '.join(str(part or '') for part in parts).casefold()
    return tuple(int(piece) if piece.isdigit() else piece
                 for piece in re.split(r'(\d+)', text_value))


def _dashboard_incident_payload(record, cfg, now_utc):
    if record.delay_reason_id and record.delay_reason:
        reason = record.delay_reason.reason
    else:
        reason = record.delay_reason_text or ''
    pending_until = None
    pending_remaining = 0
    if record.is_pending:
        pending_until = record.created_at + timedelta(minutes=cfg.commit_delay_min or 0)
        pending_remaining = max(0, int((pending_until - now_utc).total_seconds()))
    return {
        'id': record.id,
        'type_id': record.incident_type_id,
        'type': record.incident_type.name,
        'color': record.incident_type.color,
        'icon': record.incident_type.icon,
        'priority': int(record.incident_type.operational_priority or 0),
        'is_default': bool(record.incident_type.is_default),
        'delay': int(record.delay_minutes or 0),
        'eta': record.eta or '',
        'eta_label': fmt_time(record.eta, cfg.time_format),
        'reason': reason,
        'reason_id': record.delay_reason_id,
        'reason_text': record.delay_reason_text or '',
        'notes': record.notes or '',
        'schedule': record.schedule_type.name if record.schedule_type else '',
        'schedule_id': record.schedule_type_id,
        'pending': bool(record.is_pending),
        'pending_remaining_seconds': pending_remaining,
        'pending_until_utc': (
            pending_until.replace(tzinfo=timezone.utc).isoformat()
            if pending_until else ''),
        'created_label': format_district_datetime(record.created_at, cfg, '%I:%M %p'),
        'created_by': record.created_by.username if record.created_by else '',
        'version': f'{record.id}:{record.updated_at.isoformat() if record.updated_at else ""}',
    }


def _build_dashboard_snapshot(period, date_from, date_to, *, can_view_buses,
                              can_view_statistics, can_view_notifications):
    """Build one permission-aware read model for the operational dashboard."""
    cfg = get_config()
    today = district_today(cfg)
    now_local = district_now(cfg)
    now_utc = _utcnow()
    d_from, d_to = _parse_period(period, date_from, date_to, today)
    current_period = get_current_period(cfg)
    snapshot = {
        'cfg': cfg,
        'today': today,
        'now_local': now_local,
        'timezone': district_timezone(cfg).zone,
        'period': period,
        'date_from': d_from,
        'date_to': d_to,
        'current_period': current_period,
        'buses': [],
        'attention_buses': [],
        'on_time_buses': [],
        'total_buses': 0,
        'on_time_count': 0,
        'attention_count': 0,
        'pending_today': 0,
        'pending_queue': [],
        'schedule_warning_count': 0,
        'period_incidents': 0,
        'period_pending': 0,
        'period_delay_minutes': 0,
        'period_average_delay': 0,
        'by_type': {},
        'by_type_colors': [],
        'by_bus': {},
        'by_day': {},
        'recent': [],
    }

    if can_view_buses:
        buses = Bus.query.options(
            selectinload(Bus.schedule_assignments).selectinload(
                BusScheduleAssignment.schedule_type),
        ).filter_by(active=True).order_by(Bus.identifier, Bus.name).all()
        bus_ids = [bus.id for bus in buses]
        today_records = []
        if bus_ids:
            today_records = BusIncidentRecord.query.options(
                joinedload(BusIncidentRecord.incident_type),
                joinedload(BusIncidentRecord.schedule_type),
                joinedload(BusIncidentRecord.delay_reason),
                joinedload(BusIncidentRecord.created_by),
            ).filter(
                BusIncidentRecord.bus_id.in_(bus_ids),
                BusIncidentRecord.incident_date == today,
            ).order_by(BusIncidentRecord.created_at.desc()).all()
        records_by_bus = {}
        for record in today_records:
            records_by_bus.setdefault(record.bus_id, []).append(record)

        default_type = IncidentType.query.filter_by(is_default=True).first()
        group_counts = {}
        group_metadata = {}
        group_schools = {}
        if can_view_notifications and bus_ids:
            group_rows = db.session.query(
                GroupBusAssignment.bus_id, SubscriberGroup.id, SubscriberGroup.name,
            ).join(SubscriberGroup).filter(
                GroupBusAssignment.bus_id.in_(bus_ids)).order_by(
                    SubscriberGroup.name).all()
            group_ids = {group_id for _bus_id, group_id, _name in group_rows}
            for bus_id, group_id, group_name in group_rows:
                group_metadata.setdefault(bus_id, {})[group_id] = group_name
            group_counts = {
                bus_id: len(groups) for bus_id, groups in group_metadata.items()
            }
            if group_ids:
                for group_id, school in db.session.query(
                        NotificationSubscriber.group_id,
                        NotificationSubscriber.school).filter(
                            NotificationSubscriber.active.is_(True),
                            NotificationSubscriber.group_id.in_(group_ids),
                            NotificationSubscriber.school.isnot(None),
                        ).distinct().all():
                    if school:
                        group_schools.setdefault(group_id, set()).add(school)

        bus_payloads = []
        for bus in buses:
            records = records_by_bus.get(bus.id, [])
            period_records = records
            if current_period:
                period_records = [record for record in records
                                  if record.schedule_type_id == current_period.id]
            latest = period_records[0] if period_records else None
            status = latest.incident_type if latest else default_type
            status_is_default = not status or bool(status.is_default)
            assignments = sorted(
                bus.schedule_assignments,
                key=lambda assignment: (
                    assignment.schedule_type.sort_order,
                    assignment.schedule_type.name,
                ),
            )
            warnings = [schedule_assignment_warning(assignment)
                        for assignment in assignments]
            warnings = [warning for warning in warnings if warning]
            incident_payloads = [
                _dashboard_incident_payload(record, cfg, now_utc)
                for record in records
            ]
            latest_payload = (_dashboard_incident_payload(latest, cfg, now_utc)
                              if latest else None)
            bus_groups = group_metadata.get(bus.id, {})
            bus_schools = sorted({
                school for group_id in bus_groups
                for school in group_schools.get(group_id, set())
            }, key=_dashboard_natural_key)
            priority = int(status.operational_priority or 0) if status else 0
            payload = {
                'id': bus.id,
                'identifier': bus.identifier,
                'name': bus.name,
                'display_name': f'{bus.identifier} - {bus.name}',
                'route': bus.route or '',
                'capacity': bus.capacity,
                'description': bus.description or '',
                'status': status.name if status else 'On Time',
                'status_color': status.color if status else '#10b981',
                'status_icon': status.icon if status else 'fa-check-circle',
                'status_priority': priority,
                'is_attention': not status_is_default,
                'delay': int(latest.delay_minutes or 0) if latest else 0,
                'eta': latest.eta if latest else '',
                'eta_label': fmt_time(latest.eta, cfg.time_format) if latest else '',
                'reason': latest_payload['reason'] if latest_payload else '',
                'latest': latest_payload,
                'version': latest_payload['version'] if latest_payload else '0',
                'incidents': incident_payloads,
                'schedules': [{
                    'id': assignment.schedule_type_id,
                    'name': assignment.schedule_type.name,
                    'departure_time': assignment.departure_time or '',
                    'departure_label': fmt_time(assignment.departure_time, cfg.time_format),
                } for assignment in assignments],
                'schedule_names': ', '.join(
                    assignment.schedule_type.name for assignment in assignments),
                'schedule_warnings': warnings,
                'group_count': int(group_counts.get(bus.id, 0)),
                'group_ids': list(bus_groups),
                'group_names': list(bus_groups.values()),
                'school_names': bus_schools,
            }
            bus_payloads.append(payload)

        attention = sorted(
            (item for item in bus_payloads if item['is_attention']),
            key=lambda item: (-item['status_priority'], -item['delay'],
                              _dashboard_natural_key(item['identifier'], item['name'])),
        )
        on_time = sorted(
            (item for item in bus_payloads if not item['is_attention']),
            key=lambda item: _dashboard_natural_key(item['identifier'], item['name']),
        )
        pending_queue = sorted([
            {
                **_dashboard_incident_payload(record, cfg, now_utc),
                'bus_id': record.bus_id,
                'bus_label': next((
                    item['display_name'] for item in bus_payloads
                    if item['id'] == record.bus_id), ''),
            }
            for record in today_records
            if record.is_pending
        ], key=lambda item: (
            -item['priority'], item['pending_remaining_seconds'],
            _dashboard_natural_key(item['bus_label'])))
        snapshot.update({
            'buses': attention + on_time,
            'attention_buses': attention,
            'on_time_buses': on_time,
            'total_buses': len(bus_payloads),
            'attention_count': len(attention),
            'on_time_count': len(on_time),
            'pending_today': len(pending_queue),
            'schedule_warning_count': sum(
                len(item['schedule_warnings']) for item in bus_payloads),
            'pending_queue': pending_queue,
        })

    if can_view_statistics:
        actual_filter = IncidentType.is_default.is_(False)
        base_filters = (
            BusIncidentRecord.incident_date >= d_from,
            BusIncidentRecord.incident_date <= d_to,
        )
        totals = db.session.query(
            func.count(BusIncidentRecord.id),
            func.coalesce(func.sum(case(
                (BusIncidentRecord.is_pending.is_(True), 1), else_=0)), 0),
            func.coalesce(func.sum(BusIncidentRecord.delay_minutes), 0),
            func.avg(case(
                (BusIncidentRecord.delay_minutes > 0,
                 BusIncidentRecord.delay_minutes), else_=None)),
        ).join(IncidentType).filter(*base_filters, actual_filter).one()
        by_type_rows = db.session.query(
            IncidentType.name, IncidentType.color,
            func.count(BusIncidentRecord.id),
        ).join(BusIncidentRecord).filter(*base_filters, actual_filter).group_by(
            IncidentType.id, IncidentType.name, IncidentType.color,
            IncidentType.operational_priority).order_by(
                IncidentType.operational_priority.desc(), IncidentType.name).all()
        by_bus_rows = db.session.query(
            Bus.id, Bus.identifier, Bus.name, func.count(BusIncidentRecord.id),
        ).join(BusIncidentRecord).join(IncidentType).filter(
            *base_filters, actual_filter).group_by(
                Bus.id, Bus.identifier, Bus.name).order_by(
                    func.count(BusIncidentRecord.id).desc(),
                    Bus.identifier, Bus.name).limit(12).all()
        by_day_rows = db.session.query(
            BusIncidentRecord.incident_date, func.count(BusIncidentRecord.id),
        ).join(IncidentType).filter(*base_filters, actual_filter).group_by(
            BusIncidentRecord.incident_date).order_by(
                BusIncidentRecord.incident_date).all()
        recent = BusIncidentRecord.query.options(
            joinedload(BusIncidentRecord.bus),
            joinedload(BusIncidentRecord.incident_type),
            joinedload(BusIncidentRecord.schedule_type),
            joinedload(BusIncidentRecord.created_by),
        ).join(IncidentType).filter(
            *base_filters, actual_filter).order_by(
                BusIncidentRecord.created_at.desc()).limit(15).all()
        snapshot.update({
            'period_incidents': int(totals[0] or 0),
            'period_pending': int(totals[1] or 0),
            'period_delay_minutes': int(totals[2] or 0),
            'period_average_delay': round(float(totals[3] or 0), 1),
            'by_type': {name: int(count) for name, _color, count in by_type_rows},
            'by_type_colors': [color for _name, color, _count in by_type_rows],
            'by_bus': {
                f'{identifier} - {name}': int(count)
                for _bus_id, identifier, name, count in by_bus_rows
            },
            'by_day': {day.isoformat(): int(count) for day, count in by_day_rows},
            'recent': recent,
        })
    return snapshot


def _dashboard_operations_revision(snapshot):
    material = {
        'period_id': snapshot['current_period'].id if snapshot['current_period'] else None,
        'buses': [{
            'id': bus['id'], 'version': bus['version'], 'status': bus['status'],
            'schedule_names': bus['schedule_names'],
            'group_ids': bus['group_ids'], 'school_names': bus['school_names'],
        } for bus in snapshot['buses']],
        'pending': [item['version'] for item in snapshot['pending_queue']],
    }
    return hashlib.sha256(json.dumps(
        material, sort_keys=True, separators=(',', ':')).encode('utf-8')).hexdigest()


def _dashboard_operations_payload(snapshot):
    return {
        'revision': _dashboard_operations_revision(snapshot),
        'generated_at': _utcnow().replace(tzinfo=timezone.utc).isoformat(),
        'generated_label': snapshot['now_local'].strftime('%-I:%M:%S %p'),
        'timezone': snapshot['timezone'],
        'current_period_id': (
            snapshot['current_period'].id if snapshot['current_period'] else None),
        'current_period_name': (
            snapshot['current_period'].name if snapshot['current_period'] else ''),
        'attention_count': snapshot['attention_count'],
        'pending_count': snapshot['pending_today'],
        'on_time_count': snapshot['on_time_count'],
        'total_buses': snapshot['total_buses'],
        'schedule_warning_count': snapshot['schedule_warning_count'],
        'buses': snapshot['buses'],
        'pending_queue': snapshot['pending_queue'],
    }


def _dashboard_recipient_preview(bus_ids, schedule_type_id=None):
    """Return aggregate, non-PII delivery scope for selected buses."""
    bus_ids = {int(bus_id) for bus_id in bus_ids}
    active_bus_ids = {
        bus_id for (bus_id,) in db.session.query(Bus.id).filter(
            Bus.active.is_(True), Bus.id.in_(bus_ids)).all()
    }
    if not active_bus_ids:
        return {
            'bus_count': 0, 'subscriber_count': 0, 'contact_count': 0,
            'email_count': 0, 'sms_count': 0, 'roles': {}, 'languages': {},
            'schools': {}, 'groups': [], 'buses_without_recipients': 0,
        }

    group_query = db.session.query(
        GroupBusAssignment.bus_id, SubscriberGroup.id, SubscriberGroup.name,
    ).join(SubscriberGroup).filter(
        GroupBusAssignment.bus_id.in_(active_bus_ids))
    if schedule_type_id:
        group_query = group_query.filter(or_(
            GroupBusAssignment.schedule_type_id.is_(None),
            GroupBusAssignment.schedule_type_id == schedule_type_id,
        ))
    group_rows = group_query.all()
    group_ids = {group_id for _bus_id, group_id, _name in group_rows}
    group_names = {group_id: name for _bus_id, group_id, name in group_rows}
    subscribers = []
    if group_ids:
        subscribers.extend(NotificationSubscriber.query.filter(
            NotificationSubscriber.active.is_(True),
            NotificationSubscriber.group_id.in_(group_ids),
        ).all())
    direct_subscriber_ids = {
        subscriber_id for _bus_id, subscriber_id in db.session.query(
            NotificationBusAssignment.bus_id,
            NotificationBusAssignment.subscriber_id,
        ).filter(NotificationBusAssignment.bus_id.in_(active_bus_ids)).all()
    }
    if direct_subscriber_ids:
        subscribers.extend(NotificationSubscriber.query.filter(
            NotificationSubscriber.active.is_(True),
            NotificationSubscriber.id.in_(direct_subscriber_ids),
        ).all())
    subscriber_by_id = {subscriber.id: subscriber for subscriber in subscribers}
    active_group_ids = {
        subscriber.group_id for subscriber in subscriber_by_id.values()
        if subscriber.group_id
    }
    buses_with_scope = {
        bus_id for bus_id, group_id, _name in group_rows
        if group_id in active_group_ids
    }
    if subscriber_by_id:
        buses_with_scope.update(
            bus_id for bus_id, _subscriber_id in db.session.query(
                NotificationBusAssignment.bus_id,
                NotificationBusAssignment.subscriber_id,
            ).filter(
                NotificationBusAssignment.bus_id.in_(active_bus_ids),
                NotificationBusAssignment.subscriber_id.in_(subscriber_by_id),
            ).all()
        )
    contacts = (SubscriberContact.query.filter(
        SubscriberContact.subscriber_id.in_(subscriber_by_id)).all()
        if subscriber_by_id else [])
    contacts_by_subscriber = {}
    for contact in contacts:
        contacts_by_subscriber.setdefault(contact.subscriber_id, []).append(contact)

    emails, phones = set(), set()
    roles, languages, schools = {}, {}, {}
    contact_count = 0
    for subscriber in subscriber_by_id.values():
        school = subscriber.school or 'Unspecified'
        schools[school] = schools.get(school, 0) + 1
        subscriber_contacts = contacts_by_subscriber.get(subscriber.id, [])
        if not subscriber_contacts:
            subscriber_contacts = [SimpleNamespace(
                role='parent', preferred_language='en',
                email=subscriber.email, phone=subscriber.phone)]
        for contact in subscriber_contacts:
            contact_count += 1
            role = contact.role if contact.role in {'parent', 'student'} else 'parent'
            language = _normalize_language(contact.preferred_language)
            roles[role] = roles.get(role, 0) + 1
            languages[language] = languages.get(language, 0) + 1
            for address in (contact.email or '').split(','):
                normalized = address.strip().lower()
                if normalized:
                    emails.add(normalized)
            phone = _normalize_phone(contact.phone)
            if phone:
                phones.add(phone)

    group_counts = dict(db.session.query(
        NotificationSubscriber.group_id,
        func.count(NotificationSubscriber.id),
    ).filter(
        NotificationSubscriber.active.is_(True),
        NotificationSubscriber.group_id.in_(group_ids),
    ).group_by(NotificationSubscriber.group_id).all()) if group_ids else {}
    return {
        'bus_count': len(active_bus_ids),
        'subscriber_count': len(subscriber_by_id),
        'contact_count': contact_count,
        'email_count': len(emails),
        'sms_count': len(phones),
        'roles': dict(sorted(roles.items())),
        'languages': dict(sorted(languages.items())),
        'schools': dict(sorted(schools.items(), key=lambda item: (-item[1], item[0]))),
        'groups': [{
            'id': group_id, 'name': group_names[group_id],
            'subscribers': int(group_counts.get(group_id, 0)),
        } for group_id in sorted(group_ids, key=lambda value: group_names[value].casefold())],
        'buses_without_recipients': len(active_bus_ids - buses_with_scope),
    }

@app.route('/admin/')
@app.route('/admin/dashboard')
@login_required
def dashboard():
    cfg = get_config()
    today = district_today(cfg)
    period = request.args.get('period', 'today')
    if period not in {'today', 'week', 'month', 'year', 'custom'}:
        period = 'today'
    date_from = request.args.get('date_from', today.isoformat())
    date_to = request.args.get('date_to', today.isoformat())
    can_view_buses = current_user.has_access('buses', 'limited')
    can_write_buses = current_user.has_access('buses', 'full')
    can_view_statistics = current_user.has_access('statistics', 'limited')
    can_view_notifications = current_user.has_access('notifications', 'limited')
    snapshot = _build_dashboard_snapshot(
        period, date_from, date_to,
        can_view_buses=can_view_buses,
        can_view_statistics=can_view_statistics,
        can_view_notifications=can_view_notifications,
    )
    incident_types = []
    schedule_types = []
    delay_reasons = []
    incident_request_token = None
    bulk_request_token = None
    if can_view_buses:
        incident_types = IncidentType.query.order_by(
            IncidentType.operational_priority.desc(),
            IncidentType.sort_order, IncidentType.name).all()
        schedule_types = BusScheduleType.query.order_by(
            BusScheduleType.sort_order).all()
    if can_write_buses:
        delay_reasons = DelayReason.query.order_by(DelayReason.sort_order).all()
        incident_request_token = _issue_incident_request_token()
        bulk_request_token = _issue_incident_request_token()
    selected_bus_id = request.args.get('bus', type=int)
    return render_template(
        'admin/dashboard.html', snapshot=snapshot,
        period=period, date_from=snapshot['date_from'].isoformat(),
        date_to=snapshot['date_to'].isoformat(), today=today,
        can_view_buses=can_view_buses,
        can_write_buses=can_write_buses,
        can_view_statistics=can_view_statistics,
        can_view_notifications=can_view_notifications,
        incident_types=incident_types, schedule_types=schedule_types,
        delay_reasons=delay_reasons,
        incident_request_token=incident_request_token,
        bulk_request_token=bulk_request_token,
        operations_revision=_dashboard_operations_revision(snapshot),
        selected_bus_id=selected_bus_id,
    )


@app.route('/admin/dashboard/operations.json')
@login_required
@require_module('buses', 'limited')
def dashboard_operations():
    cfg = get_config()
    today = district_today(cfg)
    snapshot = _build_dashboard_snapshot(
        'today', today.isoformat(), today.isoformat(),
        can_view_buses=True, can_view_statistics=False,
        can_view_notifications=current_user.has_access('notifications', 'limited'),
    )
    payload = _dashboard_operations_payload(snapshot)
    response = jsonify(payload)
    response.set_etag(payload['revision'])
    response.cache_control.private = True
    response.cache_control.no_cache = True
    return response.make_conditional(request)


@app.route('/admin/dashboard/recipients/preview', methods=['POST'])
@login_required
@require_module('buses', 'limited')
@require_module('notifications', 'limited')
def dashboard_recipient_preview():
    payload = request.get_json(silent=True) or {}
    raw_bus_ids = payload.get('bus_ids', [])
    if not isinstance(raw_bus_ids, list) or not 1 <= len(raw_bus_ids) <= 250:
        return jsonify({'ok': False, 'message': 'Select between 1 and 250 buses.'}), 400
    try:
        bus_ids = {int(value) for value in raw_bus_ids}
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'The bus selection is invalid.'}), 400
    schedule_type_id = payload.get('schedule_type_id')
    try:
        schedule_type_id = int(schedule_type_id) if schedule_type_id else None
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'The schedule period is invalid.'}), 400
    if schedule_type_id and not db.session.get(BusScheduleType, schedule_type_id):
        return jsonify({'ok': False, 'message': 'The schedule period is unavailable.'}), 400
    return jsonify({'ok': True, 'preview': _dashboard_recipient_preview(
        bus_ids, schedule_type_id)})


# ── BUSES MODULE ──────────────────────────────────────────────────────────────

@app.route('/admin/buses')
@login_required
@require_module('buses')
def buses():
    today          = district_today()
    current_period = get_current_period()
    buses_data     = bus_list_today(admin=True)   # show all buses regardless of schedule period
    incident_types = IncidentType.query.order_by(IncidentType.sort_order).all()
    schedule_types = BusScheduleType.query.order_by(BusScheduleType.sort_order).all()
    delay_reasons  = DelayReason.query.order_by(DelayReason.sort_order).all()
    can_write      = current_user.has_access('buses', 'full')
    return render_template('admin/buses.html',
                           buses_data=buses_data, incident_types=incident_types,
                           schedule_types=schedule_types, delay_reasons=delay_reasons,
                           current_period=current_period,
                           can_write=can_write, today=today,
                           incident_request_token=(
                               _issue_incident_request_token() if can_write else None))

@app.route('/admin/buses/add', methods=['POST'])
@login_required
@require_module('buses', 'full')
@_serialized_roster_mutation('html')
def add_bus():
    identifier = request.form.get('identifier', '').strip().upper()
    name       = request.form.get('name', '').strip()
    if not identifier or not name:
        flash('Identifier and name are required.', 'error')
        return redirect(url_for('buses'))
    if Bus.query.filter_by(identifier=identifier, name=name).first():
        flash(f'A bus with identifier "{identifier}" and name "{name}" already exists.', 'error')
        return redirect(url_for('buses'))
    assignments, assignment_error = _schedule_assignments_from_form()
    if assignment_error:
        flash(assignment_error, 'error')
        return redirect(url_for('buses'))
    bus = Bus(identifier=identifier, name=name,
              route=request.form.get('route','').strip() or None,
              capacity=request.form.get('capacity', type=int),
              description=request.form.get('description','').strip() or None)
    db.session.add(bus)
    db.session.flush()
    for schedule_id, departure_time in assignments:
        db.session.add(BusScheduleAssignment(
            bus_id=bus.id, schedule_type_id=schedule_id, departure_time=departure_time))
    db.session.commit()
    _audit('add_bus', 'buses', bus.display_name)
    flash(f'Bus {bus.display_name} registered successfully.', 'success')
    return redirect(url_for('buses'))

@app.route('/admin/buses/<int:bus_id>/edit', methods=['POST'])
@login_required
@require_module('buses', 'full')
@_serialized_roster_mutation('html')
def edit_bus(bus_id):
    bus = Bus.query.get_or_404(bus_id)
    new_identifier = request.form.get('identifier', bus.identifier).strip().upper()
    new_name       = request.form.get('name', bus.name).strip()
    # Check duplicate only if identifier+name changed
    if (new_identifier != bus.identifier or new_name != bus.name):
        dup = Bus.query.filter_by(identifier=new_identifier, name=new_name).first()
        if dup and dup.id != bus_id:
            flash(f'A bus with identifier "{new_identifier}" and name "{new_name}" already exists.', 'error')
            return redirect(url_for('buses'))
    assignments, assignment_error = _schedule_assignments_from_form()
    if assignment_error:
        flash(assignment_error, 'error')
        return redirect(url_for('buses'))
    bus.identifier  = new_identifier
    bus.name        = new_name
    bus.route       = request.form.get('route', '').strip() or None
    bus.capacity    = request.form.get('capacity', type=int)
    bus.description = request.form.get('description', '').strip() or None
    bus.active      = 'active' in request.form
    # Update schedules
    BusScheduleAssignment.query.filter_by(bus_id=bus_id).delete()
    for schedule_id, departure_time in assignments:
        db.session.add(BusScheduleAssignment(
            bus_id=bus_id, schedule_type_id=schedule_id, departure_time=departure_time))
    db.session.commit()
    _audit('edit_bus', 'buses', bus.display_name)
    flash(f'Bus {bus.display_name} updated.', 'success')
    return redirect(url_for('buses'))

@app.route('/admin/buses/<int:bus_id>/delete', methods=['POST'])
@login_required
@require_module('buses', 'full')
@_serialized_roster_mutation('html')
def delete_bus(bus_id):
    bus = Bus.query.get_or_404(bus_id)
    bus.active = False
    db.session.commit()
    _audit('delete_bus', 'buses', bus.display_name)
    flash(f'Bus {bus.identifier} deactivated.', 'success')
    return redirect(url_for('buses'))


def _mapping_integer(values, key, default=None):
    raw = values.get(key, default)
    if raw in (None, ''):
        return default
    try:
        return int(raw)
    except (TypeError, ValueError):
        return None


def _validated_incident_input(values):
    incident_type_id = _mapping_integer(values, 'incident_type_id')
    incident_type = (db.session.get(IncidentType, incident_type_id)
                     if incident_type_id else None)
    if not incident_type:
        return None, 'Select a valid incident type.'
    schedule_type_id = _mapping_integer(values, 'schedule_type_id')
    if values.get('schedule_type_id') not in (None, '') and schedule_type_id is None:
        return None, 'Select a valid schedule period.'
    if schedule_type_id and not db.session.get(BusScheduleType, schedule_type_id):
        return None, 'Select a valid schedule period.'
    delay_minutes = _mapping_integer(values, 'delay_minutes', 0)
    if delay_minutes is None or not 0 <= delay_minutes <= 999:
        return None, 'Delay must be between 0 and 999 minutes.'
    eta = _normalize_text(values.get('eta'), 5) or None
    if eta and _parse_clock_value(eta) is None:
        return None, 'Enter a valid ETA.'
    reason_raw = _normalize_text(values.get('delay_reason_id'), 20)
    reason_id = int(reason_raw) if reason_raw.isdigit() else None
    if reason_id and not db.session.get(DelayReason, reason_id):
        return None, 'Select a valid delay reason.'
    reason_text = _normalize_text(values.get('delay_reason_text'), 201) or None
    if reason_text and len(reason_text) > 200:
        return None, 'The custom delay reason cannot exceed 200 characters.'
    if reason_id:
        reason_text = None
    notes = _normalize_text(values.get('notes'), 2001) or None
    if notes and len(notes) > 2000:
        return None, 'Notes cannot exceed 2,000 characters.'
    return {
        'incident_type_id': incident_type.id,
        'schedule_type_id': schedule_type_id,
        'delay_minutes': delay_minutes,
        'eta': eta,
        'delay_reason_id': reason_id,
        'delay_reason_text': reason_text,
        'notes': notes,
    }, None


def _latest_bus_incident(bus_id, incident_date, schedule_type_id):
    query = BusIncidentRecord.query.filter_by(
        bus_id=bus_id, incident_date=incident_date)
    if schedule_type_id:
        query = query.filter_by(schedule_type_id=schedule_type_id)
    else:
        query = query.filter(BusIncidentRecord.schedule_type_id.is_(None))
    return query.order_by(
        BusIncidentRecord.created_at.desc(), BusIncidentRecord.id.desc()).first()

@app.route('/admin/buses/<int:bus_id>/incident', methods=['POST'])
@login_required
@require_module('buses', 'full')
def add_bus_incident(bus_id):
    bus = Bus.query.get_or_404(bus_id)
    next_url = _safe_local_redirect(request.form.get('next', '')) or url_for('buses')
    request_token = request.form.get('request_token', '').strip()
    existing = (BusIncidentRecord.query.filter_by(request_token=request_token).first()
                if request_token else None)
    if existing:
        flash('This status update was already recorded.', 'warning')
        return redirect(next_url)
    if not _consume_incident_request_token(request_token):
        flash('This status form expired. Open the bus and try again.', 'error')
        return redirect(next_url)
    if not bus.active:
        flash('This bus is inactive and cannot receive a status update.', 'error')
        return redirect(next_url)

    incident_input, validation_error = _validated_incident_input(request.form)
    if validation_error:
        flash(validation_error, 'error')
        return redirect(next_url)
    expected_latest_raw = request.form.get('expected_latest_id', '').strip()
    current_latest = _latest_bus_incident(
        bus.id, district_today(), incident_input['schedule_type_id'])
    if expected_latest_raw:
        try:
            expected_latest_id = int(expected_latest_raw)
        except ValueError:
            flash('The displayed bus state is invalid. Refresh and try again.', 'error')
            return redirect(next_url)
        if expected_latest_id != (current_latest.id if current_latest else 0):
            flash('Another operator changed this bus. Review the latest status before retrying.',
                  'warning')
            return redirect(next_url)
    replace_id = request.form.get('replace_incident_id', type=int)
    replaced = None
    if replace_id:
        replaced = BusIncidentRecord.query.filter_by(
            id=replace_id, bus_id=bus.id, is_pending=True).first()
        if not replaced or (current_latest and replaced.id != current_latest.id):
            flash('The pending update changed and can no longer be corrected.', 'warning')
            return redirect(next_url)
        db.session.delete(replaced)
        db.session.flush()
    rec = BusIncidentRecord(
        bus_id=bus_id, **incident_input,
        incident_date=district_today(), is_pending=True,
        created_by_id=current_user.id,
        request_token=request_token,
    )
    db.session.add(rec)
    try:
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        if BusIncidentRecord.query.filter_by(request_token=request_token).first():
            flash('This status update was already recorded.', 'warning')
            return redirect(next_url)
        raise
    _audit('add_bus_incident', 'buses', bus.display_name, json.dumps({
        'incident_record_id': rec.id,
        'incident_type_id': rec.incident_type_id,
        'schedule_type_id': rec.schedule_type_id,
        'delay_minutes': rec.delay_minutes,
        'replaced_incident_id': replaced.id if replaced else None,
    }, sort_keys=True))
    flash(f'Incident recorded for {bus.identifier}.', 'success')
    return redirect(next_url)

@app.route('/admin/bus-incidents/<int:rec_id>/delete', methods=['POST'])
@login_required
@require_module('buses', 'full')
def delete_bus_incident(rec_id):
    rec = BusIncidentRecord.query.get_or_404(rec_id)
    target = rec.bus.display_name
    details = json.dumps({
        'incident_record_id': rec.id,
        'incident_type_id': rec.incident_type_id,
        'was_pending': bool(rec.is_pending),
    }, sort_keys=True)
    db.session.delete(rec)
    db.session.commit()
    _audit('delete_bus_incident', 'buses', target, details)
    flash('Incident removed.', 'success')
    return redirect(request.referrer or url_for('buses'))


@app.route('/admin/dashboard/incidents/<int:rec_id>/confirm', methods=['POST'])
@login_required
@require_module('buses', 'full')
def dashboard_confirm_incident(rec_id):
    payload = request.get_json(silent=True) or {}
    record = db.session.get(BusIncidentRecord, rec_id)
    if not record:
        return jsonify({'ok': False, 'message': 'The pending update no longer exists.'}), 404
    expected_version = str(payload.get('version') or '')
    current_version = f'{record.id}:{record.updated_at.isoformat() if record.updated_at else ""}'
    if not record.is_pending or (expected_version and not hmac.compare_digest(
            expected_version, current_version)):
        return jsonify({'ok': False, 'message':
                        'The pending update changed. Refresh before continuing.'}), 409
    bus_label = record.bus.display_name
    committed = _commit_pending_incident_once(record.id)
    if not committed:
        return jsonify({'ok': False, 'message':
                        'Another operator already processed this update.'}), 409
    _audit('confirm_bus_incident', 'buses', bus_label, json.dumps({
        'incident_record_id': rec_id, 'manual_confirmation': True,
    }, sort_keys=True))
    return jsonify({'ok': True, 'message': f'{bus_label} was confirmed.',
                    'request_token': _issue_incident_request_token()})


@app.route('/admin/dashboard/incidents/<int:rec_id>/cancel', methods=['POST'])
@login_required
@require_module('buses', 'full')
def dashboard_cancel_incident(rec_id):
    payload = request.get_json(silent=True) or {}
    record = db.session.get(BusIncidentRecord, rec_id)
    if not record:
        return jsonify({'ok': False, 'message': 'The pending update no longer exists.'}), 404
    expected_version = str(payload.get('version') or '')
    current_version = f'{record.id}:{record.updated_at.isoformat() if record.updated_at else ""}'
    if not record.is_pending or (expected_version and not hmac.compare_digest(
            expected_version, current_version)):
        return jsonify({'ok': False, 'message':
                        'The pending update changed. Refresh before continuing.'}), 409
    bus_label = record.bus.display_name
    details = json.dumps({
        'incident_record_id': record.id,
        'incident_type_id': record.incident_type_id,
        'schedule_type_id': record.schedule_type_id,
        'delay_minutes': int(record.delay_minutes or 0),
    }, sort_keys=True)
    claimed = BusIncidentRecord.query.filter(
        BusIncidentRecord.id == rec_id,
        BusIncidentRecord.is_pending.is_(True),
        BusIncidentRecord.updated_at == record.updated_at,
    ).delete(synchronize_session=False)
    if claimed != 1:
        db.session.rollback()
        return jsonify({'ok': False, 'message':
                        'Another operator already processed this update.'}), 409
    db.session.commit()
    _audit('cancel_bus_incident', 'buses', bus_label, details)
    return jsonify({'ok': True, 'message': f'Pending update for {bus_label} was cancelled.',
                    'request_token': _issue_incident_request_token()})


@app.route('/admin/dashboard/incidents/bulk', methods=['POST'])
@login_required
@require_module('buses', 'full')
def dashboard_bulk_incidents():
    payload = request.get_json(silent=True) or {}
    raw_bus_ids = payload.get('bus_ids', [])
    if not isinstance(raw_bus_ids, list) or not 1 <= len(raw_bus_ids) <= 250:
        return jsonify({'ok': False, 'message': 'Select between 1 and 250 buses.'}), 400
    try:
        bus_ids = sorted({int(value) for value in raw_bus_ids})
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'The bus selection is invalid.'}), 400
    if payload.get('confirmed') is not True:
        return jsonify({'ok': False, 'message':
                        'Bulk operations require explicit confirmation.'}), 400
    request_token = _normalize_text(payload.get('request_token'), 64)
    if not re.fullmatch(r'[0-9a-f]{48}', request_token):
        return jsonify({'ok': False, 'message': 'The bulk operation form expired.'}), 409
    row_tokens = {
        bus_id: hashlib.sha256(
            f'bulk:{request_token}:{bus_id}'.encode('utf-8')).hexdigest()
        for bus_id in bus_ids
    }
    existing_count = BusIncidentRecord.query.filter(
        BusIncidentRecord.request_token.in_(row_tokens.values())).count()
    if existing_count:
        if existing_count == len(bus_ids):
            return jsonify({'ok': True, 'duplicate': True,
                            'message': 'This bulk operation was already recorded.',
                            'request_token': _issue_incident_request_token()})
        return jsonify({'ok': False, 'message':
                        'The bulk operation is partially recorded and requires review.'}), 409
    if not _consume_incident_request_token(request_token):
        return jsonify({'ok': False, 'message': 'The bulk operation form expired.'}), 409
    incident_input, validation_error = _validated_incident_input(
        payload.get('incident') or {})
    if validation_error:
        return jsonify({'ok': False, 'message': validation_error,
                        'request_token': _issue_incident_request_token()}), 400
    buses = Bus.query.filter(Bus.id.in_(bus_ids), Bus.active.is_(True)).all()
    if len(buses) != len(bus_ids):
        return jsonify({'ok': False, 'message':
                        'One or more selected buses are no longer active.',
                        'request_token': _issue_incident_request_token()}), 409
    expected_versions = payload.get('expected_latest_ids') or {}
    if not isinstance(expected_versions, dict):
        return jsonify({'ok': False, 'message': 'The displayed bus versions are invalid.',
                        'request_token': _issue_incident_request_token()}), 400
    conflicts = []
    today = district_today()
    for bus in buses:
        latest = _latest_bus_incident(
            bus.id, today, incident_input['schedule_type_id'])
        try:
            expected_id = int(expected_versions.get(str(bus.id), 0))
        except (TypeError, ValueError):
            expected_id = -1
        if expected_id != (latest.id if latest else 0):
            conflicts.append(bus.display_name)
    if conflicts:
        return jsonify({'ok': False, 'message':
                        'Another operator changed one or more selected buses.',
                        'conflict_count': len(conflicts),
                        'request_token': _issue_incident_request_token()}), 409
    records = []
    for bus in buses:
        record = BusIncidentRecord(
            bus_id=bus.id, **incident_input, incident_date=today,
            is_pending=True, created_by_id=current_user.id,
            request_token=row_tokens[bus.id])
        db.session.add(record)
        records.append(record)
    try:
        db.session.commit()
    except IntegrityError:
        db.session.rollback()
        return jsonify({'ok': False, 'message':
                        'The bulk operation raced with another request; refresh and review.',
                        'request_token': _issue_incident_request_token()}), 409
    _audit('bulk_add_bus_incident', 'buses', f'{len(records)} buses', json.dumps({
        'incident_record_ids': [record.id for record in records],
        'incident_type_id': incident_input['incident_type_id'],
        'schedule_type_id': incident_input['schedule_type_id'],
        'delay_minutes': incident_input['delay_minutes'],
    }, sort_keys=True))
    return jsonify({
        'ok': True, 'created': len(records),
        'message': f'{len(records)} bus status updates were staged.',
        'request_token': _issue_incident_request_token(),
    })

@app.route('/admin/delay-reasons/add', methods=['POST'])
@login_required
@require_module('buses', 'full')
def add_delay_reason():
    reason = request.form.get('reason', '').strip()
    if not reason:
        return jsonify({'success': False, 'error': 'Reason text required'})
    existing = DelayReason.query.filter_by(reason=reason).first()
    if existing:
        return jsonify({'success': True, 'id': existing.id, 'reason': existing.reason})
    dr = DelayReason(reason=reason, sort_order=99)
    db.session.add(dr)
    db.session.commit()
    return jsonify({'success': True, 'id': dr.id, 'reason': dr.reason})


# ── INCIDENT TYPES MODULE ─────────────────────────────────────────────────────

@app.route('/admin/incidents')
@login_required
@require_module('incidents')
def incidents():
    types     = IncidentType.query.order_by(IncidentType.sort_order, IncidentType.name).all()
    can_write = current_user.has_access('incidents', 'full')
    return render_template('admin/incidents.html', incident_types=types, can_write=can_write)

@app.route('/admin/incidents/add', methods=['POST'])
@login_required
@require_module('incidents', 'full')
def add_incident_type():
    name = request.form.get('name', '').strip()
    if not name:
        flash('Name is required.', 'error')
        return redirect(url_for('incidents'))
    if IncidentType.query.filter_by(name=name).first():
        flash(f'"{name}" already exists.', 'error')
        return redirect(url_for('incidents'))
    color = request.form.get('color', '#6b7280').strip()
    icon = request.form.get('icon', 'fa-circle').strip()
    description = request.form.get('description', '').strip() or None
    priority = request.form.get('operational_priority', 50, type=int)
    if (len(name) > 100 or not re.fullmatch(r'#[0-9A-Fa-f]{6}', color) or
            not re.fullmatch(r'fa-[a-z0-9-]{1,47}', icon) or
            priority is None or not 0 <= priority <= 100 or
            (description and len(description) > 255)):
        abort(400)
    it = IncidentType(name=name, color=color, icon=icon, description=description,
                      operational_priority=priority)
    db.session.add(it)
    db.session.commit()
    flash(f'Status type "{name}" created.', 'success')
    return redirect(url_for('incidents'))

@app.route('/admin/incidents/<int:type_id>/edit', methods=['POST'])
@login_required
@require_module('incidents', 'full')
def edit_incident_type(type_id):
    it = IncidentType.query.get_or_404(type_id)
    name = request.form.get('name', it.name).strip()
    color = request.form.get('color', it.color).strip()
    icon = request.form.get('icon', it.icon).strip()
    description = request.form.get('description', '').strip() or None
    priority = request.form.get(
        'operational_priority', it.operational_priority, type=int)
    if (not name or len(name) > 100 or not re.fullmatch(r'#[0-9A-Fa-f]{6}', color) or
            not re.fullmatch(r'fa-[a-z0-9-]{1,47}', icon) or
            priority is None or not 0 <= priority <= 100 or
            (description and len(description) > 255)):
        abort(400)
    it.name = name
    it.color = color
    it.icon = icon
    it.description = description
    it.operational_priority = priority
    db.session.commit()
    flash('Status type updated.', 'success')
    return redirect(url_for('incidents'))

@app.route('/admin/incidents/<int:type_id>/delete', methods=['POST'])
@login_required
@require_module('incidents', 'full')
def delete_incident_type(type_id):
    it = IncidentType.query.get_or_404(type_id)
    if it.is_system:
        flash('Cannot delete a system status type.', 'error')
        return redirect(url_for('incidents'))
    if BusIncidentRecord.query.filter_by(incident_type_id=type_id).first():
        flash('Cannot delete: this type has incident records.', 'error')
        return redirect(url_for('incidents'))
    db.session.delete(it)
    db.session.commit()
    flash(f'"{it.name}" deleted.', 'success')
    return redirect(url_for('incidents'))


# ── STATISTICS MODULE ─────────────────────────────────────────────────────────

@app.route('/admin/statistics')
@login_required
@require_module('statistics')
def statistics():
    today    = district_today()
    period   = request.args.get('period', 'today')
    d_from_s = request.args.get('date_from', today.isoformat())
    d_to_s   = request.args.get('date_to',   today.isoformat())
    bus_id   = request.args.get('bus_id', type=int)
    type_id  = request.args.get('type_id', type=int)

    d_from, d_to = _parse_period(period, d_from_s, d_to_s, today)

    q = BusIncidentRecord.query.filter(
        BusIncidentRecord.is_pending == False,
        BusIncidentRecord.incident_date >= d_from,
        BusIncidentRecord.incident_date <= d_to,
    )
    if bus_id:  q = q.filter_by(bus_id=bus_id)
    if type_id: q = q.filter_by(incident_type_id=type_id)
    records = q.order_by(BusIncidentRecord.incident_date.desc(),
                         BusIncidentRecord.created_at.desc()).all()

    by_type   = {}; by_type_colors = {}
    by_bus    = {}; by_day = {}; avg_delay = {}
    for r in records:
        n = r.incident_type.name
        by_type[n]        = by_type.get(n, 0) + 1
        by_type_colors[n] = r.incident_type.color
        b = f"{r.bus.identifier} · {r.bus.name}"
        by_bus[b] = by_bus.get(b, 0) + 1
        d = r.incident_date.isoformat()
        by_day[d] = by_day.get(d, 0) + 1
        if r.delay_minutes:
            avg_delay[b] = avg_delay.get(b, [])
            avg_delay[b].append(r.delay_minutes)
    avg_delay_final = {k: round(sum(v)/len(v), 1) for k, v in avg_delay.items()}

    # Period × Bus breakdown for the period chart
    schedule_periods = BusScheduleType.query.order_by(BusScheduleType.sort_order).all()
    period_bus_data  = {p.name: {} for p in schedule_periods}
    for r in records:
        pname = r.schedule_type.name if r.schedule_type else None
        if pname and pname in period_bus_data:
            b = f"{r.bus.identifier} · {r.bus.name}"
            period_bus_data[pname][b] = period_bus_data[pname].get(b, 0) + 1
    # Only keep periods that actually have data
    period_bus_data = {k: v for k, v in period_bus_data.items() if v}
    record_buses = sorted({f"{r.bus.identifier} · {r.bus.name}" for r in records})

    # ── Bus Audit ─────────────────────────────────────────────────────────
    default_type = IncidentType.query.filter_by(is_default=True).first()
    audit_buses_q = Bus.query.filter_by(active=True).order_by(Bus.identifier)
    if bus_id:
        audit_buses_q = audit_buses_q.filter_by(id=bus_id)
    audit_buses_list = audit_buses_q.all()

    total_days_in_range = (d_to - d_from).days + 1

    # Days per bus that had at least one non-default incident
    bus_incident_dates = {}
    for r in records:
        if not r.incident_type.is_default:
            bus_incident_dates.setdefault(r.bus_id, set()).add(r.incident_date)

    on_time_by_bus = {}
    bus_audit      = {}
    for bus in audit_buses_list:
        inc_days = len(bus_incident_dates.get(bus.id, set()))
        ot_days  = max(0, total_days_in_range - inc_days)
        blabel = f"{bus.identifier} · {bus.name}"
        on_time_by_bus[blabel] = ot_days
        bus_delays = [r.delay_minutes for r in records
                      if r.bus_id == bus.id and r.delay_minutes and r.delay_minutes > 0]
        avg_d = round(sum(bus_delays) / len(bus_delays), 1) if bus_delays else 0.0
        tot_d = sum(bus_delays)
        rate  = round(ot_days / total_days_in_range * 100, 1) if total_days_in_range else 100.0
        bus_audit[blabel] = {
            'id': bus.identifier, 'name': bus.name, 'route': bus.route or '',
            'total_days': total_days_in_range,
            'on_time_days': ot_days, 'incident_days': inc_days,
            'on_time_rate': rate, 'avg_delay': avg_d, 'total_delay': tot_d,
        }

    # Include On Time in the by_type chart (only when not filtered to a specific type)
    on_time_total = sum(on_time_by_bus.values())
    if on_time_total > 0 and default_type and not type_id:
        by_type[default_type.name]        = on_time_total
        by_type_colors[default_type.name] = default_type.color

    # Stacked datasets for audit chart: {status_name: {data:[...], color:hex}}
    audit_bus_order = [f"{b.identifier} · {b.name}" for b in audit_buses_list]
    audit_datasets  = {}
    if default_type:
        audit_datasets[default_type.name] = {
            'data':  [on_time_by_bus.get(bid, 0) for bid in audit_bus_order],
            'color': default_type.color,
        }
    for r in records:
        if not r.incident_type.is_default:
            n   = r.incident_type.name
            blb = f"{r.bus.identifier} · {r.bus.name}"
            if n not in audit_datasets:
                audit_datasets[n] = {
                    'data':  [0] * len(audit_bus_order),
                    'color': r.incident_type.color,
                }
            if blb in audit_bus_order:
                audit_datasets[n]['data'][audit_bus_order.index(blb)] += 1

    all_buses  = Bus.query.filter_by(active=True).order_by(Bus.identifier).all()
    all_types  = IncidentType.query.order_by(IncidentType.sort_order).all()
    can_export = current_user.has_access('statistics', 'limited')

    # Notification delivery data contains parent/contact PII and belongs to the
    # notifications capability even when displayed alongside statistics.
    can_view_notifications = current_user.has_capability('notifications.pii')
    can_delete_notifications = current_user.has_capability('notifications.write')
    notif_logs = []
    if can_view_notifications:
        notif_q = NotificationLog.query.filter(
            NotificationLog.sent_at >= datetime.combine(d_from, datetime.min.time()),
            NotificationLog.sent_at <= datetime.combine(d_to,   datetime.max.time()),
        )
        if bus_id:
            notif_q = notif_q.filter_by(bus_id=bus_id)
        notif_logs = notif_q.order_by(NotificationLog.sent_at.desc()).all()

    notif_by_channel = {}; notif_by_day = {}; notif_by_group = {}
    notif_sent = 0; notif_failed = 0; notif_total_cost = 0.0
    notif_email_sent = 0; notif_sms_sent = 0
    for nl in notif_logs:
        ch = nl.channel
        notif_by_channel[ch] = notif_by_channel.get(ch, 0) + 1
        d = nl.sent_at.strftime('%Y-%m-%d')
        notif_by_day[d] = notif_by_day.get(d, 0) + 1
        if nl.group_name:
            notif_by_group[nl.group_name] = notif_by_group.get(nl.group_name, 0) + 1
        if nl.status == 'sent':
            notif_sent += 1
            if ch == 'email': notif_email_sent += 1
            if ch == 'sms':   notif_sms_sent   += 1
        else:
            notif_failed += 1
        if nl.sms_cost_usd:
            notif_total_cost += nl.sms_cost_usd
    notif_total_cost = round(notif_total_cost, 4)

    return render_template('admin/statistics.html',
        records=records, period=period,
        date_from=d_from.isoformat(), date_to=d_to.isoformat(),
        bus_id=bus_id, type_id=type_id,
        by_type_json=by_type,
        by_type_colors_json=list(by_type_colors.values()),
        by_bus_json=by_bus,
        by_day_json=by_day,
        avg_delay_json=avg_delay_final,
        period_bus_json=period_bus_data,
        record_buses_json=record_buses,
        bus_audit_json=bus_audit,
        audit_datasets_json=audit_datasets,
        audit_bus_order_json=audit_bus_order,
        default_type_name=(default_type.name if default_type else 'On Time'),
        total_days_in_range=total_days_in_range,
        total=len(records), all_buses=all_buses, all_types=all_types,
        can_export=can_export, can_write=current_user.has_access('statistics', 'full'),
        today=today,
        notif_logs=notif_logs,
        notif_by_channel_json=notif_by_channel,
        notif_by_day_json=notif_by_day,
        notif_by_group_json=notif_by_group,
        notif_sent=notif_sent, notif_failed=notif_failed,
        notif_email_sent=notif_email_sent, notif_sms_sent=notif_sms_sent,
        notif_total_cost=notif_total_cost,
        can_view_notifications=can_view_notifications,
        can_delete_notifications=can_delete_notifications,
    )

def _parse_period(period, d_from_s, d_to_s, today):
    if period == 'today':    return today, today
    if period == 'week':     return today - timedelta(days=today.weekday()), today
    if period == 'month':    return today.replace(day=1), today
    if period == 'year':     return today.replace(month=1, day=1), today
    try:
        selected_from = date.fromisoformat(d_from_s)
        selected_to = date.fromisoformat(d_to_s)
        if selected_from > selected_to:
            selected_from, selected_to = selected_to, selected_from
        return selected_from, selected_to
    except Exception:
        return today, today

@app.route('/admin/statistics/export/<fmt>')
@login_required
@require_module('statistics')
def export_statistics(fmt):
    today    = district_today()
    period   = request.args.get('period', 'today')
    d_from_s = request.args.get('date_from', today.isoformat())
    d_to_s   = request.args.get('date_to',   today.isoformat())
    bus_id   = request.args.get('bus_id', type=int)
    type_id  = request.args.get('type_id', type=int)
    d_from, d_to = _parse_period(period, d_from_s, d_to_s, today)

    q = BusIncidentRecord.query.filter(
        BusIncidentRecord.is_pending == False,
        BusIncidentRecord.incident_date >= d_from,
        BusIncidentRecord.incident_date <= d_to,
    )
    if bus_id:  q = q.filter_by(bus_id=bus_id)
    if type_id: q = q.filter_by(incident_type_id=type_id)
    records = q.order_by(BusIncidentRecord.incident_date, BusIncidentRecord.created_at).all()

    cfg = get_config()
    title = f"{cfg.app_name} — Incident Report ({d_from} to {d_to})"
    headers = ['Date','Bus ID','Bus Name','Route','Status','Delay (min)','Schedule','Notes','Recorded By']
    rows = [[
        r.incident_date.strftime('%Y-%m-%d'), r.bus.identifier, r.bus.name,
        r.bus.route or '', r.incident_type.name, r.delay_minutes,
        r.schedule_type.name if r.schedule_type else '',
        r.notes or '', r.created_by.username if r.created_by else '',
    ] for r in records]

    # ── Bus Audit for export ──────────────────────────────────────────────
    default_type_exp = IncidentType.query.filter_by(is_default=True).first()
    exp_buses_q = Bus.query.filter_by(active=True).order_by(Bus.identifier)
    if bus_id:
        exp_buses_q = exp_buses_q.filter_by(id=bus_id)
    exp_buses_list = exp_buses_q.all()
    total_days_exp = (d_to - d_from).days + 1
    bus_inc_dates_exp = {}
    for r in records:
        if not r.incident_type.is_default:
            bus_inc_dates_exp.setdefault(r.bus_id, set()).add(r.incident_date)
    audit_headers = ['Bus ID','Bus Name','Route','Total Days','On-Time Days',
                     'Incident Days','On-Time Rate (%)','Avg Delay (min)','Total Delay (min)']
    audit_rows = []
    for bus in exp_buses_list:
        inc_d = len(bus_inc_dates_exp.get(bus.id, set()))
        ot_d  = max(0, total_days_exp - inc_d)
        bdel  = [r.delay_minutes for r in records
                 if r.bus_id == bus.id and r.delay_minutes and r.delay_minutes > 0]
        avg_d = round(sum(bdel)/len(bdel), 1) if bdel else 0.0
        rate  = round(ot_d / total_days_exp * 100, 1) if total_days_exp else 100.0
        audit_rows.append([bus.identifier, bus.name, bus.route or '',
                           total_days_exp, ot_d, inc_d, rate, avg_d, sum(bdel)])

    # ── Notification log for general export ──────────────────────────────────
    notif_exp = []
    if current_user.has_capability('notifications.pii'):
        notif_exp_q = NotificationLog.query.filter(
            NotificationLog.sent_at >= datetime.combine(d_from, datetime.min.time()),
            NotificationLog.sent_at <= datetime.combine(d_to,   datetime.max.time()),
        )
        if bus_id:
            notif_exp_q = notif_exp_q.filter_by(bus_id=bus_id)
        notif_exp = notif_exp_q.order_by(NotificationLog.sent_at.desc()).all()
    notif_headers = ['Sent At (UTC)', 'Channel', 'Bus', 'Recipient', 'Address',
                     'Group', 'Status', 'SMS SID', 'Segments', 'Cost (USD)', 'Error']
    notif_rows = [[
        nl.sent_at.strftime('%Y-%m-%d %H:%M:%S'), nl.channel, nl.bus_label or '',
        nl.recipient_name or '', nl.recipient_address or '',
        nl.group_name or '', nl.status, nl.sms_sid or '',
        nl.sms_segments or '',
        f'{nl.sms_cost_usd:.6f}' if nl.sms_cost_usd else '',
        nl.error_message or '',
    ] for nl in notif_exp]

    if fmt == 'csv':
        output = io.StringIO()
        w = csv.writer(output)
        w.writerow(headers)
        w.writerows(_csv_safe_row(row) for row in rows)
        w.writerow([])
        w.writerow(['Bus Audit Summary'])
        w.writerow(audit_headers)
        w.writerows(_csv_safe_row(row) for row in audit_rows)
        w.writerow([])
        w.writerow(['Notification Log'])
        w.writerow(notif_headers)
        w.writerows(_csv_safe_row(row) for row in notif_rows)
        resp = make_response(output.getvalue())
        resp.headers['Content-Type'] = 'text/csv'
        resp.headers['Content-Disposition'] = f'attachment; filename="bus_report_{d_from}_{d_to}.csv"'
        return resp

    elif fmt == 'pdf' and PDF_AVAILABLE:
        def _pdf_safe(text):
            return (str(text)
                    .replace('\u2014', '--').replace('\u2013', '-')
                    .replace('\u2018', "'").replace('\u2019', "'")
                    .replace('\u201c', '"').replace('\u201d', '"')
                    .encode('latin-1', errors='replace').decode('latin-1'))

        class _BusReportPDF(FPDF):
            def footer(self):
                self.set_y(-12)
                self.set_font('Helvetica', 'I', 7)
                self.set_text_color(150, 150, 150)
                self.cell(0, 5, 'Powered by Avidity Technologies Inc', align='C')
                self.set_text_color(0, 0, 0)

        pdf = _BusReportPDF(orientation='L', unit='mm', format='A4')
        pdf.set_auto_page_break(auto=True, margin=16)
        pdf.set_margins(10, 10, 10)
        pdf.add_page()

        # ── Header: logo + app name + report title ──────────────────────────
        logo_x = 10
        logo_fs = None
        if cfg.logo_path:
            candidate = os.path.join(BASE_DIR, cfg.logo_path.lstrip('/').replace('/', os.sep))
            if os.path.exists(candidate):
                logo_fs = candidate
        if logo_fs:
            try:
                pdf.image(logo_fs, x=logo_x, y=10, h=14)
                text_x = logo_x + 18
            except Exception:
                text_x = logo_x
        else:
            text_x = logo_x

        pdf.set_xy(text_x, 10)
        pdf.set_font('Helvetica', 'B', 15)
        pdf.set_text_color(30, 64, 175)
        pdf.cell(0, 8, _pdf_safe(cfg.app_name or 'Bus Tracker'), ln=True)
        pdf.set_x(text_x)
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(100, 116, 139)
        pdf.cell(0, 5, _pdf_safe(title), ln=True)
        pdf.set_text_color(0, 0, 0)

        # Separator line
        pdf.set_y(max(pdf.get_y(), 26))
        pdf.set_draw_color(226, 232, 240)
        pdf.line(10, pdf.get_y(), 287, pdf.get_y())
        pdf.ln(3)

        # ── Table header ────────────────────────────────────────────────────
        # A4 landscape usable width ≈ 277mm (297 - 2×10)
        col_widths = [28, 20, 42, 38, 34, 20, 27, 40, 28]
        pdf.set_font('Helvetica', 'B', 7)
        pdf.set_fill_color(30, 64, 175)
        pdf.set_text_color(255, 255, 255)
        for h, w in zip(headers, col_widths):
            pdf.cell(w, 7, _pdf_safe(h), border=0, fill=True, align='C')
        pdf.ln()

        # ── Table rows ──────────────────────────────────────────────────────
        pdf.set_font('Helvetica', '', 7)
        pdf.set_text_color(15, 23, 42)
        alt = False
        for row in rows:
            pdf.set_fill_color(241, 245, 249) if alt else pdf.set_fill_color(255, 255, 255)
            for val, w in zip(row, col_widths):
                pdf.cell(w, 6, _pdf_safe(str(val))[:35], border=0, fill=True)
            pdf.ln()
            alt = not alt

        # ── Bus Audit table ──────────────────────────────────────────────
        pdf.ln(6)
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_text_color(30, 64, 175)
        pdf.cell(0, 6, 'Bus Audit Summary', ln=True)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)
        a_widths = [22, 45, 38, 22, 22, 22, 28, 28, 28]
        pdf.set_font('Helvetica', 'B', 7)
        pdf.set_fill_color(30, 64, 175)
        pdf.set_text_color(255, 255, 255)
        for h, w in zip(audit_headers, a_widths):
            pdf.cell(w, 7, _pdf_safe(h), border=0, fill=True, align='C')
        pdf.ln()
        pdf.set_font('Helvetica', '', 7)
        pdf.set_text_color(15, 23, 42)
        alt = False
        for row in audit_rows:
            pdf.set_fill_color(241, 245, 249) if alt else pdf.set_fill_color(255, 255, 255)
            for val, w in zip(row, a_widths):
                pdf.cell(w, 6, _pdf_safe(str(val))[:30], border=0, fill=True)
            pdf.ln()
            alt = not alt

        # Notification section in PDF
        if notif_exp:
            pdf.ln(4)
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(30, 64, 175)
            pdf.cell(0, 6, 'Notification Log', ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)
            n_widths = [32, 16, 32, 38, 38, 30, 16, 30, 16, 22]
            n_hdrs   = ['Sent At', 'Channel', 'Bus', 'Recipient', 'Address',
                        'Group', 'Status', 'SMS SID', 'Segs', 'Cost']
            pdf.set_font('Helvetica', 'B', 6)
            pdf.set_fill_color(30, 64, 175)
            pdf.set_text_color(255, 255, 255)
            for h, w in zip(n_hdrs, n_widths):
                pdf.cell(w, 7, _pdf_safe(h), border=0, fill=True, align='C')
            pdf.ln()
            pdf.set_font('Helvetica', '', 6)
            pdf.set_text_color(15, 23, 42)
            alt = False
            for nl in notif_exp:
                pdf.set_fill_color(241, 245, 249) if alt else pdf.set_fill_color(255, 255, 255)
                vals = [nl.sent_at.strftime('%Y-%m-%d %H:%M'), nl.channel,
                        nl.bus_label or '', nl.recipient_name or '', nl.recipient_address or '',
                        nl.group_name or '', nl.status, nl.sms_sid or '',
                        str(nl.sms_segments or ''),
                        f'${nl.sms_cost_usd:.4f}' if nl.sms_cost_usd else '']
                for val, w in zip(vals, n_widths):
                    pdf.cell(w, 5, _pdf_safe(str(val))[:28], border=0, fill=True)
                pdf.ln()
                alt = not alt

        resp = make_response(bytes(pdf.output()))
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = f'attachment; filename="bus_report_{d_from}_{d_to}.pdf"'
        return resp

    elif fmt == 'docx' and DOCX_AVAILABLE:
        doc = DocxDocument()
        doc.add_heading(title, 0)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        doc.add_heading('Bus Audit Summary', level=2)
        a_table = doc.add_table(rows=1, cols=len(audit_headers))
        a_table.style = 'Table Grid'
        for i, h in enumerate(audit_headers):
            a_table.rows[0].cells[i].text = h
        for row in audit_rows:
            cells = a_table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        if notif_exp:
            doc.add_heading('Notification Log', level=2)
            n_table = doc.add_table(rows=1, cols=len(notif_headers))
            n_table.style = 'Table Grid'
            for i, h in enumerate(notif_headers):
                n_table.rows[0].cells[i].text = h
            for row in notif_rows:
                cells = n_table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name=f'bus_report_{d_from}_{d_to}.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    flash(f'Export format "{fmt}" not available. Try CSV.', 'error')
    return redirect(url_for('statistics'))

@app.route('/admin/statistics/email', methods=['POST'])
@login_required
@require_module('statistics')
def email_statistics():
    to_email = request.form.get('email', '').strip()
    if not to_email:
        flash('Enter a recipient email.', 'error')
        return redirect(url_for('statistics'))
    cfg = get_config()
    try:
        to_email = _validated_email(to_email, 'recipient email address', required=True)
        settings = _smtp_settings_from_config(cfg)
        today = district_today(cfg)
        records = BusIncidentRecord.query.filter(
            BusIncidentRecord.is_pending == False,
            BusIncidentRecord.incident_date == today,
        ).all()
        body = f"{cfg.app_name} — Daily Bus Report ({today})\n\n"
        body += f"{'Bus':<12} {'Status':<22} {'Delay':>6}  Schedule\n"
        body += '-' * 60 + '\n'
        for r in records:
            body += f"{r.bus.identifier:<12} {r.incident_type.name:<22} {r.delay_minutes:>5}m  {r.schedule_type.name if r.schedule_type else ''}\n"
        send_email(
            settings,
            subject=f"Bus Report — {today}",
            recipients=[to_email],
            body=body,
        )
        flash(f'Report sent to {to_email}.', 'success')
    except (EmailTransportError, ValueError, RuntimeError) as exc:
        failure = exc if isinstance(exc, EmailTransportError) else EmailTransportError(
            'configuration_invalid', str(exc))
        flash(f'Could not send email: {failure.safe_message}', 'error')
    return redirect(url_for('statistics'))


@app.route('/admin/statistics/reset', methods=['POST'])
@login_required
@require_module('statistics', 'full')
def reset_statistics():
    preset      = request.form.get('preset', '')
    date_from_s = request.form.get('rs_date_from', '')
    date_to_s   = request.form.get('rs_date_to', '')
    today = district_today()
    try:
        if preset == 'today':
            d_from = d_to = today
        elif preset == 'week':
            d_from = today - timedelta(days=today.weekday())
            d_to   = today
        elif preset == 'month':
            d_from = today.replace(day=1)
            d_to   = today
        elif preset == 'year':
            d_from = today.replace(month=1, day=1)
            d_to   = today
        elif preset == 'custom':
            d_from = date.fromisoformat(date_from_s)
            d_to   = date.fromisoformat(date_to_s)
            if d_from > d_to:
                flash('Start date must be before end date.', 'error')
                return redirect(url_for('statistics'))
        else:
            flash('Invalid selection.', 'error')
            return redirect(url_for('statistics'))
    except (ValueError, TypeError):
        flash('Invalid date.', 'error')
        return redirect(url_for('statistics'))

    include_notifs = request.form.get('include_notifications') == '1'
    if include_notifs and not current_user.has_capability('notifications.write'):
        abort(403)

    # Collect the incident IDs that will be deleted
    incident_ids = [r.id for r in BusIncidentRecord.query.filter(
        BusIncidentRecord.incident_date >= d_from,
        BusIncidentRecord.incident_date <= d_to
    ).with_entities(BusIncidentRecord.id).all()]

    notif_deleted = 0
    if incident_ids:
        if include_notifs:
            # Delete notification logs tied to those incidents OR sent in that date range
            dt_from = datetime.combine(d_from, datetime.min.time())
            dt_to   = datetime.combine(d_to,   datetime.max.time())
            notif_deleted = NotificationLog.query.filter(
                db.or_(
                    NotificationLog.incident_record_id.in_(incident_ids),
                    db.and_(
                        NotificationLog.sent_at >= dt_from,
                        NotificationLog.sent_at <= dt_to
                    )
                )
            ).delete(synchronize_session='fetch')
        else:
            # Null-out FK so the incident delete doesn't violate the constraint
            NotificationLog.query.filter(
                NotificationLog.incident_record_id.in_(incident_ids)
            ).update({NotificationLog.incident_record_id: None}, synchronize_session='fetch')

    deleted = BusIncidentRecord.query.filter(
        BusIncidentRecord.incident_date >= d_from,
        BusIncidentRecord.incident_date <= d_to
    ).delete(synchronize_session='fetch')
    db.session.commit()

    detail = f'Deleted {deleted} incident record{"s" if deleted != 1 else ""}'
    if include_notifs:
        detail += f' + {notif_deleted} notification log{"s" if notif_deleted != 1 else ""}'
    _audit('reset_statistics', 'statistics', f'{d_from} → {d_to}', detail)
    flash(f'{detail} ({d_from} → {d_to}).', 'success')
    return redirect(url_for('statistics'))


@app.route('/admin/statistics/export/notifications')
@login_required
@require_module('statistics', 'limited')
@require_module('notifications', 'limited')
@require_capability('notifications.export_pii')
def export_notification_stats():
    cfg = get_config()
    today    = district_today(cfg)
    period   = request.args.get('period', 'today')
    d_from_s = request.args.get('date_from', today.isoformat())
    d_to_s   = request.args.get('date_to',   today.isoformat())
    bus_id   = request.args.get('bus_id', type=int)
    d_from, d_to = _parse_period(period, d_from_s, d_to_s, today)

    utc_from, utc_until = district_date_utc_bounds(d_from, d_to, cfg)
    q = NotificationLog.query.filter(
        NotificationLog.sent_at >= utc_from,
        NotificationLog.sent_at < utc_until,
    )
    if bus_id:
        q = q.filter_by(bus_id=bus_id)
    logs = q.order_by(NotificationLog.sent_at.desc()).all()

    output = io.StringIO()
    w = csv.writer(output)
    w.writerow(['Sent At (UTC)', 'Channel', 'Bus', 'Recipient', 'Address',
                'Group', 'Status', 'SMS SID', 'Segments', 'Cost (USD)', 'Error'])
    for nl in logs:
        w.writerow(_csv_safe_row([
            nl.sent_at.strftime('%Y-%m-%d %H:%M:%S'),
            nl.channel, nl.bus_label or '',
            nl.recipient_name or '', nl.recipient_address or '',
            nl.group_name or '', nl.status,
            nl.sms_sid or '', nl.sms_segments or '',
            f'{nl.sms_cost_usd:.6f}' if nl.sms_cost_usd else '',
            nl.error_message or '',
        ]))
    resp = make_response(output.getvalue())
    resp.headers['Content-Type'] = 'text/csv'
    resp.headers['Content-Disposition'] = (
        f'attachment; filename="notification_stats_{d_from}_{d_to}.csv"'
    )
    return resp


# ── USERS MODULE ──────────────────────────────────────────────────────────────

_ACCESS_RANK = {'none': 0, 'limited': 1, 'full': 2}


def _group_access_level(group, module_key):
    if group.is_admin:
        return 'full'
    perm = GroupPermission.query.filter_by(group_id=group.id, module_key=module_key).first()
    return perm.access_level if perm and perm.access_level in _ACCESS_RANK else 'none'


def _can_assign_group(actor, group):
    if not group:
        return True
    if actor.has_capability('user.assign_admin'):
        return True
    if group.is_admin:
        return False
    return all(_ACCESS_RANK[_group_access_level(group, mod['key'])] <=
               _ACCESS_RANK['full' if actor.has_access(mod['key'], 'full') else
                            'limited' if actor.has_access(mod['key'], 'limited') else 'none']
               for mod in MODULES)


def _requested_group(group_id):
    if not group_id:
        return None
    group = db.session.get(UserGroup, group_id)
    if not group:
        abort(400)
    if not _can_assign_group(current_user, group):
        abort(403)
    return group


def _requested_permission_level(module_key):
    level = request.form.get(f'perm_{module_key}', 'none')
    if level not in _ACCESS_RANK:
        abort(400)
    if current_user.is_admin:
        return level
    actor_level = ('full' if current_user.has_access(module_key, 'full') else
                   'limited' if current_user.has_access(module_key, 'limited') else 'none')
    if _ACCESS_RANK[level] > _ACCESS_RANK[actor_level]:
        abort(403)
    return level

@app.route('/admin/users')
@login_required
@require_module('users')
def users():
    all_users  = User.query.order_by(User.username).all()
    all_groups = UserGroup.query.order_by(UserGroup.name).all()
    assignable_groups = [group for group in all_groups if _can_assign_group(current_user, group)]
    can_write  = current_user.has_access('users', 'full')
    return render_template('admin/users.html', users=all_users, groups=all_groups,
                           assignable_groups=assignable_groups,
                           MODULES=MODULES, can_write=can_write)

@app.route('/admin/users/add', methods=['POST'])
@login_required
@require_module('users', 'full')
def add_user():
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')
    email    = request.form.get('email', '').strip() or None
    if not username or not password:
        flash('Username and password are required.', 'error')
        return redirect(url_for('users'))
    if User.query.filter_by(username=username).first():
        flash(f'Username "{username}" already exists.', 'error')
        return redirect(url_for('users'))
    if email and User.query.filter_by(email=email).first():
        flash('Email already in use.', 'error')
        return redirect(url_for('users'))
    password_error = _password_error(password)
    if password_error:
        flash(password_error, 'error')
        return redirect(url_for('users'))
    group = _requested_group(request.form.get('group_id', type=int))
    u = User(username=username, email=email,
             first_name=request.form.get('first_name', '').strip() or None,
             last_name=request.form.get('last_name', '').strip() or None,
             phone=request.form.get('phone', '').strip() or None,
             workplace=request.form.get('workplace', '').strip() or None,
             job_title=request.form.get('job_title', '').strip() or None,
             group_id=group.id if group else None,
             use_email_auth='use_email_auth' in request.form,
             receive_notifications='receive_notifications' in request.form,
             active=True)
    u.set_password(password)
    db.session.add(u)
    db.session.commit()
    _audit('add_user', 'users', username)
    flash(f'User "{username}" created.', 'success')
    return redirect(url_for('users'))

@app.route('/admin/users/<int:uid>/edit', methods=['POST'])
@login_required
@require_module('users', 'full')
def edit_user(uid):
    u = User.query.get_or_404(uid)
    # Only admins can edit other users' group; everyone can edit own profile
    if uid != current_user.id and not current_user.is_admin:
        flash('Permission denied.', 'error')
        return redirect(url_for('users'))
    u.first_name  = request.form.get('first_name', '').strip() or None
    u.last_name   = request.form.get('last_name', '').strip() or None
    u.email       = request.form.get('email', '').strip() or None
    u.phone       = request.form.get('phone', '').strip() or None
    u.workplace   = request.form.get('workplace', '').strip() or None
    u.job_title   = request.form.get('job_title', '').strip() or None
    u.use_email_auth        = 'use_email_auth' in request.form
    u.receive_notifications = 'receive_notifications' in request.form
    security_changed = False
    if current_user.is_admin:
        group = _requested_group(request.form.get('group_id', type=int))
        new_group_id = group.id if group else None
        new_active = 'active' in request.form
        if uid == current_user.id and not new_active:
            abort(400)
        security_changed = u.group_id != new_group_id or u.active != new_active
        u.group_id = new_group_id
        u.active = new_active
    pwd = request.form.get('new_password', '').strip()
    if pwd:
        password_error = _password_error(pwd)
        if password_error:
            flash(password_error, 'error')
            return redirect(url_for('users'))
        u.set_password(pwd)
        security_changed = True
    if security_changed:
        u.session_version = int(u.session_version or 1) + 1
        if uid == current_user.id:
            session['session_version'] = u.session_version
    db.session.commit()
    _audit('edit_user', 'users', u.username)
    flash(f'User "{u.username}" updated.', 'success')
    return redirect(url_for('users'))

@app.route('/admin/users/<int:uid>/delete', methods=['POST'])
@login_required
@require_module('users', 'full')
def delete_user(uid):
    if not current_user.is_admin:
        flash('Only administrators can delete users.', 'error')
        return redirect(url_for('users'))
    u = User.query.get_or_404(uid)
    if uid == current_user.id:
        flash('You cannot delete your own signed-in account.', 'error')
        return redirect(url_for('users'))
    if u.is_admin and User.query.join(UserGroup).filter(
            UserGroup.is_admin == True, User.active == True).count() <= 1:
        flash('Cannot delete the last active administrator.', 'error')
        return redirect(url_for('users'))
    uname = u.username
    db.session.delete(u)
    db.session.commit()
    _audit('delete_user', 'users', uname)
    flash(f'User "{uname}" deleted.', 'success')
    return redirect(url_for('users'))

@app.route('/admin/groups/add', methods=['POST'])
@login_required
@require_module('users', 'full')
def add_group():
    name = request.form.get('name', '').strip()
    if not name:
        flash('Name is required.', 'error')
        return redirect(url_for('users'))
    if UserGroup.query.filter_by(name=name).first():
        flash('Group already exists.', 'error')
        return redirect(url_for('users'))
    g = UserGroup(name=name, description=request.form.get('description','').strip() or None)
    db.session.add(g)
    db.session.flush()
    for mod in MODULES:
        level = _requested_permission_level(mod['key'])
        db.session.add(GroupPermission(group_id=g.id, module_key=mod['key'], access_level=level))
    db.session.flush()
    _sync_group_capabilities(g.id, overwrite_existing=True)
    db.session.commit()
    _audit('permission_group_created', 'users', name)
    flash(f'Group "{name}" created.', 'success')
    return redirect(url_for('users'))

@app.route('/admin/groups/<int:gid>/edit', methods=['POST'])
@login_required
@require_module('users', 'full')
def edit_group(gid):
    g = UserGroup.query.get_or_404(gid)
    if g.is_admin and not current_user.is_admin:
        abort(403)
    g.name        = request.form.get('name', g.name).strip()
    g.description = request.form.get('description', '').strip() or None
    changed_permissions = False
    if not g.is_admin:
        for mod in MODULES:
            level = _requested_permission_level(mod['key'])
            perm  = GroupPermission.query.filter_by(group_id=gid, module_key=mod['key']).first()
            if perm:
                changed_permissions = changed_permissions or perm.access_level != level
                perm.access_level = level
            else:
                changed_permissions = changed_permissions or level != 'none'
                db.session.add(GroupPermission(group_id=gid, module_key=mod['key'], access_level=level))
        db.session.flush()
        _sync_group_capabilities(gid, overwrite_existing=True)
    if changed_permissions:
        for member in g.users:
            member.session_version = int(member.session_version or 1) + 1
    db.session.commit()
    _audit('permissions_changed', 'users', g.name,
           'Existing sessions revoked.' if changed_permissions else 'Metadata only.')
    flash(f'Group "{g.name}" updated.', 'success')
    return redirect(url_for('users'))

@app.route('/admin/groups/<int:gid>/delete', methods=['POST'])
@login_required
@require_module('users', 'full')
def delete_group(gid):
    if not current_user.is_admin:
        flash('Only administrators can delete groups.', 'error')
        return redirect(url_for('users'))
    g = UserGroup.query.get_or_404(gid)
    if g.is_admin:
        flash('Cannot delete the Administrator group.', 'error')
        return redirect(url_for('users'))
    if g.users:
        flash('Cannot delete: group has assigned users.', 'error')
        return redirect(url_for('users'))
    gname = g.name
    db.session.delete(g)
    db.session.commit()
    _audit('delete_group', 'users', gname)
    flash(f'Group "{gname}" deleted.', 'success')
    return redirect(url_for('users'))


# ── NOTIFICATIONS MODULE ──────────────────────────────────────────────────────

def _mask_name(value):
    value = (value or '').strip()
    return (value[:1] + '***') if value else ''


def _mask_email(value):
    value = (value or '').strip()
    if '@' not in value:
        return '***' if value else ''
    local, domain = value.rsplit('@', 1)
    return f'{local[:1]}***@{domain}'


def _mask_phone(value):
    value = (value or '').strip()
    digits = re.sub(r'\D', '', value)
    return f'***-***-{digits[-4:]}' if len(digits) >= 4 else ('***' if value else '')


def _masked_subscriber(subscriber):
    contacts = [SimpleNamespace(
        first_name=_mask_name(contact.first_name),
        last_name=_mask_name(contact.last_name),
        email=', '.join(_mask_email(item) for item in (contact.email or '').split(',') if item.strip()),
        phone=_mask_phone(contact.phone),
        role=contact.role,
        preferred_language=contact.preferred_language or 'en',
        full_name=' '.join(filter(None, (
            _mask_name(contact.first_name), _mask_name(contact.last_name)))),
    ) for contact in subscriber.contacts]
    return SimpleNamespace(
        id=subscriber.id, notes=_mask_name(subscriber.notes), active=subscriber.active,
        group_id=subscriber.group_id, group=subscriber.group,
        school=subscriber.school or '',
        contacts=contacts, first_name=_mask_name(subscriber.first_name),
        last_name=_mask_name(subscriber.last_name), email=_mask_email(subscriber.email),
        phone=_mask_phone(subscriber.phone), full_name=(
            contacts[0].full_name if contacts else _mask_name(subscriber.full_name)),
    )


def _masked_user(user):
    return SimpleNamespace(
        username=user.username, first_name=_mask_name(user.first_name),
        last_name=_mask_name(user.last_name), email=_mask_email(user.email),
        receive_notifications=user.receive_notifications, group=user.group,
    )

@app.route('/admin/notifications')
@login_required
@require_module('notifications')
def notifications():
    subs           = NotificationSubscriber.query.order_by(NotificationSubscriber.last_name).all()
    groups         = SubscriberGroup.query.order_by(SubscriberGroup.name).all()
    all_buses      = Bus.query.filter_by(active=True).order_by(Bus.identifier).all()
    admin_users    = User.query.filter_by(active=True).order_by(User.username).all()
    schedule_types = BusScheduleType.query.order_by(BusScheduleType.sort_order).all()
    can_write      = current_user.has_capability('notifications.write')
    can_view_pii   = current_user.has_capability('notifications.pii')
    if not can_view_pii:
        subs = [_masked_subscriber(subscriber) for subscriber in subs]
        admin_users = [_masked_user(user) for user in admin_users]
    return render_template('admin/notifications.html',
                           subscribers=subs, groups=groups,
                           all_buses=all_buses, admin_users=admin_users,
                           schedule_types=schedule_types,
                           can_write=can_write, can_view_pii=can_view_pii,
                           can_export_pii=current_user.has_capability(
                               'notifications.export_pii'),
                           powerschool_enabled=app.config['POWERSCHOOL_IMPORT_ENABLED'],
                           can_powerschool_import=current_user.has_capability(
                               'import.powerschool'))

def _save_contacts(subscriber_id, form):
    """Read contact_{i}_* fields from form and create SubscriberContact records."""
    count = int(form.get('contact_count', 0) or 0)
    for i in range(min(count, 20)):
        fn = form.get(f'contact_{i}_first_name', '').strip()
        ln = form.get(f'contact_{i}_last_name',  '').strip()
        em = form.get(f'contact_{i}_email',       '').strip()
        ph = form.get(f'contact_{i}_phone',       '').strip()
        rl = form.get(f'contact_{i}_role',        'parent').strip()
        language = _normalize_language(
            form.get(f'contact_{i}_preferred_language', 'en'))
        if fn or em:
            db.session.add(SubscriberContact(
                subscriber_id=subscriber_id,
                first_name=fn or None, last_name=ln or None,
                email=em or None,      phone=ph or None,
                role=rl, preferred_language=language, sort_order=i,
            ))


def _delete_contact_external_identities(contact_ids):
    """Remove mappings for contacts that are about to lose their row identity."""
    contact_ids = {contact_id for contact_id in contact_ids if contact_id}
    if not contact_ids:
        return
    ExternalIdentity.query.filter(
        ExternalIdentity.local_table == 'subscriber_contact',
        ExternalIdentity.local_id.in_(contact_ids),
    ).delete(synchronize_session=False)


def _delete_subscriber_external_identities(subscriber):
    """Remove non-FK identity mappings before a subscriber can be deleted."""
    ExternalIdentity.query.filter(and_(
        ExternalIdentity.local_table == 'notification_subscriber',
        ExternalIdentity.local_id == subscriber.id,
    )).delete(synchronize_session=False)
    _delete_contact_external_identities(
        contact.id for contact in subscriber.contacts)

@app.route('/admin/notifications/add', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def add_subscriber():
    s = NotificationSubscriber(
        notes=request.form.get('notes', '').strip() or None,
        group_id=request.form.get('group_id', type=int) or None,
        school=_normalize_text(request.form.get('school'), 100) or None,
    )
    db.session.add(s)
    db.session.flush()
    _save_contacts(s.id, request.form)
    db.session.flush()
    db.session.expire(s, ['contacts'])
    _record_manual_subscriber_provenance(s, current_user)
    db.session.commit()
    _audit('add_subscriber', 'notifications', s.full_name)
    flash(f'Enrollment "{s.full_name}" added.', 'success')
    return redirect(url_for('notifications'))

@app.route('/admin/notifications/<int:sid>/edit', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def edit_subscriber(sid):
    s = NotificationSubscriber.query.get_or_404(sid)
    s.notes    = request.form.get('notes', '').strip() or None
    s.active   = 'active' in request.form
    s.group_id = request.form.get('group_id', type=int) or None
    s.school   = _normalize_text(request.form.get('school'), 100) or None
    _delete_contact_external_identities(
        contact.id for contact in s.contacts)
    SubscriberContact.query.filter_by(subscriber_id=sid).delete()
    _save_contacts(sid, request.form)
    db.session.commit()
    _audit('edit_subscriber', 'notifications', s.full_name)
    flash('Enrollment updated.', 'success')
    return redirect(url_for('notifications'))

@app.route('/admin/notifications/<int:sid>/delete', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def delete_subscriber(sid):
    s = NotificationSubscriber.query.get_or_404(sid)
    name = s.full_name
    _delete_subscriber_external_identities(s)
    db.session.delete(s)
    db.session.commit()
    _audit('delete_subscriber', 'notifications', name)
    flash('Subscriber removed.', 'success')
    return redirect(url_for('notifications'))


@app.route('/admin/notifications/bulk-delete', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def bulk_delete_subscribers():
    ids = request.form.getlist('subscriber_ids')
    count = 0
    for sid in ids:
        try:
            s = NotificationSubscriber.query.get(int(sid))
            if s:
                _delete_subscriber_external_identities(s)
                db.session.delete(s)
                count += 1
        except (ValueError, TypeError):
            pass
    if count:
        db.session.commit()
        _audit('bulk_delete_subscribers', 'notifications', f'{count} subscribers deleted')
        flash(f'{count} subscriber(s) deleted.', 'success')
    return redirect(url_for('notifications'))


@app.route('/admin/notifications/export-csv')
@login_required
@require_capability('notifications.export_pii')
def export_subscribers_csv():
    import csv, io
    subs = (NotificationSubscriber.query
            .order_by(NotificationSubscriber.id).all())
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['schema_version', 'subscriber_id', 'household_label', 'group', 'active',
                     'role', 'first_name', 'last_name', 'email', 'phone'])
    for sub in subs:
        group_name = sub.group.name if sub.group else ''
        active_str = 'yes' if sub.active else 'no'
        if sub.contacts:
            for c in sub.contacts:
                writer.writerow(_csv_safe_row([
                    'Legacy CSV v1', sub.id, sub.notes or '', group_name, active_str,
                    c.role or 'parent', c.first_name or '', c.last_name or '',
                    c.email or '', c.phone or '',
                ]))
        else:
            writer.writerow(_csv_safe_row([
                'Legacy CSV v1', sub.id, sub.notes or '', group_name, active_str,
                'parent', sub.first_name or '', sub.last_name or '',
                sub.email or '', sub.phone or '',
            ]))
    output.seek(0)
    from flask import Response
    return Response(
        '\ufeff' + output.getvalue(),   # BOM for Excel UTF-8 compatibility
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment;filename=subscribers.csv'})


def _auto_create_group_from_name(group_name):
    """Parse a group name like 'PC 01 AM PM' to auto-create a SubscriberGroup
    with the appropriate GroupBusAssignment records.

    Token rules:
      - Tokens before the first AM/MD/PM keyword → bus identifier
      - AM/MD/PM tokens → schedule periods (Morning/Midday/Afternoon)
    Returns the created SubscriberGroup, or None if the bus can't be found.
    """
    PERIOD_MAP = {'AM': 'Morning', 'MD': 'Midday', 'PM': 'Afternoon'}
    tokens = group_name.strip().split()

    bus_tokens = []
    period_tokens = []
    for tok in tokens:
        if tok.upper() in PERIOD_MAP and bus_tokens:   # only after we have some bus tokens
            period_tokens.append(tok.upper())
        elif not period_tokens:
            bus_tokens.append(tok)

    bus_identifier = ' '.join(bus_tokens).strip()
    period_names   = [PERIOD_MAP[t] for t in period_tokens]

    if not bus_identifier:
        return None

    # Bus identifier from group name is "PREFIX NUMBER" (e.g. "PC 01", "TT 55").
    # The DB stores these as separate fields: Bus.identifier="PC", Bus.name="1".
    # Strategy: split into prefix + number, try exact name match first, then
    # strip leading zeros (CSV may zero-pad numbers that the DB stores without zeros).
    bus = None
    parts = bus_identifier.rsplit(' ', 1)
    if len(parts) == 2:
        prefix, number = parts
        bus = Bus.query.filter(
            db.func.lower(Bus.identifier) == prefix.lower(),
            db.func.lower(Bus.name) == number.lower()
        ).first()
        if not bus:
            stripped = number.lstrip('0') or '0'
            if stripped != number:
                bus = Bus.query.filter(
                    db.func.lower(Bus.identifier) == prefix.lower(),
                    db.func.lower(Bus.name) == stripped.lower()
                ).first()
    if not bus:
        # Fallback: match identifier field exactly (single-token group names)
        bus = Bus.query.filter(
            db.func.lower(Bus.identifier) == bus_identifier.lower()
        ).first()
    if not bus:
        return None

    # Create the group
    grp = SubscriberGroup(name=group_name)
    db.session.add(grp)
    db.session.flush()   # get grp.id

    # Determine which schedule_type records to link
    bus_period_ids = {bsa.schedule_type_id for bsa in bus.schedule_assignments}

    if period_names:
        sched_types = BusScheduleType.query.filter(
            BusScheduleType.name.in_(period_names)
        ).all()
        assigned_any = False
        for st in sched_types:
            if st.id in bus_period_ids:
                db.session.add(GroupBusAssignment(
                    group_id=grp.id, bus_id=bus.id, schedule_type_id=st.id))
                assigned_any = True
        if not assigned_any:
            # Periods specified but none matched → assign all (NULL)
            db.session.add(GroupBusAssignment(
                group_id=grp.id, bus_id=bus.id, schedule_type_id=None))
    else:
        # No period tokens → assign all (NULL = all periods)
        db.session.add(GroupBusAssignment(
            group_id=grp.id, bus_id=bus.id, schedule_type_id=None))

    db.session.flush()
    return grp


def _store_import_stage(file_name, payload, content, headers, normalized_rows, report):
    public_id = secrets.token_urlsafe(32)
    file_sha = hashlib.sha256(payload).hexdigest()
    canonical_rows = json.dumps(normalized_rows, ensure_ascii=False,
                                sort_keys=True, separators=(',', ':'))
    plan_hash = hashlib.sha256(
        f'legacy_csv:1:{file_sha}:{canonical_rows}'.encode('utf-8')).hexdigest()
    now = _utcnow()
    batch = ImportBatch(
        public_id=public_id, source_type='legacy_csv', schema_version='1',
        status='staged' if not report['critical'] else 'blocked',
        snapshot_type='delta', uploaded_by_id=current_user.id,
        file_sha256=file_sha, plan_hash=plan_hash,
        total_rows=report['total_rows'],
        selected_rows=sum(1 for row in normalized_rows if row['classification'] == 'new'),
        rejected_rows=sum(1 for row in normalized_rows if row['classification'] == 'rejected'),
        excluded_rows=0, metadata_json=json.dumps(report, sort_keys=True),
        created_at=now,
        expires_at=now + timedelta(hours=app.config['IMPORT_STAGE_TTL_HOURS']))
    db.session.add(batch)
    db.session.flush()
    storage_path = os.path.join(IMPORT_STAGE_DIR, f'{public_id}.csv')
    _write_private_file(storage_path, payload, binary=True)
    db.session.add(ImportFile(
        batch_id=batch.id, file_type='legacy_csv',
        original_name=re.sub(r'[^A-Za-z0-9._ -]', '_', os.path.basename(file_name))[:255],
        sha256=file_sha, byte_size=len(payload), storage_path=storage_path,
        headers_json=json.dumps(headers)))
    for row in normalized_rows:
        normalized = row['normalized']
        normalized_json = json.dumps(normalized, ensure_ascii=False, sort_keys=True)
        db.session.add(ImportRow(
            batch_id=batch.id, row_number=row['row_number'],
            external_key=None, classification=row['classification'],
            selected=row['classification'] == 'new', normalized_json=normalized_json,
            errors_json=json.dumps(row.get('errors', [])),
            row_hash=hashlib.sha256(normalized_json.encode('utf-8')).hexdigest()))
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
        if os.path.isfile(storage_path):
            os.remove(storage_path)
        raise
    _audit('import_staged', 'notifications', public_id,
           f'Legacy CSV v1; {batch.total_rows} rows; sha256={file_sha[:12]}')
    return batch


def _purge_import_raw_files(batch):
    """Remove raw staged bytes while preserving normalized audit/change evidence."""
    stage_root = os.path.realpath(IMPORT_STAGE_DIR) + os.sep
    failures = []
    for staged_file in ImportFile.query.filter_by(batch_id=batch.id).all():
        path = os.path.realpath(staged_file.storage_path)
        try:
            if not path.startswith(stage_root):
                failures.append(staged_file.original_name)
                continue
            if os.path.isfile(path):
                os.remove(path)
            db.session.delete(staged_file)
        except OSError:
            failures.append(staged_file.original_name)
    return failures


def _expire_powerschool_stage(batch, processing_owner=False):
    """Expire one owned PowerSchool stage before it can mutate the roster."""
    expirable_statuses = {'staged'}
    if processing_owner:
        expirable_statuses.update({'selecting', 'applying'})
    if (batch.source_type != 'powerschool'
            or batch.status not in expirable_statuses
            or batch.expires_at > _utcnow()):
        return False
    cleanup_failures = _purge_import_raw_files(batch)
    ImportRow.query.filter_by(batch_id=batch.id).delete(
        synchronize_session=False)
    batch.status = 'expired'
    db.session.commit()
    _audit(
        'powerschool_import_expired', 'notifications', batch.public_id,
        f'Staged plan expired before mutation; cleanup_warnings='
        f'{len(cleanup_failures)}')
    return True


def _active_powerschool_roster_exists():
    """Return whether PowerSchool identities own any active subscriber."""
    return db.session.query(ExternalIdentity.id).join(
        NotificationSubscriber,
        NotificationSubscriber.id == ExternalIdentity.local_id,
    ).filter(
        ExternalIdentity.source_type == 'powerschool',
        ExternalIdentity.entity_type == 'student',
        ExternalIdentity.local_table == 'notification_subscriber',
        NotificationSubscriber.active.is_(True),
        NotificationSubscriber.created_at <= ExternalIdentity.created_at,
    ).first() is not None


def _recover_abandoned_import_processing(now):
    """Close expired worker claims only while the roster mutex is free.

    Apply and rollback hold the same cross-worker mutex for their full
    request. Acquiring it here proves that no live mutation still owns the
    claim. Ambiguous Apply evidence remains untouched for manual recovery.
    """
    processing_grace = now - timedelta(minutes=15)
    candidate_ids = [batch_id for (batch_id,) in db.session.query(
        ImportBatch.id,
    ).filter(
        ImportBatch.expires_at <= processing_grace,
        ImportBatch.status.in_(['selecting', 'applying', 'rolling_back']),
    ).all()]
    if not candidate_ids:
        return
    with _roster_import_lock() as acquired:
        if not acquired:
            return
        candidates = ImportBatch.query.filter(
            ImportBatch.id.in_(candidate_ids),
            ImportBatch.expires_at <= processing_grace,
            ImportBatch.status.in_(['selecting', 'applying', 'rolling_back']),
        ).with_for_update().all()
        changed = False
        for batch in candidates:
            metadata = _import_metadata(batch)
            if batch.status == 'selecting':
                batch.status = 'expired'
                action = 'powerschool_import_selection_abandoned'
                details = 'Expired selection claim closed after worker exit.'
            elif batch.status == 'rolling_back':
                # A successful compensating commit stores `rolled_back` in
                # the same transaction. A durable `rolling_back` claim means
                # the compensating transaction did not commit.
                batch.status = 'rollback_failed'
                metadata['rollback_failure'] = (
                    'Worker exited before the atomic rollback commit.')
                batch.metadata_json = json.dumps(metadata, sort_keys=True)
                action = 'powerschool_import_rollback_abandoned'
                details = 'Rollback transaction did not commit; retry remains available.'
            else:
                durable_state, _, _ = _powerschool_apply_durable_state(
                    batch.id, batch.selected_rows)
                if durable_state == 'clean_interruption':
                    batch.status = 'failed'
                    metadata['failure'] = (
                        'Worker exited before the atomic Apply commit.')
                    batch.metadata_json = json.dumps(metadata, sort_keys=True)
                    action = 'powerschool_import_apply_abandoned'
                    details = 'Apply transaction did not commit; reanalysis is required.'
                else:
                    if metadata.get('processing_cleanup_alerted_at'):
                        continue
                    metadata['processing_cleanup_alerted_at'] = (
                        now.isoformat() + 'Z')
                    batch.metadata_json = json.dumps(metadata, sort_keys=True)
                    action = 'powerschool_import_processing_inconsistent'
                    details = (
                        'Expired Apply claim has inconsistent durable evidence; '
                        'no status or roster data was changed.')
            db.session.add(AuditLog(
                username='system', action=action, module='notifications',
                target=batch.public_id, details=details,
                ip_address='background'))
            changed = True
        if changed:
            db.session.commit()


def _cleanup_import_stages():
    now = _utcnow()
    _recover_abandoned_import_processing(now)
    expired = ImportBatch.query.filter(
        ImportBatch.expires_at <= now,
        ImportBatch.status.in_([
            'staged', 'blocked', 'applied', 'failed', 'expired',
            'rollback_failed', 'rolled_back', 'retention_closed',
        ]),
    ).all()
    for batch in expired:
        failures = _purge_import_raw_files(batch)
        if batch.status not in {
                'applied', 'rolled_back', 'rollback_failed', 'retention_closed'}:
            ImportRow.query.filter_by(batch_id=batch.id).delete(synchronize_session=False)
            batch.status = 'expired'
        if failures:
            db.session.add(AuditLog(
                username='system', action='import_stage_cleanup_failed',
                module='notifications', target=batch.public_id,
                details=f'{len(failures)} raw file(s) could not be removed.',
                ip_address='background'))
    if expired:
        db.session.commit()

    retention_cutoff = now - timedelta(
        days=app.config['POWERSCHOOL_ROLLBACK_RETENTION_DAYS'])
    retained = ImportBatch.query.filter(
        ImportBatch.source_type == 'powerschool',
        ImportBatch.status.in_(['applied', 'rollback_failed', 'rolled_back']),
        func.coalesce(ImportBatch.applied_at, ImportBatch.created_at) <= retention_cutoff,
    ).all()
    for batch in retained:
        metadata = _import_metadata(batch)
        if metadata.get('pii_purged_at'):
            continue
        failures = _purge_import_raw_files(batch)
        if failures:
            db.session.add(AuditLog(
                username='system', action='powerschool_import_retention_pending',
                module='notifications', target=batch.public_id,
                details=(f'{len(failures)} raw file(s) could not be removed; '
                         'PII retention closure will be retried.'),
                ip_address='background'))
            continue
        for row in ImportRow.query.filter_by(batch_id=batch.id).all():
            external_key_sha256 = (
                hashlib.sha256(row.external_key.encode()).hexdigest()
                if row.external_key else None
            )
            minimal = {
                'retained': True, 'classification': row.classification,
                'external_key_sha256': external_key_sha256,
            }
            row.external_key = None
            row.normalized_json = json.dumps(minimal, sort_keys=True)
            row.errors_json = '[]'
        for change in ImportChange.query.filter_by(batch_id=batch.id).all():
            change.before_json = None
            change.after_json = None
        metadata['pii_purged_at'] = now.isoformat() + 'Z'
        metadata['rollback_retention_days'] = app.config[
            'POWERSCHOOL_ROLLBACK_RETENTION_DAYS']
        batch.metadata_json = json.dumps(metadata, sort_keys=True)
        if batch.status in {'applied', 'rollback_failed'}:
            batch.status = 'retention_closed'
        db.session.add(AuditLog(
            username='system', action='powerschool_import_pii_purged',
            module='notifications', target=batch.public_id,
            details='Normalized PII and compensating snapshots purged after retention.',
            ip_address='background'))
    if retained:
        db.session.commit()


@app.route('/admin/notifications/import-csv/preview', methods=['POST'])
@login_required
@require_capability('import.legacy')
def preview_import_csv():
    """Normalize once into immutable staging and return a review report."""
    import csv, io
    from collections import OrderedDict

    _cleanup_import_stages()
    if _active_powerschool_roster_exists():
        return jsonify({
            'ok': False,
            'message': (
                'Legacy CSV import is disabled because an active PowerSchool '
                'roster is authoritative. Use PowerSchool Import for roster '
                'updates.'),
        }), 409
    file = request.files.get('csv_file')
    if not _valid_csv_upload(file):
        return jsonify({'ok': False, 'message':
                        'Select a CSV file with an approved content type.'}), 400
    try:
        payload = file.read()
        content = payload.decode('utf-8-sig')
    except UnicodeDecodeError:
        return jsonify({'ok': False, 'message': 'File must be UTF-8 encoded.'})

    PERIOD_MAP = {'AM': 'Morning', 'MD': 'Midday', 'PM': 'Afternoon'}

    def _resolve_bus(group_name):
        """Return (bus, period_names) for a group name without touching the DB."""
        tokens = group_name.strip().split()
        bus_tokens, period_tokens = [], []
        for tok in tokens:
            if tok.upper() in PERIOD_MAP and bus_tokens:
                period_tokens.append(tok.upper())
            elif not period_tokens:
                bus_tokens.append(tok)
        bus_identifier = ' '.join(bus_tokens).strip()
        period_names   = [PERIOD_MAP[t] for t in period_tokens]
        if not bus_identifier:
            return None, []
        bus = None
        parts = bus_identifier.rsplit(' ', 1)
        if len(parts) == 2:
            prefix, number = parts
            bus = Bus.query.filter(
                db.func.lower(Bus.identifier) == prefix.lower(),
                db.func.lower(Bus.name) == number.lower()
            ).first()
            if not bus:
                stripped = number.lstrip('0') or '0'
                if stripped != number:
                    bus = Bus.query.filter(
                        db.func.lower(Bus.identifier) == prefix.lower(),
                        db.func.lower(Bus.name) == stripped.lower()
                    ).first()
        if not bus:
            bus = Bus.query.filter(
                db.func.lower(Bus.identifier) == bus_identifier.lower()
            ).first()
        return bus, period_names

    existing_groups = {g.name.strip().lower(): g for g in SubscriberGroup.query.all()}

    # Per-group analysis (resolved once per unique group name)
    groups_info   = {}   # lower_name → dict
    households    = OrderedDict()
    total_rows    = 0
    skipped_blank = 0
    warn_no_email = 0
    warn_no_phone = 0
    skipped_bus   = 0
    normalized_rows = []

    reader = csv.DictReader(io.StringIO(content))
    headers = reader.fieldnames or []
    if (not headers or len(headers) > app.config['IMPORT_MAX_COLUMNS'] or
            len(headers) != len(set(headers)) or 'group' not in headers or
            not ({'first_name', 'email'} & set(headers))):
        return jsonify({'ok': False, 'message':
                        'CSV headers are missing, duplicated, or do not match Legacy CSV v1.'}), 400
    for i, row in enumerate(reader, 2):
        if total_rows >= app.config['IMPORT_MAX_ROWS']:
            return jsonify({'ok': False, 'message':
                            f'CSV exceeds the {app.config["IMPORT_MAX_ROWS"]} row limit.'}), 400
        total_rows += 1
        group_name = (row.get('group') or '').strip()
        household  = (row.get('household_label') or '').strip()
        first_name = (row.get('first_name') or '').strip()
        last_name  = (row.get('last_name') or '').strip()
        email      = _normalize_email(row.get('email'))
        phone      = _normalize_phone(row.get('phone'))
        role_raw   = (row.get('role') or 'parent').strip().lower()
        role       = role_raw if role_raw in ('parent', 'student') else 'parent'
        school     = _normalize_text(row.get('school'), 100)
        language   = _normalize_language(row.get('preferred_language'))
        normalized = {
            'group': group_name, 'household_label': household,
            'first_name': first_name, 'last_name': last_name,
            'email': email, 'phone': phone, 'role': role,
            'school': school, 'preferred_language': language,
        }

        if not first_name and not email:
            skipped_blank += 1
            normalized_rows.append({'row_number': i, 'classification': 'rejected',
                                    'normalized': normalized,
                                    'errors': ['A name or email is required.']})
            continue

        gkey = group_name.lower() if group_name else ''

        # Resolve group once
        if group_name and gkey not in groups_info:
            if gkey in existing_groups:
                groups_info[gkey] = {'name': group_name, 'status': 'existing',
                                     'bus': None, 'periods': [], 'rows': 0}
            else:
                bus, period_names = _resolve_bus(group_name)
                if bus:
                    bus_period_ids = {bsa.schedule_type_id for bsa in bus.schedule_assignments}
                    matched = []
                    if period_names:
                        sts = BusScheduleType.query.filter(BusScheduleType.name.in_(period_names)).all()
                        matched = [st.name for st in sts if st.id in bus_period_ids]
                    groups_info[gkey] = {
                        'name':    group_name,
                        'status':  'create',
                        'bus':     f'{bus.identifier} - {bus.name}',
                        'periods': matched or ['All periods'],
                        'rows':    0,
                    }
                else:
                    groups_info[gkey] = {'name': group_name, 'status': 'error',
                                         'bus': None, 'periods': [], 'rows': 0}

        # Skip rows whose group has an unresolvable bus
        if group_name and groups_info.get(gkey, {}).get('status') == 'error':
            skipped_bus += 1
            groups_info[gkey]['rows'] += 1
            normalized_rows.append({'row_number': i, 'classification': 'rejected',
                                    'normalized': normalized,
                                    'errors': ['Group bus could not be resolved.']})
            continue

        if group_name and gkey in groups_info:
            groups_info[gkey]['rows'] += 1

        if not email:
            warn_no_email += 1
        if not phone:
            warn_no_phone += 1

        hh_key = (gkey or '__none__', household if household else f'__row_{i}__')
        if hh_key not in households:
            households[hh_key] = 0
        households[hh_key] += 1
        normalized_rows.append({'row_number': i, 'classification': 'new',
                                'normalized': normalized, 'errors': []})

    critical = [
        f'Group "{v["name"]}": bus not found — {v["rows"]} row(s) will be skipped'
        for v in groups_info.values() if v['status'] == 'error'
    ]

    report = {
        'ok':            True,
        'total_rows':    total_rows,
        'skipped_blank': skipped_blank,
        'skipped_bus':   skipped_bus,
        'households':    len(households),
        'contacts':      sum(households.values()),
        'groups_create': [v for v in groups_info.values() if v['status'] == 'create'],
        'groups_existing': [v['name'] for v in groups_info.values() if v['status'] == 'existing'],
        'critical':      critical,
        'warn_no_email': warn_no_email,
        'warn_no_phone': warn_no_phone,
        'can_import':    len(critical) == 0,
    }
    batch = _store_import_stage(
        file.filename, payload, content, headers, normalized_rows, report)
    report.update({'batch_id': batch.public_id, 'plan_hash': batch.plan_hash,
                   'schema_version': 'Legacy CSV v1'})
    return jsonify(report)


@app.route('/admin/notifications/import-csv', methods=['POST'])
@login_required
@require_capability('import.legacy')
@_serialized_roster_mutation('html')
def import_subscribers_csv():
    batch_id = (request.form.get('batch_id') or '').strip()
    plan_hash = (request.form.get('plan_hash') or '').strip()
    if not batch_id or not plan_hash:
        flash('Analyze the CSV before confirming the import.', 'error')
        return redirect(url_for('notifications'))

    batch = ImportBatch.query.filter_by(public_id=batch_id).first()
    if not batch or batch.source_type != 'legacy_csv':
        abort(404)
    if batch.uploaded_by_id != current_user.id and not current_user.is_admin:
        abort(403)
    if batch.expires_at <= _utcnow():
        batch.status = 'expired'
        db.session.commit()
        flash('The analyzed import expired. Analyze the file again.', 'error')
        return redirect(url_for('notifications'))
    if not secrets.compare_digest(batch.plan_hash, plan_hash):
        abort(409)
    if batch.status == 'applied':
        flash('This analyzed batch was already imported; no records were duplicated.', 'warning')
        return redirect(url_for('notifications'))
    if batch.status != 'staged':
        flash('This import batch is not eligible to be applied.', 'error')
        return redirect(url_for('notifications'))
    if _active_powerschool_roster_exists():
        flash(
            'Legacy CSV import is disabled because an active PowerSchool '
            'roster is authoritative. Use PowerSchool Import for roster updates.',
            'error')
        return redirect(url_for('notifications'))

    # A conditional UPDATE is atomic on both SQLite and PostgreSQL. It closes
    # the double-submit race that SELECT ... FOR UPDATE cannot close on SQLite.
    claimed = ImportBatch.query.filter_by(
        id=batch.id, status='staged').update(
            {'status': 'applying'}, synchronize_session=False)
    db.session.commit()
    if claimed != 1:
        flash('This analyzed batch is already being applied.', 'warning')
        return redirect(url_for('notifications'))
    batch = db.session.get(ImportBatch, batch.id)

    rows = ImportRow.query.filter_by(
        batch_id=batch.id, selected=True, classification='new',
    ).order_by(ImportRow.row_number).all()
    households = {}
    groups_cache = {group.name.strip().lower(): group
                    for group in SubscriberGroup.query.all()}
    created_groups = []
    try:
        for row in rows:
            normalized = json.loads(row.normalized_json)
            group_name = normalized.get('group', '').strip()
            group_obj = groups_cache.get(group_name.lower()) if group_name else None
            if group_name and not group_obj:
                group_obj = _auto_create_group_from_name(group_name)
                if not group_obj:
                    raise ValueError(f'Group "{group_name}" can no longer be resolved.')
                groups_cache[group_name.lower()] = group_obj
                created_groups.append(group_name)
            household = normalized.get('household_label', '').strip()
            school = _normalize_text(normalized.get('school'), 100)
            key = (group_obj.id if group_obj else None, school,
                   household if household else f'__row_{row.row_number}__')
            households.setdefault(key, {
                'group_id': group_obj.id if group_obj else None,
                'school': school or None, 'notes': household or None,
                'contacts': [], 'rows': [],
            })
            households[key]['contacts'].append({
                'first_name': normalized.get('first_name') or None,
                'last_name': normalized.get('last_name') or None,
                'email': normalized.get('email') or None,
                'phone': normalized.get('phone') or None,
                'role': normalized.get('role') or 'parent',
                'preferred_language': _normalize_language(
                    normalized.get('preferred_language')),
            })
            households[key]['rows'].append(row)

        imported = 0
        for household in households.values():
            subscriber = NotificationSubscriber(
                notes=household['notes'], group_id=household['group_id'],
                school=household['school'], active=True)
            db.session.add(subscriber)
            db.session.flush()
            for index, contact in enumerate(household['contacts']):
                db.session.add(SubscriberContact(
                    subscriber_id=subscriber.id, sort_order=index, **contact))
            for row in household['rows']:
                db.session.add(ImportChange(
                    batch_id=batch.id, row_id=row.id, operation='create',
                    target_table='notification_subscriber', target_id=subscriber.id,
                    after_json=json.dumps({
                        'subscriber_id': subscriber.id,
                        'group_id': subscriber.group_id,
                    }, sort_keys=True)))
            imported += 1

        batch.status = 'applied'
        batch.applied_at = _utcnow()
        batch.selected_rows = len(rows)
        db.session.commit()
        cleanup_failures = _purge_import_raw_files(batch)
        db.session.commit()
        _audit('import_applied', 'notifications', batch.public_id,
               f'Legacy CSV v1; {imported} subscribers; {len(rows)} rows')
        if cleanup_failures:
            _audit('import_stage_cleanup_failed', 'notifications', batch.public_id,
                   f'{len(cleanup_failures)} raw file(s) could not be removed.')
        message = f'{imported} subscriber(s) imported from the analyzed batch.'
        if created_groups:
            message += f' {len(created_groups)} group(s) auto-created.'
        flash(message, 'success')
    except Exception:
        db.session.rollback()
        failed = ImportBatch.query.filter_by(public_id=batch_id).first()
        if failed and failed.status != 'applied':
            failed.status = 'failed'
            db.session.commit()
        _audit('import_failed', 'notifications', batch_id,
               'Atomic transaction rolled back.')
        flash('Import failed and no subscriber records were changed.', 'error')
    return redirect(url_for('notifications'))


# ── POWERSCHOOL STAGED IMPORT ───────────────────────────────────────────────

def _powerschool_enabled():
    if not app.config['POWERSCHOOL_IMPORT_ENABLED']:
        abort(404)


@app.route('/admin/notifications/powerschool-guide')
@login_required
@require_module('notifications')
def powerschool_import_guide():
    return render_template(
        'admin/powerschool_guide.html',
        powerschool_enabled=app.config['POWERSCHOOL_IMPORT_ENABLED'],
        can_powerschool_import=current_user.has_capability('import.powerschool'))


def _import_metadata(batch):
    try:
        value = json.loads(batch.metadata_json or '{}')
        return value if isinstance(value, dict) else {}
    except (TypeError, ValueError):
        return {}


def _powerschool_analysis_context_hash(profile, mapping, combined_file_sha256,
                                       school_year, snapshot_type):
    """Bind duplicate detection to every input that can change normalization."""
    material = {
        'source_type': 'powerschool',
        'schema_version': profile.schema_version,
        'mapping_profile_id': profile.id,
        'mapping': mapping,
        'normalizer_revision': NORMALIZER_REVISION,
        'combined_file_sha256': combined_file_sha256,
        'school_year': school_year,
        'snapshot_type': snapshot_type,
    }
    canonical = json.dumps(
        material, ensure_ascii=False, sort_keys=True,
        separators=(',', ':')).encode('utf-8')
    return hashlib.sha256(canonical).hexdigest()


def _powerschool_preflight_failure(parsed, snapshot_type='delta'):
    """Return a safe blocking diagnosis, or None when staging may continue."""
    preflight = parsed.get('preflight') or {}
    valid_transport = int(preflight.get('valid_transport_rows') or 0)
    errors = preflight.get('errors') or []
    if valid_transport < 1:
        return {
            'code': 'no_valid_transportation_rows',
            'message': (
                'Transportation contains no valid bus assignments. '
                'Re-export D205 BusRoute - Transportation v2 from PowerSchool '
                'and verify that route contains usable bus routes. No batch '
                'was created.'),
            'errors': errors or ['No valid transportation assignments were found.'],
        }
    if preflight.get('ok') is False:
        return {
            'code': 'powerschool_preflight_failed',
            'message': (
                'The PowerSchool exports failed preflight validation. '
                'No batch was created.'),
            'errors': errors,
        }
    if (snapshot_type == 'full_district'
            and preflight.get('transportation_contract')
            != TRANSPORTATION_V2_CONTRACT):
        return {
            'code': 'powerschool_full_snapshot_requires_transportation_v2',
            'message': (
                'Full Snapshot requires the approved Transportation v2 '
                'dual-route export. Use Delta for legacy/v1 files or '
                're-export template 941; no batch was created.'),
            'errors': [
                'Transportation v1 cannot prove district-wide AM/PM coverage.'
            ],
        }
    if snapshot_type == 'full_district':
        transport = (parsed.get('metrics') or {}).get('transportation') or {}
        missing_periods = [
            period for period, key in (
                ('AM', 'period_am_rows'), ('PM', 'period_pm_rows'))
            if int(transport.get(key) or 0) < 1
        ]
        anomaly_count = sum(int(transport.get(key) or 0) for key in (
            'invalid_route_am_rows', 'invalid_route_pm_rows',
            'route_am_period_conflict_rows',
            'route_pm_period_conflict_rows',
        ))
        if missing_periods or anomaly_count:
            details = []
            if missing_periods:
                details.append(
                    'The export contains no usable '
                    + '/'.join(missing_periods)
                    + ' route assignments.')
            if anomaly_count:
                details.append(
                    f'{anomaly_count} directional route value(s) are invalid '
                    'or contradict their AM/PM column.')
            return {
                'code': 'transportation_v2_full_snapshot_not_proven',
                'message': (
                    'Transportation v2 did not prove a clean AM and PM '
                    'district snapshot. Select Delta or correct the reported '
                    'route anomalies and analyze again; no batch was created.'),
                'errors': details,
            }
    return None


def _powerschool_apply_contract_error(batch):
    """Reject staged Full Snapshot plans that predate the v2-only gate."""
    if batch.snapshot_type != 'full_district':
        return None
    preflight = (_import_metadata(batch).get('preflight') or {})
    if preflight.get('transportation_contract') == TRANSPORTATION_V2_CONTRACT:
        return None
    return (
        'This Full Snapshot was not analyzed from the approved '
        'Transportation v2 dual-route export. Analyze the three files again; '
        'the stored plan cannot be applied.')


def _snapshot_datetime(value):
    """Serialize an audit timestamp consistently across supported databases."""
    if value is None:
        return None
    if value.tzinfo is not None:
        value = value.astimezone(timezone.utc).replace(tzinfo=None)
    return value.isoformat(timespec='microseconds')


def _row_incarnation_matches(created_at, expected, change_created_at):
    """Bind an audit target to the row incarnation that produced the change.

    New snapshots carry the exact creation timestamp. Older snapshots did
    not, so their safe compatibility boundary is that the target row must have
    existed no later than the immutable ImportChange record. Missing
    timestamps cannot prove identity and therefore fail closed.
    """
    if (created_at is None or change_created_at is None
            or created_at > change_created_at):
        return False
    if 'created_at' not in expected:
        return True
    expected_created_at = expected.get('created_at')
    return (expected_created_at is not None
            and _snapshot_datetime(created_at) == expected_created_at)


def _subscriber_snapshot(subscriber, contact_identities=None):
    contacts = []
    for contact in sorted(subscriber.contacts, key=lambda item: item.id):
        if contact_identities is None:
            identities = ExternalIdentity.query.filter_by(
                source_type='powerschool', local_table='subscriber_contact',
                local_id=contact.id).order_by(ExternalIdentity.entity_type,
                                              ExternalIdentity.external_key).all()
        else:
            identities = contact_identities.get(contact.id, [])
        identities = sorted(
            identities,
            key=lambda item: (item.entity_type, item.external_key),
        )
        contacts.append({
            'id': contact.id, 'first_name': contact.first_name,
            'last_name': contact.last_name, 'email': contact.email,
            'phone': contact.phone, 'role': contact.role,
            'preferred_language': contact.preferred_language or 'en',
            'sort_order': contact.sort_order,
            'identities': [[item.entity_type, item.external_key]
                           for item in identities],
        })
    return {
        'id': subscriber.id, 'notes': subscriber.notes,
        'active': bool(subscriber.active), 'group_id': subscriber.group_id,
        'school': subscriber.school,
        'created_at': _snapshot_datetime(subscriber.created_at),
        'contacts': contacts,
    }


def _subscriber_matches_snapshot(subscriber, expected,
                                 contact_identities=None,
                                 change_created_at=None):
    """Compare a subscriber while accepting pre-incarnation snapshots."""
    if not subscriber or not isinstance(expected, dict):
        return False
    if not _row_incarnation_matches(
            subscriber.created_at, expected, change_created_at):
        return False
    current = _subscriber_snapshot(subscriber, contact_identities)
    if 'school' not in expected:
        current.pop('school', None)
    expected_contacts = expected.get('contacts') or []
    current_contacts = current.get('contacts') or []
    if all('preferred_language' not in contact for contact in expected_contacts):
        for contact in current_contacts:
            contact.pop('preferred_language', None)
    if 'created_at' not in expected:
        current.pop('created_at', None)
    return _snapshot_hash(current) == _snapshot_hash(expected)


def _contact_identity_map(subscribers, for_update=False):
    contact_ids = {
        contact.id
        for subscriber in subscribers
        for contact in subscriber.contacts
    }
    if not contact_ids:
        return {}
    query = ExternalIdentity.query.filter(
        ExternalIdentity.source_type == 'powerschool',
        ExternalIdentity.local_table == 'subscriber_contact',
        ExternalIdentity.local_id.in_(contact_ids),
    ).order_by(
        ExternalIdentity.local_id,
        ExternalIdentity.entity_type,
        ExternalIdentity.external_key,
    )
    if for_update:
        query = query.with_for_update()
    identities = query.all()
    result = {}
    for identity in identities:
        result.setdefault(identity.local_id, []).append(identity)
    return result


def _snapshot_hash(snapshot):
    if snapshot is None:
        return 'absent'
    payload = json.dumps(snapshot, ensure_ascii=False, sort_keys=True,
                         separators=(',', ':'))
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()


def _subscriber_group_snapshot(group):
    if group is None:
        return None
    return {
        'id': group.id,
        'name': group.name,
        'description': group.description or '',
        'color': group.color or '',
        'created_at': _snapshot_datetime(group.created_at),
        'assignments': sorted([
            [item.bus_id, item.schedule_type_id]
            for item in group.bus_assignments
        ], key=lambda item: (item[0], item[1] or 0)),
    }


def _subscriber_group_matches_snapshot(group, expected,
                                       change_created_at=None):
    """Compare a group while accepting older snapshots with fewer fields."""
    if not group or not isinstance(expected, dict):
        return False
    expected = dict(expected)
    if not _row_incarnation_matches(
            group.created_at, expected, change_created_at):
        return False
    # PowerSchool-created groups have always used these fixed values.  Filling
    # them for older audit snapshots avoids ignoring later edits merely because
    # those two fields were not serialized by the original recorder.
    expected.setdefault('description', 'Created by PowerSchool Import v1')
    expected.setdefault('color', 'blue')
    current = _subscriber_group_snapshot(group)
    if 'created_at' not in expected:
        current.pop('created_at', None)
    return all(current.get(key) == value for key, value in expected.items())


def _powerschool_identity(entity_type, external_key, identity_cache=None):
    if identity_cache is not None:
        return identity_cache.get((entity_type, external_key))
    return ExternalIdentity.query.filter_by(
        source_type='powerschool', entity_type=entity_type,
        external_key=external_key).first()


def _powerschool_subscriber_identity_current(identity, subscriber):
    """Prove that a student identity belongs to this row incarnation."""
    return bool(
        identity
        and subscriber
        and identity.local_table == 'notification_subscriber'
        and identity.local_id == subscriber.id
        and identity.created_at is not None
        and subscriber.created_at is not None
        and subscriber.created_at <= identity.created_at
    )


def _powerschool_bus_for_route(prefix, number, period=None):
    """Resolve a source route against the canonical identity of active buses.

    Numeric buses commonly store ``identifier=TT, name=55``. Some district
    buses instead store compound names such as ``identifier=TR, name=ALG1`` or
    encode a schedule token in the name (``identifier=MCK1, name=AM``). Parsing
    the combined local identity prevents those valid routes from being missed
    without introducing district-specific aliases.
    """
    desired = normalize_route(f'{prefix or ""} {number or ""}')
    if not desired:
        return None
    desired_key = (desired['prefix'], desired['number'])
    cache_key = '_powerschool_active_bus_index'
    index = getattr(g, cache_key, None) if has_request_context() else None
    if index is None:
        index = {}
        for bus in Bus.query.options(
                selectinload(Bus.schedule_assignments),
        ).filter(Bus.active.is_(True)).order_by(Bus.id).all():
            parsed = normalize_route(f'{bus.identifier or ""} {bus.name or ""}')
            if not parsed:
                continue
            key = (parsed['prefix'], parsed['number'])
            index.setdefault(key, []).append(bus)
        if has_request_context():
            setattr(g, cache_key, index)
    matches = index.get(desired_key, [])
    if len(matches) == 1:
        return matches[0]
    if period and matches:
        period_matches = [
            bus for bus in matches
            if str(bus.name or '').strip().upper() == period
        ]
        if len(period_matches) == 1:
            return period_matches[0]
    return None


def _powerschool_group(proposal, create=False, assignment_index=None,
                       schedule_types_by_name=None):
    assignments = proposal.get('assignments') or []
    if not assignments:
        return None, False, 'The proposal contains no bus assignments.'
    period_map = {'AM': 'Morning', 'MD': 'Midday', 'PM': 'Afternoon'}
    raw_periods = {item.get('period') for item in assignments}
    unknown_periods = raw_periods - {'AM', 'MD', 'PM', 'ALL'}
    if unknown_periods:
        return None, False, 'One or more normalized periods are not configured.'
    desired_names = {period_map[item] for item in raw_periods if item in period_map}
    if schedule_types_by_name is None:
        types = BusScheduleType.query.filter(
            BusScheduleType.name.in_(desired_names)).all() if desired_names else []
        types_by_name = {item.name: item for item in types}
    else:
        types_by_name = {
            name: schedule_types_by_name[name]
            for name in desired_names if name in schedule_types_by_name
        }
    if len(types_by_name) != len(desired_names):
        return None, False, 'One or more normalized periods are not configured.'

    desired_assignments = set()
    resolved = []
    for item in assignments:
        period = item.get('period')
        bus = _powerschool_bus_for_route(
            item.get('route_prefix'), item.get('route_number'), period)
        if not bus:
            return None, False, (
                f'The normalized {period or "unknown"} route does not match '
                'an active bus.')
        schedule_type_id = None
        if period != 'ALL':
            schedule_type = types_by_name[period_map[period]]
            bus_period_ids = {
                configured.schedule_type_id
                for configured in bus.schedule_assignments
            }
            if schedule_type.id not in bus_period_ids:
                return None, False, (
                    f'The matched bus is not configured for the proposed '
                    f'{period} period.')
            schedule_type_id = schedule_type.id
        desired_assignments.add((bus.id, schedule_type_id))
        resolved.append((period, bus))

    all_period_bus_ids = {
        bus_id for bus_id, schedule_type_id in desired_assignments
        if schedule_type_id is None
    }
    if all_period_bus_ids:
        desired_assignments = {
            (bus_id, schedule_type_id)
            for bus_id, schedule_type_id in desired_assignments
            if schedule_type_id is None or bus_id not in all_period_bus_ids
        }

    assignment_key = tuple(sorted(
        desired_assignments, key=lambda value: (value[0], value[1] or 0)))
    if assignment_index is not None:
        group = assignment_index.get(assignment_key)
        if group:
            return group, False, None
    else:
        for group in SubscriberGroup.query.order_by(SubscriberGroup.id).all():
            group_assignments = {
                (item.bus_id, item.schedule_type_id)
                for item in group.bus_assignments
            }
            if group_assignments and group_assignments == desired_assignments:
                return group, False, None

    buses = {bus.id: bus for _, bus in resolved}
    if len(buses) == 1:
        bus = next(iter(buses.values()))
        compact_bus = f'{bus.identifier}{bus.name}'
        period_tokens = ([] if 'ALL' in raw_periods else [
            item for item in ('AM', 'MD', 'PM') if item in raw_periods
        ])
        base_name = ' '.join([compact_bus] + period_tokens)
    else:
        order = {'AM': 0, 'MD': 1, 'PM': 2, 'ALL': 3}
        tokens = []
        for period, bus in sorted(
                resolved,
                key=lambda pair: (order.get(pair[0], 9),
                                  pair[1].identifier, pair[1].name)):
            token = f'{bus.identifier}{bus.name}'
            if period != 'ALL':
                token += f' {period}'
            if token not in tokens:
                tokens.append(token)
        base_name = ' / '.join(tokens)
    if not create:
        return SimpleNamespace(id=None, name=base_name[:100]), False, None

    name = base_name[:100]
    existing = SubscriberGroup.query.filter(func.lower(
        SubscriberGroup.name) == name.lower()).first()
    if existing:
        return None, False, 'A group with the proposed name has different assignments.'
    group = SubscriberGroup(
        name=name, description='Created by PowerSchool Import v1', color='blue')
    db.session.add(group)
    db.session.flush()
    for bus_id, schedule_type_id in sorted(
            desired_assignments,
            key=lambda value: (value[0], value[1] or 0)):
        db.session.add(GroupBusAssignment(
            group_id=group.id, bus_id=bus_id,
            schedule_type_id=schedule_type_id))
    db.session.flush()
    if assignment_index is not None:
        assignment_index[assignment_key] = group
    return group, True, None


def _powerschool_household_label(proposal):
    household = proposal.get('household_id') or proposal['student_number']
    suffix = hashlib.sha256(
        f'powerschool:{household}'.encode('utf-8')).hexdigest()[:8].upper()
    name = ' '.join(filter(None, [proposal.get('first_name'),
                                  proposal.get('last_name')])).strip()
    return f'{name or "Student"} [PS-{suffix}]'[:200]


def _powerschool_contact_specs(proposal):
    student_number = proposal['student_number']
    student_spec = {
        'first_name': proposal.get('first_name') or '',
        'last_name': proposal.get('last_name') or '',
        'email': '', 'phone': '', 'role': 'student',
        'preferred_language': _normalize_language(
            proposal.get('preferred_language')),
        'identities': [('student_contact', student_number)],
    }
    specs = []
    student_relationships = {'student', 'self', 'child', 'pupil'}
    for contact in proposal.get('contacts', []):
        relationship = (contact.get('relationship') or '').lower()
        identity = ('contact', f'{student_number}|{contact["contact_id"]}')
        if relationship in student_relationships:
            for field in ('first_name', 'last_name', 'email', 'phone'):
                if contact.get(field):
                    student_spec[field] = contact[field]
            student_spec['identities'].append(identity)
        else:
            specs.append({
                'first_name': contact.get('first_name') or '',
                'last_name': contact.get('last_name') or '',
                'email': contact.get('email') or '',
                'phone': contact.get('phone') or '', 'role': 'parent',
                'preferred_language': _normalize_language(
                    contact.get('preferred_language')),
                'identities': [identity],
            })
    if student_spec['first_name'] or student_spec['email'] or student_spec['phone']:
        specs.insert(0, student_spec)
    return specs


def _powerschool_compare_proposal(proposal):
    conflicts = list(dict.fromkeys(proposal.get('conflicts') or []))
    group, _, group_error = _powerschool_group(proposal, create=False)
    if group_error:
        conflicts.append(group_error)
    student_identity = _powerschool_identity('student', proposal['student_number'])
    subscriber = None
    if student_identity:
        if student_identity.local_table != 'notification_subscriber':
            conflicts.append('Student identity points to an unexpected local table.')
        else:
            identity_subscriber = db.session.get(
                NotificationSubscriber, student_identity.local_id)
            if not identity_subscriber:
                conflicts.append('Student identity points to a missing enrollment.')
            elif not _powerschool_subscriber_identity_current(
                    student_identity, identity_subscriber):
                conflicts.append(
                    'Student identity points to a stale enrollment incarnation.')
            else:
                subscriber = identity_subscriber

    specs = _powerschool_contact_specs(proposal)
    changes = []
    if subscriber and group:
        expected_label = _powerschool_household_label(proposal)
        current_group_name = subscriber.group.name if subscriber.group else ''
        if ((group.id and subscriber.group_id != group.id)
                or (not group.id and current_group_name != group.name)):
            changes.append({'field': 'group',
                            'current': current_group_name,
                            'proposed': group.name})
        if not subscriber.active:
            changes.append({'field': 'active', 'current': 'no', 'proposed': 'yes'})
        if subscriber.notes != expected_label:
            changes.append({'field': 'household_label',
                            'current': subscriber.notes or '',
                            'proposed': expected_label})
        proposed_school = _normalize_text(proposal.get('school'), 100)
        if (subscriber.school or '') != proposed_school:
            changes.append({'field': 'school', 'current': subscriber.school or '',
                            'proposed': proposed_school})
        for spec in specs:
            mapped = []
            for entity_type, external_key in spec['identities']:
                identity = _powerschool_identity(entity_type, external_key)
                if identity:
                    if identity.local_table != 'subscriber_contact':
                        conflicts.append('Contact identity points to an unexpected table.')
                        continue
                    contact = db.session.get(SubscriberContact, identity.local_id)
                    if not contact or contact.subscriber_id != subscriber.id:
                        conflicts.append('Contact identity conflicts with this enrollment.')
                    else:
                        mapped.append(contact)
            if mapped and len({item.id for item in mapped}) > 1:
                conflicts.append('Equivalent contact identities point to different contacts.')
            contact = mapped[0] if mapped else None
            if not contact:
                changes.append({'field': 'contact', 'current': '',
                                'proposed': ' '.join(filter(None, [
                                    spec['first_name'], spec['last_name']]))})
            else:
                for field in ('first_name', 'last_name', 'email', 'phone', 'role',
                              'preferred_language'):
                    current = getattr(contact, field) or ''
                    proposed = spec[field] or ''
                    if current != proposed:
                        changes.append({'field': f'contact.{field}',
                                        'current': current, 'proposed': proposed})

    snapshot = _subscriber_snapshot(subscriber) if subscriber else None
    proposal['group_name'] = group.name if group else ''
    proposal['changes'] = changes
    proposal['expected_state_hash'] = _snapshot_hash(snapshot)
    proposal['target_subscriber_id'] = subscriber.id if subscriber else None
    if conflicts:
        return 'conflict', False, list(dict.fromkeys(conflicts))
    if not subscriber:
        return 'new', True, []
    return ('update', True, []) if changes else ('unchanged', False, [])


LEGACY_BASELINE_SCHEMA_VERSION = 'baseline-1'
LEGACY_BASELINE_KIND = 'legacy_roster_provenance_baseline'
MANUAL_PROVENANCE_SOURCE_TYPE = 'manual'
MANUAL_PROVENANCE_SCHEMA_VERSION = 'subscriber-1'
MANUAL_PROVENANCE_KIND = 'manual_subscriber_provenance'


def _legacy_baseline_integrity():
    """Validate the one-time baseline without reopening its historical CSV."""
    batches = _legacy_baseline_existing_batches()
    if not batches:
        return set(), False
    if len(batches) != 1:
        return set(), True
    batch = batches[0]
    metadata = _import_metadata(batch)
    source_sha = str(metadata.get('source_sha256') or '')
    manifest_sha = str(metadata.get('manifest_sha256') or '')
    valid = (
        batch.status == 'applied'
        and batch.applied_at is not None
        and metadata.get('kind') == LEGACY_BASELINE_KIND
        and metadata.get('version') == LEGACY_BASELINE_SCHEMA_VERSION
        and re.fullmatch(r'[0-9a-f]{64}', source_sha)
        and re.fullmatch(r'[0-9a-f]{64}', manifest_sha)
        and batch.file_sha256 == source_sha
        and batch.analysis_context_sha256 == manifest_sha
        and batch.plan_hash == manifest_sha
        and batch.rejected_rows == 0
        and batch.excluded_rows == 0
        and ImportRow.query.filter_by(batch_id=batch.id).count() == 0
        and ImportFile.query.filter_by(batch_id=batch.id).count() == 0
    )
    changes = ImportChange.query.filter_by(
        batch_id=batch.id,
        target_table='notification_subscriber',
    ).order_by(ImportChange.target_id).all()
    all_change_count = ImportChange.query.filter_by(batch_id=batch.id).count()
    entries = []
    target_ids = set()
    operation_counts = {
        'adopt_legacy_ownership': 0,
        'preserve_manual': 0,
    }
    for change in changes:
        try:
            recorded = json.loads(change.after_json or '{}')
        except (TypeError, ValueError, json.JSONDecodeError):
            valid = False
            continue
        if (change.operation not in operation_counts
                or change.target_id is None
                or change.target_id in target_ids
                or not isinstance(recorded, dict)
                or recorded.get('subscriber_id') != change.target_id
                or recorded.get('baseline_version')
                != LEGACY_BASELINE_SCHEMA_VERSION
                or recorded.get('created_at') is None
                or not re.fullmatch(
                    r'[0-9a-f]{64}', str(recorded.get('state_hash') or ''))):
            valid = False
            continue
        target_ids.add(change.target_id)
        operation_counts[change.operation] += 1
        entries.append({
            'subscriber_id': change.target_id,
            'created_at': recorded['created_at'],
            'state_hash': recorded['state_hash'],
            'operation': change.operation,
        })
    entries.sort(key=lambda item: item['subscriber_id'])
    recalculated_manifest = hashlib.sha256(json.dumps({
        'version': LEGACY_BASELINE_SCHEMA_VERSION,
        'source_sha256': source_sha,
        'entries': entries,
    }, ensure_ascii=False, sort_keys=True,
        separators=(',', ':')).encode('utf-8')).hexdigest()
    expected_total = operation_counts['adopt_legacy_ownership'] + operation_counts[
        'preserve_manual']
    valid = bool(
        valid
        and len(entries) == len(changes) == all_change_count == expected_total
        and batch.total_rows == batch.selected_rows == expected_total
        and metadata.get('candidate_count')
        == operation_counts['adopt_legacy_ownership']
        and metadata.get('preserved_count')
        == operation_counts['preserve_manual']
        and isinstance(metadata.get('contact_count'), int)
        and metadata.get('contact_count') >= 0
        and isinstance(metadata.get('group_count'), int)
        and metadata.get('group_count') >= 0
        and recalculated_manifest == manifest_sha
    )
    return ({batch.id} if valid else set()), not valid


def _legacy_provenance_for_subscribers(subscribers):
    """Classify current row incarnations from immutable audit provenance.

    A target ID is not enough because SQLite may reuse deleted integer primary
    keys.  Baseline records carry the exact creation timestamp; older Legacy
    CSV ``create`` records retain the conservative created-before-change
    compatibility boundary used by the original cutover guard.
    """
    by_id = {subscriber.id: subscriber for subscriber in subscribers}
    result = {subscriber_id: set() for subscriber_id in by_id}
    incarnation_excluded = set()
    if not by_id:
        _, invalid_baseline = _legacy_baseline_integrity()
        return result, incarnation_excluded, invalid_baseline

    valid_baseline_ids, invalid_baseline = _legacy_baseline_integrity()
    records = db.session.query(ImportChange, ImportBatch).join(
        ImportBatch, ImportBatch.id == ImportChange.batch_id,
    ).filter(
        ImportBatch.status == 'applied',
        ImportChange.target_table == 'notification_subscriber',
        ImportChange.target_id.in_(set(by_id)),
        or_(
            and_(
                ImportBatch.source_type == 'legacy_csv',
                ImportChange.operation == 'create',
            ),
            and_(
                ImportBatch.source_type == 'legacy_csv',
                ImportBatch.schema_version == LEGACY_BASELINE_SCHEMA_VERSION,
                ImportChange.operation.in_([
                    'adopt_legacy_ownership', 'preserve_manual',
                ]),
            ),
            and_(
                ImportBatch.source_type == MANUAL_PROVENANCE_SOURCE_TYPE,
                ImportBatch.schema_version == MANUAL_PROVENANCE_SCHEMA_VERSION,
                ImportChange.operation == 'preserve_manual',
            ),
        ),
    ).order_by(ImportChange.id).all()
    for change, batch in records:
        subscriber = by_id.get(change.target_id)
        if not subscriber:
            continue
        try:
            recorded = json.loads(change.after_json or '{}')
        except (TypeError, ValueError, json.JSONDecodeError):
            recorded = {}
        if not isinstance(recorded, dict):
            recorded = {}
        if batch.schema_version == LEGACY_BASELINE_SCHEMA_VERSION:
            if batch.id not in valid_baseline_ids:
                incarnation_excluded.add(subscriber.id)
                continue
            strict_record = (
                change.operation in {
                    'adopt_legacy_ownership', 'preserve_manual'}
                and recorded.get('subscriber_id') == subscriber.id
                and recorded.get('baseline_version')
                == LEGACY_BASELINE_SCHEMA_VERSION
                and recorded.get('created_at') is not None
                and re.fullmatch(
                    r'[0-9a-f]{64}', str(recorded.get('state_hash') or ''))
            )
            if not strict_record:
                incarnation_excluded.add(subscriber.id)
                continue
        elif batch.source_type == MANUAL_PROVENANCE_SOURCE_TYPE:
            if not _manual_provenance_record_valid(
                    batch, change, subscriber):
                incarnation_excluded.add(subscriber.id)
                continue
        if not _row_incarnation_matches(
                subscriber.created_at, recorded, change.created_at):
            incarnation_excluded.add(subscriber.id)
            continue
        disposition = (
            'manual' if change.operation == 'preserve_manual' else 'legacy'
        )
        result[subscriber.id].add(disposition)
    return result, incarnation_excluded, invalid_baseline


def _active_non_powerschool_roster_state():
    """Return PII-free ownership state for every active non-PowerSchool row."""
    powerschool_ids = {
        local_id for (local_id,) in db.session.query(
            ExternalIdentity.local_id,
        ).join(
            NotificationSubscriber,
            NotificationSubscriber.id == ExternalIdentity.local_id,
        ).filter(
            ExternalIdentity.source_type == 'powerschool',
            ExternalIdentity.entity_type == 'student',
            ExternalIdentity.local_table == 'notification_subscriber',
            NotificationSubscriber.active.is_(True),
            NotificationSubscriber.created_at <= ExternalIdentity.created_at,
        ).distinct().all()
    }
    subscribers = NotificationSubscriber.query.options(
        selectinload(NotificationSubscriber.contacts),
        selectinload(NotificationSubscriber.group),
    ).filter(
        NotificationSubscriber.active.is_(True),
        ~NotificationSubscriber.id.in_(powerschool_ids),
    ).order_by(NotificationSubscriber.id).all()
    provenance, incarnation_excluded, invalid_baseline = (
        _legacy_provenance_for_subscribers(
            subscribers)
    )
    legacy = []
    manual = []
    unmanaged = []
    conflicts = []
    for subscriber in subscribers:
        dispositions = provenance.get(subscriber.id) or set()
        if invalid_baseline:
            conflicts.append(subscriber)
        elif subscriber.id in incarnation_excluded and not dispositions:
            unmanaged.append(subscriber)
        elif dispositions == {'legacy'}:
            legacy.append(subscriber)
        elif dispositions == {'manual'}:
            manual.append(subscriber)
        elif not dispositions:
            unmanaged.append(subscriber)
        else:
            conflicts.append(subscriber)
    return {
        'subscribers': subscribers,
        'legacy': legacy,
        'manual': manual,
        'unmanaged': unmanaged,
        'conflicts': conflicts,
        'incarnation_excluded_count': len(incarnation_excluded),
        'powerschool_active_count': len(powerschool_ids),
    }


def _legacy_baseline_source_plan(source_path):
    """Reconstruct exactly what the pre-audit Legacy CSV importer created.

    This intentionally does not call the modern normalizers.  The August 5
    importer stripped text fields, lower-cased only ``role``, resolved groups
    case-insensitively, and kept contact order.  Any row that importer would
    have skipped makes the baseline unverifiable and therefore unusable.
    """
    from collections import OrderedDict

    try:
        with open(source_path, 'rb') as handle:
            payload = handle.read(app.config['MAX_CONTENT_LENGTH'] + 1)
    except OSError as exc:
        raise ValueError('The Legacy source CSV could not be read.') from exc
    if len(payload) > app.config['MAX_CONTENT_LENGTH']:
        raise ValueError('The Legacy source CSV exceeds the configured size limit.')
    try:
        content = payload.decode('utf-8-sig')
    except UnicodeDecodeError as exc:
        raise ValueError('The Legacy source CSV must be UTF-8 encoded.') from exc

    reader = csv.DictReader(io.StringIO(content))
    headers = reader.fieldnames or []
    expected_headers = [
        'subscriber_id', 'household_label', 'group', 'active', 'role',
        'first_name', 'last_name', 'email', 'phone',
    ]
    if headers != expected_headers:
        raise ValueError('The Legacy source CSV headers do not match the original contract.')

    groups = {}
    duplicate_group_keys = set()
    for group in SubscriberGroup.query.order_by(SubscriberGroup.id).all():
        key = group.name.strip().lower()
        if key in groups:
            duplicate_group_keys.add(key)
        groups[key] = group
    if duplicate_group_keys:
        raise ValueError('Subscriber groups are ambiguous under Legacy case-insensitive matching.')

    households = OrderedDict()
    row_count = 0
    for row_number, row in enumerate(reader, 2):
        row_count += 1
        if row_count > app.config['IMPORT_MAX_ROWS']:
            raise ValueError('The Legacy source CSV exceeds the configured row limit.')
        group_name = (row.get('group') or '').strip()
        household = (row.get('household_label') or '').strip()
        first_name = (row.get('first_name') or '').strip()
        last_name = (row.get('last_name') or '').strip()
        email = (row.get('email') or '').strip()
        phone = (row.get('phone') or '').strip()
        role_raw = (row.get('role') or 'parent').strip().lower()
        role = role_raw if role_raw in {'parent', 'student'} else 'parent'
        if not first_name and not email:
            raise ValueError(
                'The Legacy source contains a row the historical importer would skip.')
        group = groups.get(group_name.lower()) if group_name else None
        if group_name and not group:
            raise ValueError(
                'The Legacy source references a group that cannot be resolved exactly.')
        group_id = group.id if group else None
        key = (group_id, household if household else f'__row_{row_number}__')
        households.setdefault(key, {
            'group_id': group_id,
            'notes': household or None,
            'contacts': [],
        })['contacts'].append({
            'first_name': first_name or None,
            'last_name': last_name or None,
            'email': email or None,
            'phone': phone or None,
            'role': role,
        })
    if not households:
        raise ValueError('The Legacy source CSV contains no importable households.')
    return {
        'source_sha256': hashlib.sha256(payload).hexdigest(),
        'households': list(households.values()),
        'candidate_count': len(households),
        'contact_count': sum(
            len(household['contacts']) for household in households.values()),
        'group_count': len({
            household['group_id'] for household in households.values()
            if household['group_id'] is not None
        }),
    }


def _legacy_baseline_household_signature(group_id, notes, contacts):
    material = {
        'group_id': group_id,
        'notes': notes,
        'contacts': [{
            'sort_order': index,
            'first_name': contact.get('first_name') or None,
            'last_name': contact.get('last_name') or None,
            'email': contact.get('email') or None,
            'phone': contact.get('phone') or None,
            'role': contact.get('role') or 'parent',
        } for index, contact in enumerate(contacts)],
    }
    return json.dumps(
        material, ensure_ascii=False, sort_keys=True, separators=(',', ':'))


def _legacy_baseline_subscriber_signature(subscriber):
    contacts = sorted(
        subscriber.contacts,
        key=lambda contact: (contact.sort_order, contact.id),
    )
    if any(contact.sort_order != index
           for index, contact in enumerate(contacts)):
        return None
    return _legacy_baseline_household_signature(
        subscriber.group_id,
        subscriber.notes,
        [{
            'first_name': contact.first_name,
            'last_name': contact.last_name,
            'email': contact.email,
            'phone': contact.phone,
            'role': contact.role,
        } for contact in contacts],
    )


def _legacy_baseline_state_hash(snapshot):
    """Key the audit fingerprint so stored hashes do not expose PII oracles."""
    canonical = json.dumps(
        snapshot, ensure_ascii=False, sort_keys=True,
        separators=(',', ':'),
    ).encode('utf-8')
    return hmac.new(
        app.config['SECRET_KEY'].encode('utf-8'),
        b'legacy-roster-baseline-state-v1\x00' + canonical,
        hashlib.sha256,
    ).hexdigest()


def _manual_provenance_manifest(entry):
    material = {
        'version': MANUAL_PROVENANCE_SCHEMA_VERSION,
        'entry': entry,
    }
    return hashlib.sha256(json.dumps(
        material, ensure_ascii=False, sort_keys=True,
        separators=(',', ':'),
    ).encode('utf-8')).hexdigest()


def _manual_provenance_record_valid(batch, change, subscriber):
    metadata = _import_metadata(batch)
    try:
        recorded = json.loads(change.after_json or '{}')
    except (TypeError, ValueError, json.JSONDecodeError):
        return False
    if not isinstance(recorded, dict):
        return False
    entry = {
        'subscriber_id': recorded.get('subscriber_id'),
        'created_at': recorded.get('created_at'),
        'state_hash': recorded.get('state_hash'),
        'operation': 'preserve_manual',
    }
    manifest_sha = _manual_provenance_manifest(entry)
    return bool(
        batch.status == 'applied'
        and batch.applied_at is not None
        and metadata.get('kind') == MANUAL_PROVENANCE_KIND
        and metadata.get('version') == MANUAL_PROVENANCE_SCHEMA_VERSION
        and metadata.get('manifest_sha256') == manifest_sha
        and batch.file_sha256 == manifest_sha
        and batch.analysis_context_sha256 == manifest_sha
        and batch.plan_hash == manifest_sha
        and batch.total_rows == batch.selected_rows == 1
        and batch.rejected_rows == batch.excluded_rows == 0
        and ImportChange.query.filter_by(batch_id=batch.id).count() == 1
        and ImportRow.query.filter_by(batch_id=batch.id).count() == 0
        and ImportFile.query.filter_by(batch_id=batch.id).count() == 0
        and change.operation == 'preserve_manual'
        and change.target_table == 'notification_subscriber'
        and change.target_id == subscriber.id
        and entry['subscriber_id'] == subscriber.id
        and recorded.get('provenance_version')
        == MANUAL_PROVENANCE_SCHEMA_VERSION
        and entry['created_at'] == _snapshot_datetime(subscriber.created_at)
        and re.fullmatch(r'[0-9a-f]{64}', str(entry['state_hash'] or ''))
    )


def _record_manual_subscriber_provenance(subscriber, operator):
    """Create PII-free ownership evidence in the subscriber transaction."""
    if not subscriber.id or subscriber.created_at is None:
        raise ValueError('A manual subscriber needs a stable row before audit.')
    contact_identities = _contact_identity_map([subscriber])
    entry = {
        'subscriber_id': subscriber.id,
        'created_at': _snapshot_datetime(subscriber.created_at),
        'state_hash': _legacy_baseline_state_hash(
            _subscriber_snapshot(subscriber, contact_identities)),
        'operation': 'preserve_manual',
    }
    manifest_sha = _manual_provenance_manifest(entry)
    now = max(_utcnow(), subscriber.created_at)
    metadata = {
        'kind': MANUAL_PROVENANCE_KIND,
        'version': MANUAL_PROVENANCE_SCHEMA_VERSION,
        'manifest_sha256': manifest_sha,
        'approved_by_id': operator.id,
        'applied_at': now.isoformat() + 'Z',
    }
    batch = ImportBatch(
        public_id=secrets.token_urlsafe(32),
        source_type=MANUAL_PROVENANCE_SOURCE_TYPE,
        schema_version=MANUAL_PROVENANCE_SCHEMA_VERSION,
        status='applied', snapshot_type='delta', school_year=None,
        uploaded_by_id=operator.id,
        file_sha256=manifest_sha,
        analysis_context_sha256=manifest_sha,
        plan_hash=manifest_sha,
        total_rows=1, selected_rows=1, rejected_rows=0, excluded_rows=0,
        metadata_json=json.dumps(metadata, sort_keys=True),
        created_at=now, applied_at=now,
        expires_at=now + timedelta(
            hours=app.config['IMPORT_STAGE_TTL_HOURS']),
    )
    db.session.add(batch)
    db.session.flush()
    db.session.add(ImportChange(
        batch_id=batch.id,
        operation='preserve_manual',
        target_table='notification_subscriber',
        target_id=subscriber.id,
        after_json=json.dumps({
            'subscriber_id': entry['subscriber_id'],
            'created_at': entry['created_at'],
            'state_hash': entry['state_hash'],
            'provenance_version': MANUAL_PROVENANCE_SCHEMA_VERSION,
        }, sort_keys=True),
        created_at=now,
    ))
    return batch


def _legacy_baseline_existing_batches():
    return ImportBatch.query.filter_by(
        source_type='legacy_csv',
        schema_version=LEGACY_BASELINE_SCHEMA_VERSION,
    ).order_by(ImportBatch.id).all()


def _legacy_baseline_build_manifest(source_path, existing_batch=None):
    source = _legacy_baseline_source_plan(source_path)
    state = _active_non_powerschool_roster_state()
    if state['powerschool_active_count']:
        raise ValueError(
            'A PowerSchool roster is already active; baseline adoption is no longer safe.')
    applied_powerschool = ImportBatch.query.filter(
        ImportBatch.source_type == 'powerschool',
        ImportBatch.status.in_([
            'applied', 'rollback_failed', 'retention_closed',
        ]),
    ).first()
    if applied_powerschool:
        raise ValueError(
            'A PowerSchool batch has already been applied; baseline adoption is no longer safe.')
    if state['conflicts'] or state['incarnation_excluded_count']:
        raise ValueError('Existing roster provenance is conflicting or unverifiable.')

    subscribers = state['subscribers']
    if existing_batch is None:
        if state['legacy'] or state['manual']:
            raise ValueError(
                'The active roster already has partial explicit provenance; '
                'baseline adoption is blocked.')
    else:
        metadata = _import_metadata(existing_batch)
        if (metadata.get('kind') != LEGACY_BASELINE_KIND
                or metadata.get('version') != LEGACY_BASELINE_SCHEMA_VERSION
                or existing_batch.file_sha256 != source['source_sha256']
                or metadata.get('source_sha256') != source['source_sha256']
                or existing_batch.analysis_context_sha256
                != metadata.get('manifest_sha256')
                or existing_batch.plan_hash != metadata.get('manifest_sha256')):
            raise ValueError('The existing baseline metadata is incomplete or inconsistent.')
        baseline_changes = ImportChange.query.filter_by(
            batch_id=existing_batch.id,
            target_table='notification_subscriber',
        ).order_by(ImportChange.id).all()
        recorded = {}
        for change in baseline_changes:
            if (change.operation not in {
                    'adopt_legacy_ownership', 'preserve_manual'}
                    or change.target_id is None
                    or change.target_id in recorded):
                raise ValueError('The existing baseline audit is incomplete or duplicated.')
            try:
                change_record = json.loads(change.after_json or '{}')
            except (TypeError, ValueError, json.JSONDecodeError) as exc:
                raise ValueError(
                    'The existing baseline audit record is malformed.') from exc
            if (not isinstance(change_record, dict)
                    or change_record.get('subscriber_id') != change.target_id
                    or change_record.get('baseline_version')
                    != LEGACY_BASELINE_SCHEMA_VERSION
                    or change_record.get('created_at') is None
                    or not re.fullmatch(
                        r'[0-9a-f]{64}',
                        str(change_record.get('state_hash') or ''))):
                raise ValueError('The existing baseline audit record is incomplete.')
            recorded[change.target_id] = change.operation
        if set(recorded) != {subscriber.id for subscriber in subscribers}:
            raise ValueError('The active roster no longer matches the applied baseline.')

    source_signatures = []
    for household in source['households']:
        source_signatures.append(_legacy_baseline_household_signature(
            household['group_id'], household['notes'], household['contacts']))
    if len(source_signatures) != len(set(source_signatures)):
        raise ValueError(
            'The Legacy source contains duplicate household signatures and '
            'cannot be matched one-to-one.')

    database_matches = {}
    for subscriber in subscribers:
        signature = _legacy_baseline_subscriber_signature(subscriber)
        if signature is None:
            raise ValueError('A roster contact order is not historically reproducible.')
        database_matches.setdefault(signature, []).append(subscriber)
    matched = []
    matched_ids = set()
    for signature in source_signatures:
        candidates = database_matches.get(signature, [])
        if len(candidates) != 1 or candidates[0].id in matched_ids:
            raise ValueError(
                'The Legacy source does not have an exact one-to-one active roster match.')
        matched.append(candidates[0])
        matched_ids.add(candidates[0].id)
    preserved = [
        subscriber for subscriber in subscribers
        if subscriber.id not in matched_ids
    ]
    if existing_batch is not None:
        actual_operations = {
            subscriber.id: 'adopt_legacy_ownership' for subscriber in matched
        }
        actual_operations.update({
            subscriber.id: 'preserve_manual' for subscriber in preserved
        })
        recorded_operations = {
            change.target_id: change.operation
            for change in ImportChange.query.filter_by(
                batch_id=existing_batch.id,
                target_table='notification_subscriber',
            ).all()
        }
        if actual_operations != recorded_operations:
            raise ValueError('The source disposition no longer matches the applied baseline.')

    contact_identities = _contact_identity_map(subscribers)
    entries = []
    for operation, collection in (
            ('adopt_legacy_ownership', matched),
            ('preserve_manual', preserved)):
        for subscriber in collection:
            if subscriber.created_at is None:
                raise ValueError('A roster row has no verifiable creation timestamp.')
            snapshot = _subscriber_snapshot(subscriber, contact_identities)
            entries.append({
                'subscriber_id': subscriber.id,
                'created_at': _snapshot_datetime(subscriber.created_at),
                'state_hash': _legacy_baseline_state_hash(snapshot),
                'operation': operation,
            })
    entries.sort(key=lambda item: item['subscriber_id'])
    manifest = {
        'version': LEGACY_BASELINE_SCHEMA_VERSION,
        'source_sha256': source['source_sha256'],
        'entries': entries,
    }
    manifest_sha256 = hashlib.sha256(json.dumps(
        manifest, ensure_ascii=False, sort_keys=True,
        separators=(',', ':'),
    ).encode('utf-8')).hexdigest()
    if existing_batch is not None:
        metadata = _import_metadata(existing_batch)
        expected_counts = {
            'candidate_count': len(matched),
            'contact_count': source['contact_count'],
            'group_count': source['group_count'],
            'preserved_count': len(preserved),
        }
        if (manifest_sha256 != metadata.get('manifest_sha256')
                or existing_batch.total_rows != len(entries)
                or existing_batch.selected_rows != len(entries)
                or existing_batch.rejected_rows != 0
                or existing_batch.excluded_rows != 0
                or any(metadata.get(key) != value
                       for key, value in expected_counts.items())):
            raise ValueError('The applied baseline no longer matches its manifest.')
        recorded_by_id = {}
        for change in ImportChange.query.filter_by(
                batch_id=existing_batch.id,
                target_table='notification_subscriber').all():
            recorded_by_id[change.target_id] = json.loads(change.after_json)
        for entry in entries:
            recorded_entry = recorded_by_id.get(entry['subscriber_id']) or {}
            if (recorded_entry.get('created_at') != entry['created_at']
                    or recorded_entry.get('state_hash') != entry['state_hash']):
                raise ValueError(
                    'The active roster state changed after baseline adoption.')
    return {
        'source_sha256': source['source_sha256'],
        'manifest_sha256': manifest_sha256,
        'candidate_count': len(matched),
        'contact_count': source['contact_count'],
        'group_count': source['group_count'],
        'preserved_count': len(preserved),
        'entries': entries,
    }


def _legacy_baseline_public_summary(plan, mode, already_applied=False):
    return {
        'ok': True,
        'mode': mode,
        'already_applied': bool(already_applied),
        'candidate_count': plan['candidate_count'],
        'contact_count': plan['contact_count'],
        'group_count': plan['group_count'],
        'preserved_count': plan['preserved_count'],
        'source_sha256': plan['source_sha256'],
        'manifest_sha256': plan['manifest_sha256'],
    }


@app.cli.command('adopt-legacy-baseline')
@click.argument('source_csv', type=click.Path(
    exists=True, dir_okay=False, readable=True, path_type=str))
@click.option('--apply', 'apply_changes', is_flag=True,
              help='Persist the exactly approved baseline manifest.')
@click.option('--source-sha', default=None,
              help='Expected SHA-256 of the original Legacy combined CSV.')
@click.option('--manifest-sha', default=None,
              help='Exact manifest SHA-256 printed by the dry-run.')
@click.option('--expected-candidates', type=click.IntRange(min=0), default=None)
@click.option('--expected-contacts', type=click.IntRange(min=0), default=None)
@click.option('--expected-groups', type=click.IntRange(min=0), default=None)
@click.option('--expected-preserved', type=click.IntRange(min=0), default=None)
@click.option('--approved-by', default=None,
              help='Username of the active administrator approving adoption.')
def adopt_legacy_baseline_command(
        source_csv, apply_changes, source_sha, manifest_sha,
        expected_candidates, expected_contacts, expected_groups,
        expected_preserved, approved_by):
    """Dry-run or adopt provenance for one historical Legacy CSV roster."""
    with _roster_import_lock() as acquired:
        if not acquired:
            raise click.ClickException(
                'Another roster mutation is in progress; try again later.')
        try:
            # Lock all relevant rows before deriving or applying a manifest.
            active = NotificationSubscriber.query.options(
                selectinload(NotificationSubscriber.contacts),
                selectinload(NotificationSubscriber.group),
            ).filter(NotificationSubscriber.active.is_(True)).with_for_update().all()
            active_ids = {subscriber.id for subscriber in active}
            if active_ids:
                SubscriberContact.query.filter(
                    SubscriberContact.subscriber_id.in_(active_ids),
                ).with_for_update().all()
            ExternalIdentity.query.with_for_update().all()
            ImportBatch.query.filter(
                or_(
                    ImportBatch.source_type == 'powerschool',
                    and_(
                        ImportBatch.source_type == 'legacy_csv',
                        ImportBatch.schema_version == LEGACY_BASELINE_SCHEMA_VERSION,
                    ),
                ),
            ).with_for_update().all()
            existing_batches = _legacy_baseline_existing_batches()
            if len(existing_batches) > 1:
                raise ValueError('Multiple provenance baselines exist; manual review is required.')
            existing = existing_batches[0] if existing_batches else None
            if existing and existing.status != 'applied':
                raise ValueError('A partial provenance baseline exists; manual review is required.')
            plan = _legacy_baseline_build_manifest(source_csv, existing)

            if not apply_changes:
                db.session.rollback()
                click.echo(json.dumps(_legacy_baseline_public_summary(
                    plan, 'dry-run', already_applied=bool(existing)), sort_keys=True))
                return

            required_options = {
                '--source-sha': source_sha,
                '--manifest-sha': manifest_sha,
                '--expected-candidates': expected_candidates,
                '--expected-contacts': expected_contacts,
                '--expected-groups': expected_groups,
                '--expected-preserved': expected_preserved,
                '--approved-by': approved_by,
            }
            missing = [name for name, value in required_options.items()
                       if value is None or value == '']
            if missing:
                raise ValueError(
                    'Apply requires every dry-run checksum, aggregate, and approver option.')
            if not re.fullmatch(r'[0-9a-fA-F]{64}', str(source_sha)):
                raise ValueError('The expected source SHA-256 is invalid.')
            if not re.fullmatch(r'[0-9a-fA-F]{64}', str(manifest_sha)):
                raise ValueError('The expected manifest SHA-256 is invalid.')
            expected = {
                'candidate_count': expected_candidates,
                'contact_count': expected_contacts,
                'group_count': expected_groups,
                'preserved_count': expected_preserved,
            }
            if not hmac.compare_digest(
                    plan['source_sha256'], str(source_sha).lower()):
                raise ValueError('The Legacy source checksum does not match approval.')
            if not hmac.compare_digest(
                    plan['manifest_sha256'], str(manifest_sha).lower()):
                raise ValueError('The current roster manifest does not match approval.')
            if any(plan[key] != value for key, value in expected.items()):
                raise ValueError('One or more approved aggregate counts changed.')
            approver = User.query.filter_by(
                username=str(approved_by), active=True,
            ).with_for_update().first()
            approver_group = None
            if approver and approver.group_id:
                approver_group = UserGroup.query.filter_by(
                    id=approver.group_id, is_admin=True,
                ).with_for_update().first()
            if not approver or not approver_group:
                raise ValueError('The approving username is not an active administrator.')
            if existing:
                db.session.rollback()
                click.echo(json.dumps(_legacy_baseline_public_summary(
                    plan, 'apply', already_applied=True), sort_keys=True))
                return

            now = _utcnow()
            metadata = {
                'kind': LEGACY_BASELINE_KIND,
                'version': LEGACY_BASELINE_SCHEMA_VERSION,
                'source_sha256': plan['source_sha256'],
                'manifest_sha256': plan['manifest_sha256'],
                'candidate_count': plan['candidate_count'],
                'contact_count': plan['contact_count'],
                'group_count': plan['group_count'],
                'preserved_count': plan['preserved_count'],
                'approved_by_id': approver.id,
                'applied_at': now.isoformat() + 'Z',
            }
            batch = ImportBatch(
                public_id=secrets.token_urlsafe(32),
                source_type='legacy_csv',
                schema_version=LEGACY_BASELINE_SCHEMA_VERSION,
                status='applied', snapshot_type='delta', school_year=None,
                uploaded_by_id=approver.id,
                file_sha256=plan['source_sha256'],
                analysis_context_sha256=plan['manifest_sha256'],
                plan_hash=plan['manifest_sha256'],
                total_rows=len(plan['entries']),
                selected_rows=len(plan['entries']),
                rejected_rows=0, excluded_rows=0,
                metadata_json=json.dumps(metadata, sort_keys=True),
                created_at=now, applied_at=now,
                expires_at=now + timedelta(
                    hours=app.config['IMPORT_STAGE_TTL_HOURS']),
            )
            db.session.add(batch)
            db.session.flush()
            for entry in plan['entries']:
                db.session.add(ImportChange(
                    batch_id=batch.id,
                    operation=entry['operation'],
                    target_table='notification_subscriber',
                    target_id=entry['subscriber_id'],
                    after_json=json.dumps({
                        'subscriber_id': entry['subscriber_id'],
                        'created_at': entry['created_at'],
                        'state_hash': entry['state_hash'],
                        'baseline_version': LEGACY_BASELINE_SCHEMA_VERSION,
                    }, sort_keys=True),
                    created_at=now,
                ))
            db.session.add(AuditLog(
                user_id=approver.id,
                username=approver.username,
                action='legacy_baseline_adopted',
                module='notifications',
                target=batch.public_id,
                details=(
                    f'candidates={plan["candidate_count"]}; '
                    f'contacts={plan["contact_count"]}; '
                    f'groups={plan["group_count"]}; '
                    f'preserved={plan["preserved_count"]}; '
                    f'source_sha256={plan["source_sha256"]}; '
                    f'manifest_sha256={plan["manifest_sha256"]}'
                ),
                ip_address='cli',
                created_at=now,
            ))
            db.session.commit()
            click.echo(json.dumps(_legacy_baseline_public_summary(
                plan, 'apply'), sort_keys=True))
        except ValueError as exc:
            db.session.rollback()
            raise click.ClickException(str(exc)) from exc
        except Exception:
            db.session.rollback()
            raise


def _legacy_cutover_inventory():
    """Return provenance-bound legacy candidates and reincarnation exclusions.

    ImportChange is the authoritative provenance boundary.  Names, emails,
    phones and household labels are deliberately not used to infer identity,
    and manually-created subscribers are therefore never swept into a
    PowerSchool cutover.  The creation timestamps also bind the audit target
    to the same row incarnation because SQLite may reuse a deleted integer PK.
    """
    state = _active_non_powerschool_roster_state()
    return state['legacy'], (
        state['incarnation_excluded_count'] + len(state['conflicts'])
    )


def _active_legacy_cutover_subscribers():
    subscribers, _ = _legacy_cutover_inventory()
    return subscribers


def _currently_applied_powerschool_batch_exists():
    return ImportBatch.query.filter(
        ImportBatch.source_type == 'powerschool',
        ImportBatch.status.in_([
            'applied', 'rollback_failed', 'retention_closed',
        ]),
    ).first() is not None


def _unmanaged_roster_block_message(baseline_available):
    if baseline_available:
        return (
            'Active subscribers have no explicit Legacy or manual '
            'provenance. Reconcile the exact original Legacy CSV with the '
            'one-time baseline command, then analyze all three PowerSchool '
            'files again.')
    return (
        'Active subscribers have no explicit roster provenance while '
        'PowerSchool ownership or an applied PowerSchool roster is already '
        'present. Automatic baseline adoption is unavailable; keep Apply '
        'blocked and perform manual provenance review before creating a new '
        'analysis.')


def _legacy_cutover_rows(batch):
    result = []
    for row in ImportRow.query.filter_by(
            batch_id=batch.id, classification='deactivate_candidate').all():
        try:
            data = json.loads(row.normalized_json)
        except (TypeError, ValueError, json.JSONDecodeError):
            continue
        if data.get('cutover_source') == 'legacy_csv':
            result.append((row, data))
    return result


def _legacy_cutover_payload(batch, metadata):
    stored = dict(metadata.get('legacy_cutover') or {})
    defaults = {
        'required': False,
        'candidate_count': 0,
        'incarnation_excluded_count': 0,
        'baseline_required': False,
        'baseline_available': False,
        'unmanaged_count': 0,
        'approved': False,
        'blocked': False,
        'requires_reanalysis': False,
        'message': '',
    }
    result = {**defaults, **stored}
    if batch.status != 'staged':
        return result

    roster_state = _active_non_powerschool_roster_state()
    unmanaged_count = (
        len(roster_state['unmanaged']) + len(roster_state['conflicts'])
    )
    if unmanaged_count:
        baseline_available = bool(
            not roster_state['powerschool_active_count']
            and not _currently_applied_powerschool_batch_exists())
        return {
            **defaults,
            'required': True,
            'baseline_required': True,
            'baseline_available': baseline_available,
            'unmanaged_count': unmanaged_count,
            'incarnation_excluded_count': roster_state[
                'incarnation_excluded_count'],
            'blocked': True,
            'requires_reanalysis': True,
            'message': _unmanaged_roster_block_message(baseline_available),
        }

    live_ids = {
        subscriber.id for subscriber in _active_legacy_cutover_subscribers()
    }
    counts = metadata.get('counts') or {}
    has_replacement_roster = any(
        int(counts.get(classification) or 0) > 0
        for classification in ('new', 'update', 'unchanged', 'conflict')
    )
    if not result['required'] and live_ids and has_replacement_roster:
        return {
            **defaults,
            'required': True,
            'candidate_count': len(live_ids),
            'blocked': True,
            'requires_reanalysis': True,
            'message': (
                'This staged batch predates the Legacy CSV cutover guard. '
                'Analyze the three files again before approval.'),
        }

    if result['required'] and not result['blocked']:
        staged_ids = {
            int(data.get('target_subscriber_id'))
            for _, data in _legacy_cutover_rows(batch)
            if data.get('target_subscriber_id') is not None
        }
        if staged_ids != live_ids:
            result.update({
                'approved': False,
                'blocked': True,
                'requires_reanalysis': True,
                'candidate_count': len(live_ids),
                'message': (
                    'The active Legacy CSV roster changed after analysis. '
                    'Analyze the three files again.'),
            })
    return result


def _legacy_cutover_apply_error(batch):
    roster_state = _active_non_powerschool_roster_state()
    unmanaged_count = (
        len(roster_state['unmanaged']) + len(roster_state['conflicts'])
    )
    if unmanaged_count:
        return _unmanaged_roster_block_message(bool(
            not roster_state['powerschool_active_count']
            and not _currently_applied_powerschool_batch_exists()))
    regular_rows = ImportRow.query.filter(
        ImportRow.batch_id == batch.id,
        ImportRow.classification.in_(['new', 'update']),
    ).all()
    regular_ids = {row.id for row in regular_rows}
    selected_regular_ids = {row.id for row in regular_rows if row.selected}
    live_ids = {
        subscriber.id for subscriber in _active_legacy_cutover_subscribers()
    }
    metadata = _import_metadata(batch)
    cutover = _legacy_cutover_payload(batch, metadata)
    selected_legacy_ids = {
        int(data['target_subscriber_id'])
        for row, data in _legacy_cutover_rows(batch)
        if row.selected and data.get('target_subscriber_id') is not None
    }
    if not cutover.get('required') and not selected_legacy_ids:
        return None
    if (not cutover.get('required') or not cutover.get('approved')
            or cutover.get('blocked') or cutover.get('requires_reanalysis')):
        return (cutover.get('message') or
                'Approve the complete Legacy CSV cutover before applying this batch.')
    if not hmac.compare_digest(
            str(cutover.get('approved_plan_hash') or ''),
            str(batch.plan_hash or '')):
        return ('The Legacy CSV cutover approval is not bound to the current '
                'plan. Save the complete selection and approve it again.')
    preflight = metadata.get('preflight') or {}
    if (batch.snapshot_type != 'full_district'
            or preflight.get('transportation_contract')
            != TRANSPORTATION_V2_CONTRACT):
        return ('Legacy CSV cutover requires a current Transportation v2 '
                'district-wide Full Snapshot. Analyze the three files again.')
    counts = metadata.get('counts') or {}
    replacement_count = sum(
        int(counts.get(classification) or 0)
        for classification in ('new', 'update', 'unchanged')
    )
    if replacement_count < 1 or selected_regular_ids != regular_ids:
        return ('Legacy CSV cutover requires every importable New and Update '
                'row to remain selected. Restore the complete selection and '
                'save it before applying.')
    staged_legacy_ids = {
        int(data['target_subscriber_id'])
        for _, data in _legacy_cutover_rows(batch)
        if data.get('target_subscriber_id') is not None
    }
    if (selected_legacy_ids != staged_legacy_ids
            or staged_legacy_ids != live_ids):
        return ('The active Legacy CSV roster no longer matches the approved '
                'cutover. Analyze the three files again.')
    return None


def _powerschool_row_records(batch):
    return [{
        'id': row.id, 'row_hash': row.row_hash,
        'classification': row.classification, 'selected': row.selected,
    } for row in ImportRow.query.filter_by(batch_id=batch.id).all()]


def _refresh_import_counts_and_hash(batch):
    rows = ImportRow.query.filter_by(batch_id=batch.id).all()
    batch.total_rows = len(rows)
    batch.selected_rows = sum(bool(row.selected) for row in rows)
    batch.rejected_rows = sum(row.classification == 'rejected' for row in rows)
    batch.excluded_rows = batch.total_rows - batch.selected_rows - batch.rejected_rows
    batch.plan_hash = canonical_plan_hash(
        batch.public_id, batch.schema_version, _powerschool_row_records(batch))


def _store_powerschool_file(batch, file_type, uploaded, payload, headers):
    safe_name = re.sub(r'[^A-Za-z0-9._ -]', '_',
                       os.path.basename(uploaded.filename))[:255]
    path = os.path.join(IMPORT_STAGE_DIR, f'{batch.public_id}-{file_type}.csv')
    _write_private_file(path, payload, binary=True)
    db.session.add(ImportFile(
        batch_id=batch.id, file_type=file_type, original_name=safe_name,
        sha256=hashlib.sha256(payload).hexdigest(), byte_size=len(payload),
        storage_path=path, headers_json=json.dumps(headers)))
    return path


def _powerschool_batch_payload(batch, include_rows=True):
    metadata = _import_metadata(batch)
    result = {
        'ok': True, 'batch_id': batch.public_id, 'status': batch.status,
        'schema_version': batch.schema_version, 'school_year': batch.school_year,
        'snapshot_type': batch.snapshot_type, 'plan_hash': batch.plan_hash,
        'total': batch.total_rows, 'selected': batch.selected_rows,
        'excluded': batch.excluded_rows, 'rejected': batch.rejected_rows,
        'counts': metadata.get('counts', {}), 'issues': metadata.get('issues', []),
        'metrics': metadata.get('metrics', {}),
        'preflight': metadata.get('preflight', {}),
        'normalizer_revision': metadata.get('normalizer_revision'),
        'reanalyzed_from': metadata.get('reanalyzed_from'),
        'created_at': batch.created_at.isoformat() + 'Z',
        'applied_at': batch.applied_at.isoformat() + 'Z' if batch.applied_at else None,
        'rolled_back_at': metadata.get('rolled_back_at'),
        'legacy_cutover': _legacy_cutover_payload(batch, metadata),
    }
    if include_rows:
        rows = ImportRow.query.filter_by(batch_id=batch.id).order_by(
            ImportRow.row_number).all()
        can_view_pii = current_user.has_capability('notifications.pii')
        result['rows'] = []
        for row in rows:
            data = json.loads(row.normalized_json)
            external_key = row.external_key
            if not can_view_pii:
                external_key = ('••••' + external_key[-4:]
                                if external_key else None)
                data = _masked_powerschool_data(data)
            result['rows'].append({
                'id': row.id, 'row_number': row.row_number,
                'external_key': external_key,
                'classification': row.classification, 'selected': bool(row.selected),
                'data': data, 'errors': json.loads(row.errors_json or '[]'),
            })
    return result


def _masked_powerschool_data(data):
    value = json.loads(json.dumps(data))
    for field in ('student_number', 'student_id', 'household_id', 'source_id'):
        if field in value:
            raw = str(value.get(field) or '').strip()
            value[field] = ('••••' + raw[-4:]) if raw else ''
    if 'student_numbers' in value:
        value['student_numbers'] = [
            ('••••' + str(item or '').strip()[-4:])
            if str(item or '').strip() else ''
            for item in value.get('student_numbers') or []
        ]
    if value.get('stop'):
        value['stop'] = '***'
    for field in ('first_name', 'last_name'):
        if field in value:
            value[field] = _mask_name(value.get(field))
    for contact in value.get('contacts', []):
        raw_contact_id = str(contact.get('contact_id') or '').strip()
        contact['contact_id'] = ('••••' + raw_contact_id[-4:]
                                 if raw_contact_id else '')
        contact['first_name'] = _mask_name(contact.get('first_name'))
        contact['last_name'] = _mask_name(contact.get('last_name'))
        contact['email'] = _mask_email(contact.get('email'))
        contact['phone'] = _mask_phone(contact.get('phone'))
    for change in value.get('changes', []):
        field = change.get('field', '')
        masker = (_mask_email if 'email' in field else
                  _mask_phone if 'phone' in field else
                  (lambda item: ('••••' + str(item or '')[-4:])
                   if item else '') if any(item in field for item in (
                       'student', 'identity', 'external_key')) else
                  _mask_name if any(item in field for item in (
                      'name', 'household_label')) else
                  (lambda item: '***' if item else '')
                  if 'stop' in field else None)
        if masker:
            change['current'] = masker(change.get('current'))
            change['proposed'] = masker(change.get('proposed'))
    return value


@app.route('/admin/notifications/powerschool')
@login_required
@require_capability('import.powerschool')
def powerschool_import_page():
    _powerschool_enabled()
    profiles = ImportMappingProfile.query.filter_by(
        source_type='powerschool', active=True).order_by(
        ImportMappingProfile.schema_version.desc()).all()
    recent = ImportBatch.query.filter_by(
        source_type='powerschool', uploaded_by_id=current_user.id).order_by(
        ImportBatch.created_at.desc()).limit(20).all()
    return render_template('admin/powerschool_import.html', profiles=profiles,
                           recent_batches=recent,
                           can_rollback=current_user.has_capability('import.rollback'))


@app.route('/admin/notifications/powerschool/preview', methods=['POST'])
@login_required
@require_capability('import.powerschool')
def powerschool_import_preview():
    _powerschool_enabled()
    _cleanup_import_stages()
    transportation = request.files.get('transportation_file')
    contacts = request.files.get('contacts_file')
    student_contacts = request.files.get('student_contacts_file')
    guardian_contacts = request.files.get('guardian_contacts_file')
    has_combined_contacts = bool(contacts and contacts.filename)
    has_student_contacts = bool(student_contacts and student_contacts.filename)
    has_guardian_contacts = bool(guardian_contacts and guardian_contacts.filename)
    has_any_split_contacts = has_student_contacts or has_guardian_contacts
    if not _valid_csv_upload(transportation):
        return jsonify({'ok': False, 'message':
                        'Select the approved UTF-8 Transportation CSV export.'}), 400
    if has_combined_contacts and has_any_split_contacts:
        return jsonify({'ok': False, 'message':
                        'Choose either the combined Contacts CSV or both PowerSchool contact exports, not both.'}), 400
    if has_combined_contacts:
        if not _valid_csv_upload(contacts):
            return jsonify({'ok': False, 'message':
                            'Select a valid UTF-8 combined Contacts CSV.'}), 400
    elif not (has_student_contacts and has_guardian_contacts):
        return jsonify({'ok': False, 'message':
                        'Select both the Student Contacts and Guardian Contacts UTF-8 CSV exports.'}), 400
    elif (not _valid_csv_upload(student_contacts)
          or not _valid_csv_upload(guardian_contacts)):
        return jsonify({'ok': False, 'message':
                        'Both PowerSchool contact exports must be valid UTF-8 CSV files.'}), 400
    school_year = _normalize_text(request.form.get('school_year'), 20)
    if not re.fullmatch(r'20\d{2}-\d{2}', school_year):
        return jsonify({'ok': False, 'message':
                        'School year must use YYYY-YY format.'}), 400
    snapshot_type = request.form.get('snapshot_type', 'delta')
    if snapshot_type not in {'delta', 'full_district'}:
        return jsonify({'ok': False, 'message': 'Invalid snapshot policy.'}), 400
    force_reanalyze = request.form.get('force_reanalyze') == '1'
    try:
        profile_id = int(request.form.get('mapping_profile_id', ''))
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'Select a mapping profile.'}), 400
    profile = ImportMappingProfile.query.filter_by(
        id=profile_id, source_type='powerschool', active=True).first()
    if not profile:
        return jsonify({'ok': False, 'message': 'Mapping profile is unavailable.'}), 400
    try:
        mapping = json.loads(profile.mapping_json)
        transport_payload = transportation.read()
        contact_uploads = {}
        if has_combined_contacts:
            contacts_payload = contacts.read()
            contact_sources = None
            contact_uploads['contacts'] = (contacts, contacts_payload)
        else:
            contacts_payload = None
            student_contacts_payload = student_contacts.read()
            guardian_contacts_payload = guardian_contacts.read()
            contact_sources = [
                {
                    'key': 'student_contacts',
                    'payload': student_contacts_payload,
                    'force_relationship': 'student',
                },
                {
                    'key': 'guardian_contacts',
                    'payload': guardian_contacts_payload,
                    'default_relationship': 'guardian',
                },
            ]
            contact_uploads.update({
                'student_contacts': (student_contacts, student_contacts_payload),
                'guardian_contacts': (guardian_contacts, guardian_contacts_payload),
            })
        parsed = build_normalized_plan(
            transport_payload, contacts_payload, mapping,
            app.config['IMPORT_MAX_ROWS'], app.config['IMPORT_MAX_COLUMNS'],
            contact_sources=contact_sources)
    except (ImportValidationError, json.JSONDecodeError) as exc:
        return jsonify({'ok': False, 'message': str(exc)}), 400

    preflight_failure = _powerschool_preflight_failure(parsed, snapshot_type)
    if preflight_failure:
        return jsonify({
            'ok': False,
            **preflight_failure,
            'preflight': parsed.get('preflight') or {},
            'metrics': parsed.get('metrics') or {},
        }), 400

    analysis_context_sha256 = _powerschool_analysis_context_hash(
        profile, mapping, parsed['combined_sha256'], school_year, snapshot_type)
    same_context = ImportBatch.query.filter(
        ImportBatch.source_type == 'powerschool',
        ImportBatch.uploaded_by_id == current_user.id,
        ImportBatch.analysis_context_sha256 == analysis_context_sha256,
    )
    busy_duplicate = same_context.filter(
        ImportBatch.status.in_(['selecting', 'applying'])
    ).order_by(ImportBatch.created_at.desc()).first()
    duplicate = busy_duplicate or same_context.filter(
        ImportBatch.status.in_(['staged', 'applied'])
    ).order_by(ImportBatch.created_at.desc()).first()
    if duplicate:
        reanalyze_allowed = (
            busy_duplicate is None
            and duplicate.status in {'staged', 'applied'})
        if not force_reanalyze or not reanalyze_allowed:
            return jsonify({
                'ok': False,
                'message': (
                    'A matching analysis is currently changing selection or '
                    'being applied. Wait for it to finish before re-analyzing.'
                    if busy_duplicate else
                    'These files were already analyzed with the same year, '
                    'policy, mapping and normalizer revision.'),
                'existing_batch_id': duplicate.public_id,
                'existing_status': duplicate.status,
                'can_open': True,
                'reanalyze_allowed': reanalyze_allowed,
            }), 409

    now = _utcnow()
    batch = ImportBatch(
        public_id=secrets.token_urlsafe(32), source_type='powerschool',
        schema_version=profile.schema_version, status='staged',
        snapshot_type=snapshot_type, school_year=school_year,
        uploaded_by_id=current_user.id,
        file_sha256=parsed['combined_sha256'], plan_hash='pending',
        analysis_context_sha256=analysis_context_sha256,
        created_at=now,
        expires_at=now + timedelta(hours=app.config['IMPORT_STAGE_TTL_HOURS']))
    db.session.add(batch)
    db.session.flush()
    paths = []
    try:
        paths.append(_store_powerschool_file(
            batch, 'transportation', transportation, transport_payload,
            parsed['files']['transportation']['headers']))
        for file_type, (uploaded, payload) in contact_uploads.items():
            paths.append(_store_powerschool_file(
                batch, file_type, uploaded, payload,
                parsed['files'][file_type]['headers']))
        row_number = 1
        uploaded_students = set()
        counts = {name: 0 for name in (
            'new', 'update', 'unchanged', 'duplicate', 'conflict',
            'rejected', 'ignored', 'warning', 'deactivate_candidate')}
        for proposal in parsed['students']:
            uploaded_students.add(proposal['student_number'])
            if proposal.get('school_year') and proposal['school_year'] != school_year:
                proposal.setdefault('conflicts', []).append(
                    'source school_year does not match the selected batch year')
            classification, selected, errors = _powerschool_compare_proposal(proposal)
            counts[classification] += 1
            normalized_json = json.dumps(
                proposal, ensure_ascii=False, sort_keys=True)
            db.session.add(ImportRow(
                batch_id=batch.id, row_number=row_number,
                external_key=proposal['student_number'],
                classification=classification, selected=selected,
                normalized_json=normalized_json,
                errors_json=json.dumps(errors),
                row_hash=hashlib.sha256(normalized_json.encode()).hexdigest()))
            row_number += 1

        for issue in parsed['issues']:
            classification = issue['classification']
            counts[classification] = counts.get(classification, 0) + 1
            normalized_json = json.dumps(issue, sort_keys=True)
            db.session.add(ImportRow(
                batch_id=batch.id, row_number=row_number, external_key=None,
                classification=classification, selected=False,
                normalized_json=normalized_json,
                errors_json=json.dumps(issue.get('errors', [])),
                row_hash=hashlib.sha256(normalized_json.encode()).hexdigest()))
            row_number += 1

        roster_state = _active_non_powerschool_roster_state()
        legacy_subscribers = roster_state['legacy']
        legacy_incarnation_excluded = (
            roster_state['incarnation_excluded_count']
            + len(roster_state['conflicts'])
        )
        unmanaged_count = (
            len(roster_state['unmanaged']) + len(roster_state['conflicts'])
        )
        baseline_required = bool(
            unmanaged_count and parsed['students']
        )
        baseline_available = bool(
            baseline_required
            and not roster_state['powerschool_active_count']
            and not _currently_applied_powerschool_batch_exists())
        legacy_contact_identities = _contact_identity_map(legacy_subscribers)
        legacy_cutover_required = bool(
            (legacy_subscribers or baseline_required) and parsed['students'])
        legacy_cutover_contract_blocked = bool(
            legacy_cutover_required
            and (parsed.get('preflight') or {}).get(
                'transportation_contract') != TRANSPORTATION_V2_CONTRACT)
        legacy_cutover_scope_blocked = bool(
            legacy_cutover_required
            and (snapshot_type != 'full_district'
                 or legacy_cutover_contract_blocked))
        legacy_cutover_blocked = bool(
            legacy_cutover_required
            and (baseline_required or legacy_cutover_scope_blocked
                 or counts['conflict'] or counts['rejected']))
        if legacy_cutover_required and not legacy_cutover_blocked:
            for subscriber in legacy_subscribers:
                candidate = {
                    'target_subscriber_id': subscriber.id,
                    'cutover_source': 'legacy_csv',
                    'group_name': subscriber.group.name if subscriber.group else '',
                    'expected_state_hash': _snapshot_hash(
                        _subscriber_snapshot(
                            subscriber, legacy_contact_identities)),
                    'changes': [{'field': 'active', 'current': 'yes',
                                 'proposed': 'no'}],
                }
                normalized_json = json.dumps(candidate, sort_keys=True)
                db.session.add(ImportRow(
                    batch_id=batch.id, row_number=row_number,
                    external_key=None,
                    classification='deactivate_candidate', selected=False,
                    normalized_json=normalized_json, errors_json='[]',
                    row_hash=hashlib.sha256(
                        normalized_json.encode()).hexdigest()))
                counts['deactivate_candidate'] += 1
                row_number += 1

        transportation_rejections = any(
            issue.get('file') == 'transportation'
            and issue.get('classification') in {'rejected', 'conflict'}
            for issue in parsed['issues'])
        snapshot_complete_for_deactivation = bool(
            snapshot_type == 'full_district' and parsed['students']
            and not transportation_rejections and counts['conflict'] == 0)
        if snapshot_complete_for_deactivation:
            identities = ExternalIdentity.query.filter_by(
                source_type='powerschool', entity_type='student',
                local_table='notification_subscriber').all()
            by_subscriber = {}
            for identity in identities:
                identity_subscriber = db.session.get(
                    NotificationSubscriber, identity.local_id)
                if not _powerschool_subscriber_identity_current(
                        identity, identity_subscriber):
                    continue
                by_subscriber.setdefault(identity.local_id, set()).add(identity.external_key)
            for subscriber_id, student_numbers in sorted(by_subscriber.items()):
                if student_numbers & uploaded_students:
                    continue
                subscriber = db.session.get(NotificationSubscriber, subscriber_id)
                if not subscriber or not subscriber.active:
                    continue
                candidate = {
                    'target_subscriber_id': subscriber.id,
                    'student_numbers': sorted(student_numbers),
                    'group_name': subscriber.group.name if subscriber.group else '',
                    'expected_state_hash': _snapshot_hash(
                        _subscriber_snapshot(subscriber)),
                    'changes': [{'field': 'active', 'current': 'yes',
                                 'proposed': 'no'}],
                }
                normalized_json = json.dumps(candidate, sort_keys=True)
                db.session.add(ImportRow(
                    batch_id=batch.id, row_number=row_number,
                    external_key='|'.join(sorted(student_numbers))[:160],
                    classification='deactivate_candidate', selected=False,
                    normalized_json=normalized_json, errors_json='[]',
                    row_hash=hashlib.sha256(normalized_json.encode()).hexdigest()))
                counts['deactivate_candidate'] += 1
                row_number += 1

        db.session.flush()
        metadata = {
            'mapping_profile_id': profile.id, 'mapping_profile_name': profile.name,
            'files': parsed['files'], 'issues': parsed['issues'], 'counts': counts,
            'metrics': parsed.get('metrics') or {},
            'preflight': parsed.get('preflight') or {},
            'normalizer_revision': NORMALIZER_REVISION,
            'analysis_context_sha256': analysis_context_sha256,
            'deactivation_policy': 'separate_explicit_approval',
            'snapshot_complete_for_deactivation': snapshot_complete_for_deactivation,
            'legacy_cutover': {
                'required': legacy_cutover_required,
                'candidate_count': len(legacy_subscribers),
                'incarnation_excluded_count': legacy_incarnation_excluded,
                'baseline_required': baseline_required,
                'baseline_available': baseline_available,
                'unmanaged_count': unmanaged_count if baseline_required else 0,
                'approved': False,
                'blocked': legacy_cutover_blocked,
                'requires_reanalysis': legacy_cutover_blocked,
                'message': (
                    _unmanaged_roster_block_message(baseline_available)
                    if baseline_required else
                    'Legacy CSV cutover requires the approved Transportation '
                    'v2 export and a district-wide Full Snapshot. Select Full '
                    'Snapshot and analyze the three files again before '
                    'replacing the active legacy roster.'
                    if legacy_cutover_scope_blocked else
                    'Resolve every conflict and rejected row, then analyze '
                    'again before replacing the active Legacy CSV roster.'
                    if legacy_cutover_blocked else
                    'Approve the atomic Legacy CSV to PowerSchool roster '
                    'cutover before applying this first PowerSchool batch.'
                    if legacy_cutover_required else ''
                ),
            },
        }
        if duplicate and force_reanalyze:
            metadata['reanalyzed_from'] = duplicate.public_id
        batch.metadata_json = json.dumps(metadata, sort_keys=True)
        _refresh_import_counts_and_hash(batch)
        db.session.commit()
    except Exception:
        db.session.rollback()
        for path in paths:
            if os.path.isfile(path):
                os.remove(path)
        raise
    _audit('powerschool_import_staged', 'notifications', batch.public_id,
           f'PowerSchool v{batch.schema_version}; {batch.total_rows} review rows; '
           f'sha256={batch.file_sha256[:12]}; legacy_reincarnation_excluded='
           f'{legacy_incarnation_excluded}')
    if duplicate and force_reanalyze:
        _audit('powerschool_import_reanalyzed', 'notifications', batch.public_id,
               f'New immutable analysis from {duplicate.public_id[:12]}; '
               f'normalizer={NORMALIZER_REVISION}')
    return jsonify(_powerschool_batch_payload(batch))


def _owned_powerschool_batch(public_id, statuses=None):
    query = ImportBatch.query.filter_by(
        public_id=public_id, source_type='powerschool')
    if not current_user.is_admin:
        query = query.filter_by(uploaded_by_id=current_user.id)
    if statuses:
        query = query.filter(ImportBatch.status.in_(statuses))
    return query.first_or_404()


@app.route('/admin/notifications/powerschool/batch/<public_id>')
@login_required
@require_capability('import.powerschool')
def powerschool_import_batch(public_id):
    _powerschool_enabled()
    batch = _owned_powerschool_batch(public_id)
    if batch.status == 'expired' or _expire_powerschool_stage(batch):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired. Analyze the three '
                        'files again.'}), 410
    return jsonify(_powerschool_batch_payload(batch))


@app.route('/admin/notifications/powerschool/batch/<public_id>/selection',
           methods=['POST'])
@login_required
@require_capability('import.powerschool')
@_serialized_roster_mutation('json')
def powerschool_import_selection(public_id):
    _powerschool_enabled()
    batch = _owned_powerschool_batch(public_id)
    if batch.status == 'expired' or _expire_powerschool_stage(batch):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired. Analyze the three '
                        'files again.'}), 410
    if batch.status != 'staged':
        return jsonify({'ok': False, 'message':
                        'The batch is not available for selection changes.'}), 409
    payload = request.get_json(silent=True) or {}
    expected_plan_hash = str(payload.get('plan_hash', ''))
    if not hmac.compare_digest(batch.plan_hash, expected_plan_hash):
        return jsonify({'ok': False, 'message':
                        'The displayed batch changed; reload it before selecting rows.'}), 409
    selected_ids = payload.get('selected_row_ids', [])
    deactivation_ids = payload.get('deactivation_row_ids', [])
    if not isinstance(selected_ids, list) or not isinstance(deactivation_ids, list):
        return jsonify({'ok': False, 'message': 'Selection must be a list.'}), 400
    try:
        selected_ids = {int(value) for value in selected_ids}
        deactivation_ids = {int(value) for value in deactivation_ids}
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'Selection contains an invalid row.'}), 400
    metadata = _import_metadata(batch)
    cutover = _legacy_cutover_payload(batch, metadata)
    legacy_cutover_approved = payload.get('legacy_cutover_approved') is True
    if cutover.get('baseline_required') or cutover.get('requires_reanalysis'):
        return jsonify({'ok': False, 'message':
                        cutover.get('message') or
                        'This roster state requires a new analysis.'}), 409
    claimed = ImportBatch.query.filter_by(
        id=batch.id, status='staged', plan_hash=expected_plan_hash).update(
            {'status': 'selecting'}, synchronize_session=False)
    db.session.commit()
    if claimed != 1:
        return jsonify({'ok': False, 'message':
                        'The batch is being changed by another request.'}), 409
    batch = db.session.get(ImportBatch, batch.id)
    if _expire_powerschool_stage(batch, processing_owner=True):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired during selection. '
                        'Analyze the three files again.'}), 410
    cutover = _legacy_cutover_payload(batch, _import_metadata(batch))
    if cutover.get('baseline_required') or cutover.get('requires_reanalysis'):
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message':
                        cutover.get('message') or
                        'This roster state requires a new analysis.'}), 409
    rows = ImportRow.query.filter_by(batch_id=batch.id).all()
    legacy_rows = _legacy_cutover_rows(batch)
    legacy_ids = {row.id for row, _ in legacy_rows}
    requested_legacy_ids = deactivation_ids & legacy_ids
    deactivation_ids -= legacy_ids
    if requested_legacy_ids and not legacy_cutover_approved:
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message':
                        'Legacy CSV cutover rows require the separate cutover approval.'}), 409
    if deactivation_ids and not payload.get('confirm_deactivations'):
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message':
                        'Deactivations require separate explicit approval.'}), 409
    valid_regular = {row.id for row in rows
                     if row.classification in {'new', 'update'}}
    valid_deactivation = {
        row.id for row in rows
        if row.classification == 'deactivate_candidate'
        and row.id not in legacy_ids
    }
    if not selected_ids <= valid_regular or not deactivation_ids <= valid_deactivation:
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message':
                        'Selection contains a non-importable row.'}), 409
    if legacy_cutover_approved:
        if not cutover.get('required'):
            batch.status = 'staged'
            db.session.commit()
            return jsonify({'ok': False, 'message':
                            'This batch does not require a Legacy CSV cutover.'}), 409
        if cutover.get('blocked') or cutover.get('requires_reanalysis'):
            batch.status = 'staged'
            db.session.commit()
            return jsonify({'ok': False, 'message':
                            cutover.get('message') or
                            'Analyze the files again before approving the cutover.'}), 409
        if len(legacy_ids) != int(cutover.get('candidate_count') or 0):
            batch.status = 'staged'
            db.session.commit()
            return jsonify({'ok': False, 'message':
                            'The staged Legacy CSV cutover is incomplete; analyze again.'}), 409
        replacement_count = sum(
            int((metadata.get('counts') or {}).get(classification) or 0)
            for classification in ('new', 'update', 'unchanged')
        )
        if replacement_count < 1 or selected_ids != valid_regular:
            batch.status = 'staged'
            db.session.commit()
            return jsonify({'ok': False, 'message':
                            'Legacy CSV cutover requires every importable New '
                            'and Update row to remain selected.'}), 409
    selected_legacy_ids = legacy_ids if legacy_cutover_approved else set()
    for row in rows:
        row.selected = (
            row.id in selected_ids
            or row.id in deactivation_ids
            or row.id in selected_legacy_ids
        )
    stored_cutover = dict(metadata.get('legacy_cutover') or {})
    if stored_cutover:
        stored_cutover['approved'] = bool(legacy_cutover_approved)
        if legacy_cutover_approved:
            stored_cutover['approved_by_id'] = current_user.id
            stored_cutover['approved_at'] = _utcnow().isoformat() + 'Z'
        else:
            for field in ('approved_by_id', 'approved_at',
                          'approved_plan_hash'):
                stored_cutover.pop(field, None)
        metadata['legacy_cutover'] = stored_cutover
    _refresh_import_counts_and_hash(batch)
    if stored_cutover and legacy_cutover_approved:
        stored_cutover['approved_plan_hash'] = batch.plan_hash
        metadata['legacy_cutover'] = stored_cutover
    batch.metadata_json = json.dumps(metadata, sort_keys=True)
    batch.status = 'staged'
    db.session.commit()
    cutover_audit = (
        f'; legacy_cutover=approved; candidates='
        f'{int(stored_cutover.get("candidate_count") or 0)}; '
        f'approver_id={current_user.id}; plan_hash={batch.plan_hash}'
        if stored_cutover and legacy_cutover_approved else
        '; legacy_cutover=not_approved'
        if stored_cutover.get('required') else '')
    _audit('powerschool_import_selection', 'notifications', batch.public_id,
           f'{batch.selected_rows} selected; {batch.excluded_rows} excluded; '
           f'{batch.rejected_rows} rejected{cutover_audit}')
    return jsonify(_powerschool_batch_payload(batch, include_rows=False))


def _ensure_external_identity(batch, row, entity_type, external_key,
                              local_table, local_id, identity_cache=None,
                              contact_identities=None,
                              pending_changes=None):
    identity = _powerschool_identity(
        entity_type, external_key, identity_cache)
    if identity:
        if identity.local_table != local_table or identity.local_id != local_id:
            raise ValueError('An external identity changed after preview.')
        if local_table == 'notification_subscriber':
            subscriber = db.session.get(NotificationSubscriber, local_id)
            if not _powerschool_subscriber_identity_current(
                    identity, subscriber):
                raise ValueError(
                    'An external identity belongs to a stale enrollment incarnation.')
        return identity
    identity = ExternalIdentity(
        source_type='powerschool', entity_type=entity_type,
        external_key=external_key, local_table=local_table, local_id=local_id)
    db.session.add(identity)
    if identity_cache is not None:
        identity_cache[(entity_type, external_key)] = identity
    if contact_identities is not None and local_table == 'subscriber_contact':
        contact_identities.setdefault(local_id, []).append(identity)
    if pending_changes is not None:
        pending_changes.append((identity, entity_type, external_key))
    else:
        db.session.flush()
        db.session.add(ImportChange(
            batch_id=batch.id, row_id=row.id, operation='create',
            target_table='external_identity', target_id=identity.id,
            after_json=json.dumps({'entity_type': entity_type,
                                   'external_key': external_key}, sort_keys=True)))
    return identity


def _apply_powerschool_proposal(
        batch, row, proposal, created_groups, identity_cache=None,
        contact_identities=None, assignment_index=None,
        schedule_types_by_name=None):
    """Prepare one proposal without issuing per-row identity flushes.

    Apply prepares every selected proposal first, flushes all subscribers and
    contacts together, then attaches identities and records immutable changes.
    Keeping those phases inside the request's single transaction preserves the
    existing all-or-nothing behavior while avoiding thousands of database
    round trips for district-sized imports.
    """
    identity = _powerschool_identity(
        'student', proposal['student_number'], identity_cache)
    subscriber = None
    if identity:
        identity_subscriber = db.session.get(
            NotificationSubscriber, identity.local_id)
        if not _powerschool_subscriber_identity_current(
                identity, identity_subscriber):
            raise ValueError(
                'A student identity belongs to a stale enrollment incarnation.')
        subscriber = identity_subscriber
    current_snapshot = (
        _subscriber_snapshot(subscriber, contact_identities)
        if subscriber else None)
    if _snapshot_hash(current_snapshot) != proposal.get('expected_state_hash'):
        raise ValueError('An enrollment changed after preview; analyze the files again.')
    group, created, group_error = _powerschool_group(
        proposal, create=True, assignment_index=assignment_index,
        schedule_types_by_name=schedule_types_by_name)
    if group_error:
        raise ValueError(group_error)
    if created and group.id not in created_groups:
        created_groups.add(group.id)
        db.session.add(ImportChange(
            batch_id=batch.id, row_id=row.id, operation='create',
            target_table='subscriber_group', target_id=group.id,
            after_json=json.dumps(
                _subscriber_group_snapshot(group), sort_keys=True)))
    operation = 'update' if subscriber else 'create'
    if not subscriber:
        subscriber = NotificationSubscriber(active=True)
        db.session.add(subscriber)
    subscriber.group_id = group.id
    subscriber.notes = _powerschool_household_label(proposal)
    subscriber.school = _normalize_text(proposal.get('school'), 100) or None
    subscriber.active = True
    specs = _powerschool_contact_specs(proposal)
    prepared_contacts = []
    for index, spec in enumerate(specs):
        mapped = []
        for entity_type, external_key in spec['identities']:
            contact_identity = _powerschool_identity(
                entity_type, external_key, identity_cache)
            if contact_identity:
                contact = db.session.get(SubscriberContact,
                                         contact_identity.local_id)
                if not contact or contact.subscriber_id != subscriber.id:
                    raise ValueError('A contact identity changed after preview.')
                mapped.append(contact)
        contact = mapped[0] if mapped else None
        if mapped and len({item.id for item in mapped}) != 1:
            raise ValueError('Contact identities no longer resolve consistently.')
        if not contact:
            contact = SubscriberContact(subscriber=subscriber)
            db.session.add(contact)
        contact.first_name = spec['first_name'] or None
        contact.last_name = spec['last_name'] or None
        contact.email = spec['email'] or None
        contact.phone = spec['phone'] or None
        contact.role = spec['role']
        contact.preferred_language = spec['preferred_language'] or 'en'
        contact.sort_order = index
        prepared_contacts.append((spec, contact))

    return {
        'row': row,
        'proposal': proposal,
        'subscriber': subscriber,
        'current_snapshot': current_snapshot,
        'operation': operation,
        'contacts': prepared_contacts,
        'pending_identity_changes': [],
    }


def _attach_powerschool_proposal_identities(
        batch, prepared, identity_cache, contact_identities):
    """Attach identities after the global subscriber/contact flush."""
    row = prepared['row']
    proposal = prepared['proposal']
    pending = prepared['pending_identity_changes']
    for spec, contact in prepared['contacts']:
        for entity_type, external_key in spec['identities']:
            _ensure_external_identity(
                batch, row, entity_type, external_key,
                'subscriber_contact', contact.id,
                identity_cache=identity_cache,
                contact_identities=contact_identities,
                pending_changes=pending)
    _ensure_external_identity(
        batch, row, 'student', proposal['student_number'],
        'notification_subscriber', prepared['subscriber'].id,
        identity_cache=identity_cache,
        contact_identities=contact_identities,
        pending_changes=pending)


def _record_powerschool_proposal_changes(
        batch, prepared, contact_identities):
    """Record immutable changes after every new identity has an ID."""
    row = prepared['row']
    for external_identity, entity_type, external_key in prepared[
            'pending_identity_changes']:
        db.session.add(ImportChange(
            batch_id=batch.id, row_id=row.id, operation='create',
            target_table='external_identity', target_id=external_identity.id,
            after_json=json.dumps({'entity_type': entity_type,
                                   'external_key': external_key}, sort_keys=True)))
    subscriber = prepared['subscriber']
    after_snapshot = _subscriber_snapshot(subscriber, contact_identities)
    db.session.add(ImportChange(
        batch_id=batch.id, row_id=row.id, operation=prepared['operation'],
        target_table='notification_subscriber', target_id=subscriber.id,
        before_json=(json.dumps(prepared['current_snapshot'], sort_keys=True)
                     if prepared['current_snapshot'] else None),
        after_json=json.dumps(after_snapshot, sort_keys=True)))


def _powerschool_apply_durable_state(batch_id, expected_selected=None):
    """Classify the durable result after an Apply transaction error."""
    batch = db.session.get(ImportBatch, batch_id)
    if not batch:
        return 'inconsistent', None, {'changes': 0}
    changes = ImportChange.query.filter_by(batch_id=batch.id).count()
    metadata = _import_metadata(batch)
    summary = metadata.get('applied_summary')
    selected_matches = (
        expected_selected is None
        or (isinstance(summary, dict)
            and summary.get('selected') == expected_selected)
    )
    if (batch.status == 'applied'
            and batch.applied_at is not None
            and isinstance(summary, dict)
            and summary.get('changes') == changes
            and selected_matches):
        return 'applied', batch, {'changes': changes}
    if (batch.status == 'applying'
            and batch.applied_at is None
            and changes == 0
            and summary is None):
        return 'clean_interruption', batch, {'changes': 0}
    return 'inconsistent', batch, {'changes': changes}


def _powerschool_apply_recovery_summary(batch):
    """Build an aggregate, PII-safe proof for an interrupted Apply."""
    durable_state, _, durable = _powerschool_apply_durable_state(
        batch.id, batch.selected_rows)
    rows = ImportRow.query.filter_by(batch_id=batch.id).order_by(
        ImportRow.row_number).all()
    selected = [row for row in rows if row.selected]
    recalculated_plan_hash = canonical_plan_hash(
        batch.public_id, batch.schema_version,
        _powerschool_row_records(batch))
    plan_intact = hmac.compare_digest(
        str(batch.plan_hash or ''), recalculated_plan_hash)
    counts_intact = (
        batch.total_rows == len(rows)
        and batch.selected_rows == len(selected)
        and batch.rejected_rows
        == sum(row.classification == 'rejected' for row in rows)
        and batch.excluded_rows
        == len(rows) - len(selected)
        - sum(row.classification == 'rejected' for row in rows)
    )

    files = ImportFile.query.filter_by(batch_id=batch.id).order_by(
        ImportFile.file_type).all()
    file_evidence_intact = bool(files)
    for stored in files:
        if not os.path.isfile(stored.storage_path):
            file_evidence_intact = False
            continue
        digest = hashlib.sha256()
        with open(stored.storage_path, 'rb') as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b''):
                digest.update(chunk)
        if (digest.hexdigest() != stored.sha256
                or os.path.getsize(stored.storage_path) != stored.byte_size):
            file_evidence_intact = False

    identities = ExternalIdentity.query.filter_by(
        source_type='powerschool').with_for_update().all()
    identity_cache = {
        (identity.entity_type, identity.external_key): identity
        for identity in identities
    }
    target_ids = {
        identity.local_id for identity in identities
        if identity.entity_type == 'student'
        and identity.local_table == 'notification_subscriber'
    }
    proposals = []
    malformed_rows = 0
    for row in selected:
        try:
            proposal = json.loads(row.normalized_json)
            proposals.append((row, proposal))
            target_id = proposal.get('target_subscriber_id')
            if target_id is not None:
                target_ids.add(int(target_id))
        except (TypeError, ValueError, json.JSONDecodeError):
            malformed_rows += 1

    subscribers = (
        NotificationSubscriber.query.options(
            selectinload(NotificationSubscriber.contacts),
        ).filter(NotificationSubscriber.id.in_(target_ids)).with_for_update().all()
        if target_ids else []
    )
    if target_ids:
        SubscriberContact.query.filter(
            SubscriberContact.subscriber_id.in_(target_ids),
        ).with_for_update().all()
    subscribers_by_id = {subscriber.id: subscriber for subscriber in subscribers}
    contact_identities = {}
    for identity in identities:
        if identity.local_table == 'subscriber_contact':
            contact_identities.setdefault(identity.local_id, []).append(identity)

    prestate_mismatches = malformed_rows
    for row, proposal in proposals:
        subscriber = None
        if row.classification in {'new', 'update'}:
            identity = identity_cache.get(
                ('student', str(proposal.get('student_number') or '')))
            if identity:
                subscriber = subscribers_by_id.get(identity.local_id)
                if not _powerschool_subscriber_identity_current(
                        identity, subscriber):
                    prestate_mismatches += 1
                    continue
        elif row.classification == 'deactivate_candidate':
            try:
                subscriber = subscribers_by_id.get(
                    int(proposal.get('target_subscriber_id')))
            except (TypeError, ValueError):
                prestate_mismatches += 1
                continue
        else:
            prestate_mismatches += 1
            continue
        current = (_subscriber_snapshot(subscriber, contact_identities)
                   if subscriber else None)
        if _snapshot_hash(current) != proposal.get('expected_state_hash'):
            prestate_mismatches += 1

    applied_audits = AuditLog.query.filter_by(
        action='powerschool_import_applied', target=batch.public_id).count()
    cutover_error = _legacy_cutover_apply_error(batch)
    recoverable = bool(
        durable_state == 'clean_interruption'
        and durable['changes'] == 0
        and applied_audits == 0
        and plan_intact
        and counts_intact
        and file_evidence_intact
        and len(selected) > 0
        and prestate_mismatches == 0
        and cutover_error is None
    )
    proof = {
        'batch_id': batch.public_id,
        'status': batch.status,
        'durable_state': durable_state,
        'plan_hash': batch.plan_hash,
        'file_sha256': batch.file_sha256,
        'total_rows': len(rows),
        'selected_rows': len(selected),
        'changes': durable['changes'],
        'applied_audits': applied_audits,
        'prestate_mismatches': prestate_mismatches,
        'plan_intact': plan_intact,
        'counts_intact': counts_intact,
        'file_evidence_intact': file_evidence_intact,
        'cutover_intact': cutover_error is None,
        'recoverable': recoverable,
    }
    proof['manifest_sha256'] = hashlib.sha256(json.dumps(
        proof, ensure_ascii=False, sort_keys=True,
        separators=(',', ':')).encode('utf-8')).hexdigest()
    return proof


@app.cli.command('recover-powerschool-apply')
@click.argument('batch_public_id')
@click.option('--apply', 'apply_changes', is_flag=True,
              help='Mark a proven clean interruption failed so it can be reanalyzed.')
@click.option('--manifest-sha', default=None,
              help='Exact recovery manifest SHA-256 printed by the dry-run.')
@click.option('--expected-plan-hash', default=None)
@click.option('--expected-file-sha', default=None)
@click.option('--expected-selected', type=click.IntRange(min=1), default=None)
@click.option('--approved-by', default=None,
              help='Username of the active administrator approving recovery.')
@click.option('--confirm-worker-stopped', is_flag=True,
              help='Confirm the interrupted worker/transaction no longer exists.')
def recover_powerschool_apply_command(
        batch_public_id, apply_changes, manifest_sha, expected_plan_hash,
        expected_file_sha, expected_selected, approved_by,
        confirm_worker_stopped):
    """Prove and close one Apply interrupted before its atomic commit."""
    with _roster_import_lock() as acquired:
        if not acquired:
            raise click.ClickException(
                'Another roster mutation is in progress; try again later.')
        try:
            batch = ImportBatch.query.filter_by(
                public_id=batch_public_id,
                source_type='powerschool',
            ).with_for_update().first()
            if not batch:
                raise ValueError('The PowerSchool batch was not found.')
            proof = _powerschool_apply_recovery_summary(batch)
            if not apply_changes:
                db.session.rollback()
                click.echo(json.dumps(proof, sort_keys=True))
                return

            required = {
                '--manifest-sha': manifest_sha,
                '--expected-plan-hash': expected_plan_hash,
                '--expected-file-sha': expected_file_sha,
                '--expected-selected': expected_selected,
                '--approved-by': approved_by,
            }
            missing = [name for name, value in required.items()
                       if value is None or value == '']
            if missing or not confirm_worker_stopped:
                raise ValueError(
                    'Apply recovery requires every dry-run pin, an approver, '
                    'and --confirm-worker-stopped.')
            for label, value in (
                    ('manifest', manifest_sha),
                    ('plan', expected_plan_hash),
                    ('file', expected_file_sha)):
                if not re.fullmatch(r'[0-9a-fA-F]{64}', str(value)):
                    raise ValueError(f'The expected {label} SHA-256 is invalid.')
            if not proof['recoverable']:
                raise ValueError(
                    'The batch is not a proven clean interruption; leave it '
                    'unchanged for manual review.')
            pins_match = (
                hmac.compare_digest(
                    proof['manifest_sha256'], str(manifest_sha).lower())
                and hmac.compare_digest(
                    proof['plan_hash'], str(expected_plan_hash).lower())
                and hmac.compare_digest(
                    proof['file_sha256'], str(expected_file_sha).lower())
                and proof['selected_rows'] == expected_selected
            )
            if not pins_match:
                raise ValueError('The recovery proof changed after approval.')
            approver = User.query.filter_by(
                username=str(approved_by), active=True,
            ).with_for_update().first()
            approver_group = None
            if approver and approver.group_id:
                approver_group = UserGroup.query.filter_by(
                    id=approver.group_id, is_admin=True,
                ).with_for_update().first()
            if not approver or not approver_group:
                raise ValueError(
                    'The approving username is not an active administrator.')

            now = _utcnow()
            metadata = _import_metadata(batch)
            metadata['apply_recovery'] = {
                'reason': 'worker_interrupted_before_atomic_commit',
                'manifest_sha256': proof['manifest_sha256'],
                'approved_by_id': approver.id,
                'recovered_at': now.isoformat() + 'Z',
            }
            batch.status = 'failed'
            batch.metadata_json = json.dumps(metadata, sort_keys=True)
            db.session.add(AuditLog(
                user_id=approver.id, username=approver.username,
                action='powerschool_import_apply_recovered',
                module='notifications', target=batch.public_id,
                details=(
                    f'clean_interruption; selected={proof["selected_rows"]}; '
                    f'manifest_sha256={proof["manifest_sha256"]}'
                ),
                ip_address='cli', created_at=now,
            ))
            db.session.commit()
            click.echo(json.dumps({
                **proof, 'status': 'failed', 'recovered': True,
            }, sort_keys=True))
        except ValueError as exc:
            db.session.rollback()
            raise click.ClickException(str(exc)) from exc
        except Exception:
            db.session.rollback()
            raise


@app.route('/admin/notifications/powerschool/batch/<public_id>/apply',
           methods=['POST'])
@login_required
@require_capability('import.powerschool')
@_serialized_roster_mutation('json')
def powerschool_import_apply(public_id):
    _powerschool_enabled()
    actor_id = current_user.id
    actor_username = current_user.username
    actor_ip = request.remote_addr or '0.0.0.0'
    payload = request.get_json(silent=True) or request.form
    plan_hash = str(payload.get('plan_hash', ''))
    batch = _owned_powerschool_batch(public_id)
    if batch.status == 'applied':
        return jsonify({'ok': True, 'already_applied': True,
                        'batch': _powerschool_batch_payload(
                            batch, include_rows=False)})
    if batch.status == 'expired' or _expire_powerschool_stage(batch):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired. Analyze the three '
                        'files again.'}), 410
    if batch.status != 'staged':
        return jsonify({'ok': False, 'message':
                        'The batch is not available to apply.'}), 409
    if not hmac.compare_digest(batch.plan_hash, plan_hash):
        return jsonify({'ok': False, 'message':
                        'The approved plan no longer matches the staged selection.'}), 409
    if batch.selected_rows < 1:
        return jsonify({'ok': False, 'message': 'Select at least one change.'}), 409
    contract_error = _powerschool_apply_contract_error(batch)
    if contract_error:
        return jsonify({'ok': False, 'message': contract_error}), 409
    cutover_error = _legacy_cutover_apply_error(batch)
    if cutover_error:
        return jsonify({'ok': False, 'message': cutover_error}), 409
    claimed = ImportBatch.query.filter_by(
        id=batch.id, status='staged', plan_hash=plan_hash).update(
            {'status': 'applying'}, synchronize_session=False)
    db.session.commit()
    if claimed != 1:
        return jsonify({'ok': False, 'message':
                        'The batch is already being processed.'}), 409
    batch = db.session.get(ImportBatch, batch.id)
    if _expire_powerschool_stage(batch, processing_owner=True):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired before Apply could '
                        'mutate the roster. Analyze the three files again.'}), 410
    contract_error = _powerschool_apply_contract_error(batch)
    if contract_error:
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message': contract_error}), 409
    selected = ImportRow.query.filter_by(
        batch_id=batch.id, selected=True).order_by(ImportRow.row_number).all()
    cutover_error = _legacy_cutover_apply_error(batch)
    if cutover_error:
        batch.status = 'staged'
        db.session.commit()
        return jsonify({'ok': False, 'message': cutover_error}), 409
    created_groups = set()
    try:
        selected_items = [
            (row, json.loads(row.normalized_json)) for row in selected
        ]
        target_ids = {
            int(proposal['target_subscriber_id'])
            for _, proposal in selected_items
            if proposal.get('target_subscriber_id') is not None
        }
        locked_subscribers = (
            NotificationSubscriber.query.options(
                selectinload(NotificationSubscriber.contacts),
            ).filter(
                NotificationSubscriber.id.in_(target_ids),
            ).with_for_update().all()
            if target_ids else []
        )
        if target_ids:
            SubscriberContact.query.filter(
                SubscriberContact.subscriber_id.in_(target_ids),
            ).with_for_update().all()
        locked_by_id = {
            subscriber.id: subscriber for subscriber in locked_subscribers
        }
        locked_identities = ExternalIdentity.query.filter_by(
            source_type='powerschool').with_for_update().all()
        identity_cache = {
            (identity.entity_type, identity.external_key): identity
            for identity in locked_identities
        }
        contact_identities = {}
        for identity in locked_identities:
            if identity.local_table == 'subscriber_contact':
                contact_identities.setdefault(
                    identity.local_id, []).append(identity)
        schedule_types_by_name = {
            schedule_type.name: schedule_type
            for schedule_type in BusScheduleType.query.all()
        }
        assignment_index = {}
        groups = SubscriberGroup.query.options(
            selectinload(SubscriberGroup.bus_assignments),
        ).order_by(SubscriberGroup.id).all()
        for group in groups:
            assignment_key = tuple(sorted(
                {
                    (assignment.bus_id, assignment.schedule_type_id)
                    for assignment in group.bus_assignments
                },
                key=lambda value: (value[0], value[1] or 0),
            ))
            if assignment_key:
                assignment_index.setdefault(assignment_key, group)

        prepared_proposals = []
        prepared_deactivations = []
        for row, proposal in selected_items:
            if row.classification in {'new', 'update'}:
                prepared_proposals.append(
                    _apply_powerschool_proposal(
                        batch, row, proposal, created_groups,
                        identity_cache=identity_cache,
                        contact_identities=contact_identities,
                        assignment_index=assignment_index,
                        schedule_types_by_name=schedule_types_by_name))
            elif row.classification == 'deactivate_candidate':
                subscriber = locked_by_id.get(
                    int(proposal['target_subscriber_id']))
                current = (_subscriber_snapshot(
                    subscriber, contact_identities)
                    if subscriber else None)
                if not subscriber or _snapshot_hash(current) != proposal.get(
                        'expected_state_hash'):
                    raise ValueError('A deactivation candidate changed after preview.')
                subscriber.active = False
                prepared_deactivations.append((row, subscriber, current))
            else:
                raise ValueError('The staged selection contains a non-importable row.')

        # Phase 1 assigns every new subscriber/contact ID in one unit-of-work
        # flush. Phase 2 then creates every external identity and assigns those
        # IDs in a second global flush. No operational row is committed until
        # the immutable audit changes below are complete.
        db.session.flush()
        for prepared in prepared_proposals:
            _attach_powerschool_proposal_identities(
                batch, prepared, identity_cache, contact_identities)
        db.session.flush()
        for prepared in prepared_proposals:
            _record_powerschool_proposal_changes(
                batch, prepared, contact_identities)
        for row, subscriber, current in prepared_deactivations:
            db.session.add(ImportChange(
                batch_id=batch.id, row_id=row.id, operation='deactivate',
                target_table='notification_subscriber', target_id=subscriber.id,
                before_json=json.dumps(current, sort_keys=True),
                after_json=json.dumps(_subscriber_snapshot(
                    subscriber, contact_identities), sort_keys=True)))
        batch.status = 'applied'
        batch.applied_at = _utcnow()
        metadata = _import_metadata(batch)
        metadata['applied_summary'] = {
            'selected': len(selected),
            'changes': ImportChange.query.filter_by(batch_id=batch.id).count(),
        }
        batch.metadata_json = json.dumps(metadata, sort_keys=True)
        db.session.commit()
    except Exception as exc:
        batch_id = batch.id
        db.session.rollback()
        db.session.remove()
        try:
            durable_state, durable_batch, _ = _powerschool_apply_durable_state(
                batch_id, len(selected))
        except Exception:
            app.logger.exception(
                'Could not classify the durable PowerSchool Apply result.')
            return jsonify({'ok': False, 'message':
                            'Apply ended in an unverified database state. '
                            'Do not retry until an administrator runs the '
                            'recovery proof.'}), 503
        if durable_state == 'applied':
            # PostgreSQL committed, but the client connection lost the commit
            # acknowledgement. Continue the idempotent cleanup/audit path and
            # report success instead of corrupting the durable status.
            batch = durable_batch
            app.logger.warning(
                'PowerSchool Apply commit was confirmed after a request error.')
        elif durable_state == 'clean_interruption':
            failed = durable_batch
            failed.status = 'failed'
            metadata = _import_metadata(failed)
            metadata['failure'] = (
                f'{type(exc).__name__}: {str(exc)[:240]}')
            failed.metadata_json = json.dumps(metadata, sort_keys=True)
            db.session.add(AuditLog(
                user_id=actor_id,
                username=actor_username,
                action='powerschool_import_failed',
                module='notifications', target=public_id,
                details='Atomic transaction rolled back.',
                ip_address=actor_ip))
            db.session.commit()
            return jsonify({'ok': False, 'message':
                            'Import failed atomically; no selected changes '
                            'were applied. Analyze the files again.'}), 409
        else:
            app.logger.exception(
                'PowerSchool Apply ended in an inconsistent durable state.')
            return jsonify({'ok': False, 'message':
                            'Apply ended in an unverified database state. '
                            'Do not retry until an administrator runs the '
                            'recovery proof.'}), 503
    cleanup_failures = _purge_import_raw_files(batch)
    db.session.add(AuditLog(
        user_id=actor_id, username=actor_username,
        action='powerschool_import_applied', module='notifications',
        target=public_id,
        details=(f'{len(selected)} selected rows; '
                 f'{batch.excluded_rows} excluded; '
                 f'{batch.rejected_rows} rejected'),
        ip_address=actor_ip))
    db.session.commit()
    return jsonify({
        'ok': True, 'cleanup_warnings': len(cleanup_failures),
        'batch': _powerschool_batch_payload(batch, include_rows=False),
    })


def _restore_subscriber_snapshot(subscriber, snapshot):
    subscriber.notes = snapshot.get('notes')
    subscriber.active = bool(snapshot.get('active'))
    subscriber.group_id = snapshot.get('group_id')
    if 'school' in snapshot:
        subscriber.school = snapshot.get('school')
    desired = {item['id']: item for item in snapshot.get('contacts', [])}
    for contact in list(subscriber.contacts):
        if contact.id not in desired:
            db.session.delete(contact)
    for contact_id, item in desired.items():
        contact = db.session.get(SubscriberContact, contact_id)
        if not contact:
            raise ValueError('Rollback cannot recreate a removed pre-import contact safely.')
        for field in ('first_name', 'last_name', 'email', 'phone', 'role',
                      'preferred_language', 'sort_order'):
            if field not in item:
                continue
            setattr(contact, field, item.get(field))


def _detach_notification_history(*, subscriber_id=None, group_id=None):
    """Preserve delivery history while a rollback removes imported targets.

    PostgreSQL enforces the nullable notification/outbox foreign keys with
    restrictive constraints in existing installations.  Clear only those
    references before deleting the imported subscriber or group; the stored
    recipient, group name, bus label and delivery result remain available for
    auditing.
    """
    detached = {'notification_log': 0, 'email_outbox': 0}
    if subscriber_id is not None:
        detached['notification_log'] += NotificationLog.query.filter_by(
            subscriber_id=subscriber_id).update(
                {NotificationLog.subscriber_id: None},
                synchronize_session=False)
        detached['email_outbox'] += EmailOutbox.query.filter_by(
            subscriber_id=subscriber_id).update(
                {EmailOutbox.subscriber_id: None},
                synchronize_session=False)
    if group_id is not None:
        detached['notification_log'] += NotificationLog.query.filter_by(
            group_id=group_id).update(
                {NotificationLog.group_id: None},
                synchronize_session=False)
        detached['email_outbox'] += EmailOutbox.query.filter_by(
            group_id=group_id).update(
                {EmailOutbox.group_id: None},
                synchronize_session=False)
    return detached


@app.route('/admin/notifications/powerschool/batch/<public_id>/rollback',
           methods=['POST'])
@login_required
@require_capability('import.rollback')
@_serialized_roster_mutation('json')
def powerschool_import_rollback(public_id):
    _powerschool_enabled()
    batch = _owned_powerschool_batch(public_id)
    if batch.status == 'rolled_back':
        return jsonify({'ok': True, 'already_rolled_back': True,
                        'batch': _powerschool_batch_payload(
                            batch, include_rows=False)})
    if batch.status not in {'applied', 'rollback_failed'}:
        return jsonify({'ok': False, 'message':
                        'The batch is not available to roll back.'}), 409
    rollback_from_status = batch.status
    claimed = ImportBatch.query.filter(
        ImportBatch.id == batch.id,
        ImportBatch.status.in_(['applied', 'rollback_failed'])).update(
            {'status': 'rolling_back'}, synchronize_session=False)
    db.session.commit()
    if claimed != 1:
        return jsonify({'ok': False, 'message': 'Rollback is already running.'}), 409
    try:
        batch = db.session.get(ImportBatch, batch.id)
        changes = ImportChange.query.filter_by(batch_id=batch.id).order_by(
            ImportChange.id.desc()).all()
        subscriber_ids = {
            change.target_id for change in changes
            if change.target_table == 'notification_subscriber'
            and change.target_id is not None
        }
        group_ids = {
            change.target_id for change in changes
            if change.target_table == 'subscriber_group'
            and change.operation == 'create'
            and change.target_id is not None
        }
        locked_subscribers = (
            NotificationSubscriber.query.options(
                selectinload(NotificationSubscriber.contacts),
            ).filter(
                NotificationSubscriber.id.in_(subscriber_ids),
            ).with_for_update().all()
            if subscriber_ids else []
        )
        if subscriber_ids:
            SubscriberContact.query.filter(
                SubscriberContact.subscriber_id.in_(subscriber_ids),
            ).with_for_update().all()
        subscribers_by_id = {
            subscriber.id: subscriber for subscriber in locked_subscribers
        }
        locked_contact_identities = _contact_identity_map(
            locked_subscribers, for_update=True)
        locked_groups = (
            SubscriberGroup.query.options(
                selectinload(SubscriberGroup.bus_assignments),
            ).filter(
                SubscriberGroup.id.in_(group_ids),
            ).with_for_update().all()
            if group_ids else []
        )
        if group_ids:
            GroupBusAssignment.query.filter(
                GroupBusAssignment.group_id.in_(group_ids),
            ).with_for_update().all()
        groups_by_id = {group.id: group for group in locked_groups}

        # Validate after the batch claim and while PostgreSQL row locks are
        # held, so a later manual edit cannot slip between preflight and restore.
        for change in changes:
            if change.target_table != 'notification_subscriber':
                continue
            subscriber = subscribers_by_id.get(change.target_id)
            expected = json.loads(change.after_json) if change.after_json else None
            if not _subscriber_matches_snapshot(
                    subscriber, expected, locked_contact_identities,
                    change.created_at):
                batch.status = rollback_from_status
                db.session.commit()
                return jsonify({'ok': False, 'message':
                                'Rollback blocked because imported data was '
                                'edited later.'}), 409
        for change in changes:
            if (change.target_table != 'subscriber_group'
                    or change.operation != 'create'):
                continue
            group = groups_by_id.get(change.target_id)
            expected = json.loads(change.after_json) if change.after_json else None
            if not _subscriber_group_matches_snapshot(
                    group, expected, change.created_at):
                batch.status = rollback_from_status
                db.session.commit()
                return jsonify({'ok': False, 'message':
                                'Rollback blocked because an imported group '
                                'was edited later.'}), 409

        detached_history = {'notification_log': 0, 'email_outbox': 0}
        for change in changes:
            if change.target_table == 'external_identity':
                identity = db.session.get(ExternalIdentity, change.target_id)
                if identity:
                    db.session.delete(identity)
            elif change.target_table == 'notification_subscriber':
                subscriber = subscribers_by_id.get(change.target_id)
                if change.operation == 'create':
                    if subscriber:
                        detached = _detach_notification_history(
                            subscriber_id=subscriber.id)
                        for key, value in detached.items():
                            detached_history[key] += value
                        db.session.delete(subscriber)
                elif subscriber and change.before_json:
                    _restore_subscriber_snapshot(
                        subscriber, json.loads(change.before_json))
            elif change.target_table == 'subscriber_group' and change.operation == 'create':
                group = groups_by_id.get(change.target_id)
                if group:
                    db.session.flush()
                    if NotificationSubscriber.query.filter_by(
                            group_id=group.id).count():
                        raise ValueError('A created group is now used by another enrollment.')
                    detached = _detach_notification_history(group_id=group.id)
                    for key, value in detached.items():
                        detached_history[key] += value
                    db.session.delete(group)
        batch = db.session.get(ImportBatch, batch.id)
        batch.status = 'rolled_back'
        metadata = _import_metadata(batch)
        metadata['rolled_back_at'] = _utcnow().isoformat() + 'Z'
        metadata['rolled_back_by_id'] = current_user.id
        metadata['detached_notification_history'] = detached_history
        batch.metadata_json = json.dumps(metadata, sort_keys=True)
        db.session.commit()
    except Exception:
        db.session.rollback()
        failed = db.session.get(ImportBatch, batch.id)
        failed.status = 'rollback_failed'
        db.session.commit()
        _audit('powerschool_import_rollback_failed', 'notifications', public_id,
               'Compensating transaction rolled back.')
        return jsonify({'ok': False, 'message':
                        'Rollback failed atomically; imported data remains unchanged.'}), 409
    _audit('powerschool_import_rolled_back', 'notifications', public_id,
           f'{len(changes)} recorded change(s) reversed')
    return jsonify({'ok': True, 'batch': _powerschool_batch_payload(
        batch, include_rows=False)})


@app.route('/admin/notifications/powerschool/batch/<public_id>/report.csv')
@login_required
@require_capability('import.powerschool')
def powerschool_import_report(public_id):
    _powerschool_enabled()
    batch = _owned_powerschool_batch(public_id)
    if batch.status == 'expired' or _expire_powerschool_stage(batch):
        return jsonify({'ok': False, 'status': 'expired', 'message':
                        'This staged analysis expired. Analyze the three '
                        'files again before downloading a report.'}), 410
    cutover = _legacy_cutover_payload(batch, _import_metadata(batch))
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['schema_version', 'batch_id', 'school_year', 'row_number',
                     'external_key', 'classification', 'selected', 'status',
                     'school', 'grade', 'group', 'errors',
                     'legacy_cutover_approved',
                     'legacy_cutover_candidate_count',
                     'legacy_cutover_approved_by_id',
                     'legacy_cutover_approved_at',
                     'legacy_cutover_approved_plan_hash'])
    can_view_pii = current_user.has_capability('notifications.pii')
    for row in ImportRow.query.filter_by(batch_id=batch.id).order_by(
            ImportRow.row_number).all():
        data = json.loads(row.normalized_json)
        external_key = row.external_key or ''
        if not can_view_pii:
            external_key = '••••' + external_key[-4:] if external_key else ''
            data = _masked_powerschool_data(data)
        writer.writerow([safe_csv_cell(value) for value in (
            f'PowerSchool Import v{batch.schema_version}', batch.public_id,
            batch.school_year, row.row_number, external_key,
            row.classification, 'yes' if row.selected else 'no', batch.status,
            data.get('school', ''), data.get('grade', ''),
            data.get('group_name', ''),
            '; '.join(json.loads(row.errors_json or '[]')),
            'yes' if cutover.get('approved') else 'no',
            int(cutover.get('candidate_count') or 0),
            cutover.get('approved_by_id', ''),
            cutover.get('approved_at', ''),
            cutover.get('approved_plan_hash', ''),
        )])
    response = make_response('\ufeff' + output.getvalue())
    response.headers['Content-Type'] = 'text/csv; charset=utf-8'
    response.headers['Content-Disposition'] = (
        f'attachment; filename=powerschool-{batch.public_id[:12]}-report.csv')
    return response

# ── SUBSCRIBER GROUPS ──────────────────────────────────────────────────────────

@app.route('/admin/notifications/groups/add', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def add_subscriber_group():
    name  = request.form.get('name', '').strip()
    color = request.form.get('color', 'blue').strip()
    desc  = request.form.get('description', '').strip()
    if not name:
        flash('Group name is required.', 'error')
        return redirect(url_for('notifications', tab='groups'))
    if SubscriberGroup.query.filter_by(name=name).first():
        flash(f'Group "{name}" already exists.', 'error')
        return redirect(url_for('notifications', tab='groups'))
    g = SubscriberGroup(name=name, color=color, description=desc)
    db.session.add(g)
    db.session.flush()
    for val in request.form.getlist('bus_sched'):
        try:
            parts  = val.split('_', 1)
            bid    = int(parts[0])
            sid    = int(parts[1]) if len(parts) > 1 and parts[1] else None
            db.session.add(GroupBusAssignment(group_id=g.id, bus_id=bid, schedule_type_id=sid))
        except Exception:
            pass
    db.session.commit()
    flash(f'Group "{name}" created.', 'success')
    return redirect(url_for('notifications', tab='groups'))


@app.route('/admin/notifications/groups/<int:gid>/delete', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def delete_subscriber_group(gid):
    g = SubscriberGroup.query.get_or_404(gid)
    # Unassign subscribers from this group before deleting
    NotificationSubscriber.query.filter_by(group_id=gid).update({'group_id': None})
    db.session.delete(g)
    db.session.commit()
    flash(f'Group "{g.name}" deleted.', 'success')
    return redirect(url_for('notifications', tab='groups'))


@app.route('/admin/notifications/groups/bulk-delete', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def bulk_delete_groups():
    ids = request.form.getlist('group_ids')
    count = 0
    for gid in ids:
        try:
            g = SubscriberGroup.query.get(int(gid))
            if g:
                NotificationSubscriber.query.filter_by(group_id=g.id).update({'group_id': None})
                db.session.delete(g)
                count += 1
        except (ValueError, TypeError):
            pass
    if count:
        db.session.commit()
        _audit('bulk_delete_groups', 'notifications', f'{count} groups deleted')
        flash(f'{count} group(s) deleted.', 'success')
    return redirect(url_for('notifications', tab='groups'))


@app.route('/admin/notifications/groups/<int:gid>/edit', methods=['POST'])
@login_required
@require_module('notifications', 'full')
@_serialized_roster_mutation('html')
def edit_subscriber_group(gid):
    g = SubscriberGroup.query.get_or_404(gid)
    name = request.form.get('name', '').strip()
    if not name:
        flash('Group name is required.', 'error')
        return redirect(url_for('notifications', tab='groups'))
    existing = SubscriberGroup.query.filter_by(name=name).first()
    if existing and existing.id != gid:
        flash(f'Group "{name}" already exists.', 'error')
        return redirect(url_for('notifications', tab='groups'))
    g.name        = name
    g.color       = request.form.get('color', g.color)
    g.description = request.form.get('description', '').strip()
    GroupBusAssignment.query.filter_by(group_id=gid).delete()
    for val in request.form.getlist('bus_sched'):
        try:
            parts  = val.split('_', 1)
            bid    = int(parts[0])
            sid    = int(parts[1]) if len(parts) > 1 and parts[1] else None
            db.session.add(GroupBusAssignment(group_id=gid, bus_id=bid, schedule_type_id=sid))
        except Exception:
            pass
    db.session.commit()
    flash('Group updated.', 'success')
    return redirect(url_for('notifications', tab='groups'))


# ── BROADCAST ─────────────────────────────────────────────────────────────────

def _build_recipient_list(target, group_ids, subscriber_id, user_id):
    """Return list of (name, email) tuples based on target selection."""
    recipients = []
    seen = set()

    def add(name, email):
        if email and email not in seen:
            seen.add(email)
            recipients.append((name, email))

    def add_sub(s):
        """Add all contacts of a subscriber (all emails per contact), or legacy email."""
        if s.contacts:
            for c in s.contacts:
                if c.email:
                    for em in [e.strip() for e in c.email.split(',') if e.strip()]:
                        add(c.full_name, em)
        else:
            if s.email:
                for em in [e.strip() for e in s.email.split(',') if e.strip()]:
                    add(s.full_name, em)

    if target in ('all', 'subscribers', 'group'):
        query = NotificationSubscriber.query.filter_by(active=True)
        if target == 'group' and group_ids:
            query = query.filter(NotificationSubscriber.group_id.in_(group_ids))
        for s in query.all():
            add_sub(s)

    if target in ('all', 'admins'):
        for u in User.query.filter_by(active=True, receive_notifications=True).all():
            add(u.username, u.email)

    if target == 'individual_subscriber' and subscriber_id:
        s = NotificationSubscriber.query.get(subscriber_id)
        if s:
            add_sub(s)

    if target == 'individual_user' and user_id:
        u = User.query.get(user_id)
        if u:
            add(u.username, u.email)

    return recipients


@app.route('/admin/notifications/broadcast', methods=['POST'])
@login_required
@require_capability('notifications.broadcast')
def send_broadcast():
    data         = request.get_json(silent=True) or {}
    target       = data.get('target', 'all')
    group_ids    = [int(x) for x in data.get('group_ids', []) if x]
    sub_id       = data.get('subscriber_id')
    user_id      = data.get('user_id')
    subject      = (data.get('subject') or '').strip()
    body         = (data.get('body') or '').strip()
    try:
        interval_sec = max(0, min(3600, int(data.get('interval', 0))))
    except (TypeError, ValueError):
        return jsonify({'ok': False, 'message': 'Send interval must be numeric.'}), 400

    if not subject or not body:
        return jsonify({'ok': False, 'message': 'Subject and body are required.'})

    recipients = _build_recipient_list(target, group_ids, sub_id, user_id)
    if not recipients:
        return jsonify({'ok': False, 'message': 'No valid recipients found.'})

    job_id = secrets.token_urlsafe(32)
    now = _utcnow()
    job = BroadcastJob(
        public_id=job_id, owner_id=current_user.id, status='queued',
        total=len(recipients), sent=0, failed=0, errors_json='[]',
        created_at=now, updated_at=now,
        expires_at=now + timedelta(seconds=max(
            app.config['BROADCAST_JOB_TTL_SECONDS'],
            interval_sec * max(0, len(recipients) - 1) + 3600,
        )))
    db.session.add(job)
    db.session.flush()
    for index, (name, email) in enumerate(recipients):
        try:
            _enqueue_email(
                dedupe_key=_email_dedupe_key(
                    'broadcast', job_id, email.strip().lower()),
                kind='broadcast',
                recipient_name=name,
                recipient_address=email,
                subject=subject,
                body=f"Hi {name or 'there'},\n\n{body}",
                available_at=now + timedelta(seconds=index * interval_sec),
                broadcast_job_id=job_id,
            )
        except ValueError:
            job.failed += 1
            errors = job.errors
            if len(errors) < 100:
                errors.append(f'{_mask_email(email)}: invalid_recipient')
            job.errors_json = json.dumps(errors)
    _complete_broadcast_if_ready(job_id)
    db.session.commit()
    _audit('broadcast_started', 'notifications', job_id,
           f'{len(recipients)} recipients')
    return jsonify({'ok': True, 'job_id': job_id, 'total': len(recipients)})


@app.route('/admin/notifications/broadcast/<job_id>/status')
@login_required
@require_capability('notifications.broadcast')
def broadcast_status(job_id):
    job = db.session.get(BroadcastJob, job_id)
    if not job:
        abort(404)
    if job.expires_at <= _utcnow():
        if job.status not in {'completed', 'failed', 'expired'}:
            job.status = 'expired'
            db.session.commit()
        abort(404)
    if job.owner_id != current_user.id and not current_user.is_admin:
        abort(403)
    return jsonify({
        'total': job.total, 'sent': job.sent, 'failed': job.failed,
        'done': job.done, 'status': job.status, 'errors': job.errors,
    })


# ── CONFIGURATION MODULE ──────────────────────────────────────────────────────

@app.route('/admin/config', methods=['GET', 'POST'])
@login_required
@require_module('config')
def config_page():
    cfg = get_config()
    if request.method == 'POST':
        if not current_user.has_access('config', 'full'):
            abort(403)
        section = request.form.get('section', 'general')
        allowed_sections = {'general', 'theme', 'operational', 'schedule_windows',
                            'language', 'email', 'sms'}
        if section not in allowed_sections:
            abort(400)
        if section == 'general':
            cfg.app_name     = request.form.get('app_name', cfg.app_name).strip()
            cfg.app_subtitle = request.form.get('app_subtitle', cfg.app_subtitle).strip()
            cfg.time_format  = request.form.get('time_format', '12h')
        elif section == 'theme':
            mode = request.form.get('theme_mode', 'light')
            if mode not in {'light', 'dark'}:
                abort(400)
            colors = {field: request.form.get(field, getattr(cfg, field)) for field in
                      ('color_bg', 'color_nav', 'color_card', 'color_text',
                       'color_accent', 'color_nav_text')}
            if any(not re.fullmatch(r'#[0-9A-Fa-f]{6}', value or '') for value in colors.values()):
                abort(400)
            cfg.theme_mode = mode
            for field, value in colors.items():
                setattr(cfg, field, value)
        elif section == 'operational':
            timezone = request.form.get('timezone', cfg.timezone)
            reset_time = request.form.get('daily_reset_time', cfg.daily_reset_time)
            delay = request.form.get('commit_delay_min', cfg.commit_delay_min, type=int)
            if timezone not in pytz.all_timezones or not re.fullmatch(r'(?:[01]\d|2[0-3]):[0-5]\d', reset_time or ''):
                abort(400)
            if delay is None or not 0 <= delay <= 1440:
                abort(400)
            cfg.timezone = timezone
            cfg.daily_reset_time = reset_time
            cfg.commit_delay_min = delay
            cfg.offline_message  = request.form.get('offline_message', cfg.offline_message)
            cfg.show_always      = 'show_always' in request.form
        elif section == 'schedule_windows':
            for p in BusScheduleType.query.all():
                w_start = request.form.get(f'window_start_{p.id}', '').strip()
                w_end   = request.form.get(f'window_end_{p.id}',   '').strip()
                p.window_start = w_start or None
                p.window_end   = w_end   or None
        elif section == 'language':
            frontend = request.form.get('lang_frontend', 'en')
            backend = request.form.get('lang_backend', 'en')
            if frontend not in TRANSLATIONS or backend not in TRANSLATIONS:
                abort(400)
            cfg.lang_frontend = frontend
            cfg.lang_backend = backend
        elif section == 'email':
            provider = request.form.get('mail_provider', 'custom').strip().lower()
            server = request.form.get('mail_server', '').strip()
            tls = 'mail_use_tls' in request.form
            ssl = 'mail_use_ssl' in request.form
            username = request.form.get('mail_username', '').strip()
            from_email = request.form.get('mail_from_email', '').strip()
            from_name = request.form.get('mail_from_name', '').strip()
            try:
                validated_server, port, tls, ssl = _canonical_smtp_transport(
                    provider, server, request.form.get('mail_port', 587), tls, ssl)
                from_email = _validated_email(
                    from_email, 'From email address', required=True)
                if len(username) > 320 or len(from_name) > 100:
                    raise ValueError('Email account fields exceed the allowed length.')
            except (ValueError, RuntimeError) as exc:
                flash(str(exc), 'error')
                return redirect(url_for('config_page', tab='email'))
            new_pwd = request.form.get('mail_password', '')
            try:
                old_identity = _current_smtp_identity(cfg)
            except ValueError:
                old_identity = ''
            new_identity = _smtp_identity(
                provider, validated_server, port, tls, ssl, username)
            if old_identity != new_identity and cfg.mail_password and not new_pwd:
                flash('Re-enter the SMTP password when changing connection settings.', 'error')
                return redirect(url_for('config_page', tab='email'))
            try:
                if new_pwd:
                    encrypted_password = _encrypt_mail_password(new_pwd)
                elif cfg.mail_password and not _mail_password_is_encrypted(cfg):
                    encrypted_password = _encrypt_mail_password(cfg.mail_password)
                else:
                    encrypted_password = cfg.mail_password or ''
            except RuntimeError as exc:
                flash(str(exc), 'error')
                return redirect(url_for('config_page', tab='email'))
            cfg.mail_provider = provider
            cfg.mail_server = validated_server
            cfg.mail_port = port
            cfg.mail_use_tls = tls
            cfg.mail_use_ssl = ssl
            cfg.mail_username = username
            cfg.mail_password = encrypted_password
            cfg.mail_from_email = from_email
            cfg.mail_from_name = from_name or 'Bus Tracker'
            new_verification_identity = _mail_verification_identity(
                new_identity, cfg.mail_from_email)
            if cfg.mail_last_verified_identity != new_verification_identity:
                cfg.mail_last_verification_status = 'unverified'
                cfg.mail_last_error_code = ''
        elif section == 'sms':
            enabled = 'twilio_enabled' in request.form
            account_sid = request.form.get('twilio_account_sid', '').strip()
            from_number = request.form.get('twilio_from_number', '').strip()
            new_tok = request.form.get('twilio_auth_token', '').strip()
            if ((account_sid != (cfg.twilio_account_sid or '') or
                 from_number != (cfg.twilio_from_number or '')) and
                    cfg.twilio_auth_token and not new_tok):
                flash('Re-enter the Twilio Auth Token when changing account settings.', 'error')
                return redirect(url_for('config_page', tab='sms'))
            cfg.twilio_enabled = enabled
            cfg.twilio_account_sid = account_sid
            cfg.twilio_from_number = from_number
            if new_tok:
                cfg.twilio_auth_token = new_tok
            try:
                cfg.twilio_sms_cost_per_seg = float(request.form.get('twilio_sms_cost_per_seg', 0.0079))
            except (ValueError, TypeError):
                pass
        requeued_email_count = _requeue_configuration_failures() if section == 'email' else 0
        db.session.commit()
        if section == 'email':
            _audit(
                'email_config_saved', 'config', cfg.mail_provider,
                f'credential_updated={bool(new_pwd)}; requeued={requeued_email_count}',
            )
        flash('Configuration saved.', 'success')
        return redirect(url_for('config_page', tab=section))

    # Operational schedules and holidays
    schedules      = OperationalSchedule.query.order_by(OperationalSchedule.name).all()
    holidays       = Holiday.query.order_by(Holiday.holiday_date.desc()).all()
    schedule_types = BusScheduleType.query.order_by(BusScheduleType.sort_order).all()
    timezones      = ['America/New_York','America/Chicago','America/Denver',
                      'America/Los_Angeles','America/Anchorage','Pacific/Honolulu',
                      'America/Puerto_Rico','Europe/London','Europe/Madrid']
    active_tab = request.args.get('tab', 'general')
    can_write  = current_user.has_access('config', 'full')
    return render_template('admin/config.html', cfg=cfg, schedules=schedules,
                           holidays=holidays, schedule_types=schedule_types,
                           timezones=timezones, active_tab=active_tab, can_write=can_write,
                           mail_status=_mail_configuration_status(cfg),
                           smtp_presets=SMTP_PROVIDER_PRESETS,
                           smtp_current=_smtp_public_settings(cfg))

@app.route('/admin/config/upload-logo', methods=['POST'])
@login_required
@require_module('config', 'full')
def upload_logo():
    cfg  = get_config()
    field = request.form.get('field', 'logo')
    f = request.files.get('file')
    if field not in {'logo', 'icon'}:
        abort(400)
    if f and allowed_file(f.filename):
        import warnings
        from PIL import Image, UnidentifiedImageError
        expected_formats = {
            'png': {'PNG'}, 'jpg': {'JPEG'}, 'jpeg': {'JPEG'},
            'gif': {'GIF'}, 'ico': {'ICO'},
        }
        extension = f.filename.rsplit('.', 1)[1].lower()
        payload = f.read()
        try:
            with warnings.catch_warnings():
                warnings.simplefilter('error', Image.DecompressionBombWarning)
                image = Image.open(io.BytesIO(payload))
                image.verify()
            if image.format not in expected_formats[extension]:
                raise ValueError('Image format does not match its extension.')
        except (Image.DecompressionBombError, Image.DecompressionBombWarning,
                UnidentifiedImageError, OSError, ValueError):
            flash('Upload rejected: the file is not a valid approved image.', 'error')
            return redirect(url_for('config_page', tab='general'))
        fn = f'app_{field}_{secrets.token_hex(12)}.{extension}'
        destination = os.path.join(app.config['UPLOAD_FOLDER'], fn)
        with open(destination, 'xb') as output:
            output.write(payload)
        os.chmod(destination, 0o600)
        if field == 'logo':
            cfg.logo_path = f'/static/uploads/{fn}'
        else:
            cfg.icon_path = f'/static/uploads/{fn}'
        db.session.commit()
        flash('File uploaded.', 'success')
    else:
        flash('Upload rejected: select a PNG, JPEG, GIF, or ICO image.', 'error')
    return redirect(url_for('config_page', tab='general'))

@app.route('/admin/config/test-email', methods=['POST'])
@login_required
@require_capability('smtp.diagnose')
def test_email():
    cfg = get_config()
    to = request.form.get('test_email', current_user.email or '')
    try:
        to = _validated_email(to, 'recipient email address', required=True)
        settings = _smtp_settings_from_config(cfg)
        send_email(
            settings,
            subject=f'Test Email — {cfg.app_name}',
            recipients=[to],
            body=f'This is a test email from {cfg.app_name}.',
        )
        _record_mail_verification(
            cfg, _current_smtp_identity(cfg), settings.from_email,
            'delivery_verified')
        _audit('smtp_test_succeeded', 'config', 'saved approved destination')
        flash(f'Test email sent to {to}.', 'success')
    except (EmailTransportError, ValueError, RuntimeError) as exc:
        failure = exc if isinstance(exc, EmailTransportError) else EmailTransportError(
            'configuration_invalid', str(exc))
        try:
            _record_mail_verification(
                cfg, _current_smtp_identity(cfg), cfg.mail_from_email,
                'failed', failure.code)
        except (ValueError, RuntimeError):
            db.session.rollback()
        _audit('smtp_test_failed', 'config', failure.code)
        flash(f'Email test failed: {failure.safe_message}', 'error')
    return redirect(url_for('config_page', tab='email'))


@app.route('/admin/config/test-email-live', methods=['POST'])
@login_required
@require_capability('smtp.diagnose')
def test_email_live():
    """AJAX endpoint: test SMTP with current form values (does not save to DB)."""
    data = request.get_json(silent=True) or {}
    cfg = get_config()
    settings = None
    identity = None
    try:
        test_to = _validated_email(
            data.get('test_to', ''), 'recipient email address', required=True)
        settings, identity = _smtp_settings_from_payload(
            cfg, data, allow_saved_password=True)
        send_email(
            settings,
            subject=f'Test Email — {cfg.app_name}',
            recipients=[test_to],
            body=(f'This is a test email from {cfg.app_name}.\n\n'
                  f'SMTP: {settings.server}:{settings.port}\n'
                  f'TLS: {settings.use_tls}  SSL: {settings.use_ssl}\n'
                  f'From: {settings.from_name} <{settings.from_email}>'),
        )
        if _matches_saved_mail_identity(cfg, identity, settings.from_email):
            _record_mail_verification(
                cfg, identity, settings.from_email, 'delivery_verified')
        _audit('smtp_live_test_succeeded', 'config', settings.server)
        return jsonify({'ok': True, 'message': f'Test email sent successfully to {test_to}.'})
    except (EmailTransportError, ValueError, RuntimeError) as exc:
        failure = exc if isinstance(exc, EmailTransportError) else EmailTransportError(
            'configuration_invalid', str(exc))
        if settings and _matches_saved_mail_identity(cfg, identity, settings.from_email):
            _record_mail_verification(
                cfg, identity, settings.from_email, 'failed', failure.code)
        _audit('smtp_live_test_failed', 'config', failure.code)
        return jsonify({'ok': False, 'message': failure.safe_message,
                        'code': failure.code}), 400

@app.route('/admin/config/check-smtp', methods=['POST'])
@login_required
@require_capability('smtp.diagnose')
def check_smtp():
    """Verify the current form using an isolated SMTP connection."""
    data = request.get_json(silent=True) or {}
    cfg  = get_config()
    settings = None
    identity = None
    try:
        settings, identity = _smtp_settings_from_payload(
            cfg, data, allow_saved_password=True)
        _audit('smtp_diagnostic_started', 'config', f'{settings.server}:{settings.port}')
        verify_connection(settings)
        if _matches_saved_mail_identity(cfg, identity, settings.from_email):
            _record_mail_verification(
                cfg, identity, settings.from_email, 'connection_verified')
        security_label = 'SSL/TLS' if settings.use_ssl else (
            'STARTTLS' if settings.use_tls else 'unencrypted')
        steps = [
            {'ok': True, 'label': f'TCP and SMTP handshake — {settings.server}:{settings.port}'},
            {'ok': True, 'label': f'Connection security — {security_label}'},
            {'ok': True, 'label': 'Authentication accepted' if settings.username
             else 'Authentication not configured'},
        ]
        _audit('smtp_diagnostic_succeeded', 'config', f'{settings.server}:{settings.port}')
        return jsonify({'ok': True, 'steps': steps})
    except (EmailTransportError, ValueError, RuntimeError) as exc:
        failure = exc if isinstance(exc, EmailTransportError) else EmailTransportError(
            'configuration_invalid', str(exc))
        if settings and _matches_saved_mail_identity(cfg, identity, settings.from_email):
            _record_mail_verification(
                cfg, identity, settings.from_email, 'failed', failure.code)
        _audit('smtp_diagnostic_failed', 'config', failure.code)
        return jsonify({'ok': False, 'code': failure.code, 'steps': [{
            'ok': False, 'label': 'SMTP diagnostic failed',
            'detail': failure.safe_message,
        }]}), 400


@app.route('/admin/config/check-twilio', methods=['POST'])
@login_required
@require_module('config', 'full')
def check_twilio():
    """AJAX: verify Twilio credentials without sending a message."""
    if not TWILIO_AVAILABLE:
        return jsonify({'ok': False, 'message': 'Twilio library not installed. Run: pip install twilio'})
    data = request.get_json(silent=True) or {}
    cfg  = get_config()
    sid  = data.get('account_sid', '') or ''
    tok  = data.get('auth_token',  '') or ''
    if not sid or not tok:
        return jsonify({'ok': False, 'message': 'Account SID and Auth Token are required.'})
    try:
        tw = TwilioClient(sid, tok)
        # Use incoming phone numbers list — works on both trial and paid accounts
        numbers = tw.incoming_phone_numbers.list(limit=1)
        label = numbers[0].phone_number if numbers else '(no numbers purchased yet)'
        return jsonify({'ok': True,
                        'message': f'Connected! Credentials valid. From number on account: {label}'})
    except TwilioRestException as e:
        # 20003 = authentication error (bad SID/token)
        if e.code == 20003:
            return jsonify({'ok': False, 'message': 'Authentication failed — check your Account SID and Auth Token.'})
        return jsonify({'ok': False, 'message': f'Twilio error {e.code}: {e.msg}'})
    except Exception:
        return jsonify({'ok': False, 'message': 'Twilio diagnostic failed.'}), 400


@app.route('/admin/config/test-sms', methods=['POST'])
@login_required
@require_module('config', 'full')
def test_sms():
    """AJAX: send a test SMS using current form values."""
    if not TWILIO_AVAILABLE:
        return jsonify({'ok': False, 'message': 'Twilio library not installed. Run: pip install twilio'})
    data = request.get_json(silent=True) or {}
    cfg  = get_config()
    sid      = data.get('account_sid', '') or ''
    tok      = data.get('auth_token',  '') or ''
    from_num = data.get('from_number', '') or ''
    to_num   = data.get('to_number',   '').strip()
    if not sid or not tok:
        return jsonify({'ok': False, 'message': 'Account SID and Auth Token are required.'})
    if not from_num:
        return jsonify({'ok': False, 'message': 'From Number is required.'})
    if not to_num:
        return jsonify({'ok': False, 'message': 'Destination phone number is required.'})
    try:
        tw  = TwilioClient(sid, tok)
        msg = tw.messages.create(
            to=to_num, from_=from_num,
            body=f'[{get_config().app_name}] Test SMS — configuration verified successfully.'
        )
        return jsonify({'ok': True, 'message': f'SMS sent! SID: {msg.sid} — Status: {msg.status}'})
    except TwilioRestException as e:
        return jsonify({'ok': False, 'message': f'Twilio error {e.code}: {e.msg}'})
    except Exception:
        return jsonify({'ok': False, 'message': 'Twilio test failed.'}), 400


@app.route('/admin/config/schedules/add', methods=['POST'])
@login_required
@require_module('config', 'full')
def add_schedule():
    s = OperationalSchedule(
        name=request.form.get('name','').strip(),
        days=request.form.get('days','mon-fri'),
        start_time=request.form.get('start_time','07:00'),
        end_time=request.form.get('end_time','17:00'),
    )
    db.session.add(s)
    db.session.commit()
    flash('Schedule added.', 'success')
    return redirect(url_for('config_page', tab='operational'))

@app.route('/admin/config/schedules/<int:sid>/delete', methods=['POST'])
@login_required
@require_module('config', 'full')
def delete_schedule(sid):
    db.session.delete(OperationalSchedule.query.get_or_404(sid))
    db.session.commit()
    flash('Schedule removed.', 'success')
    return redirect(url_for('config_page', tab='operational'))

@app.route('/admin/config/holidays/add', methods=['POST'])
@login_required
@require_module('config', 'full')
def add_holiday():
    try:
        h = Holiday(
            name=request.form.get('name','').strip(),
            holiday_type=request.form.get('holiday_type','school'),
            holiday_date=date.fromisoformat(request.form.get('holiday_date','')),
            is_recurring='is_recurring' in request.form,
            custom_message=request.form.get('custom_message','').strip() or None,
        )
        db.session.add(h)
        db.session.commit()
        flash('Holiday added.', 'success')
    except Exception:
        flash('Invalid date.', 'error')
    return redirect(url_for('config_page', tab='operational'))

@app.route('/admin/config/holidays/<int:hid>/edit', methods=['POST'])
@login_required
@require_module('config', 'full')
def edit_holiday(hid):
    h = Holiday.query.get_or_404(hid)
    name = request.form.get('name', '').strip()
    if name:
        h.name = name
    h.holiday_type   = request.form.get('holiday_type', h.holiday_type)
    h.custom_message = request.form.get('custom_message', '').strip() or None
    try:
        new_date = request.form.get('holiday_date', '').strip()
        if new_date:
            h.holiday_date = date.fromisoformat(new_date)
    except Exception:
        pass
    db.session.commit()
    flash('Holiday updated.', 'success')
    return redirect(url_for('config_page', tab='operational'))

@app.route('/admin/config/holidays/<int:hid>/delete', methods=['POST'])
@login_required
@require_module('config', 'full')
def delete_holiday(hid):
    db.session.delete(Holiday.query.get_or_404(hid))
    db.session.commit()
    flash('Holiday removed.', 'success')
    return redirect(url_for('config_page', tab='operational'))

@app.route('/admin/config/export-db')
@login_required
@require_capability('backup.export_sensitive')
def export_db():
    db_path = db.engine.url.database if db.engine.url.get_backend_name() == 'sqlite' else None
    if db_path and os.path.exists(db_path):
        try:
            with open(db_path, 'rb') as stream:
                response = _encrypted_download(stream.read(),
                                               f'bustrack_sqlite_{district_today()}.bustrack-db')
            _audit('export_full_backup', 'config', 'encrypted SQLite database')
            return response
        except RuntimeError as exc:
            flash(str(exc), 'error')
            return redirect(url_for('config_page', tab='data'))
    flash('Database file not found.', 'error')
    return redirect(url_for('config_page', tab='data'))


@app.route('/admin/config/system-status')
@login_required
@require_module('config')
def system_status():
    import platform, sys
    result = {}
    # DB info
    db_url = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    is_pg = db_url.startswith('postgresql')
    result['db_type'] = 'PostgreSQL' if is_pg else 'SQLite'
    try:
        if is_pg:
            row = db.session.execute(db.text('SELECT version()')).fetchone()
            result['db_version'] = row[0].split('\n')[0] if row else 'Unknown'
            size_row = db.session.execute(db.text(
                "SELECT pg_size_pretty(pg_database_size(current_database()))"
            )).fetchone()
            result['db_size'] = size_row[0] if size_row else 'Unknown'
        else:
            db_path = os.path.join(BASE_DIR, 'bustrack.db')
            if os.path.exists(db_path):
                sz = os.path.getsize(db_path)
                result['db_size'] = f'{sz/1024/1024:.2f} MB' if sz > 1024*1024 else f'{sz/1024:.1f} KB'
            else:
                result['db_size'] = 'N/A'
            result['db_version'] = 'SQLite ' + db.session.execute(db.text('SELECT sqlite_version()')).fetchone()[0]
        result['db_ok'] = True
        result['db_error'] = None
    except Exception as e:
        result['db_ok'] = False
        result['db_error'] = str(e)
        result['db_version'] = 'N/A'
        result['db_size'] = 'N/A'

    # Table row counts
    try:
        from sqlalchemy import inspect as sa_inspect
        inspector = sa_inspect(db.engine)
        tables = inspector.get_table_names()
        counts = {}
        for t in sorted(tables):
            try:
                row = db.session.execute(db.text(f'SELECT COUNT(*) FROM "{t}"')).fetchone()
                counts[t] = row[0]
            except Exception:
                counts[t] = '?'
        result['tables'] = counts
    except Exception as e:
        result['tables'] = {}

    # Server / process stats
    try:
        import shutil
        du = shutil.disk_usage(BASE_DIR)
        result['disk_total'] = f'{du.total/1024**3:.1f} GB'
        result['disk_used']  = f'{du.used/1024**3:.1f} GB'
        result['disk_free']  = f'{du.free/1024**3:.1f} GB'
        result['disk_pct']   = round(du.used / du.total * 100, 1)
    except Exception:
        result['disk_total'] = result['disk_used'] = result['disk_free'] = 'N/A'
        result['disk_pct'] = 0

    try:
        import psutil
        result['cpu_pct']   = psutil.cpu_percent(interval=0.5)
        mem = psutil.virtual_memory()
        result['mem_total'] = f'{mem.total/1024**3:.1f} GB'
        result['mem_used']  = f'{mem.used/1024**3:.1f} GB'
        result['mem_pct']   = mem.percent
        boot = datetime.fromtimestamp(psutil.boot_time())
        delta = datetime.now() - boot
        d, rem = divmod(int(delta.total_seconds()), 86400)
        h, rem = divmod(rem, 3600)
        m = rem // 60
        result['uptime'] = f'{d}d {h}h {m}m'
        result['psutil'] = True
    except ImportError:
        result['psutil'] = False
        result['cpu_pct'] = result['mem_pct'] = 'N/A'
        result['mem_total'] = result['mem_used'] = 'N/A'
        result['uptime'] = 'N/A (psutil not installed)'

    result['python'] = sys.version.split(' ')[0]
    result['platform'] = platform.system() + ' ' + platform.release()
    return jsonify(result)


@app.route('/admin/config/check-deps')
@login_required
@require_module('config')
def check_deps():
    import importlib.metadata, urllib.request
    from concurrent.futures import ThreadPoolExecutor, as_completed

    packages = [
        'Flask', 'Flask-SQLAlchemy', 'Flask-Login', 'Flask-Mail',
        'Werkzeug', 'APScheduler', 'fpdf2', 'python-docx',
        'pytz', 'psycopg2-binary', 'python-dotenv', 'gunicorn',
    ]

    def _installed_version(pkg):
        # normalise: psycopg2-binary → psycopg2-binary, try both hyphen/underscore
        for name in (pkg, pkg.replace('-', '_'), pkg.lower()):
            try:
                return importlib.metadata.version(name)
            except importlib.metadata.PackageNotFoundError:
                continue
        return None

    def _pypi_latest(pkg):
        try:
            url = f'https://pypi.org/pypi/{pkg}/json'
            with urllib.request.urlopen(url, timeout=5) as r:
                data = json.loads(r.read())
            return data['info']['version']
        except Exception:
            return None

    results = []
    with ThreadPoolExecutor(max_workers=8) as ex:
        fut_installed = {ex.submit(_installed_version, p): p for p in packages}
        installed = {fut_installed[f]: f.result() for f in as_completed(fut_installed)}

    with ThreadPoolExecutor(max_workers=8) as ex:
        fut_latest = {ex.submit(_pypi_latest, p): p for p in packages}
        latest = {fut_latest[f]: f.result() for f in as_completed(fut_latest)}

    def _parse_ver(v):
        if not v: return (0,)
        try:
            return tuple(int(x) for x in v.split('.')[:3])
        except Exception:
            return (0,)

    for pkg in packages:
        inst = installed.get(pkg)
        lat  = latest.get(pkg)
        iv = _parse_ver(inst)
        lv = _parse_ver(lat)
        status = 'ok'
        if not inst:
            status = 'missing'
        elif lat and lv > iv:
            status = 'major_update' if lv[0] > iv[0] else 'update'
        results.append({
            'package': pkg,
            'installed': inst or 'Not installed',
            'latest': lat or 'Unknown',
            'status': status,
        })
    return jsonify(results)


_IMPORT_TABLE_ORDER_V1 = [
    'user_group', 'group_permission', 'user', 'configuration',
    'operational_schedule', 'bus_schedule_type', 'incident_type',
    'delay_reason', 'bus', 'bus_schedule_assignment',
    'bus_incident_record', 'subscriber_group', 'group_bus_assignment',
    'notification_subscriber', 'subscriber_contact',
    'notification_bus_assignment', 'holiday', 'audit_log', 'login_throttle',
]

_IMPORT_TABLE_ORDER_V2 = [
    'user_group', 'group_permission', 'group_capability', 'user', 'configuration',
    'operational_schedule', 'bus_schedule_type', 'incident_type',
    'delay_reason', 'bus', 'bus_schedule_assignment',
    'bus_incident_record', 'subscriber_group', 'group_bus_assignment',
    'notification_subscriber', 'subscriber_contact',
    'notification_bus_assignment', 'notification_log', 'holiday',
    'audit_log', 'login_throttle', 'broadcast_job',
    'import_mapping_profile', 'import_batch', 'import_file', 'import_row',
    'external_identity', 'import_change',
]

_IMPORT_TABLE_ORDER = [
    'user_group', 'group_permission', 'group_capability', 'user', 'configuration',
    'operational_schedule', 'bus_schedule_type', 'incident_type',
    'delay_reason', 'bus', 'bus_schedule_assignment',
    'bus_incident_record', 'subscriber_group', 'group_bus_assignment',
    'notification_subscriber', 'subscriber_contact',
    'notification_bus_assignment', 'notification_log', 'holiday',
    'audit_log', 'login_throttle', 'broadcast_job', 'email_outbox',
    'import_mapping_profile', 'import_batch', 'import_file', 'import_row',
    'external_identity', 'import_change',
]


_OPERATIONAL_EXPORT_TABLES = [
    'operational_schedule', 'bus_schedule_type', 'holiday',
]
_SAFE_CONFIGURATION_COLUMNS = [
    'app_name', 'app_subtitle', 'logo_path', 'icon_path', 'theme_mode',
    'color_bg', 'color_nav', 'color_card', 'color_text', 'color_accent',
    'color_nav_text', 'timezone', 'daily_reset_time', 'commit_delay_min',
    'offline_message', 'show_always', 'lang_frontend', 'lang_backend',
    'time_format', 'mail_provider', 'mail_server', 'mail_port', 'mail_use_tls',
    'mail_use_ssl', 'mail_from_email', 'mail_from_name', 'twilio_enabled',
    'twilio_from_number', 'twilio_sms_cost_per_seg',
]
_BACKUP_FORMAT = 'bustrack-full-backup'
_BACKUP_VERSION = 3
_SUPPORTED_BACKUP_VERSIONS = {1, 2, 3}


def _json_default(value):
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    raise TypeError(f'Unsupported backup value type: {type(value).__name__}')


def _backup_fernet():
    from cryptography.fernet import Fernet
    key = os.environ.get('BACKUP_ENCRYPTION_KEY', '').strip().encode('ascii', 'strict')
    if not key:
        raise RuntimeError('BACKUP_ENCRYPTION_KEY is not configured.')
    try:
        return Fernet(key)
    except (ValueError, TypeError) as exc:
        raise RuntimeError('BACKUP_ENCRYPTION_KEY is invalid.') from exc


def _database_dump(table_names=None):
    from sqlalchemy import inspect as sa_inspect, text as sa_text
    inspector = sa_inspect(db.engine)
    existing = set(inspector.get_table_names())
    tables = sorted(existing) if table_names is None else [t for t in table_names if t in existing]
    dump = {}
    for table in tables:
        rows = db.session.execute(sa_text(f'SELECT * FROM "{table}"')).mappings().all()
        dump[table] = [dict(row) for row in rows]
    return dump


def _full_backup_document():
    return {
        'format': _BACKUP_FORMAT,
        'version': _BACKUP_VERSION,
        'created_at': _utcnow().isoformat() + 'Z',
        'database': 'postgresql' if str(db.engine.url).startswith('postgresql') else 'sqlite',
        'tables': _database_dump(_IMPORT_TABLE_ORDER),
    }


def _encrypted_download(payload, filename):
    token = _backup_fernet().encrypt(payload)
    response = make_response(token)
    response.headers['Content-Type'] = 'application/octet-stream'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    response.headers['Cache-Control'] = 'no-store'
    return response


@app.route('/admin/config/export-operational-json')
@login_required
@require_capability('backup.export_operational')
def export_operational_json():
    tables = list(_OPERATIONAL_EXPORT_TABLES)
    if current_user.has_access('buses'):
        tables.extend(['bus', 'bus_schedule_assignment'])
    if current_user.has_access('incidents'):
        tables.extend(['incident_type', 'delay_reason'])
    data = _database_dump(tables)
    cfg = get_config()
    data['configuration'] = [{column: getattr(cfg, column) for column in
                              _SAFE_CONFIGURATION_COLUMNS}]
    payload = json.dumps({
        'format': 'bustrack-operational-export',
        'version': 1,
        'created_at': _utcnow().isoformat() + 'Z',
        'tables': data,
    }, default=_json_default, indent=2, ensure_ascii=False)
    response = make_response(payload)
    response.headers['Content-Type'] = 'application/json; charset=utf-8'
    response.headers['Content-Disposition'] = (
        f'attachment; filename="bustrack_operational_{district_today()}.json"')
    response.headers['Cache-Control'] = 'no-store'
    _audit('export_operational', 'config', 'redacted operational data')
    return response


@app.route('/admin/config/export-json')
@login_required
@require_capability('backup.export_sensitive')
def export_json():
    try:
        payload = json.dumps(_full_backup_document(), default=_json_default,
                             separators=(',', ':'), ensure_ascii=False).encode('utf-8')
        response = _encrypted_download(payload, f'bustrack_full_{district_today()}.bustrack')
        _audit('export_full_backup', 'config', 'encrypted JSON backup')
        return response
    except RuntimeError as exc:
        flash(str(exc), 'error')
        return redirect(url_for('config_page', tab='data'))


@app.route('/admin/config/export-sql')
@login_required
@require_capability('backup.export_sensitive')
def export_sql():
    from sqlalchemy import inspect as sa_inspect, text as sa_text
    lines = [f'-- BusTrack encrypted SQL export source — {_utcnow().isoformat()}Z']
    for table in sorted(sa_inspect(db.engine).get_table_names()):
        rows = db.session.execute(sa_text(f'SELECT * FROM "{table}"')).mappings().all()
        for row in rows:
            data = dict(row)
            columns = ', '.join(f'"{key}"' for key in data)
            values = []
            for value in data.values():
                if value is None:
                    values.append('NULL')
                elif isinstance(value, bool):
                    values.append('TRUE' if value else 'FALSE')
                elif isinstance(value, (int, float)):
                    values.append(str(value))
                else:
                    values.append("'" + str(value).replace("'", "''") + "'")
            lines.append(f'INSERT INTO "{table}" ({columns}) VALUES ({", ".join(values)});')
    try:
        response = _encrypted_download('\n'.join(lines).encode('utf-8'),
                                       f'bustrack_sql_{district_today()}.bustrack-sql')
        _audit('export_full_backup', 'config', 'encrypted SQL backup')
        return response
    except RuntimeError as exc:
        flash(str(exc), 'error')
        return redirect(url_for('config_page', tab='data'))


RESTORE_JOB_DIR = os.path.join(INSTANCE_DIR, 'restore_jobs')
RESTORE_SNAPSHOT_DIR = os.path.join(INSTANCE_DIR, 'restore_snapshots')


def _job_path(job_id: str) -> str:
    if not re.fullmatch(r'[A-Za-z0-9_-]{20,100}', job_id):
        abort(404)
    return os.path.join(RESTORE_JOB_DIR, f'{job_id}.json')


def _write_private_file(path, payload, binary=False):
    os.makedirs(os.path.dirname(path), mode=0o700, exist_ok=True)
    flags = os.O_WRONLY | os.O_CREAT | os.O_EXCL
    fd = os.open(path, flags, 0o600)
    mode = 'wb' if binary else 'w'
    kwargs = {} if binary else {'encoding': 'utf-8'}
    with os.fdopen(fd, mode, **kwargs) as stream:
        stream.write(payload)


def _cleanup_restore_jobs():
    if os.path.isdir(RESTORE_JOB_DIR):
        cutoff = time.time() - app.config['RESTORE_JOB_TTL_SECONDS']
        for name in os.listdir(RESTORE_JOB_DIR):
            path = os.path.join(RESTORE_JOB_DIR, name)
            try:
                if os.path.isfile(path) and os.path.getmtime(path) < cutoff:
                    os.remove(path)
            except OSError:
                continue
    if os.path.isdir(RESTORE_SNAPSHOT_DIR):
        snapshot_cutoff = time.time() - (
            app.config['RESTORE_SNAPSHOT_RETENTION_DAYS'] * 86400)
        for name in os.listdir(RESTORE_SNAPSHOT_DIR):
            path = os.path.join(RESTORE_SNAPSHOT_DIR, name)
            try:
                if (name.endswith('.bustrack') and os.path.isfile(path) and
                        os.path.getmtime(path) < snapshot_cutoff):
                    os.remove(path)
            except OSError:
                continue


def _validate_backup_document(document):
    from sqlalchemy import inspect as sa_inspect
    if not isinstance(document, dict) or document.get('format') != _BACKUP_FORMAT:
        raise ValueError('Unsupported backup format.')
    version = document.get('version')
    if version not in _SUPPORTED_BACKUP_VERSIONS:
        raise ValueError('Unsupported backup version.')
    tables = document.get('tables')
    if not isinstance(tables, dict) or not tables:
        raise ValueError('Backup does not contain tables.')
    inspector = sa_inspect(db.engine)
    existing = set(inspector.get_table_names())
    import_order = (_IMPORT_TABLE_ORDER if version == 3 else
                    _IMPORT_TABLE_ORDER_V2 if version == 2 else
                    _IMPORT_TABLE_ORDER_V1)
    allowed = set(import_order) & existing
    unknown = set(tables) - allowed
    if unknown:
        raise ValueError('Backup contains tables that are not approved for restore.')
    missing = allowed - set(tables)
    if missing:
        raise ValueError('Full backup is incomplete and cannot be restored safely.')
    ordered = []
    for table in import_order:
        if table not in tables:
            continue
        rows = tables[table]
        if not isinstance(rows, list):
            raise ValueError(f'Backup table {table} is not a row list.')
        allowed_columns = {column['name'] for column in inspector.get_columns(table)}
        for row in rows:
            if not isinstance(row, dict) or not row or not set(row).issubset(allowed_columns):
                raise ValueError(f'Backup table {table} contains an invalid row schema.')
        ordered.append((table, rows))

    groups = tables.get('user_group', [])
    users = tables.get('user', [])
    truthy = {True, 1, '1', 'true', 'True'}
    admin_group_ids = {row.get('id') for row in groups if row.get('is_admin') in truthy}
    has_active_admin = any(
        row.get('active') in truthy and row.get('group_id') in admin_group_ids and
        bool(row.get('password_hash')) for row in users)
    if not has_active_admin:
        raise ValueError('Full backup does not contain an active administrator.')
    if len(tables.get('configuration', [])) != 1:
        raise ValueError('Full backup must contain exactly one configuration row.')
    return ordered


@app.route('/admin/config/import-db', methods=['POST'])
@login_required
@require_capability('restore.identity')
def import_db():
    """Validate an encrypted, versioned backup and stage an owner-bound restore job."""
    _cleanup_restore_jobs()
    f = request.files.get('backup_file')
    if not f or not f.filename.lower().endswith('.bustrack'):
        return jsonify({'ok': False, 'error': 'Upload an encrypted .bustrack backup.'}), 400
    try:
        decrypted = _backup_fernet().decrypt(f.read())
        document = json.loads(decrypted.decode('utf-8'))
        ordered = _validate_backup_document(document)
    except RuntimeError as exc:
        return jsonify({'ok': False, 'error': str(exc)}), 503
    except Exception:
        return jsonify({'ok': False, 'error':
                        'Backup could not be authenticated or validated.'}), 400

    job_id = secrets.token_urlsafe(32)
    job = {
        'owner_id': current_user.id,
        'created_at': _utcnow().isoformat() + 'Z',
        'backup_version': document.get('version'),
        'tables': ordered,
        'is_pg': str(db.engine.url).startswith('postgresql'),
    }
    _write_private_file(_job_path(job_id), json.dumps(job, default=_json_default))
    _audit('import_db_start', 'config', f'{len(ordered)} tables', f'File: {f.filename}')
    return jsonify({'ok': True, 'job_id': job_id, 'total': len(ordered)})


@app.route('/admin/config/import-run/<job_id>', methods=['POST'])
@login_required
@require_capability('restore.identity')
def import_run(job_id):
    """Apply the validated restore atomically after creating an encrypted snapshot."""
    _cleanup_restore_jobs()
    jpath = _job_path(job_id)
    if not os.path.exists(jpath):
        return jsonify({'ok': False, 'error': 'Restore job not found or expired.'}), 404
    with open(jpath, 'r', encoding='utf-8') as fp:
        job = json.load(fp)
    if job.get('owner_id') != current_user.id:
        abort(403)
    try:
        created = datetime.fromisoformat(job['created_at'].rstrip('Z'))
    except (KeyError, TypeError, ValueError):
        return jsonify({'ok': False, 'error': 'Restore job metadata is invalid.'}), 400
    if _utcnow() - created > timedelta(seconds=app.config['RESTORE_JOB_TTL_SECONDS']):
        os.remove(jpath)
        return jsonify({'ok': False, 'error': 'Restore job has expired.'}), 410

    try:
        tables = _validate_backup_document({
            'format': _BACKUP_FORMAT,
            'version': job.get('backup_version', _BACKUP_VERSION),
            'tables': dict(job['tables']),
        })
        for table_name, rows in tables:
            if table_name != 'configuration' or not rows:
                continue
            restored_secret = rows[0].get('mail_password') or ''
            if restored_secret.startswith(_ENCRYPTED_SECRET_PREFIX):
                _decrypt_mail_password(restored_secret)
            elif restored_secret:
                rows[0]['mail_password'] = _encrypt_mail_password(restored_secret)
        snapshot = json.dumps(_full_backup_document(), default=_json_default,
                              separators=(',', ':'), ensure_ascii=False).encode('utf-8')
        encrypted_snapshot = _backup_fernet().encrypt(snapshot)
        snapshot_name = f'pre_restore_{_utcnow().strftime("%Y%m%dT%H%M%SZ")}_{job_id}.bustrack'
        _write_private_file(os.path.join(RESTORE_SNAPSHOT_DIR, snapshot_name),
                            encrypted_snapshot, binary=True)

        from sqlalchemy import text as sa_text
        with db.engine.begin() as conn:
            table_names = [table for table, _ in tables]
            legacy_phase2_tables = [
                'import_change', 'import_file', 'import_row', 'import_batch',
                'external_identity', 'broadcast_job', 'group_capability',
                'import_mapping_profile',
            ] if job.get('backup_version') == 1 else []
            legacy_phase3_tables = (
                ['email_outbox'] if job.get('backup_version') in {1, 2} else [])
            if job.get('is_pg'):
                clear_tables = table_names + legacy_phase2_tables + legacy_phase3_tables
                quoted = ', '.join(f'"{table}"' for table in clear_tables)
                if quoted:
                    conn.execute(sa_text(f'TRUNCATE TABLE {quoted} RESTART IDENTITY CASCADE'))
            else:
                # Phase 3 outbox rows reference phase 2 broadcast jobs, so clear
                # the newer child tables first when restoring an older backup.
                for table in legacy_phase3_tables + legacy_phase2_tables:
                    conn.execute(sa_text(f'DELETE FROM "{table}"'))
                for table in reversed(table_names):
                    conn.execute(sa_text(f'DELETE FROM "{table}"'))
            for table, rows in tables:
                for row in rows:
                    columns = ', '.join(f'"{key}"' for key in row)
                    placeholders = ', '.join(f':{key}' for key in row)
                    conn.execute(sa_text(
                        f'INSERT INTO "{table}" ({columns}) VALUES ({placeholders})'), row)
                if job.get('is_pg') and rows and 'id' in rows[0]:
                    conn.execute(sa_text(
                        f"SELECT setval(pg_get_serial_sequence('\"{table}\"','id'), "
                        f"COALESCE(MAX(id), 1), MAX(id) IS NOT NULL) FROM \"{table}\""))
        db.session.expire_all()
        _seed_phase2_security_and_imports()
        os.remove(jpath)
        _audit('import_db_done', 'config', f'{len(tables)} tables restored',
               f'Pre-restore snapshot: {snapshot_name}')
        return jsonify({'ok': True, 'restored': len(tables), 'total': len(tables),
                        'snapshot': snapshot_name})
    except RuntimeError as exc:
        return jsonify({'ok': False, 'error': str(exc)}), 503
    except Exception:
        db.session.rollback()
        _audit('import_db_failed', 'config', 'atomic restore rolled back')
        return jsonify({'ok': False, 'error':
                        'Restore failed and the database transaction was rolled back.'}), 400

@app.route('/admin/config/manual-commit', methods=['POST'])
@login_required
@require_module('config', 'full')
def manual_commit():
    commit_pending_incidents()
    flash('All pending incidents committed to statistics.', 'success')
    return redirect(url_for('config_page', tab='operational'))


# ── SYSTEM LOGS ───────────────────────────────────────────────────────────────

@app.route('/admin/logs')
@login_required
@require_module('logs')
def system_logs():
    page     = request.args.get('page', 1, type=int)
    module_f = request.args.get('module', '').strip()
    user_f   = request.args.get('user', '').strip()
    date_f   = request.args.get('date', '').strip()
    search_f = request.args.get('q', '').strip()

    q = AuditLog.query.order_by(AuditLog.created_at.desc())
    if module_f:
        q = q.filter(AuditLog.module == module_f)
    if user_f:
        q = q.filter(AuditLog.username.ilike(f'%{user_f}%'))
    if search_f:
        like = f'%{search_f}%'
        q = q.filter(
            db.or_(AuditLog.action.ilike(like), AuditLog.target.ilike(like),
                   AuditLog.details.ilike(like))
        )
    if date_f:
        try:
            d = date.fromisoformat(date_f)
            q = q.filter(func.date(AuditLog.created_at) == d)
        except Exception:
            pass

    logs_page = q.paginate(page=page, per_page=50, error_out=False)
    all_modules = [r[0] for r in db.session.query(AuditLog.module).distinct().order_by(AuditLog.module).all() if r[0]]
    all_users   = [r[0] for r in db.session.query(AuditLog.username).distinct().order_by(AuditLog.username).all() if r[0]]
    return render_template('admin/logs.html',
                           logs=logs_page, all_modules=all_modules, all_users=all_users,
                           module_f=module_f, user_f=user_f, date_f=date_f, search_f=search_f)


@app.route('/admin/logs/export-csv')
@login_required
@require_capability('audit.export')
def export_logs_csv():
    module_f = request.args.get('module', '').strip()
    user_f   = request.args.get('user', '').strip()
    date_f   = request.args.get('date', '').strip()
    search_f = request.args.get('q', '').strip()

    q = AuditLog.query.order_by(AuditLog.created_at.desc())
    if module_f: q = q.filter(AuditLog.module == module_f)
    if user_f:   q = q.filter(AuditLog.username.ilike(f'%{user_f}%'))
    if search_f:
        like = f'%{search_f}%'
        q = q.filter(db.or_(AuditLog.action.ilike(like), AuditLog.target.ilike(like)))
    if date_f:
        try:
            d = date.fromisoformat(date_f)
            q = q.filter(func.date(AuditLog.created_at) == d)
        except Exception:
            pass

    buf = io.StringIO()
    buf.write('\ufeff')  # BOM for Excel
    writer = csv.writer(buf)
    writer.writerow(['Timestamp', 'Username', 'Action', 'Module', 'Target', 'Details', 'IP'])
    for log in q.all():
        writer.writerow(_csv_safe_row([
            log.created_at.strftime('%Y-%m-%d %H:%M:%S') if log.created_at else '',
            log.username or '', log.action or '', log.module or '',
            log.target or '', log.details or '', log.ip_address or '',
        ]))
    resp = make_response(buf.getvalue())
    resp.headers['Content-Type'] = 'text/csv; charset=utf-8'
    resp.headers['Content-Disposition'] = f'attachment; filename=audit_log_{district_today()}.csv'
    return resp


# ── PROFILE ───────────────────────────────────────────────────────────────────

@app.route('/admin/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        current_user.first_name  = request.form.get('first_name','').strip() or None
        current_user.last_name   = request.form.get('last_name','').strip() or None
        current_user.email       = request.form.get('email','').strip() or None
        current_user.phone       = request.form.get('phone','').strip() or None
        current_user.workplace   = request.form.get('workplace','').strip() or None
        current_user.job_title   = request.form.get('job_title','').strip() or None
        current_user.use_email_auth = 'use_email_auth' in request.form
        current_user.receive_notifications = 'receive_notifications' in request.form
        pwd = request.form.get('new_password','').strip()
        if pwd:
            password_error = _password_error(pwd)
            if password_error:
                flash(password_error, 'error')
                return redirect(url_for('profile'))
            if not current_user.check_password(request.form.get('current_password','')):
                flash('Current password is incorrect.', 'error')
                return redirect(url_for('profile'))
            current_user.set_password(pwd)
            current_user.session_version = int(current_user.session_version or 1) + 1
            session['session_version'] = current_user.session_version
        db.session.commit()
        flash('Profile updated.', 'success')
        return redirect(url_for('profile'))
    return render_template('admin/profile.html')


# ── MAIN ─────────────────────────────────────────────────────────────────────

with app.app_context():
    init_db()

if __name__ == '__main__':
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(debug=debug, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
